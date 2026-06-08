from fastapi import FastAPI, APIRouter, HTTPException, Query, Depends, File, UploadFile
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials, HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from datetime import datetime, timedelta, date
from decimal import Decimal
import os
import uuid
import json
import io
import csv
import secrets
import pytz
import logging
import hashlib
import base64
import math
import re
import smtplib
import requests
from email.message import EmailMessage

try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
except Exception:
    service_account = None
    build = None
    MediaIoBaseDownload = None
    MediaIoBaseUpload = None

try:
    from google_drive_service import get_google_drive_service, GoogleDriveService
except Exception as drive_import_error:
    logger = logging.getLogger(__name__)
    logger.warning(f"Google Drive service import failed: {drive_import_error}")
    get_google_drive_service = None
    GoogleDriveService = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
MONGO_URL = os.environ.get('MONGO_URL')
DB_NAME = os.environ.get('DB_NAME', 'lda_timetracking')
PORT = int(os.environ.get('PORT', 8001))

# Google Drive job folder automation settings
GOOGLE_DRIVE_PARENT_FOLDER_ID = os.environ.get('GOOGLE_DRIVE_PARENT_FOLDER_ID', '').strip()
GOOGLE_DRIVE_TEMPLATE_FOLDER_ID = os.environ.get('GOOGLE_DRIVE_TEMPLATE_FOLDER_ID', '').strip()
GOOGLE_SERVICE_ACCOUNT_JSON = os.environ.get('GOOGLE_SERVICE_ACCOUNT_JSON', '').strip()
GOOGLE_DRIVE_START_JOB_NUMBER = int(os.environ.get('GOOGLE_DRIVE_START_JOB_NUMBER', '34'))
GOOGLE_DRIVE_JOB_NUMBER_PADDING = int(os.environ.get('GOOGLE_DRIVE_JOB_NUMBER_PADDING', '3'))

# UK timezone handling
UK_TZ = pytz.timezone('Europe/London')

def get_uk_time():
    """Get current UK time (handles BST/GMT automatically)"""
    return datetime.now(UK_TZ)

def utc_to_uk(utc_dt):
    """Convert UTC datetime to UK time"""
    if utc_dt is None:
        return None
    if utc_dt.tzinfo is None:
        utc_dt = pytz.utc.localize(utc_dt)
    return utc_dt.astimezone(UK_TZ)

def uk_to_utc(uk_dt):
    """Convert UK time to UTC"""
    if uk_dt is None:
        return None
    if uk_dt.tzinfo is None:
        uk_dt = UK_TZ.localize(uk_dt)
    return uk_dt.astimezone(pytz.utc)




def format_uk_date_only(value):
    """Format date-only values as DD-MM-YYYY for PO PDFs/emails/webhook payloads."""
    if not value:
        return "-"

    try:
        if isinstance(value, datetime):
            return value.strftime("%d-%m-%Y")

        value_str = str(value).strip()
        if not value_str:
            return "-"

        # Handles ISO date strings like 2026-05-21 or 2026-05-21T00:00:00
        if len(value_str) >= 10 and value_str[4] == "-" and value_str[7] == "-":
            parsed = datetime.fromisoformat(value_str[:10])
            return parsed.strftime("%d-%m-%Y")

        return value_str
    except Exception:
        return str(value)

def format_uk_datetime_for_export(value):
    """Format datetimes for CSV exports in UK local time (handles BST/GMT)."""
    if not value:
        return ""

    if isinstance(value, str):
        try:
            value = datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            return value

    if value.tzinfo is None:
        value = pytz.utc.localize(value)

    return value.astimezone(UK_TZ).strftime("%d/%m/%Y %H:%M")

# MongoDB connection with retry logic
async def get_database():
    """Get database connection with retry logic"""
    try:
        client = AsyncIOMotorClient(
            MONGO_URL,
            serverSelectionTimeoutMS=5000,
            connectTimeoutMS=10000,
            socketTimeoutMS=10000,
            maxPoolSize=10,
            retryWrites=True,
            retryReads=True,
            tlsInsecure=True  # Keep this for MongoDB Atlas compatibility
        )

        # Test connection
        await client.admin.command('ping')
        logger.info("MongoDB connection successful")

        return client[DB_NAME]
    except Exception as e:
        logger.error(f"MongoDB connection failed: {e}")
        raise HTTPException(status_code=503, detail="Database connection failed")

# Global database instance
db = None

# Security
security = HTTPBasic()
bearer_scheme = HTTPBearer()

# Admin credentials (in production, store securely)
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "ldagroup2024"

# Create the main app
app = FastAPI(title="LDA Group Time Tracking API - Production")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Health check endpoint
@app.get("/ping")
@app.post("/ping")
@app.head("/ping")
async def health_check():
    """Health check endpoint for monitoring"""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")
# Define Models
class Worker(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    email: str
    phone: str
    role: str = "worker"  # worker, admin, supervisor
    worker_type: str = "worker"  # worker or contractor
    division: str = ""  # LDA FM, LDA Building Services, LDA Construction, etc.
    trades: List[str] = []  # Multiple trades: Roofer, Plasterer, Plumber, Builder, etc.
    hourly_rate: float = 15.0  # Default £15/hour
    password: Optional[str] = None  # For admin users
    app_role: Optional[str] = None  # super_admin, admin, project_manager, accounts
    active: bool = True
    archived: bool = False
    gps_exempt: bool = False
    created_date: datetime = Field(default_factory=datetime.utcnow)

class WorkerCreate(BaseModel):
    name: str
    email: str
    phone: str
    role: str = "worker"
    worker_type: str = "worker"
    division: str = ""
    trades: List[str] = []
    hourly_rate: float = 15.0
    gps_exempt: bool = False
    password: Optional[str] = None
    app_role: Optional[str] = None

class WorkerUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    role: Optional[str] = None
    worker_type: Optional[str] = None
    division: Optional[str] = None
    trades: Optional[List[str]] = None
    hourly_rate: Optional[float] = None
    password: Optional[str] = None
    app_role: Optional[str] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None
    gps_exempt: Optional[bool] = None

class Job(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: str
    location: str
    client: str
    division: str = ""
    manager_id: str = ""
    manager_name: str = ""
    supervisor_id: str = ""
    supervisor_name: str = ""
    quoted_cost: float
    original_quoted_cost: Optional[float] = None
    approved_variations_total: float = 0.0
    pending_variations_total: float = 0.0
    rejected_variations_total: float = 0.0
    current_contract_value: Optional[float] = None
    status: str = "active"  # active, completed, cancelled
    archived: bool = False
    gps_required: bool = False  # If true, location is mandatory for clock in/out
    created_date: datetime = Field(default_factory=datetime.utcnow)
    job_number: Optional[int] = None
    display_name: Optional[str] = None
    include_in_gantt: bool = False
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    gantt_sections: List[Dict[str, Any]] = []
    drive_folder_id: Optional[str] = None
    drive_folder_link: Optional[str] = None
    drive_folder_url: Optional[str] = None
    google_drive_link: Optional[str] = None
    drive_folder_status: Optional[str] = None
    drive_folder_error: Optional[str] = None
    drive_folder_copy_stats: Optional[Dict[str, Any]] = None
    post_work_photos: List[Dict[str, Any]] = []
    commercial_markers: List[Dict[str, Any]] = []
    client_id: str = ""
    client_name: str = ""
    client_tax_snapshot: Dict[str, Any] = {}
    client_commercial_snapshot: Dict[str, Any] = {}
    payment_terms_days: int = 30
    retention_percent: float = 0.0
    # Finance / tax treatment. Values are intentionally simple strings so older records remain compatible.
    vat_treatment: str = "standard_20"  # standard_20 / reduced_5 / zero_rated / exempt / no_vat / drc
    vat_rate: float = 20.0
    drc_enabled: bool = False
    cis_enabled: bool = False
    cis_rate: float = 0.0  # 0 / 20 / 30 normally
    cis_deduction_basis: str = "labour_only"  # labour_only / full_net / none
    tax_notes: str = ""

class JobCreate(BaseModel):
    name: str
    description: str
    location: str
    client: str
    division: str = ""
    manager_id: str = ""
    manager_name: str = ""
    supervisor_id: str = ""
    supervisor_name: str = ""
    quoted_cost: float
    original_quoted_cost: Optional[float] = None
    approved_variations_total: float = 0.0
    pending_variations_total: float = 0.0
    rejected_variations_total: float = 0.0
    current_contract_value: Optional[float] = None
    gps_required: bool = True
    include_in_gantt: bool = False
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    gantt_sections: List[Dict[str, Any]] = []
    commercial_markers: List[Dict[str, Any]] = []
    drive_folder_id: Optional[str] = None
    drive_folder_link: Optional[str] = None
    drive_folder_url: Optional[str] = None
    google_drive_link: Optional[str] = None
    drive_folder_status: Optional[str] = None
    client_id: str = ""
    client_name: str = ""
    client_tax_snapshot: Dict[str, Any] = {}
    client_commercial_snapshot: Dict[str, Any] = {}
    payment_terms_days: int = 30
    retention_percent: float = 0.0
    vat_treatment: str = "standard_20"
    vat_rate: float = 20.0
    drc_enabled: bool = False
    cis_enabled: bool = False
    cis_rate: float = 0.0
    cis_deduction_basis: str = "labour_only"
    tax_notes: str = ""

class JobUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    location: Optional[str] = None
    client: Optional[str] = None
    division: Optional[str] = None
    manager_id: Optional[str] = None
    manager_name: Optional[str] = None
    supervisor_id: Optional[str] = None
    supervisor_name: Optional[str] = None
    quoted_cost: Optional[float] = None
    original_quoted_cost: Optional[float] = None
    approved_variations_total: Optional[float] = None
    pending_variations_total: Optional[float] = None
    rejected_variations_total: Optional[float] = None
    current_contract_value: Optional[float] = None
    status: Optional[str] = None
    archived: Optional[bool] = None
    gps_required: Optional[bool] = None
    include_in_gantt: Optional[bool] = None
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    gantt_sections: Optional[List[Dict[str, Any]]] = None
    drive_folder_id: Optional[str] = None
    drive_folder_link: Optional[str] = None
    drive_folder_url: Optional[str] = None
    google_drive_link: Optional[str] = None
    drive_folder_status: Optional[str] = None
    drive_folder_error: Optional[str] = None
    commercial_markers: Optional[List[Dict[str, Any]]] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    client_tax_snapshot: Optional[Dict[str, Any]] = None
    client_commercial_snapshot: Optional[Dict[str, Any]] = None
    payment_terms_days: Optional[int] = None
    retention_percent: Optional[float] = None
    vat_treatment: Optional[str] = None
    vat_rate: Optional[float] = None
    drc_enabled: Optional[bool] = None
    cis_enabled: Optional[bool] = None
    cis_rate: Optional[float] = None
    cis_deduction_basis: Optional[str] = None
    tax_notes: Optional[str] = None


# ==================== CLIENT / CUSTOMER SYSTEM MODELS ====================

class ClientContact(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str = ""
    role: str = ""
    email: str = ""
    phone: str = ""
    receives_invoices: bool = False
    receives_variations: bool = False
    receives_quotes: bool = False
    receives_project_updates: bool = False

class Client(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    client_type: str = "commercial"  # commercial / domestic / local_authority / insurance / contractor / other
    main_contact_name: str = ""
    phone: str = ""
    email: str = ""
    accounts_email: str = ""
    address: str = ""
    billing_address: str = ""
    vat_number: str = ""
    company_number: str = ""
    default_variation_email: str = ""
    default_invoice_email: str = ""
    default_quote_email: str = ""
    payment_terms_days: int = 30
    retention_percent: float = 0.0
    vat_treatment: str = "standard_20"
    vat_rate: float = 20.0
    drc_enabled: bool = False
    cis_enabled: bool = False
    cis_rate: float = 0.0
    cis_deduction_basis: str = "labour_only"
    notes: str = ""
    contacts: List[ClientContact] = []
    active: bool = True
    archived: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None

class ClientCreate(BaseModel):
    client_name: str
    client_type: str = "commercial"
    main_contact_name: str = ""
    phone: str = ""
    email: str = ""
    accounts_email: str = ""
    address: str = ""
    billing_address: str = ""
    vat_number: str = ""
    company_number: str = ""
    default_variation_email: str = ""
    default_invoice_email: str = ""
    default_quote_email: str = ""
    payment_terms_days: int = 30
    retention_percent: float = 0.0
    vat_treatment: str = "standard_20"
    vat_rate: float = 20.0
    drc_enabled: bool = False
    cis_enabled: bool = False
    cis_rate: float = 0.0
    cis_deduction_basis: str = "labour_only"
    notes: str = ""
    contacts: List[ClientContact] = []
    active: bool = True

class ClientUpdate(BaseModel):
    client_name: Optional[str] = None
    client_type: Optional[str] = None
    main_contact_name: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    accounts_email: Optional[str] = None
    address: Optional[str] = None
    billing_address: Optional[str] = None
    vat_number: Optional[str] = None
    company_number: Optional[str] = None
    default_variation_email: Optional[str] = None
    default_invoice_email: Optional[str] = None
    default_quote_email: Optional[str] = None
    payment_terms_days: Optional[int] = None
    retention_percent: Optional[float] = None
    vat_treatment: Optional[str] = None
    vat_rate: Optional[float] = None
    drc_enabled: Optional[bool] = None
    cis_enabled: Optional[bool] = None
    cis_rate: Optional[float] = None
    cis_deduction_basis: Optional[str] = None
    notes: Optional[str] = None
    contacts: Optional[List[ClientContact]] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None

def normalise_client_tax_payload(values: Dict[str, Any]) -> Dict[str, Any]:
    treatment = str(values.get("vat_treatment") or "standard_20").strip().lower()
    if treatment not in {"standard_20", "reduced_5", "zero_rated", "exempt", "no_vat", "drc"}:
        treatment = "standard_20"
    drc_enabled = treatment == "drc" or values.get("drc_enabled") is True
    if drc_enabled:
        treatment = "drc"
    rate = finance_to_number(values.get("vat_rate"), 20.0)
    if treatment == "reduced_5":
        rate = 5.0
    elif treatment in {"zero_rated", "exempt", "no_vat", "drc"}:
        rate = 0.0
    cis_enabled = bool(values.get("cis_enabled"))
    cis_rate = finance_to_number(values.get("cis_rate"), 20.0 if cis_enabled else 0.0) if cis_enabled else 0.0
    return {
        "vat_treatment": treatment,
        "vat_rate": max(0.0, rate),
        "drc_enabled": drc_enabled,
        "cis_enabled": cis_enabled,
        "cis_rate": max(0.0, min(100.0, cis_rate)),
        "cis_deduction_basis": values.get("cis_deduction_basis") or "labour_only",
    }

def client_job_defaults(client: Dict[str, Any]) -> Dict[str, Any]:
    tax = normalise_client_tax_payload(client or {})
    commercial = {
        "payment_terms_days": int(finance_to_number((client or {}).get("payment_terms_days"), 30)),
        "retention_percent": finance_to_number((client or {}).get("retention_percent"), 0.0),
    }
    return {
        "client_id": (client or {}).get("id", ""),
        "client_name": (client or {}).get("client_name", ""),
        "client_tax_snapshot": tax,
        "client_commercial_snapshot": commercial,
        **tax,
        **commercial,
    }

async def apply_client_defaults_to_job_data(job_data: Dict[str, Any], force_refresh: bool = False) -> Dict[str, Any]:
    client_id = str(job_data.get("client_id") or "").strip()
    if not client_id:
        return job_data
    client = await db.clients.find_one({"id": client_id, "archived": {"$ne": True}}, {"_id": 0})
    if not client:
        return job_data
    defaults = client_job_defaults(client)
    job_data["client_id"] = defaults["client_id"]
    job_data["client_name"] = defaults["client_name"]
    if not job_data.get("client"):
        job_data["client"] = defaults["client_name"]
    job_data["client_tax_snapshot"] = defaults["client_tax_snapshot"]
    job_data["client_commercial_snapshot"] = defaults["client_commercial_snapshot"]
    defaultable_keys = ["vat_treatment", "vat_rate", "drc_enabled", "cis_enabled", "cis_rate", "cis_deduction_basis", "payment_terms_days", "retention_percent"]
    for key in defaultable_keys:
        if force_refresh or job_data.get(key) in [None, ""]:
            job_data[key] = defaults.get(key)
    return job_data


# ==================== SUBCONTRACTOR SYSTEM MODELS ====================

class SubcontractorCompany(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    contact_name: str = ""
    email: str = ""
    phone: str = ""
    trades: List[str] = []
    cis_registered: bool = False
    utr_number: str = ""
    insurance_expiry: Optional[str] = None
    notes: str = ""
    active: bool = True
    archived: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None

class SubcontractorCompanyCreate(BaseModel):
    company_name: str
    contact_name: str = ""
    email: str = ""
    phone: str = ""
    trades: List[str] = []
    cis_registered: bool = False
    utr_number: str = ""
    insurance_expiry: Optional[str] = None
    notes: str = ""
    active: bool = True

class SubcontractorCompanyUpdate(BaseModel):
    company_name: Optional[str] = None
    contact_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    trades: Optional[List[str]] = None
    cis_registered: Optional[bool] = None
    utr_number: Optional[str] = None
    insurance_expiry: Optional[str] = None
    notes: Optional[str] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None

class SubcontractorResource(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    subcontractor_id: str
    name: str
    trade: str = ""
    capacity: int = 1
    phone: str = ""
    email: str = ""
    notes: str = ""
    active: bool = True
    archived: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None

class SubcontractorResourceCreate(BaseModel):
    subcontractor_id: str = ""
    name: str
    trade: str = ""
    capacity: int = 1
    phone: str = ""
    email: str = ""
    notes: str = ""
    active: bool = True

class SubcontractorResourceUpdate(BaseModel):
    subcontractor_id: Optional[str] = None
    name: Optional[str] = None
    trade: Optional[str] = None
    capacity: Optional[int] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    notes: Optional[str] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None

class GPSLocation(BaseModel):
    latitude: float
    longitude: float
    accuracy: Optional[float] = None
    address: Optional[str] = None  # Human readable address

class TimeEntry(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    worker_id: str
    job_id: str
    clock_in: datetime
    clock_out: Optional[datetime] = None
    duration_minutes: Optional[int] = None
    gps_location_in: Optional[GPSLocation] = None
    gps_location_out: Optional[GPSLocation] = None
    device_id_in: Optional[str] = None
    device_id_out: Optional[str] = None
    suspicious_flags: List[str] = []
    notes: str = ""
    created_date: datetime = Field(default_factory=datetime.utcnow)
    worker_name: Optional[str] = None
    job_name: Optional[str] = None
    job_client: Optional[str] = None
    cost: Optional[float] = 0

class TimeEntryClockIn(BaseModel):
    worker_id: str
    job_id: str
    gps_location: Optional[GPSLocation] = None
    device_id: Optional[str] = None
    notes: str = ""

class TimeEntryClockOut(BaseModel):
    gps_location: Optional[GPSLocation] = None
    device_id: Optional[str] = None
    notes: str = ""

class TimeEntryUpdate(BaseModel):
    worker_id: Optional[str] = None
    job_id: Optional[str] = None
    clock_in: Optional[str] = None
    clock_out: Optional[str] = None
    duration_minutes: Optional[int] = None
    notes: Optional[str] = None

class Material(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    job_id: str
    name: str
    cost: float
    quantity: int
    supplier: str = ""  # Supplier name
    reference: str = ""  # Receipt number or reference
    purchase_date: datetime = Field(default_factory=datetime.utcnow)
    notes: str = ""
    created_date: datetime = Field(default_factory=datetime.utcnow)

class MaterialCreate(BaseModel):
    job_id: str
    name: str
    cost: float
    quantity: int
    supplier: str = ""
    reference: str = ""
    notes: str = ""

class MaterialUpdate(BaseModel):
    name: Optional[str] = None
    cost: Optional[float] = None
    quantity: Optional[int] = None
    supplier: Optional[str] = None
    reference: Optional[str] = None
    notes: Optional[str] = None

class ScheduleEntry(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    worker_id: Optional[str] = ""
    resource_type: str = "worker"  # worker / subcontractor_resource
    resource_id: Optional[str] = None
    display_worker_name: str = ""
    job_id: Optional[str] = None
    scheduled_date: str  # YYYY-MM-DD
    notes: str = ""
    status: str = "scheduled"
    schedule_type: str = "job"  # job, holiday, sick, unavailable
    absence_type: Optional[str] = None
    project_section_id: str = ""
    project_section_name: str = ""
    linked_gantt_section_id: str = ""
    schedule_link_mode: str = "manual"  # manual / linked_to_section / manually_adjusted / detached
    hours: float = 8.0
    created_date: datetime = Field(default_factory=datetime.utcnow)
    updated_date: Optional[datetime] = None

class ScheduleEntryCreate(BaseModel):
    worker_id: Optional[str] = ""
    resource_type: str = "worker"
    resource_id: Optional[str] = None
    display_worker_name: str = ""
    job_id: Optional[str] = None
    scheduled_date: str  # YYYY-MM-DD
    notes: str = ""
    status: str = "scheduled"
    schedule_type: str = "job"
    absence_type: Optional[str] = None
    project_section_id: str = ""
    project_section_name: str = ""
    linked_gantt_section_id: str = ""
    schedule_link_mode: str = "manual"
    hours: float = 8.0

class ScheduleEntryUpdate(BaseModel):
    worker_id: Optional[str] = None
    resource_type: Optional[str] = None
    resource_id: Optional[str] = None
    display_worker_name: Optional[str] = None
    job_id: Optional[str] = None
    scheduled_date: Optional[str] = None
    notes: Optional[str] = None
    status: Optional[str] = None
    schedule_type: Optional[str] = None
    absence_type: Optional[str] = None
    project_section_id: Optional[str] = None
    project_section_name: Optional[str] = None
    linked_gantt_section_id: Optional[str] = None
    schedule_link_mode: Optional[str] = None
    hours: Optional[float] = None

class GanttPushToScheduleRequest(BaseModel):
    job_id: str
    section_id: str
    worker_ids: List[str]
    start_date: str
    end_date: str
    section_name: str = ""
    replace_existing_for_section: bool = True
    override_existing_allocations: bool = False

class GanttShiftProjectRequest(BaseModel):
    job_id: str
    planned_start_date: str
    planned_end_date: str
    delta_days: int
    shift_sections: bool = True
    shift_schedule: bool = True
    shift_commercial_markers: bool = True
    shift_purchase_orders: bool = True
    shift_work_orders: bool = True

class GanttShiftSectionRequest(BaseModel):
    job_id: str
    section_id: str
    start_date: str
    end_date: str
    delta_days: int
    shift_schedule: bool = True
    shift_purchase_orders: bool = True
    shift_work_orders: bool = True


class FinanceRecord(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    job_id: str
    application_marker_id: Optional[str] = None
    type: str = "application"  # application, invoice, payment, retention, adjustment
    label: str = ""
    submitted_date: Optional[str] = None
    submitted_value: float = 0.0
    certified_value: float = 0.0
    invoice_number: str = ""
    invoice_date: Optional[str] = None
    invoice_value: float = 0.0
    payment_due_date: Optional[str] = None
    paid_date: Optional[str] = None
    paid_value: float = 0.0
    retention_percent: float = 0.0
    retention_value: float = 0.0
    retention_due_date: Optional[str] = None
    retention_paid_date: Optional[str] = None
    status: str = "draft"  # draft, submitted, certified, invoiced, part_paid, paid, overdue, disputed
    notes: str = ""
    created_date: datetime = Field(default_factory=datetime.utcnow)
    updated_date: Optional[datetime] = None
    archived: bool = False

class FinanceRecordCreate(BaseModel):
    job_id: str
    application_marker_id: Optional[str] = None
    type: str = "application"
    label: str = ""
    submitted_date: Optional[str] = None
    submitted_value: float = 0.0
    certified_value: float = 0.0
    invoice_number: str = ""
    invoice_date: Optional[str] = None
    invoice_value: float = 0.0
    payment_due_date: Optional[str] = None
    paid_date: Optional[str] = None
    paid_value: float = 0.0
    retention_percent: float = 0.0
    retention_value: float = 0.0
    retention_due_date: Optional[str] = None
    retention_paid_date: Optional[str] = None
    status: str = "draft"
    notes: str = ""

class FinanceRecordUpdate(BaseModel):
    application_marker_id: Optional[str] = None
    type: Optional[str] = None
    label: Optional[str] = None
    submitted_date: Optional[str] = None
    submitted_value: Optional[float] = None
    certified_value: Optional[float] = None
    invoice_number: Optional[str] = None
    invoice_date: Optional[str] = None
    invoice_value: Optional[float] = None
    payment_due_date: Optional[str] = None
    paid_date: Optional[str] = None
    paid_value: Optional[float] = None
    retention_percent: Optional[float] = None
    retention_value: Optional[float] = None
    retention_due_date: Optional[str] = None
    retention_paid_date: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    archived: Optional[bool] = None



# ==================== CONTRACTOR WORK ORDER SYSTEM MODELS ====================

WORK_ORDER_STATUSES = [
    "draft",
    "price_requested",
    "price_received",
    "accepted",
    "instructed",
    "in_progress",
    "complete",
    "awaiting_sign_off",
    "invoice_received",
    "part_paid",
    "retention_held",
    "paid",
    "closed",
    "cancelled",
]

WORK_ORDER_COMMITTED_STATUSES = {
    "accepted",
    "instructed",
    "in_progress",
    "complete",
    "awaiting_sign_off",
    "invoice_received",
    "part_paid",
    "retention_held",
    "paid",
    "closed",
}

WORK_ORDER_CANCELLED_STATUSES = {"cancelled", "canceled", "void", "archived", "rejected"}

class WorkOrder(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    wo_number: str = ""
    job_id: str
    job_name: str = ""
    job_number: Optional[int] = None
    division: str = ""
    section_id: str = ""
    section_name: str = ""
    contractor_id: str = ""
    contractor_name: str = ""
    contractor_email: str = ""
    trade: str = ""
    description: str = ""
    pricing_type: str = "fixed_price"  # fixed_price / day_rate / hourly / item_rate
    quantity: float = 1.0
    rate: float = 0.0
    net_amount: float = 0.0
    vat_rate: float = 20.0
    vat_amount: float = 0.0
    gross_amount: float = 0.0
    cis_applicable: bool = False
    cis_rate: float = 20.0
    cis_deduction: float = 0.0
    payment_requirement: str = "credit_terms"  # credit_terms / proforma / immediate
    payment_terms_days: int = 30
    expected_start_date: Optional[str] = None
    expected_completion_date: Optional[str] = None
    payment_due_date: Optional[str] = None
    payment_schedule_mode: str = "single"  # single / weekly / monthly / manual / milestone
    retention_percent: float = 0.0
    retention_release_rule: str = "manual"  # manual / supervisor_sign_off / defects_period / client_payment_received
    retention_release_date: Optional[str] = None
    payment_schedule: List[Dict[str, Any]] = []
    linked_application_marker_id: str = ""
    status: str = "price_received"
    notes: str = ""
    requested_by_user_id: str = ""
    requested_by_name: str = ""
    requested_by_email: str = ""
    requested_by_role: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None
    archived: bool = False

class WorkOrderCreate(BaseModel):
    wo_number: str = ""
    job_id: str
    job_name: str = ""
    job_number: Optional[int] = None
    division: str = ""
    section_id: str = ""
    section_name: str = ""
    contractor_id: str = ""
    contractor_name: str = ""
    contractor_email: str = ""
    trade: str = ""
    description: str = ""
    pricing_type: str = "fixed_price"
    quantity: float = 1.0
    rate: float = 0.0
    net_amount: float = 0.0
    vat_rate: float = 20.0
    vat_amount: float = 0.0
    gross_amount: float = 0.0
    cis_applicable: bool = False
    cis_rate: float = 20.0
    cis_deduction: float = 0.0
    payment_requirement: str = "credit_terms"
    payment_terms_days: int = 30
    expected_start_date: Optional[str] = None
    expected_completion_date: Optional[str] = None
    payment_due_date: Optional[str] = None
    payment_schedule_mode: str = "single"  # single / weekly / monthly / manual / milestone
    retention_percent: float = 0.0
    retention_release_rule: str = "manual"  # manual / supervisor_sign_off / defects_period / client_payment_received
    retention_release_date: Optional[str] = None
    payment_schedule: List[Dict[str, Any]] = []
    linked_application_marker_id: str = ""
    status: str = "price_received"
    notes: str = ""
    requested_by_user_id: str = ""
    requested_by_name: str = ""
    requested_by_email: str = ""
    requested_by_role: str = ""

class WorkOrderUpdate(BaseModel):
    wo_number: Optional[str] = None
    job_id: Optional[str] = None
    job_name: Optional[str] = None
    job_number: Optional[int] = None
    division: Optional[str] = None
    section_id: Optional[str] = None
    section_name: Optional[str] = None
    contractor_id: Optional[str] = None
    contractor_name: Optional[str] = None
    contractor_email: Optional[str] = None
    trade: Optional[str] = None
    description: Optional[str] = None
    pricing_type: Optional[str] = None
    quantity: Optional[float] = None
    rate: Optional[float] = None
    net_amount: Optional[float] = None
    vat_rate: Optional[float] = None
    vat_amount: Optional[float] = None
    gross_amount: Optional[float] = None
    cis_applicable: Optional[bool] = None
    cis_rate: Optional[float] = None
    cis_deduction: Optional[float] = None
    payment_requirement: Optional[str] = None
    payment_terms_days: Optional[int] = None
    expected_start_date: Optional[str] = None
    expected_completion_date: Optional[str] = None
    payment_due_date: Optional[str] = None
    payment_schedule_mode: Optional[str] = None
    retention_percent: Optional[float] = None
    retention_release_rule: Optional[str] = None
    retention_release_date: Optional[str] = None
    payment_schedule: Optional[List[Dict[str, Any]]] = None
    linked_application_marker_id: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    requested_by_user_id: Optional[str] = None
    requested_by_name: Optional[str] = None
    requested_by_email: Optional[str] = None
    requested_by_role: Optional[str] = None
    archived: Optional[bool] = None

# ==================== PURCHASE ORDER SYSTEM MODELS ====================

class Supplier(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    contact_name: str = ""
    orders_email: str = ""
    accounts_email: str = ""
    phone: str = ""
    address: str = ""
    vat_number: str = ""
    payment_terms: str = "30 days"
    payment_terms_days: int = 30
    payment_requirement: str = "credit_terms"  # credit_terms / proforma / immediate
    lead_time_days: int = 0
    notes: str = ""
    active: bool = True
    archived: bool = False
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None

class SupplierCreate(BaseModel):
    name: str
    contact_name: str = ""
    orders_email: str = ""
    accounts_email: str = ""
    phone: str = ""
    address: str = ""
    vat_number: str = ""
    payment_terms: str = "30 days"
    payment_terms_days: int = 30
    payment_requirement: str = "credit_terms"
    lead_time_days: int = 0
    notes: str = ""

class SupplierUpdate(BaseModel):
    name: Optional[str] = None
    contact_name: Optional[str] = None
    orders_email: Optional[str] = None
    accounts_email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    vat_number: Optional[str] = None
    payment_terms: Optional[str] = None
    payment_terms_days: Optional[int] = None
    payment_requirement: Optional[str] = None
    lead_time_days: Optional[int] = None
    notes: Optional[str] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None

class PurchaseOrderLine(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    description: str
    quantity: float = 1.0
    unit_cost: float = 0.0
    vat_rate: float = 20.0
    net_total: float = 0.0
    vat_total: float = 0.0
    gross_total: float = 0.0
    prices_include_vat: bool = False
    source_line_net_total: Optional[float] = None
    source_line_vat_total: Optional[float] = None
    source_line_gross_total: Optional[float] = None
    job_section_id: str = ""
    job_section_name: str = ""
    cost_category: str = "Materials"
    received_quantity: float = 0.0
    material_status: str = "committed"
    material_id: Optional[str] = None

class PurchaseOrder(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    po_number: str = ""
    supplier_id: str
    supplier_name: str = ""
    supplier_email: str = ""
    job_id: str
    job_name: str = ""
    job_number: Optional[int] = None
    division: str = ""
    status: str = "draft"
    requested_by_user_id: str = ""
    requested_by_name: str = ""
    requested_by_email: str = ""
    requested_by_role: str = ""
    approved_by_user_id: Optional[str] = None
    approved_by_name: Optional[str] = None
    approved_at: Optional[datetime] = None
    sent_at: Optional[datetime] = None
    sent_by_user_id: Optional[str] = None
    sent_by_name: Optional[str] = None
    email_subject: str = ""
    required_date: Optional[str] = None
    payment_requirement: str = "credit_terms"
    payment_terms_days: int = 30
    lead_time_days: int = 0
    payment_due_date: Optional[str] = None
    order_by_date: Optional[str] = None
    expected_payment_date: Optional[str] = None
    delivery_address: str = ""
    notes: str = ""
    supplier_quote_number: str = ""
    source_type: str = "manual"
    source_upload_id: Optional[str] = None
    source_file_name: str = ""
    extraction_status: str = "not_required"
    extraction_confidence: str = ""
    lines: List[PurchaseOrderLine] = []
    net_total: float = 0.0
    vat_total: float = 0.0
    gross_total: float = 0.0
    materials_assigned: bool = False
    materials_assigned_at: Optional[datetime] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None

class PurchaseOrderCreate(BaseModel):
    supplier_id: str
    supplier_name: str = ""
    supplier_email: str = ""
    job_id: str
    job_name: str = ""
    job_number: Optional[int] = None
    division: str = ""
    required_date: Optional[str] = None
    payment_requirement: str = "credit_terms"
    payment_terms_days: int = 30
    lead_time_days: int = 0
    payment_due_date: Optional[str] = None
    order_by_date: Optional[str] = None
    expected_payment_date: Optional[str] = None
    delivery_address: str = ""
    notes: str = ""
    supplier_quote_number: str = ""
    source_type: str = "manual"
    source_upload_id: Optional[str] = None
    source_file_name: str = ""
    extraction_status: str = "not_required"
    extraction_confidence: str = ""
    lines: List[PurchaseOrderLine] = []
    net_total: float = 0.0
    vat_total: float = 0.0
    gross_total: float = 0.0
    requested_by_user_id: str = ""
    requested_by_name: str = ""
    requested_by_email: str = ""
    requested_by_role: str = ""

class PurchaseOrderUpdate(BaseModel):
    supplier_id: Optional[str] = None
    supplier_name: Optional[str] = None
    supplier_email: Optional[str] = None
    job_id: Optional[str] = None
    job_name: Optional[str] = None
    job_number: Optional[int] = None
    division: Optional[str] = None
    status: Optional[str] = None
    required_date: Optional[str] = None
    payment_requirement: Optional[str] = None
    payment_terms_days: Optional[int] = None
    lead_time_days: Optional[int] = None
    payment_due_date: Optional[str] = None
    order_by_date: Optional[str] = None
    expected_payment_date: Optional[str] = None
    delivery_address: Optional[str] = None
    notes: Optional[str] = None
    supplier_quote_number: Optional[str] = None
    source_type: Optional[str] = None
    source_upload_id: Optional[str] = None
    source_file_name: Optional[str] = None
    extraction_status: Optional[str] = None
    extraction_confidence: Optional[str] = None
    lines: Optional[List[PurchaseOrderLine]] = None
    net_total: Optional[float] = None
    vat_total: Optional[float] = None
    gross_total: Optional[float] = None
    requested_by_user_id: Optional[str] = None
    requested_by_name: Optional[str] = None
    requested_by_email: Optional[str] = None
    requested_by_role: Optional[str] = None

class PurchaseOrderBulkDeleteRequest(BaseModel):
    po_ids: List[str]

class PurchaseOrderApprovalResponseRequest(BaseModel):
    secret: str = ""
    decision: str = ""  # Approve / Reject / Approved / Rejected
    selected_option: str = ""  # Power Automate Send email with options response
    responder_name: str = ""
    responder_email: str = ""
    comments: str = ""
    reason: str = ""

class PurchaseOrderRequesterActionResponseRequest(BaseModel):
    secret: str = ""
    decision: str = ""  # Send email to supplier and assign materials / Assign materials to job only
    selected_option: str = ""
    responder_name: str = ""
    responder_email: str = ""
    comments: str = ""


# ==================== VARIATION SYSTEM MODELS ====================

class Variation(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    variation_number: str = ""
    job_id: str
    job_name: str = ""
    job_number: Optional[int] = None
    client_name: str = ""
    client_email: str = ""
    title: str
    client_instruction_summary: str = ""
    description: str = ""
    scope_of_works: str = ""
    material_value: float = 0.0
    labour_value: float = 0.0
    subcontractor_value: float = 0.0
    other_value: float = 0.0
    net_total: float = 0.0
    vat_rate: float = 20.0
    vat_total: float = 0.0
    gross_total: float = 0.0
    required_date: Optional[str] = None
    status: str = "pending_client_approval"  # draft, pending_client_approval, approved, rejected, cancelled
    approval_token: str = ""
    approval_link: str = ""
    approval_sent_at: Optional[datetime] = None
    approval_notification_status: str = "not_sent"
    approved_by_name: str = ""
    approved_by_email: str = ""
    approved_at: Optional[datetime] = None
    rejected_by_name: str = ""
    rejected_by_email: str = ""
    rejected_at: Optional[datetime] = None
    rejection_reason: str = ""
    client_signature: str = ""
    client_comments: str = ""
    add_to_gantt_on_approval: bool = True
    standalone_variation_invoice: bool = False  # false means included in contract/applications, not separately counted as cash-in
    added_to_gantt: bool = False
    gantt_section_id: str = ""
    created_by_user_id: str = ""
    created_by_name: str = ""
    created_by_email: str = ""
    created_by_role: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None
    archived: bool = False

class VariationCreate(BaseModel):
    job_id: str
    title: str
    client_email: str = ""
    client_instruction_summary: str = ""
    description: str = ""
    scope_of_works: str = ""
    material_value: float = 0.0
    labour_value: float = 0.0
    subcontractor_value: float = 0.0
    other_value: float = 0.0
    vat_rate: float = 20.0
    required_date: Optional[str] = None
    status: str = "pending_client_approval"
    add_to_gantt_on_approval: bool = True
    standalone_variation_invoice: bool = False
    created_by_user_id: str = ""
    created_by_name: str = ""
    created_by_email: str = ""
    created_by_role: str = ""

class VariationUpdate(BaseModel):
    title: Optional[str] = None
    client_email: Optional[str] = None
    client_instruction_summary: Optional[str] = None
    description: Optional[str] = None
    scope_of_works: Optional[str] = None
    material_value: Optional[float] = None
    labour_value: Optional[float] = None
    subcontractor_value: Optional[float] = None
    other_value: Optional[float] = None
    vat_rate: Optional[float] = None
    required_date: Optional[str] = None
    status: Optional[str] = None
    add_to_gantt_on_approval: Optional[bool] = None
    standalone_variation_invoice: Optional[bool] = None
    archived: Optional[bool] = None

class VariationApprovalResponse(BaseModel):
    decision: str = "approve"  # approve / reject
    name: str = ""
    email: str = ""
    signature: str = ""
    comments: str = ""
    reason: str = ""

class AdminLogin(BaseModel):
    username: str
    password: str

class WorkerLogin(BaseModel):
    worker_id: str
    password: Optional[str] = None  # For admin workers

# Quote System Models
class QuoteMaterial(BaseModel):
    name: str
    description: str = ""
    quantity: int
    unit_price: Decimal
    total_price: Decimal

class QuoteLabor(BaseModel):
    description: str
    estimated_hours: float
    hourly_rate: Decimal
    total_cost: Decimal

class Quote(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    quote_number: str
    surveyor_id: str
    surveyor_name: str
    client_name: str
    client_email: str
    client_phone: str = ""
    client_address: str = ""
    job_description: str
    materials: List[QuoteMaterial] = []
    labor: List[QuoteLabor] = []
    total_materials_cost: Decimal = Decimal('0.00')
    total_labor_cost: Decimal = Decimal('0.00')
    total_quote_amount: Decimal = Decimal('0.00')
    vat_rate: Decimal = Decimal('0.20')  # 20% VAT
    vat_amount: Decimal = Decimal('0.00')
    final_amount: Decimal = Decimal('0.00')
    status: str = "draft"  # draft, sent, accepted, declined, expired
    valid_until: date = Field(default_factory=lambda: (datetime.utcnow() + timedelta(days=30)).date())
    photos: List[Dict[str, Any]] = []
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: Optional[datetime] = None
    sent_at: Optional[datetime] = None
    client_response: Optional[str] = None  # accepted, declined
    client_response_date: Optional[datetime] = None
    client_message: str = ""

class QuoteCreate(BaseModel):
    client_name: str
    client_email: str
    client_phone: str = ""
    client_address: str = ""
    job_description: str
    materials: List[QuoteMaterial] = []
    labor: List[QuoteLabor] = []
    vat_rate: Decimal = Decimal('0.20')

class QuoteUpdate(BaseModel):
    client_name: Optional[str] = None
    client_email: Optional[str] = None
    client_phone: Optional[str] = None
    client_address: Optional[str] = None
    job_description: Optional[str] = None
    materials: Optional[List[QuoteMaterial]] = None
    labor: Optional[List[QuoteLabor]] = None
    vat_rate: Optional[Decimal] = None
    status: Optional[str] = None

class SurveyorCreate(BaseModel):
    name: str
    email: str
    phone: str
    password: str

class SurveyorLogin(BaseModel):
    email: str
    password: str

class ClientResponse(BaseModel):
    response: str  # "accepted" or "declined"
    message: str = ""

# Helper functions for quote services (placeholders)
EmailService = None
QuotePDFGenerator = None

def convert_decimals_to_float(obj):
    """Convert Decimal objects to float and date objects to datetime for MongoDB serialization"""
    if isinstance(obj, dict):
        return {key: convert_decimals_to_float(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_decimals_to_float(item) for item in obj]
    elif isinstance(obj, Decimal):
        return float(obj)
    elif isinstance(obj, date) and not isinstance(obj, datetime):
        # Convert date to datetime for MongoDB compatibility
        return datetime.combine(obj, datetime.min.time())
    else:
        return obj

# Security functions
OFFICE_LOGIN_ROLES = ["admin", "super_admin", "project_manager", "accounts", "office_admin"]

# Users in this list always receive full Super Admin access when they log in successfully.
# This prevents the owner account from being locked out by role-based navigation.
SUPER_ADMIN_EMAILS = {
    "duke.mcintyre@ldagroup.co.uk",
    "dukemcintyre@ldagroup.co.uk",
}

def normalise_app_role(value: Optional[str]) -> str:
    role = str(value or "").strip().lower().replace(" ", "_").replace("-", "_")
    if role in ["super", "superadmin", "owner"]:
        return "super_admin"
    if role in ["pm", "project", "project_manager", "project_managers"]:
        return "project_manager"
    if role in ["account", "accounts", "finance"]:
        return "accounts"
    if role in ["office", "office_admin", "administrator", "admin"]:
        return "admin"
    return role or "admin"

def public_user_from_worker(worker: Dict[str, Any]) -> Dict[str, Any]:
    email = str(worker.get("email", "")).strip().lower()
    role = normalise_app_role(worker.get("app_role") or worker.get("role") or "admin")

    # Hard owner override: Duke must always have full app access.
    if email in SUPER_ADMIN_EMAILS:
        role = "super_admin"

    return {
        "id": worker.get("id", ""),
        "name": worker.get("name") or worker.get("email") or "Office User",
        "email": worker.get("email", ""),
        "role": role,
        "worker_role": worker.get("role", ""),
        "worker_type": worker.get("worker_type", ""),
        "division": worker.get("division", ""),
        "trades": worker.get("trades", []),
    }

def builtin_super_admin_user(username: str = ADMIN_USERNAME) -> Dict[str, Any]:
    username_value = str(username or "").strip()
    email = username_value if "@" in username_value else ""
    return {
        "id": "built_in_admin",
        "name": "Duke Mcintyre" if email.lower() in SUPER_ADMIN_EMAILS else "LDA Super Admin",
        "email": email,
        "role": "super_admin",
        "worker_role": "admin",
        "worker_type": "admin",
        "division": "Management",
        "trades": [],
    }

async def find_office_user(username: str, password: str) -> Optional[Dict[str, Any]]:
    username_clean = str(username or "").strip()
    username_lower = username_clean.lower()

    # Legacy built-in login remains Super Admin.
    if username_clean == ADMIN_USERNAME and password == ADMIN_PASSWORD:
        return builtin_super_admin_user(username_clean)

    # Duke owner fallback: allows Duke to regain full access with the legacy admin password
    # even if no database user has been upgraded to super_admin yet.
    if username_lower in SUPER_ADMIN_EMAILS and password == ADMIN_PASSWORD:
        return builtin_super_admin_user(username_clean)

    # Owner email override: if Duke has a worker/admin login record, let him in as Super Admin
    # even if the worker role is currently stored as admin/worker/project_manager/accounts.
    if username_lower in SUPER_ADMIN_EMAILS:
        worker = await db.workers.find_one({
            "email": {"$regex": f"^{re.escape(username_clean)}$", "$options": "i"},
            "password": password,
            "active": True,
            "archived": {"$ne": True},
        }, {"_id": 0})

        if worker:
            return public_user_from_worker(worker)

    # Standard office users are limited to approved office roles.
    worker = await db.workers.find_one({
        "email": {"$regex": f"^{re.escape(username_clean)}$", "$options": "i"},
        "password": password,
        "active": True,
        "archived": {"$ne": True},
        "$or": [
            {"role": {"$in": OFFICE_LOGIN_ROLES}},
            {"app_role": {"$in": OFFICE_LOGIN_ROLES}},
        ],
    }, {"_id": 0})

    if worker:
        return public_user_from_worker(worker)

    return None

async def verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
    """Verify an office/admin login.

    The frontend still uses Basic auth for existing protected endpoints.
    This now accepts named office users as well as the legacy built-in admin.
    """
    try:
        user = await find_office_user(credentials.username, credentials.password)
        if user:
            return user.get("email") or credentials.username
    except Exception as e:
        print(f"Error checking office user: {e}")

    raise HTTPException(
        status_code=401,
        detail="Invalid admin credentials",
        headers={"WWW-Authenticate": "Basic"},
    )

# Surveyor authentication functions
def hash_password(password: str) -> str:
    """Hash password for storage"""
    return hashlib.sha256(password.encode()).hexdigest()

def verify_surveyor_token(token: str) -> str:
    """Verify surveyor JWT token and return surveyor_id"""
    try:
        decoded = base64.b64decode(token).decode()
        return decoded.split(':')[0]  # Extract surveyor_id
    except:
        raise HTTPException(status_code=401, detail="Invalid surveyor token")

async def get_current_surveyor(credentials: HTTPAuthorizationCredentials = Depends(bearer_scheme)) -> str:
    """Dependency to get current surveyor from token"""
    if not credentials:
        # For testing, allow access without token
        return "test_surveyor_id"

    if credentials.scheme != "Bearer":
        raise HTTPException(status_code=401, detail="Invalid authentication scheme")

    return verify_surveyor_token(credentials.credentials)

# Helper functions
def calculate_duration(clock_in: datetime, clock_out: datetime) -> int:
    """Calculate duration in minutes between two datetime objects"""
    delta = clock_out - clock_in
    return int(delta.total_seconds() / 60)


def haversine_metres(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    """Distance between two GPS co-ordinates in metres."""
    radius = 6371000
    phi1 = math.radians(float(lat1))
    phi2 = math.radians(float(lat2))
    d_phi = math.radians(float(lat2) - float(lat1))
    d_lambda = math.radians(float(lon2) - float(lon1))
    a = math.sin(d_phi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(d_lambda / 2) ** 2
    return radius * 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

def is_gps_exempt_job(job: Dict[str, Any]) -> bool:
    """Jobs like Runner/mobile jobs are allowed to clock anywhere."""
    if not job:
        return False
    if job.get("gps_exempt") or job.get("location_exempt") or job.get("allow_remote_clocking"):
        return True
    searchable = " ".join([str(job.get("name", "")), str(job.get("location", "")), str(job.get("description", ""))]).lower()
    exempt_terms = ["runner", "mobile", "roaming", "anywhere", "various", "multiple sites"]
    return any(term in searchable for term in exempt_terms)

def job_requires_gps(job: Dict[str, Any]) -> bool:
    """Return True only when the job has GPS explicitly required and is not exempt."""
    if not job or is_gps_exempt_job(job):
        return False
    return bool(
        job.get("gps_required") is True
        or job.get("require_gps") is True
        or job.get("requires_gps") is True
        or job.get("location_required") is True
    )

def get_job_coordinates(job: Dict[str, Any]):
    """Return job GPS co-ordinates if the job record has them stored."""
    if not job:
        return None
    pairs = [(job.get("latitude"), job.get("longitude")), (job.get("lat"), job.get("lng")), (job.get("job_latitude"), job.get("job_longitude"))]
    gps = job.get("gps_location") or job.get("location_gps") or {}
    if isinstance(gps, dict):
        pairs.append((gps.get("latitude"), gps.get("longitude")))
        pairs.append((gps.get("lat"), gps.get("lng")))
    for lat, lng in pairs:
        try:
            if lat is not None and lng is not None:
                return float(lat), float(lng)
        except (TypeError, ValueError):
            continue
    return None

async def build_suspicious_flags(worker_id: Optional[str], job_id: Optional[str], device_id: Optional[str], gps_location: Optional[GPSLocation], existing_flags: Optional[List[str]] = None) -> List[str]:
    """Build fraud/protection flags without blocking clock in/out."""
    flags = set(existing_flags or [])

    worker = await db.workers.find_one({"id": worker_id}) if worker_id else None
    if worker and worker.get("gps_exempt", False):
        flags.add("WORKER_GPS_EXEMPT")

    if not gps_location:
        flags.add("MISSING_GPS")
    elif gps_location.accuracy is not None and gps_location.accuracy > 100:
        flags.add("POOR_GPS_ACCURACY")

    if device_id:
        since = datetime.utcnow() - timedelta(days=60)
        shared_device = await db.time_entries.find_one({
            "worker_id": {"$ne": worker_id},
            "clock_in": {"$gte": since},
            "$or": [{"device_id_in": device_id}, {"device_id_out": device_id}],
        })
        if shared_device:
            flags.add("SHARED_DEVICE_MULTIPLE_WORKERS")
    else:
        flags.add("MISSING_DEVICE_ID")

    if gps_location:
        day_start = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
        day_end = day_start + timedelta(days=1)
        same_position = await db.time_entries.find_one({
            "worker_id": {"$ne": worker_id},
            "clock_in": {"$gte": day_start, "$lt": day_end},
            "$or": [
                {"gps_location_in.latitude": gps_location.latitude, "gps_location_in.longitude": gps_location.longitude},
                {"gps_location_out.latitude": gps_location.latitude, "gps_location_out.longitude": gps_location.longitude},
            ],
        })
        if same_position:
            flags.add("IDENTICAL_GPS_MULTIPLE_WORKERS")

    if gps_location and job_id:
        job = await db.jobs.find_one({"id": job_id}) or {}
        if not is_gps_exempt_job(job):
            job_coords = get_job_coordinates(job)
            if job_coords:
                distance = haversine_metres(gps_location.latitude, gps_location.longitude, job_coords[0], job_coords[1])
                if distance > 1609:
                    flags.add("FAR_FROM_JOB_LOCATION")

    if worker_id:
        recent_cutoff = datetime.utcnow() - timedelta(minutes=5)
        recent_count = await db.time_entries.count_documents({"worker_id": worker_id, "clock_in": {"$gte": recent_cutoff}})
        if recent_count >= 2:
            flags.add("UNUSUAL_RAPID_CLOCK_ACTIVITY")

    return sorted(flags)

# AUTHENTICATION ENDPOINTS
@api_router.post("/admin/login")
async def admin_login(login_data: AdminLogin):
    """Office/admin login endpoint.

    Returns the actual logged-in user so the frontend can display their name
    and attach it to audit trail changes.
    """
    user = await find_office_user(login_data.username, login_data.password)

    if user:
        return {
            "success": True,
            "message": "Login successful",
            "user": user,
        }

    raise HTTPException(status_code=401, detail="Invalid admin credentials")

@api_router.get("/admin/me")
async def get_admin_me(admin: str = Depends(verify_admin)):
    """Return a lightweight current-user response for existing Basic auth sessions."""
    return {"success": True, "username": admin}

# ==================== APP USER / SETTINGS ENDPOINTS ====================

APP_USER_ROLES = ["super_admin", "admin", "project_manager", "accounts", "worker"]

class AppUserCreate(BaseModel):
    name: str
    email: str
    phone: str = ""
    password: str = ""
    app_role: str = "worker"
    division: str = ""
    trades: List[str] = []
    hourly_rate: float = 15.0
    gps_exempt: bool = False
    active: bool = True

class AppUserUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    password: Optional[str] = None
    app_role: Optional[str] = None
    division: Optional[str] = None
    trades: Optional[List[str]] = None
    hourly_rate: Optional[float] = None
    gps_exempt: Optional[bool] = None
    active: Optional[bool] = None
    archived: Optional[bool] = None

class AppUserPasswordUpdate(BaseModel):
    password: str


# ==================== ROLE PERMISSIONS SETTINGS ====================

APP_SECTIONS = [
    "dashboard",
    "workers",
    "clients",
    "jobs",
    "project-management",
    "work-orders",
    "variations",
    "schedule",
    "subcontractors",
    "purchase-orders",
    "finance",
    "price-builder",
    "time-reports",
    "materials-reports",
    "settings",
]

ROLE_PERMISSION_ROLES = ["super_admin", "admin", "project_manager", "accounts", "worker"]

DEFAULT_ROLE_PERMISSIONS = {
    "super_admin": {section: True for section in APP_SECTIONS},
    "admin": {
        "dashboard": True,
        "workers": True,
        "clients": True,
        "jobs": True,
        "project-management": False,
        "work-orders": True,
        "variations": True,
        "schedule": False,
        "subcontractors": True,
        "purchase-orders": True,
        "finance": False,
        "price-builder": True,
        "time-reports": True,
        "materials-reports": True,
        "settings": False,
    },
    "project_manager": {
        "dashboard": False,
        "workers": False,
        "clients": False,
        "jobs": False,
        "project-management": True,
        "work-orders": True,
        "variations": True,
        "schedule": True,
        "subcontractors": True,
        "purchase-orders": False,
        "finance": False,
        "price-builder": True,
        "time-reports": False,
        "materials-reports": False,
        "settings": False,
    },
    "accounts": {
        "dashboard": False,
        "workers": False,
        "clients": False,
        "jobs": False,
        "project-management": False,
        "work-orders": True,
        "variations": True,
        "schedule": False,
        "subcontractors": False,
        "purchase-orders": False,
        "finance": True,
        "price-builder": False,
        "time-reports": False,
        "materials-reports": False,
        "settings": False,
    },
    "worker": {section: False for section in APP_SECTIONS},
}
DEFAULT_ROLE_PERMISSIONS["worker"]["dashboard"] = False

DEFAULT_ROLE_LANDING_PAGES = {
    "super_admin": "dashboard",
    "admin": "dashboard",
    "project_manager": "project-management",
    "accounts": "finance",
    "worker": "dashboard",
}

class RolePermissionsUpdate(BaseModel):
    permissions: Dict[str, bool] = {}
    landing_page: Optional[str] = None


def normalise_permission_role(value: Optional[str]) -> str:
    role = normalise_app_role(value)
    if role == "office_admin":
        role = "admin"
    if role not in ROLE_PERMISSION_ROLES:
        role = "worker"
    return role


def normalise_permission_map(role: str, permissions: Optional[Dict[str, Any]] = None) -> Dict[str, bool]:
    role = normalise_permission_role(role)
    default_map = DEFAULT_ROLE_PERMISSIONS.get(role, DEFAULT_ROLE_PERMISSIONS["worker"])
    cleaned = {section: bool(default_map.get(section, False)) for section in APP_SECTIONS}

    if isinstance(permissions, dict):
        for section in APP_SECTIONS:
            if section in permissions:
                cleaned[section] = bool(permissions.get(section))

    # Safety rules. Super Admin always has everything. Settings is Super Admin only.
    if role == "super_admin":
        cleaned = {section: True for section in APP_SECTIONS}
    else:
        cleaned["settings"] = False

    return cleaned


def normalise_landing_page(role: str, landing_page: Optional[str], permissions: Dict[str, bool]) -> str:
    role = normalise_permission_role(role)
    if role == "super_admin":
        return landing_page if landing_page in APP_SECTIONS else DEFAULT_ROLE_LANDING_PAGES["super_admin"]

    requested = str(landing_page or "").strip()
    default_page = DEFAULT_ROLE_LANDING_PAGES.get(role, "dashboard")

    # The landing page must be one of the visible sections for that role.
    if requested in APP_SECTIONS and permissions.get(requested):
        return requested
    if default_page in APP_SECTIONS and permissions.get(default_page):
        return default_page

    for section in APP_SECTIONS:
        if permissions.get(section):
            return section

    return "dashboard"


async def get_role_permissions_doc(role: str) -> Dict[str, Any]:
    role = normalise_permission_role(role)
    if role == "super_admin":
        permissions = normalise_permission_map("super_admin")
        landing_page = normalise_landing_page("super_admin", DEFAULT_ROLE_LANDING_PAGES["super_admin"], permissions)
        return {
            "role": "super_admin",
            "permissions": permissions,
            "landing_page": landing_page,
            "is_default": True,
            "protected": True,
        }

    doc = await db.role_permissions.find_one({"role": role}, {"_id": 0})
    permissions = normalise_permission_map(role, (doc or {}).get("permissions"))
    landing_page = normalise_landing_page(role, (doc or {}).get("landing_page"), permissions)
    return {
        "role": role,
        "permissions": permissions,
        "landing_page": landing_page,
        "is_default": not bool(doc),
        "protected": False,
        "updated_at": (doc or {}).get("updated_at"),
        "updated_by": (doc or {}).get("updated_by", ""),
        "updated_by_email": (doc or {}).get("updated_by_email", ""),
    }


async def write_role_permissions_audit(role: str, actor: Dict[str, Any], changes: Optional[List[Dict[str, Any]]] = None, action: str = "update"):
    try:
        now = datetime.utcnow().isoformat()
        actor_name = actor.get("name") or actor.get("email") or "Super Admin"
        role_label = normalise_permission_role(role).replace("_", " ").title()
        if not changes:
            changes = [{"field": "permissions", "old_value": "", "new_value": "updated"}]

        for change in changes:
            section = str(change.get("field") or "permissions")
            section_label = section.replace("-", " ").replace("_", " ").title()
            old_value = change.get("old_value")
            new_value = change.get("new_value")
            if action == "reset":
                description = f"{actor_name} reset role permissions to defaults."
            else:
                if section == "landing_page":
                    description = f"{actor_name} changed the default landing page for {role_label} users from {old_value} to {new_value}."
                else:
                    enabled_text = "enabled" if bool(new_value) else "disabled"
                    description = f"{actor_name} {enabled_text} {section_label} access for {role_label} users."

            await db.audit_logs.insert_one({
                "id": str(uuid.uuid4()),
                "record_type": "role_permissions",
                "record_id": normalise_permission_role(role),
                "project_id": None,
                "project_name": None,
                "action": action,
                "field": section,
                "field_label": section_label,
                "old_value": old_value,
                "new_value": new_value,
                "changed_by": actor_name,
                "changed_by_email": actor.get("email", ""),
                "changed_by_role": actor.get("role", "super_admin"),
                "changed_at": now,
                "description": description,
            })
    except Exception as exc:
        logger.warning("Could not write role permission audit log: %s", exc)



def normalise_settings_role(value: Optional[str]) -> str:
    role = normalise_app_role(value)
    if role not in APP_USER_ROLES:
        return "worker"
    return role


def worker_role_for_app_role(app_role: str) -> str:
    app_role = normalise_settings_role(app_role)
    if app_role in ["super_admin", "admin", "project_manager", "accounts"]:
        return app_role
    return "worker"


def app_user_public_doc(worker: Dict[str, Any]) -> Dict[str, Any]:
    if not worker:
        return {}
    email = str(worker.get("email", "")).strip().lower()
    app_role = normalise_settings_role(worker.get("app_role") or worker.get("role") or "worker")
    if email in SUPER_ADMIN_EMAILS:
        app_role = "super_admin"
    return {
        "id": worker.get("id", ""),
        "name": worker.get("name", ""),
        "email": worker.get("email", ""),
        "phone": worker.get("phone", ""),
        "role": worker.get("role", "worker"),
        "app_role": app_role,
        "worker_type": worker.get("worker_type", "admin" if app_role != "worker" else "worker"),
        "division": worker.get("division", ""),
        "trades": worker.get("trades", []),
        "hourly_rate": finance_to_number(worker.get("hourly_rate"), 15.0),
        "gps_exempt": bool(worker.get("gps_exempt", False)),
        "active": bool(worker.get("active", True)),
        "archived": bool(worker.get("archived", False)),
        "has_password": bool(worker.get("password")),
        "created_date": worker.get("created_date"),
        "updated_at": worker.get("updated_at"),
    }


async def get_super_admin_user(credentials: HTTPBasicCredentials = Depends(security)) -> Dict[str, Any]:
    user = await find_office_user(credentials.username, credentials.password)
    if user and normalise_app_role(user.get("role")) == "super_admin":
        return user
    raise HTTPException(status_code=403, detail="Super Admin access required")


async def write_user_access_audit(action: str, target: Dict[str, Any], actor: Dict[str, Any], changes: Optional[List[Dict[str, Any]]] = None):
    try:
        now = datetime.utcnow().isoformat()
        actor_name = actor.get("name") or actor.get("email") or "Super Admin"
        target_name = target.get("name") or target.get("email") or "user"
        for change in changes or [{}]:
            field = change.get("field")
            old_value = change.get("old_value")
            new_value = change.get("new_value")
            if action == "create":
                description = f"{actor_name} created user access for {target_name}."
            elif action == "archive":
                description = f"{actor_name} archived user access for {target_name}."
            elif action == "password":
                description = f"{actor_name} reset the password for {target_name}."
            elif field:
                description = f"{actor_name} changed {target_name} {field.replace('_', ' ')} from {old_value} to {new_value}."
            else:
                description = f"{actor_name} updated user access for {target_name}."
            await db.audit_logs.insert_one({
                "id": str(uuid.uuid4()),
                "record_type": "app_user",
                "record_id": target.get("id"),
                "project_id": None,
                "project_name": None,
                "action": action,
                "field": field,
                "field_label": field.replace("_", " ").title() if field else "User Access",
                "old_value": old_value,
                "new_value": new_value,
                "changed_by": actor_name,
                "changed_by_email": actor.get("email", ""),
                "changed_by_role": actor.get("role", "super_admin"),
                "changed_at": now,
                "description": description,
            })
    except Exception as exc:
        logger.warning("Could not write user access audit log: %s", exc)


@api_router.get("/app-users")
@api_router.get("/app-users/")
async def get_app_users(super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    """List all app users/workers for Settings > User Access."""
    users = await db.workers.find({}, {"_id": 0}).sort("name", 1).to_list(5000)
    return [app_user_public_doc(user) for user in users]


@api_router.post("/app-users")
@api_router.post("/app-users/")
async def create_app_user(user: AppUserCreate, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    """Create a worker/app user from Settings > User Access."""
    existing = await db.workers.find_one({"email": {"$regex": f"^{re.escape(user.email.strip())}$", "$options": "i"}})
    if existing:
        raise HTTPException(status_code=400, detail="A user with this email already exists")

    app_role = normalise_settings_role(user.app_role)
    now = datetime.utcnow()
    doc = {
        "id": str(uuid.uuid4()),
        "name": user.name.strip(),
        "email": user.email.strip(),
        "phone": user.phone.strip(),
        "password": user.password,
        "role": worker_role_for_app_role(app_role),
        "app_role": app_role,
        "worker_type": "admin" if app_role != "worker" else "worker",
        "division": user.division,
        "trades": user.trades or [],
        "hourly_rate": user.hourly_rate,
        "gps_exempt": user.gps_exempt,
        "active": user.active,
        "archived": False,
        "created_date": now,
        "updated_at": now,
    }
    await db.workers.insert_one(doc)
    public_doc = app_user_public_doc(doc)
    await write_user_access_audit("create", public_doc, super_admin)
    return public_doc


@api_router.put("/app-users/{user_id}")
async def update_app_user(user_id: str, update: AppUserUpdate, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    existing = await db.workers.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    update_data = {k: v for k, v in update.dict().items() if v is not None}
    if not update_data:
        return app_user_public_doc(existing)

    # Owner protection: Duke must never be downgraded or locked out from Settings.
    existing_email = str(existing.get("email", "")).strip().lower()
    if existing_email in SUPER_ADMIN_EMAILS:
        if "app_role" in update_data and normalise_settings_role(update_data["app_role"]) != "super_admin":
            raise HTTPException(status_code=400, detail="The owner account cannot be downgraded from Super Admin")
        if update_data.get("active") is False or update_data.get("archived") is True:
            raise HTTPException(status_code=400, detail="The owner account cannot be deactivated or archived")

    if "email" in update_data:
        duplicate = await db.workers.find_one({
            "id": {"$ne": user_id},
            "email": {"$regex": f"^{re.escape(str(update_data['email']).strip())}$", "$options": "i"},
        })
        if duplicate:
            raise HTTPException(status_code=400, detail="Another user already has this email")
        update_data["email"] = str(update_data["email"]).strip()

    if "app_role" in update_data:
        app_role = normalise_settings_role(update_data["app_role"])
        update_data["app_role"] = app_role
        update_data["role"] = worker_role_for_app_role(app_role)
        update_data["worker_type"] = "admin" if app_role != "worker" else "worker"

    update_data["updated_at"] = datetime.utcnow()

    audit_changes = []
    for key, new_value in update_data.items():
        if key == "updated_at":
            continue
        old_value = existing.get(key)
        if str(old_value or "") != str(new_value or ""):
            audit_changes.append({"field": key, "old_value": old_value, "new_value": new_value})

    await db.workers.update_one({"id": user_id}, {"$set": update_data})
    updated = await db.workers.find_one({"id": user_id}, {"_id": 0})
    public_doc = app_user_public_doc(updated)
    await write_user_access_audit("update", public_doc, super_admin, audit_changes or None)
    return public_doc


@api_router.put("/app-users/{user_id}/password")
async def update_app_user_password(user_id: str, update: AppUserPasswordUpdate, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    if not update.password:
        raise HTTPException(status_code=400, detail="Password is required")
    existing = await db.workers.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")
    await db.workers.update_one({"id": user_id}, {"$set": {"password": update.password, "updated_at": datetime.utcnow()}})
    updated = await db.workers.find_one({"id": user_id}, {"_id": 0})
    public_doc = app_user_public_doc(updated)
    await write_user_access_audit("password", public_doc, super_admin)
    return {"success": True, "message": "Password updated", "user": public_doc}


@api_router.put("/app-users/{user_id}/archive")
async def archive_app_user(user_id: str, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    existing = await db.workers.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")
    if str(existing.get("email", "")).strip().lower() in SUPER_ADMIN_EMAILS:
        raise HTTPException(status_code=400, detail="The owner account cannot be archived")
    await db.workers.update_one({"id": user_id}, {"$set": {"archived": True, "active": False, "updated_at": datetime.utcnow()}})
    updated = await db.workers.find_one({"id": user_id}, {"_id": 0})
    public_doc = app_user_public_doc(updated)
    await write_user_access_audit("archive", public_doc, super_admin)
    return {"success": True, "message": "User archived", "user": public_doc}


@api_router.get("/role-permissions")
@api_router.get("/role-permissions/")
async def get_role_permissions(super_admin: Optional[Dict[str, Any]] = None):
    """Return the current role-to-section permission matrix.

    This endpoint is readable by the app so the sidebar can build itself from saved settings.
    Updating permissions remains Super Admin only.
    """
    result = {}
    landing_pages = {}
    for role in ROLE_PERMISSION_ROLES:
        doc = await get_role_permissions_doc(role)
        result[role] = doc["permissions"]
        landing_pages[role] = doc.get("landing_page") or normalise_landing_page(role, None, doc["permissions"])
    return {
        "sections": APP_SECTIONS,
        "roles": ROLE_PERMISSION_ROLES,
        "permissions": result,
        "landing_pages": landing_pages,
        "protected_roles": ["super_admin"],
    }


@api_router.put("/role-permissions/{role}")
async def update_role_permissions(role: str, update: RolePermissionsUpdate, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    """Update access permissions for one role. Super Admin is protected."""
    role = normalise_permission_role(role)
    if role == "super_admin":
        raise HTTPException(status_code=400, detail="Super Admin permissions are protected and cannot be restricted")

    existing_doc = await get_role_permissions_doc(role)
    existing_permissions = existing_doc.get("permissions", {})
    new_permissions = normalise_permission_map(role, update.permissions)

    # Settings must remain Super Admin only.
    new_permissions["settings"] = False
    existing_landing_page = existing_doc.get("landing_page") or normalise_landing_page(role, None, existing_permissions)
    requested_landing_page = update.landing_page if update.landing_page is not None else existing_landing_page
    new_landing_page = normalise_landing_page(role, requested_landing_page, new_permissions)

    changes = []
    for section in APP_SECTIONS:
        old_value = bool(existing_permissions.get(section, False))
        new_value = bool(new_permissions.get(section, False))
        if old_value != new_value:
            changes.append({"field": section, "old_value": old_value, "new_value": new_value})

    if existing_landing_page != new_landing_page:
        changes.append({"field": "landing_page", "old_value": existing_landing_page, "new_value": new_landing_page})

    now = datetime.utcnow()
    await db.role_permissions.update_one(
        {"role": role},
        {"$set": {
            "role": role,
            "permissions": new_permissions,
            "landing_page": new_landing_page,
            "updated_at": now,
            "updated_by": super_admin.get("name") or super_admin.get("email") or "Super Admin",
            "updated_by_email": super_admin.get("email", ""),
            "updated_by_role": super_admin.get("role", "super_admin"),
        }},
        upsert=True,
    )

    if changes:
        await write_role_permissions_audit(role, super_admin, changes, action="update")

    updated_doc = await get_role_permissions_doc(role)
    return updated_doc


@api_router.post("/role-permissions/reset-defaults")
async def reset_role_permissions(super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    """Reset editable role permissions back to defaults."""
    await db.role_permissions.delete_many({"role": {"$ne": "super_admin"}})
    await write_role_permissions_audit("all", super_admin, action="reset")
    result = {}
    landing_pages = {}
    for role in ROLE_PERMISSION_ROLES:
        doc = await get_role_permissions_doc(role)
        result[role] = doc["permissions"]
        landing_pages[role] = doc.get("landing_page")
    return {"success": True, "message": "Role permissions reset to defaults", "permissions": result, "landing_pages": landing_pages}

# SURVEYOR AUTHENTICATION ENDPOINTS
@api_router.post("/surveyors/register")
async def register_surveyor(surveyor_data: SurveyorCreate):
    """Register a new surveyor"""
    # Check if surveyor already exists
    existing_surveyor = await db.surveyors.find_one({"email": surveyor_data.email})
    if existing_surveyor:
        raise HTTPException(status_code=400, detail="Surveyor already exists")

    # Create new surveyor
    surveyor = {
        "id": str(uuid.uuid4()),
        "name": surveyor_data.name,
        "email": surveyor_data.email,
        "phone": surveyor_data.phone,
        "password": hash_password(surveyor_data.password),
        "created_at": datetime.utcnow(),
        "active": True
    }

    await db.surveyors.insert_one(surveyor)

    # Return surveyor without password
    surveyor.pop("password")
    return surveyor

@api_router.post("/surveyors/login")
async def login_surveyor(login_data: SurveyorLogin):
    """Authenticate surveyor and return token"""
    # Find surveyor
    surveyor = await db.surveyors.find_one({"email": login_data.email})
    if not surveyor or not surveyor.get("active"):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # Verify password
    if surveyor["password"] != hash_password(login_data.password):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # Generate token (simple base64 encoding for now)
    token_data = f"{surveyor['id']}:{login_data.email}"
    token = base64.b64encode(token_data.encode()).decode()

    return {
        "token": token,
        "surveyor_id": surveyor["id"],
        "surveyor": {
            "id": surveyor["id"],
            "name": surveyor["name"],
            "email": surveyor["email"],
            "phone": surveyor["phone"]
        }
    }

# QUOTE ENDPOINTS
@api_router.get("/quotes")
async def get_quotes(
    surveyor_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    current_surveyor_id: str = Depends(get_current_surveyor)
):
    """Get quotes, filtered by surveyor if not admin"""
    filter_query = {}

    # If surveyor_id is provided and it's different from current surveyor, only allow admins
    if surveyor_id and surveyor_id != current_surveyor_id:
        # Check if current user is admin (basic check)
        if current_surveyor_id == "test_surveyor_id":
            # Allow test surveyor to see all quotes
            pass
        else:
            # In production, add proper admin check here
            filter_query["surveyor_id"] = current_surveyor_id
    elif not surveyor_id:
        # If no surveyor_id specified, filter by current surveyor
        filter_query["surveyor_id"] = current_surveyor_id
    else:
        filter_query["surveyor_id"] = surveyor_id

    if status:
        filter_query["status"] = status

    quotes = await db.quotes.find(filter_query).to_list(1000)

    # Remove MongoDB ObjectId fields and handle any serialization issues
    clean_quotes = []
    for quote in quotes:
        # Remove MongoDB ObjectId
        quote.pop("_id", None)

        # Convert any Decimal objects to float
        clean_quote = convert_decimals_to_float(quote)
        clean_quotes.append(clean_quote)

    return clean_quotes

@api_router.post("/quotes")
async def create_quote(
    quote_data: QuoteCreate,
    surveyor_id: str = Depends(get_current_surveyor)
):
    """Create a new quote"""
    # Get surveyor info
    surveyor = await db.surveyors.find_one({"id": surveyor_id})
    if not surveyor:
        raise HTTPException(status_code=404, detail="Surveyor not found")

    # Generate quote number
    quote_number = f"Q-{datetime.utcnow().strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"

    # Calculate totals
    total_materials = sum(item.total_price for item in quote_data.materials)
    total_labor = sum(item.total_cost for item in quote_data.labor)
    subtotal = total_materials + total_labor
    vat_amount = subtotal * quote_data.vat_rate
    final_amount = subtotal + vat_amount

    # Create quote
    quote = Quote(
        quote_number=quote_number,
        surveyor_id=surveyor_id,
        surveyor_name=surveyor["name"],
        total_materials_cost=total_materials,
        total_labor_cost=total_labor,
        total_quote_amount=subtotal,
        vat_amount=vat_amount,
        final_amount=final_amount,
        **quote_data.dict()
    )

    # Convert to dict for MongoDB and handle Decimal serialization
    quote_dict = quote.dict()
    quote_dict["created_at"] = datetime.utcnow()

    # Convert Decimal objects to float for MongoDB
    quote_dict = convert_decimals_to_float(quote_dict)

    # Insert into MongoDB
    result = await db.quotes.insert_one(quote_dict)

    # Return the quote without the MongoDB ObjectId
    quote_dict.pop("_id", None)  # Remove _id if it exists
    return quote_dict

@api_router.get("/quotes/{quote_id}")
async def get_quote(quote_id: str):
    """Get a specific quote"""
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    
    # Remove MongoDB ObjectId and convert Decimals
    quote.pop("_id", None)
    return convert_decimals_to_float(quote)

@api_router.put("/quotes/{quote_id}")
async def update_quote(
    quote_id: str,
    quote_update: QuoteUpdate,
    surveyor_id: str = Depends(get_current_surveyor)
):
    """Update a quote"""
    # Check if quote exists and belongs to surveyor
    existing_quote = await db.quotes.find_one({"id": quote_id})
    if not existing_quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    # For non-admin users, verify ownership
    if existing_quote["surveyor_id"] != surveyor_id:
        # Check if user is admin
        try:
            verify_admin(None)  # This will fail if not admin
        except:
            raise HTTPException(status_code=403, detail="Not authorized to update this quote")

    # Update quote
    update_dict = {k: v for k, v in quote_update.dict().items() if v is not None}
    update_dict["updated_at"] = datetime.utcnow()

    # Convert decimals for MongoDB storage
    update_dict = convert_decimals_to_float(update_dict)

    result = await db.quotes.update_one({"id": quote_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Quote not found")

    # Return updated quote
    updated_quote = await db.quotes.find_one({"id": quote_id})
    updated_quote.pop("_id", None)
    return convert_decimals_to_float(updated_quote)

# WORKER ENDPOINTS
@api_router.post("/workers", response_model=Worker)
async def create_worker(worker: WorkerCreate, admin: str = Depends(verify_admin)):
    """Create a new worker (Admin only)"""
    worker_dict = worker.dict()
    if worker_dict.get("role") == "contractor":
        worker_dict["role"] = "worker"
        worker_dict["worker_type"] = "contractor"
    elif worker_dict.get("role") == "worker":
        worker_dict.setdefault("worker_type", "worker")
    worker_obj = Worker(**worker_dict)
    await db.workers.insert_one(worker_obj.dict())
    return worker_obj

@api_router.get("/workers", response_model=List[Worker])
async def get_workers(
    active_only: bool = Query(True),
    include_archived: bool = Query(False),
    worker_type: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    include_admins: bool = Query(True)
):
    filter_dict = {"active": True} if active_only else {}

    if not include_archived:
        filter_dict["archived"] = {"$ne": True}

    if worker_type and worker_type != "all":
        filter_dict["worker_type"] = worker_type

    if division and division != "all":
        filter_dict["division"] = division

    if trade and trade != "all":
        filter_dict["$or"] = [{"trades": trade}, {"trade": trade}]

    if not include_admins:
        filter_dict["role"] = {"$ne": "admin"}

    workers = await db.workers.find(filter_dict).to_list(1000)

    # Backfill defaults for older workers that were created before these fields existed.
    for worker in workers:
        if worker.get("role") == "contractor":
            worker["role"] = "worker"
            worker["worker_type"] = "contractor"
        worker.setdefault("worker_type", "worker")
        worker.setdefault("division", "")
        worker.setdefault("gps_exempt", False)
        if "trades" not in worker:
            old_trade = worker.get("trade", "")
            worker["trades"] = [item.strip() for item in old_trade.split(",") if item.strip()] if old_trade else []

    return [Worker(**worker) for worker in workers]



@api_router.get("/workers/export-csv")
async def export_workers_csv(
    include_archived: bool = Query(False),
    worker_type: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    include_admins: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    """Download workers as a CSV file for admin records/backups."""
    filter_dict = {"active": True}

    if not include_archived:
        filter_dict["archived"] = {"$ne": True}

    if worker_type and worker_type != "all":
        filter_dict["worker_type"] = worker_type

    if division and division != "all":
        filter_dict["division"] = division

    if trade and trade != "all":
        filter_dict["$or"] = [{"trades": trade}, {"trade": trade}]

    if not include_admins:
        filter_dict["role"] = {"$ne": "admin"}

    workers = await db.workers.find(filter_dict, {"_id": 0}).sort("name", 1).to_list(5000)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "name",
        "email",
        "phone",
        "role",
        "worker_type",
        "division",
        "trades",
        "hourly_rate",
        "gps_exempt",
        "active",
        "archived",
        "created_date",
    ])

    for worker in workers:
        trades = worker.get("trades") or []
        if not trades and worker.get("trade"):
            trades = [worker.get("trade")]
        if isinstance(trades, list):
            trades_text = ", ".join(str(item) for item in trades if item)
        else:
            trades_text = str(trades or "")

        writer.writerow([
            worker.get("name", ""),
            worker.get("email", ""),
            worker.get("phone", ""),
            worker.get("role", ""),
            worker.get("worker_type", "worker"),
            worker.get("division", ""),
            trades_text,
            worker.get("hourly_rate", ""),
            "true" if worker.get("gps_exempt") else "false",
            "true" if worker.get("active", True) else "false",
            "true" if worker.get("archived") else "false",
            format_uk_datetime_for_export(worker.get("created_date")),
        ])

    output.seek(0)
    filename = f"workers_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.csv"
    return StreamingResponse(
        io.BytesIO(("\ufeff" + output.getvalue()).encode("utf-8")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@api_router.get("/workers/{worker_id}", response_model=Worker)
async def get_worker(worker_id: str):
    worker = await db.workers.find_one({"id": worker_id})
    if not worker:
        raise HTTPException(status_code=404, detail="Worker not found")
    return Worker(**worker)

@api_router.put("/workers/{worker_id}", response_model=Worker)
async def update_worker(worker_id: str, worker_update: WorkerUpdate, admin: str = Depends(verify_admin)):
    """Update worker (Admin only)"""
    update_dict = {k: v for k, v in worker_update.dict().items() if v is not None}
    if update_dict.get("role") == "contractor":
        update_dict["role"] = "worker"
        update_dict["worker_type"] = "contractor"
    elif update_dict.get("role") == "worker" and update_dict.get("worker_type") is None:
        update_dict["worker_type"] = "worker"

    result = await db.workers.update_one({"id": worker_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Worker not found")

    updated_worker = await db.workers.find_one({"id": worker_id})
    return Worker(**updated_worker)

@api_router.delete("/workers/{worker_id}")
async def delete_worker(worker_id: str, admin: str = Depends(verify_admin)):
    """Delete worker (Admin only)"""
    result = await db.workers.delete_one({"id": worker_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Worker not found")
    return {"message": "Worker deleted successfully"}

@api_router.put("/workers/{worker_id}/archive")
async def archive_worker(worker_id: str, admin: str = Depends(verify_admin)):
    """Archive worker (Admin only)"""
    result = await db.workers.update_one({"id": worker_id}, {"$set": {"archived": True}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Worker not found")
    return {"message": "Worker archived successfully"}


# ==================== CLIENT / CUSTOMER ENDPOINTS ====================

@api_router.get("/clients", response_model=List[Client])
async def get_clients(
    active_only: bool = Query(True),
    include_archived: bool = Query(False),
    search: Optional[str] = Query(None),
    admin: str = Depends(verify_admin),
):
    filter_dict = {}
    if active_only:
        filter_dict["active"] = True
    if not include_archived:
        filter_dict["archived"] = {"$ne": True}
    if search:
        filter_dict["$or"] = [
            {"client_name": {"$regex": search, "$options": "i"}},
            {"main_contact_name": {"$regex": search, "$options": "i"}},
            {"email": {"$regex": search, "$options": "i"}},
            {"accounts_email": {"$regex": search, "$options": "i"}},
        ]
    docs = await db.clients.find(filter_dict, {"_id": 0}).sort("client_name", 1).to_list(5000)
    return [Client(**doc) for doc in docs]

@api_router.post("/clients", response_model=Client)
async def create_client(client: ClientCreate, admin: str = Depends(verify_admin)):
    duplicate = await db.clients.find_one({
        "client_name": {"$regex": f"^{re.escape(client.client_name.strip())}$", "$options": "i"},
        "archived": {"$ne": True},
    })
    if duplicate:
        raise HTTPException(status_code=400, detail="A client with this name already exists")
    data = client.dict()
    data["client_name"] = data["client_name"].strip()
    data.update(normalise_client_tax_payload(data))
    obj = Client(**data)
    await db.clients.insert_one(obj.dict())
    return obj

@api_router.get("/clients/{client_id}", response_model=Client)
async def get_client(client_id: str, admin: str = Depends(verify_admin)):
    client = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return Client(**client)

@api_router.put("/clients/{client_id}", response_model=Client)
async def update_client(client_id: str, update: ClientUpdate, admin: str = Depends(verify_admin)):
    existing = await db.clients.find_one({"id": client_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Client not found")
    update_data = {k: v for k, v in update.dict().items() if v is not None}
    if not update_data:
        return Client(**existing)
    if "client_name" in update_data:
        duplicate = await db.clients.find_one({
            "id": {"$ne": client_id},
            "client_name": {"$regex": f"^{re.escape(str(update_data['client_name']).strip())}$", "$options": "i"},
            "archived": {"$ne": True},
        })
        if duplicate:
            raise HTTPException(status_code=400, detail="Another client already has this name")
        update_data["client_name"] = str(update_data["client_name"]).strip()
    if any(key in update_data for key in ["vat_treatment", "vat_rate", "drc_enabled", "cis_enabled", "cis_rate", "cis_deduction_basis"]):
        merged = {**existing, **update_data}
        update_data.update(normalise_client_tax_payload(merged))
    update_data["updated_at"] = datetime.utcnow()
    await db.clients.update_one({"id": client_id}, {"$set": update_data})
    updated = await db.clients.find_one({"id": client_id}, {"_id": 0})
    return Client(**updated)

@api_router.put("/clients/{client_id}/archive")
async def archive_client(client_id: str, admin: str = Depends(verify_admin)):
    result = await db.clients.update_one({"id": client_id}, {"$set": {"archived": True, "active": False, "updated_at": datetime.utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    return {"success": True, "message": "Client archived"}

@api_router.put("/clients/{client_id}/unarchive")
async def unarchive_client(client_id: str, admin: str = Depends(verify_admin)):
    result = await db.clients.update_one({"id": client_id}, {"$set": {"archived": False, "active": True, "updated_at": datetime.utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Client not found")
    return {"success": True, "message": "Client unarchived"}

@api_router.post("/jobs/{job_id}/refresh-client-defaults", response_model=Job)
async def refresh_job_client_defaults(job_id: str, admin: str = Depends(verify_admin)):
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if not job.get("client_id"):
        raise HTTPException(status_code=400, detail="This job is not linked to a client record")
    refreshed = await apply_client_defaults_to_job_data(dict(job), force_refresh=True)
    update_keys = ["client_id", "client_name", "client", "client_tax_snapshot", "client_commercial_snapshot", "vat_treatment", "vat_rate", "drc_enabled", "cis_enabled", "cis_rate", "cis_deduction_basis", "payment_terms_days", "retention_percent"]
    update_doc = {key: refreshed.get(key) for key in update_keys if key in refreshed}
    await db.jobs.update_one({"id": job_id}, {"$set": update_doc})
    updated = await db.jobs.find_one({"id": job_id})
    return Job(**updated)

# ==================== GOOGLE DRIVE JOB FOLDER AUTOMATION ====================

def google_drive_config_missing() -> List[str]:
    """Return missing Google Drive env variable names without exposing secret values."""
    missing = []
    if not GOOGLE_DRIVE_PARENT_FOLDER_ID:
        missing.append("GOOGLE_DRIVE_PARENT_FOLDER_ID")
    if not GOOGLE_SERVICE_ACCOUNT_JSON:
        missing.append("GOOGLE_SERVICE_ACCOUNT_JSON")
    return missing


def get_google_drive_api_service():
    """Build a Google Drive API client from the service account JSON in Render env vars."""
    missing = google_drive_config_missing()
    if missing:
        logger.warning("Google Drive automation not configured. Missing: %s", ", ".join(missing))
        return None
    if service_account is None or build is None:
        logger.error("Google Drive API packages are not installed. Check requirements.txt")
        return None

    try:
        account_info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
        credentials = service_account.Credentials.from_service_account_info(
            account_info,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
        return build("drive", "v3", credentials=credentials, cache_discovery=False)
    except Exception as exc:
        logger.exception("Failed to initialise Google Drive API service: %s", exc)
        return None


def drive_safe_name(value: str) -> str:
    """Keep folder names readable while removing characters Drive/Windows dislike."""
    value = (value or "Untitled Job").strip()
    for char in ['<', '>', ':', '"', '/', '\\', '|', '?', '*']:
        value = value.replace(char, '-')
    return " ".join(value.split()) or "Untitled Job"


async def get_next_job_number() -> int:
    """Find the next job number using existing job_number values in MongoDB."""
    latest = await db.jobs.find_one(
        {"job_number": {"$exists": True, "$ne": None}},
        sort=[("job_number", -1)]
    )
    if latest and isinstance(latest.get("job_number"), int):
        return latest["job_number"] + 1
    return GOOGLE_DRIVE_START_JOB_NUMBER


def build_job_display_name(job_number: int, job_name: str) -> str:
    padded = str(job_number).zfill(GOOGLE_DRIVE_JOB_NUMBER_PADDING)
    return f"{padded}: {drive_safe_name(job_name)}"


def create_drive_folder(service, name: str, parent_id: str) -> Dict[str, Any]:
    """Create one folder in Drive and return id/webViewLink."""
    metadata = {
        "name": name,
        "mimeType": "application/vnd.google-apps.folder",
        "parents": [parent_id],
    }
    return service.files().create(
        body=metadata,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()


def list_drive_folder_children(service, folder_id: str) -> List[Dict[str, Any]]:
    """List all non-trashed children of a Drive folder, including shared drives."""
    children = []
    page_token = None
    while True:
        response = service.files().list(
            q=f"'{folder_id}' in parents and trashed = false",
            fields="nextPageToken, files(id,name,mimeType,shortcutDetails,exportLinks)",
            pageSize=1000,
            pageToken=page_token,
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
        ).execute()
        children.extend(response.get("files", []))
        page_token = response.get("nextPageToken")
        if not page_token:
            break
    children.sort(key=lambda item: item.get("name", "").lower())
    return children


def copy_drive_file_with_download_fallback(
    service,
    item_id: str,
    item_name: str,
    destination_folder_id: str,
    mime_type: str,
) -> None:
    """Copy one Drive file. If Drive copy is blocked for a binary file, download and re-upload it."""
    metadata = {"name": item_name, "parents": [destination_folder_id]}

    try:
        service.files().copy(
            fileId=item_id,
            body=metadata,
            fields="id,name",
            supportsAllDrives=True,
        ).execute()
        return
    except Exception as copy_exc:
        # Google Docs/Sheets/Slides must use files.copy. Binary files can be downloaded and re-uploaded.
        if mime_type.startswith("application/vnd.google-apps"):
            raise copy_exc
        if MediaIoBaseDownload is None or MediaIoBaseUpload is None:
            raise copy_exc

        logger.warning("Drive files.copy failed for %s; trying download/upload fallback: %s", item_name, copy_exc)
        buffer = io.BytesIO()
        request = service.files().get_media(fileId=item_id, supportsAllDrives=True)
        downloader = MediaIoBaseDownload(buffer, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        buffer.seek(0)

        media = MediaIoBaseUpload(buffer, mimetype=mime_type or "application/octet-stream", resumable=False)
        service.files().create(
            body=metadata,
            media_body=media,
            fields="id,name",
            supportsAllDrives=True,
        ).execute()


def copy_drive_template_contents_recursive(
    service,
    source_folder_id: str,
    destination_folder_id: str,
    current_path: str = "template",
    stats: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Recursively copy every folder/file from the template folder into the new job folder."""
    if stats is None:
        stats = {
            "folders_created": 0,
            "files_copied": 0,
            "shortcuts_created": 0,
            "skipped": [],
            "paths_seen": [],
        }

    children = list_drive_folder_children(service, source_folder_id)
    logger.info("Copying %s Drive item(s) from %s", len(children), current_path)

    for item in children:
        item_id = item.get("id")
        item_name = item.get("name", "Untitled")
        mime_type = item.get("mimeType", "")
        item_path = f"{current_path}/{item_name}"
        stats["paths_seen"].append(item_path)

        try:
            if mime_type == "application/vnd.google-apps.folder":
                new_folder = create_drive_folder(service, item_name, destination_folder_id)
                stats["folders_created"] += 1
                copy_drive_template_contents_recursive(
                    service,
                    item_id,
                    new_folder["id"],
                    current_path=item_path,
                    stats=stats,
                )
            elif mime_type == "application/vnd.google-apps.shortcut":
                shortcut_details = item.get("shortcutDetails") or {}
                target_id = shortcut_details.get("targetId")
                target_mime_type = shortcut_details.get("targetMimeType")
                if not target_id:
                    raise ValueError("Shortcut has no targetId")
                metadata = {
                    "name": item_name,
                    "mimeType": "application/vnd.google-apps.shortcut",
                    "parents": [destination_folder_id],
                    "shortcutDetails": {"targetId": target_id},
                }
                if target_mime_type:
                    metadata["shortcutDetails"]["targetMimeType"] = target_mime_type
                service.files().create(
                    body=metadata,
                    fields="id,name",
                    supportsAllDrives=True,
                ).execute()
                stats["shortcuts_created"] += 1
            else:
                copy_drive_file_with_download_fallback(
                    service,
                    item_id=item_id,
                    item_name=item_name,
                    destination_folder_id=destination_folder_id,
                    mime_type=mime_type,
                )
                stats["files_copied"] += 1
        except Exception as exc:
            error_message = f"{item_path}: {exc}"
            stats["skipped"].append(error_message)
            logger.warning("Skipped Drive template item during recursive copy: %s", error_message)

    return stats


async def create_google_drive_job_folder(job_number: int, job_name: str) -> Dict[str, Any]:
    """Create the new numbered job folder and recursively copy the template into it."""
    missing = google_drive_config_missing()
    if missing:
        return {
            "drive_folder_status": "not_configured",
            "drive_folder_error": "Missing env var(s): " + ", ".join(missing),
        }

    service = get_google_drive_api_service()
    if not service:
        return {
            "drive_folder_status": "failed",
            "drive_folder_error": "Google Drive API service could not be initialised. Check Render env vars and requirements.txt.",
        }

    display_name = build_job_display_name(job_number, job_name)
    try:
        folder = create_drive_folder(service, display_name, GOOGLE_DRIVE_PARENT_FOLDER_ID)
        result = {
            "drive_folder_status": "created",
            "drive_folder_id": folder.get("id"),
            "drive_folder_link": folder.get("webViewLink"),
            "drive_folder_url": folder.get("webViewLink"),
            "google_drive_link": folder.get("webViewLink"),
        }

        if GOOGLE_DRIVE_TEMPLATE_FOLDER_ID:
            stats = copy_drive_template_contents_recursive(
                service,
                GOOGLE_DRIVE_TEMPLATE_FOLDER_ID,
                folder["id"],
                current_path="template",
            )
            result["drive_folder_copy_stats"] = stats
            if stats.get("skipped"):
                result["drive_folder_status"] = "created_with_copy_warnings"
                result["drive_folder_error"] = "; ".join(stats.get("skipped", [])[:5])
            logger.info("Google Drive template copy complete for %s: %s", display_name, stats)
        else:
            result["drive_folder_error"] = "No GOOGLE_DRIVE_TEMPLATE_FOLDER_ID set, so only the main job folder was created."

        return result
    except Exception as exc:
        logger.exception("Google Drive job folder creation failed for %s: %s", display_name, exc)
        return {
            "drive_folder_status": "failed",
            "drive_folder_error": str(exc),
        }


# JOB ENDPOINTS
@api_router.post("/jobs", response_model=Job)
async def create_job(job: JobCreate, admin: str = Depends(verify_admin)):
    """Create a new job, assign a job number, and create/copy its Google Drive folder."""
    job_dict = job.dict()
    job_dict = await apply_client_defaults_to_job_data(job_dict, force_refresh=False)

    job_number = await get_next_job_number()
    display_name = build_job_display_name(job_number, job.name)

    job_dict["job_number"] = job_number
    job_dict["display_name"] = display_name

    drive_result = await create_google_drive_job_folder(job_number, job.name)
    for key in [
        "drive_folder_id",
        "drive_folder_link",
        "drive_folder_url",
        "google_drive_link",
        "drive_folder_status",
        "drive_folder_error",
        "drive_folder_copy_stats",
    ]:
        if key in drive_result:
            job_dict[key] = drive_result[key]

    job_obj = Job(**job_dict)
    mongo_doc = job_obj.dict()
    await db.jobs.insert_one(mongo_doc)
    return job_obj

@api_router.get("/jobs", response_model=List[Job])
async def get_jobs(active_only: bool = Query(False), include_archived: bool = Query(False)):
    filter_dict = {}

    if active_only:
        filter_dict["status"] = {"$ne": "cancelled"}
        filter_dict["archived"] = {"$ne": True}
    elif not include_archived:
        filter_dict["archived"] = {"$ne": True}

    jobs = await db.jobs.find(filter_dict).sort("name", 1).to_list(1000)
    return [Job(**job) for job in jobs]



# ==================== COMPACT JOB DETAIL REPORT ====================

def _job_detail_money(value):
    try:
        if value is None or value == "":
            return 0.0
        return round(float(value), 2)
    except Exception:
        return 0.0


def _job_detail_first(doc, keys, default=None):
    for key in keys:
        if isinstance(doc, dict) and key in doc and doc.get(key) not in [None, ""]:
            return doc.get(key)
    return default


def _job_detail_parse_date(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        try:
            return datetime.fromisoformat(value.replace("Z", "+00:00")[:10]).date()
        except Exception:
            try:
                return datetime.strptime(value[:10], "%Y-%m-%d").date()
            except Exception:
                return None
    return None


def _job_detail_date_to_iso(value):
    parsed = _job_detail_parse_date(value)
    return parsed.isoformat() if parsed else None


def _job_detail_days_between(start_value, end_value):
    start = _job_detail_parse_date(start_value)
    end = _job_detail_parse_date(end_value)
    if not start or not end:
        return None
    return max((end - start).days + 1, 0)



def _job_detail_is_commercial_placeholder(section: Dict[str, Any]) -> bool:
    section_type = str((section or {}).get("type") or (section or {}).get("section_type") or "").strip().lower()
    return bool(
        (section or {}).get("is_commercial_placeholder") is True
        or (section or {}).get("commercial_placeholder") is True
        or section_type in ["commercial_placeholder", "placeholder"]
    )


@api_router.get("/jobs/{job_id}/detail-report")
async def get_job_detail_report(job_id: str, admin: str = Depends(verify_admin)):
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_name = job.get("name", "Unnamed Job")
    start_date = _job_detail_first(job, ["planned_start_date", "start_date", "startDate", "project_start", "projectStart"], None)
    end_date = _job_detail_first(job, ["planned_end_date", "end_date", "endDate", "project_end", "projectEnd"], None)
    duration_days = _job_detail_first(job, ["duration_days", "durationDays", "duration"], None)
    if duration_days in [None, ""]:
        duration_days = _job_detail_days_between(start_date, end_date)

    original_value = _job_detail_money(_job_detail_first(job, ["original_quoted_cost", "original_value", "originalValue", "quoted_cost", "quotedCost", "contract_value", "value"], 0))

    variations = await db.variations.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(1000)
    approved_variations = []
    pending_variations = []
    for variation in variations:
        status = str(variation.get("status", "")).strip().lower()
        if status in ["approved", "accepted", "client approved", "client_approved"]:
            approved_variations.append(variation)
        elif status not in ["rejected", "declined", "cancelled", "canceled"]:
            pending_variations.append(variation)

    def variation_value(variation):
        return _job_detail_money(_job_detail_first(variation, ["net_total", "value", "variation_value", "total", "total_value", "amount"], 0))

    approved_variation_value = round(sum(variation_value(v) for v in approved_variations), 2)
    pending_variation_value = round(sum(variation_value(v) for v in pending_variations), 2)
    revised_value = round(original_value + approved_variation_value, 2)

    gantt_sections = job.get("gantt_sections") or []
    commercial_markers = job.get("commercial_markers") or []
    programme_sections = [section for section in gantt_sections if not _job_detail_is_commercial_placeholder(section)]
    commercial_placeholder_sections = [section for section in gantt_sections if _job_detail_is_commercial_placeholder(section)]

    def _section_value(section):
        explicit = _job_detail_money(_job_detail_first(section, ["section_value", "value", "amount", "net_value"], 0))
        if explicit > 0:
            return explicit
        return round(
            _job_detail_money(section.get("labour_value", 0))
            + _job_detail_money(section.get("material_value", 0))
            + _job_detail_money(section.get("subcontractor_value", 0))
            + _job_detail_money(section.get("other_value", 0)),
            2,
        )

    commercial_placeholder_value = round(sum(_section_value(section) for section in commercial_placeholder_sections), 2)

    programme_sections_with_dates = [section for section in programme_sections if _job_detail_first(section, ["start_date", "startDate", "start"], None) and _job_detail_first(section, ["end_date", "endDate", "end"], None)]
    if not programme_sections and not commercial_placeholder_sections:
        planning_status = "missing"
        planning_status_label = "No programme"
    elif not programme_sections and commercial_placeholder_sections:
        planning_status = "partial"
        planning_status_label = "Commercial placeholder only"
    elif len(programme_sections_with_dates) < len(programme_sections):
        planning_status = "partial"
        planning_status_label = "Programme incomplete"
    else:
        planning_status = "assigned"
        planning_status_label = "Operational programme ready"

    section_dates = []
    for section in programme_sections:
        section_start = _job_detail_parse_date(_job_detail_first(section, ["start_date", "startDate", "start"], None))
        section_end = _job_detail_parse_date(_job_detail_first(section, ["end_date", "endDate", "end"], None))
        if section_start:
            section_dates.append(section_start)
        if section_end:
            section_dates.append(section_end)

    schedule_entries = await db.schedule_entries.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    time_entries = await db.time_entries.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    material_entries = await db.materials.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    purchase_orders = await db.purchase_orders.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    work_orders = await db.work_orders.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)

    # Workforce is now a three-stage status:
    # missing   = no workforce allocated to the programme
    # allocated = workers selected/allocated to Gantt sections but not pushed/assigned to schedule
    # assigned  = workforce has been pushed/assigned to the schedule
    allocated_worker_ids = set()
    allocated_worker_names = set()
    assigned_section_count = 0

    for section in programme_sections:
        for key in ["assigned_worker_ids", "assignedWorkerIds", "worker_ids", "workerIds", "allocated_worker_ids", "allocatedWorkerIds"]:
            values = section.get(key)
            if isinstance(values, list):
                for value in values:
                    if value:
                        allocated_worker_ids.add(str(value))

        for key in ["assigned_workers", "assignedWorkers", "workers", "allocated_workers", "allocatedWorkers"]:
            values = section.get(key)
            if isinstance(values, list):
                for value in values:
                    if isinstance(value, dict):
                        worker_id = value.get("id") or value.get("worker_id") or value.get("workerId")
                        worker_name = value.get("name") or value.get("worker_name") or value.get("workerName")
                        if worker_id:
                            allocated_worker_ids.add(str(worker_id))
                        if worker_name:
                            allocated_worker_names.add(str(worker_name))
                    elif value:
                        allocated_worker_ids.add(str(value))

        schedule_status = str(section.get("schedule_status") or section.get("scheduleStatus") or "").strip().lower()
        if section.get("sent_to_schedule") is True or section.get("sentToSchedule") is True or schedule_status in ["scheduled", "partial_scheduled", "assigned"]:
            assigned_section_count += 1

    scheduled_worker_ids = {str(entry.get("worker_id")) for entry in schedule_entries if entry.get("worker_id")}
    worker_ids = sorted(allocated_worker_ids.union(scheduled_worker_ids))
    workers = await db.workers.find({"id": {"$in": worker_ids}}, {"_id": 0}).to_list(1000) if worker_ids else []
    worker_lookup = {worker.get("id"): worker for worker in workers}
    worker_names = sorted({worker_lookup.get(worker_id, {}).get("name", worker_id) for worker_id in worker_ids if worker_id}.union(allocated_worker_names))

    has_workforce_allocated = len(allocated_worker_ids) > 0 or len(allocated_worker_names) > 0
    has_workforce_assigned = len(schedule_entries) > 0 or assigned_section_count > 0
    if has_workforce_assigned:
        workforce_status = "assigned"
        workforce_status_label = "Assigned"
    elif has_workforce_allocated:
        workforce_status = "allocated"
        workforce_status_label = "Allocated, not assigned"
    else:
        workforce_status = "missing"
        workforce_status_label = "Missing"

    # Job management is also a three-stage status:
    # missing = neither Manager nor Supervisor assigned
    # partial = either Manager or Supervisor assigned
    # assigned = both Manager and Supervisor assigned
    manager_id = _job_detail_first(job, ["manager_id", "managerId", "project_manager_id", "projectManagerId"], "")
    manager_name = _job_detail_first(job, ["manager_name", "managerName", "project_manager_name", "projectManagerName"], "")
    supervisor_id = _job_detail_first(job, ["supervisor_id", "supervisorId"], "")
    supervisor_name = _job_detail_first(job, ["supervisor_name", "supervisorName"], "")

    def _assigned_person(person_id, person_name):
        name_value = str(person_name or "").strip().lower()
        return bool(str(person_id or "").strip()) or (bool(name_value) and name_value not in ["unassigned", "none", "not set", "-"])

    has_manager_assigned = _assigned_person(manager_id, manager_name)
    has_supervisor_assigned = _assigned_person(supervisor_id, supervisor_name)

    if has_manager_assigned and has_supervisor_assigned:
        management_status = "assigned"
        management_status_label = "Manager & Supervisor assigned"
    elif has_manager_assigned or has_supervisor_assigned:
        management_status = "partial"
        management_status_label = "Partially assigned"
    else:
        management_status = "missing"
        management_status_label = "Missing"

    has_management = management_status == "assigned"

    material_forecast_value = 0.0
    for marker in commercial_markers:
        marker_type = str(marker.get("type") or marker.get("category") or marker.get("label") or "").lower()
        if "material" in marker_type:
            material_forecast_value += _job_detail_money(_job_detail_first(marker, ["value", "amount", "net_value", "forecast_value"], 0))

    actual_material_value = round(sum(_job_detail_money(item.get("cost", 0)) * _job_detail_money(item.get("quantity", 1)) for item in material_entries), 2)
    if material_forecast_value == 0 and actual_material_value > 0:
        material_forecast_value = actual_material_value

    po_net_value = round(sum(_job_detail_money(_job_detail_first(po, ["net_total", "net_value", "netValue", "subtotal", "amount", "value"], 0)) for po in purchase_orders), 2)
    po_gross_value = round(sum(_job_detail_money(_job_detail_first(po, ["gross_total", "gross_value", "grossValue", "total", "grand_total"], 0)) for po in purchase_orders), 2)
    if po_gross_value == 0 and po_net_value:
        po_gross_value = round(po_net_value * 1.2, 2)

    po_status_counts = {}
    for po in purchase_orders:
        status = str(po.get("status") or "Unknown").strip() or "Unknown"
        po_status_counts[status] = po_status_counts.get(status, 0) + 1

    has_program = planning_status == "assigned"
    has_commercial_markers = len(commercial_markers) > 0 or len(commercial_placeholder_sections) > 0
    has_workforce = has_workforce_assigned
    has_materials_forecast = material_forecast_value > 0 or len(material_entries) > 0
    has_purchase_orders = len(purchase_orders) > 0
    has_work_orders = len(work_orders) > 0
    checklist = [has_management, has_program, has_commercial_markers, has_workforce, has_materials_forecast, has_purchase_orders]

    return {
        "job": {
            "id": job_id,
            "name": job_name,
            "client": job.get("client_name") or job.get("client", ""),
            "location": job.get("location", ""),
            "status": job.get("status", ""),
            "division": _job_detail_first(job, ["division", "job_division", "jobDivision"], ""),
            "manager_id": manager_id,
            "manager_name": manager_name,
            "supervisor_id": supervisor_id,
            "supervisor_name": supervisor_name,
            "start_date": _job_detail_date_to_iso(start_date),
            "end_date": _job_detail_date_to_iso(end_date),
            "duration_days": duration_days,
            "original_value": original_value,
            "approved_variation_value": approved_variation_value,
            "pending_variation_value": pending_variation_value,
            "revised_value": revised_value,
        },
        "checks": {
            "setup_score": f"{sum(1 for item in checklist if item)}/{len(checklist)}",
            "has_management": has_management,
            "has_manager_assigned": has_manager_assigned,
            "has_supervisor_assigned": has_supervisor_assigned,
            "management_status": management_status,
            "management_status_label": management_status_label,
            "has_program": has_program,
            "planning_status": planning_status,
            "planning_status_label": planning_status_label,
            "has_commercial_markers": has_commercial_markers,
            "has_workforce": has_workforce,
            "has_workforce_allocated": has_workforce_allocated,
            "has_workforce_assigned": has_workforce_assigned,
            "workforce_status": workforce_status,
            "workforce_status_label": workforce_status_label,
            "has_materials_forecast": has_materials_forecast,
            "has_purchase_orders": has_purchase_orders,
            "has_work_orders": has_work_orders,
        },
        "management": {
            "manager_id": manager_id,
            "manager_name": manager_name,
            "supervisor_id": supervisor_id,
            "supervisor_name": supervisor_name,
            "has_manager_assigned": has_manager_assigned,
            "has_supervisor_assigned": has_supervisor_assigned,
            "status": management_status,
            "status_label": management_status_label,
        },
        "programme": {
            "sections_count": len(programme_sections),
            "total_gantt_items_count": len(gantt_sections),
            "commercial_placeholder_count": len(commercial_placeholder_sections),
            "commercial_placeholder_value": commercial_placeholder_value,
            "planning_status": planning_status,
            "planning_status_label": planning_status_label,
            "start_date": min(section_dates).isoformat() if section_dates else None,
            "end_date": max(section_dates).isoformat() if section_dates else None,
        },
        "commercial": {
            "commercial_markers_count": len(commercial_markers),
            "commercial_placeholder_count": len(commercial_placeholder_sections),
            "commercial_placeholder_value": commercial_placeholder_value,
            "commercial_marker_value": round(sum(_job_detail_money(_job_detail_first(marker, ["value", "amount", "net_value", "application_value"], 0)) for marker in commercial_markers), 2),
            "approved_variations_count": len(approved_variations),
            "approved_variation_value": approved_variation_value,
            "pending_variations_count": len(pending_variations),
            "pending_variation_value": pending_variation_value,
        },
        "workforce": {
            "worker_count": len(worker_names),
            "allocated_worker_count": len(allocated_worker_ids) + len(allocated_worker_names),
            "scheduled_worker_count": len(scheduled_worker_ids),
            "scheduled_entries_count": len(schedule_entries),
            "time_entries_count": len(time_entries),
            "status": workforce_status,
            "status_label": workforce_status_label,
            "workers": worker_names,
        },
        "materials": {
            "forecast_count": len([marker for marker in commercial_markers if "material" in str(marker.get("type") or marker.get("category") or marker.get("label") or "").lower()]),
            "forecast_value": round(material_forecast_value, 2),
            "actual_material_entries_count": len(material_entries),
            "actual_material_value": actual_material_value,
        },
        "purchase_orders": {
            "po_count": len(purchase_orders),
            "po_net_value": po_net_value,
            "po_gross_value": po_gross_value,
            "status_counts": po_status_counts,
        },
        "work_orders": {
            "wo_count": len(work_orders),
            "wo_net_value": round(sum(_job_detail_money(wo.get("net_amount", 0)) for wo in work_orders), 2),
            "wo_gross_value": round(sum(_job_detail_money(wo.get("gross_amount", 0)) for wo in work_orders), 2),
            "committed_net_value": round(sum(_job_detail_money(wo.get("net_amount", 0)) for wo in work_orders if str(wo.get("status", "")).lower() in WORK_ORDER_COMMITTED_STATUSES), 2),
        },
    }


@api_router.get("/jobs/{job_id}", response_model=Job)
async def get_job(job_id: str):
    job = await db.jobs.find_one({"id": job_id})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return Job(**job)

@api_router.put("/jobs/{job_id}", response_model=Job)
async def update_job(job_id: str, job_update: JobUpdate, admin: str = Depends(verify_admin)):
    """Update job (Admin only)"""
    update_dict = {k: v for k, v in job_update.dict().items() if v is not None}

    if "vat_treatment" in update_dict or "drc_enabled" in update_dict or "vat_rate" in update_dict:
        treatment = normalise_vat_treatment(update_dict.get("vat_treatment"))
        if treatment == "drc" or update_dict.get("drc_enabled") is True:
            update_dict["vat_treatment"] = "drc"
            update_dict["drc_enabled"] = True
            update_dict["vat_rate"] = finance_to_number(update_dict.get("vat_rate"), 20.0) or 20.0
        else:
            update_dict["vat_treatment"] = treatment
            update_dict["drc_enabled"] = False
            if treatment == "reduced_5":
                update_dict["vat_rate"] = finance_to_number(update_dict.get("vat_rate"), 5.0) or 5.0
            elif treatment in ["zero_rated", "exempt", "no_vat"]:
                update_dict["vat_rate"] = 0.0
            else:
                update_dict["vat_rate"] = finance_to_number(update_dict.get("vat_rate"), 20.0) or 20.0

    if "cis_enabled" in update_dict or "cis_rate" in update_dict or "cis_deduction_basis" in update_dict:
        update_dict["cis_enabled"] = bool(update_dict.get("cis_enabled"))
        update_dict["cis_rate"] = finance_to_number(update_dict.get("cis_rate"), 0.0) if update_dict["cis_enabled"] else 0.0
        if update_dict.get("cis_deduction_basis") not in ["labour_only", "full_net", "none"]:
            update_dict["cis_deduction_basis"] = "labour_only"

    result = await db.jobs.update_one({"id": job_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job not found")

    updated_job = await db.jobs.find_one({"id": job_id})
    return Job(**updated_job)

@api_router.delete("/jobs/{job_id}")
async def delete_job(job_id: str, admin: str = Depends(verify_admin)):
    """Delete job (Admin only)"""
    result = await db.jobs.delete_one({"id": job_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"message": "Job deleted successfully"}

@api_router.put("/jobs/{job_id}/archive")
async def archive_job(job_id: str, admin: str = Depends(verify_admin)):
    """Archive job (Admin only)"""
    result = await db.jobs.update_one({"id": job_id}, {"$set": {"archived": True}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"message": "Job archived successfully"}

@api_router.put("/jobs/{job_id}/unarchive")
async def unarchive_job(job_id: str, admin: str = Depends(verify_admin)):
    """Unarchive job (Admin only)"""
    result = await db.jobs.update_one({"id": job_id}, {"$set": {"archived": False}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"message": "Job unarchived successfully"}




@api_router.get("/jobs/{job_id}/tax-settings")
async def get_job_tax_settings(job_id: str, admin: str = Depends(verify_admin)):
    """Return a job's VAT / DRC / CIS settings for Finance and Project Management."""
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return normalise_job_tax_settings(job)


@api_router.put("/jobs/{job_id}/tax-settings")
async def update_job_tax_settings(job_id: str, settings: Dict[str, Any], admin: str = Depends(verify_admin)):
    """Update VAT / DRC / CIS treatment for a job without touching programme or commercial data."""
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    treatment = normalise_vat_treatment(settings.get("vat_treatment"))
    drc_enabled = bool(settings.get("drc_enabled")) or treatment == "drc"
    if drc_enabled:
        treatment = "drc"
        vat_rate = finance_to_number(settings.get("vat_rate"), 20.0) or 20.0
    elif treatment == "reduced_5":
        vat_rate = finance_to_number(settings.get("vat_rate"), 5.0) or 5.0
    elif treatment in ["zero_rated", "exempt", "no_vat"]:
        vat_rate = 0.0
    else:
        treatment = "standard_20"
        vat_rate = finance_to_number(settings.get("vat_rate"), 20.0) or 20.0

    cis_enabled = bool(settings.get("cis_enabled"))
    cis_rate = finance_to_number(settings.get("cis_rate"), 20.0 if cis_enabled else 0.0) if cis_enabled else 0.0
    cis_basis = str(settings.get("cis_deduction_basis") or "labour_only").strip().lower()
    if cis_basis not in ["labour_only", "full_net", "none"]:
        cis_basis = "labour_only"

    patch = {
        "vat_treatment": treatment,
        "vat_rate": round(vat_rate, 2),
        "drc_enabled": drc_enabled,
        "cis_enabled": cis_enabled,
        "cis_rate": max(0.0, min(100.0, cis_rate)),
        "cis_deduction_basis": cis_basis,
        "tax_notes": str(settings.get("tax_notes") or ""),
        "updated_at": datetime.utcnow(),
    }
    await db.jobs.update_one({"id": job_id}, {"$set": patch})
    updated = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    return {"success": True, "job": updated, "tax_settings": normalise_job_tax_settings(updated)}

def get_job_drive_folder_id(job: Dict[str, Any]) -> Optional[str]:
    """Return the Google Drive folder ID stored against a job."""
    if not job:
        return None

    possible_values = [
        job.get("drive_folder_id"),
        job.get("drive_folder_link"),
        job.get("drive_folder_url"),
        job.get("google_drive_link"),
    ]

    for value in possible_values:
        if not value:
            continue
        if GoogleDriveService and hasattr(GoogleDriveService, "extract_folder_id"):
            folder_id = GoogleDriveService.extract_folder_id(value)
        else:
            folder_id = str(value).strip()
        if folder_id:
            return folder_id
    return None


@api_router.post("/jobs/{job_id}/post-work-photos")
async def upload_job_post_work_photos(
    job_id: str,
    files: List[UploadFile] = File(...),
    worker_id: Optional[str] = Query(None),
):
    """Upload worker post-works photos into the job Google Drive folder.

    Target folder inside the job folder:
    007: Site Deliverables / 04: Site Images Pre & Post / 02: Post Works Images
    """
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if not files:
        raise HTTPException(status_code=400, detail="Please select at least one photo")

    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Maximum 20 photos can be uploaded at once")

    job_folder_id = get_job_drive_folder_id(job)
    if not job_folder_id:
        raise HTTPException(
            status_code=400,
            detail="This job does not have a Google Drive folder link/id saved against it yet",
        )

    if not get_google_drive_service:
        raise HTTPException(status_code=500, detail="Google Drive service is not configured on the backend")

    drive_service = get_google_drive_service()
    if not drive_service:
        raise HTTPException(status_code=500, detail="Google Drive service could not be initialised")

    worker = None
    if worker_id:
        worker = await db.workers.find_one({"id": worker_id})

    uploaded_photos = []
    errors = []

    for file in files:
        try:
            if not file.content_type or not file.content_type.startswith("image/"):
                errors.append({"filename": file.filename, "error": "Only image files are allowed"})
                continue

            content = await file.read()
            if len(content) > 15 * 1024 * 1024:
                errors.append({"filename": file.filename, "error": "Photo is larger than 15MB"})
                continue

            uploaded = drive_service.upload_post_work_image(
                job_folder_id=job_folder_id,
                file_content=content,
                filename=file.filename,
                content_type=file.content_type,
                worker_name=(worker or {}).get("name", ""),
                job_name=job.get("name", ""),
            )

            photo_record = {
                "id": str(uuid.uuid4()),
                "job_id": job_id,
                "job_name": job.get("name", ""),
                "worker_id": worker_id,
                "worker_name": (worker or {}).get("name", ""),
                "original_filename": file.filename,
                "file_id": uploaded.get("file_id"),
                "filename": uploaded.get("filename"),
                "mime_type": uploaded.get("mime_type"),
                "share_url": uploaded.get("share_url"),
                "direct_url": uploaded.get("direct_url", ""),
                "folder_id": uploaded.get("folder_id"),
                "folder_name": uploaded.get("folder_name"),
                "folder_link": uploaded.get("folder_link"),
                "category": "post_works",
                "uploaded_at": datetime.utcnow(),
            }
            uploaded_photos.append(photo_record)

        except Exception as exc:
            logger.error(f"Error uploading post work photo {file.filename}: {exc}")
            errors.append({"filename": file.filename, "error": str(exc)})

    if not uploaded_photos and errors:
        raise HTTPException(status_code=500, detail={"message": "No photos were uploaded", "errors": errors})

    if uploaded_photos:
        await db.job_photo_uploads.insert_many(uploaded_photos)
        await db.jobs.update_one(
            {"id": job_id},
            {
                "$push": {"post_work_photos": {"$each": uploaded_photos}},
                "$set": {"updated_date": datetime.utcnow()},
            },
        )

    return {
        "message": f"Successfully uploaded {len(uploaded_photos)} photo(s)",
        "photos": uploaded_photos,
        "errors": errors,
    }
# TIME ENTRY ENDPOINTS
@api_router.post("/time-entries/clock-in", response_model=TimeEntry)
async def clock_in(entry: TimeEntryClockIn):
    """Clock a worker into a job. GPS is mandatory only when the job has gps_required=true."""
    worker = await db.workers.find_one({"id": entry.worker_id, "active": True, "archived": {"$ne": True}})
    if not worker:
        raise HTTPException(status_code=404, detail="Worker not found or inactive")

    job = await db.jobs.find_one({"id": entry.job_id, "archived": {"$ne": True}})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job_requires_gps(job) and not worker.get("gps_exempt", False) and entry.gps_location is None:
        raise HTTPException(
            status_code=400,
            detail="Location is required for this job. Please allow location access and try again."
        )

    active_entry = await db.time_entries.find_one({
        "worker_id": entry.worker_id,
        "clock_out": None
    })

    if active_entry:
        raise HTTPException(status_code=400, detail="Worker already clocked in. Must clock out first.")

    time_entry_dict = entry.dict()
    time_entry_dict["clock_in"] = datetime.utcnow()
    time_entry_dict["gps_location_in"] = entry.gps_location.dict() if entry.gps_location else None
    time_entry_dict["device_id_in"] = entry.device_id
    time_entry_dict["suspicious_flags"] = await build_suspicious_flags(entry.worker_id, entry.job_id, entry.device_id, entry.gps_location)
    time_entry_dict.pop("gps_location", None)
    time_entry_dict.pop("device_id", None)

    time_entry_obj = TimeEntry(**time_entry_dict)
    await db.time_entries.insert_one(time_entry_obj.dict())
    return time_entry_obj

@api_router.get("/time-entries")
async def get_time_entries(
    worker_id: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Return time entries, enriched defensively so missing worker/job records do not crash the app."""
    filter_dict = {}

    if worker_id:
        filter_dict["worker_id"] = worker_id
    if job_id:
        filter_dict["job_id"] = job_id

    if start_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        filter_dict["clock_in"] = {"$gte": start_dt}

    if end_date:
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        if "clock_in" in filter_dict:
            filter_dict["clock_in"]["$lte"] = end_dt
        else:
            filter_dict["clock_in"] = {"$lte": end_dt}

    entries = await db.time_entries.find(filter_dict, {"_id": 0}).sort("clock_in", -1).to_list(1000)
    workers = await db.workers.find({}, {"_id": 0}).to_list(1000)
    jobs = await db.jobs.find({}, {"_id": 0}).to_list(1000)

    worker_lookup = {w.get("id"): w for w in workers if w.get("id")}
    job_lookup = {j.get("id"): j for j in jobs if j.get("id")}

    result = []
    for entry_doc in entries:
        worker = worker_lookup.get(entry_doc.get("worker_id"), {})
        job = job_lookup.get(entry_doc.get("job_id"), {})

        if entry_doc.get("clock_in"):
            entry_doc["clock_in"] = utc_to_uk(entry_doc["clock_in"])
        if entry_doc.get("clock_out"):
            entry_doc["clock_out"] = utc_to_uk(entry_doc["clock_out"])

        duration_hours = (entry_doc.get("duration_minutes", 0) or 0) / 60
        hourly_rate = worker.get("hourly_rate", 15.0) or 15.0

        result.append({
            **entry_doc,
            "worker_name": worker.get("name", "Unknown Worker"),
            "worker_type": worker.get("worker_type") or worker.get("role") or "worker",
            "job_name": job.get("name", "Unknown Job"),
            "job_client": job.get("client", ""),
            "job_location": job.get("location", ""),
            "cost": round(duration_hours * hourly_rate, 2),
        })

    return result

@api_router.put("/time-entries/{entry_id}/clock-out", response_model=TimeEntry)
async def clock_out(entry_id: str, clock_out_data: TimeEntryClockOut):
    # Find the active time entry
    time_entry = await db.time_entries.find_one({"id": entry_id, "clock_out": None})
    if not time_entry:
        raise HTTPException(status_code=404, detail="Active time entry not found")

    job = await db.jobs.find_one({"id": time_entry.get("job_id"), "archived": {"$ne": True}}) or {}
    worker = await db.workers.find_one({"id": time_entry.get("worker_id")}) or {}
    if job_requires_gps(job) and not worker.get("gps_exempt", False) and clock_out_data.gps_location is None:
        raise HTTPException(
            status_code=400,
            detail="Location is required to clock out for this job. Please allow location access and try again."
        )

    clock_out_time = datetime.utcnow()
    clock_in_time = time_entry["clock_in"]
    duration = calculate_duration(clock_in_time, clock_out_time)

    update_dict = {
        "clock_out": clock_out_time,
        "duration_minutes": duration,
        "gps_location_out": clock_out_data.gps_location.dict() if clock_out_data.gps_location else None,
        "device_id_out": clock_out_data.device_id,
        "suspicious_flags": await build_suspicious_flags(
            time_entry.get("worker_id"),
            time_entry.get("job_id"),
            clock_out_data.device_id,
            clock_out_data.gps_location,
            time_entry.get("suspicious_flags", [])
        ),
        "notes": clock_out_data.notes
    }

    await db.time_entries.update_one({"id": entry_id}, {"$set": update_dict})

    updated_entry = await db.time_entries.find_one({"id": entry_id})
    return TimeEntry(**updated_entry)

@api_router.put("/time-entries/{entry_id}", response_model=TimeEntry)
async def update_time_entry(entry_id: str, entry_update: TimeEntryUpdate, admin: str = Depends(verify_admin)):
    """Update time entry (Admin only)"""
    # Find the existing time entry
    existing_entry = await db.time_entries.find_one({"id": entry_id})
    if not existing_entry:
        raise HTTPException(status_code=404, detail="Time entry not found")

    # Prepare update dictionary
    update_dict = {}

    # Update allowed fields
    if entry_update.worker_id:
        update_dict["worker_id"] = entry_update.worker_id
    if entry_update.job_id:
        update_dict["job_id"] = entry_update.job_id
    if entry_update.clock_in:
        update_dict["clock_in"] = datetime.fromisoformat(entry_update.clock_in.replace('Z', '+00:00'))
    if entry_update.clock_out is not None:
        if entry_update.clock_out:
            update_dict["clock_out"] = datetime.fromisoformat(entry_update.clock_out.replace('Z', '+00:00'))
        else:
            update_dict["clock_out"] = None
            update_dict["duration_minutes"] = None
    if entry_update.notes is not None:
        update_dict["notes"] = entry_update.notes

    # Recalculate duration if both clock_in and clock_out are present
    if entry_update.duration_minutes is not None:
        update_dict["duration_minutes"] = entry_update.duration_minutes
    elif "clock_in" in update_dict and "clock_out" in update_dict and update_dict["clock_out"]:
        clock_in = update_dict["clock_in"] if "clock_in" in update_dict else existing_entry["clock_in"]
        clock_out = update_dict["clock_out"]
        update_dict["duration_minutes"] = calculate_duration(clock_in, clock_out)

    # Update the time entry
    result = await db.time_entries.update_one({"id": entry_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Time entry not found")

    # Return updated entry
    updated_entry = await db.time_entries.find_one({"id": entry_id})
    return TimeEntry(**updated_entry)

@api_router.get("/workers/{worker_id}/active-entry")
async def get_active_time_entry(worker_id: str):
    active_entry = await db.time_entries.find_one({
        "worker_id": worker_id,
        "clock_out": None
    })

    if not active_entry:
        return {"active_entry": None}

    return {"active_entry": TimeEntry(**active_entry)}

# MATERIAL ENDPOINTS
@api_router.post("/materials", response_model=Material)
async def create_material(material: MaterialCreate):
    material_dict = material.dict()
    material_obj = Material(**material_dict)
    await db.materials.insert_one(material_obj.dict())
    return material_obj

@api_router.get("/materials", response_model=List[Material])
async def get_materials(job_id: Optional[str] = Query(None)):
    filter_dict = {"job_id": job_id} if job_id else {}
    materials = await db.materials.find(filter_dict).to_list(1000)
    return [Material(**material) for material in materials]

@api_router.put("/materials/{material_id}", response_model=Material)
async def update_material(material_id: str, material_update: MaterialUpdate, admin: str = Depends(verify_admin)):
    """Update material (Admin only)"""
    update_dict = {k: v for k, v in material_update.dict().items() if v is not None}

    result = await db.materials.update_one({"id": material_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Material not found")

    updated_material = await db.materials.find_one({"id": material_id})
    return Material(**updated_material)

@api_router.delete("/materials/{material_id}")
async def delete_material(material_id: str, admin: str = Depends(verify_admin)):
    """Delete material (Admin only)"""
    result = await db.materials.delete_one({"id": material_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Material not found")
    return {"message": "Material deleted successfully"}

# REPORTING ENDPOINTS
@api_router.get("/reports/time-entries", response_model=List[Dict[str, Any]])
async def get_time_entries_report(
    worker_id: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Get time entries report with filters (Admin only)"""
    # Build filter query
    filter_query = {"archived": {"$ne": True}}

    if worker_id:
        filter_query["worker_id"] = worker_id
    if job_id:
        filter_query["job_id"] = job_id

    if start_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        filter_query["clock_in"] = {"$gte": start_dt}

    if end_date:
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        if "clock_in" in filter_query:
            filter_query["clock_in"]["$lte"] = end_dt
        else:
            filter_query["clock_in"] = {"$lte": end_dt}

    # Get time entries
    time_entries = await db.time_entries.find(filter_query).to_list(1000)

    # Get all jobs and workers for lookup
    jobs = await db.jobs.find().to_list(1000)
    workers = await db.workers.find().to_list(1000)

    # Create lookup dictionaries
    job_lookup = {job["id"]: job for job in jobs if "id" in job}
    worker_lookup = {worker["id"]: worker for worker in workers if "id" in worker}

    # Process time entries for report
    result = []
    for entry in time_entries:
        worker = worker_lookup.get(entry["worker_id"])
        job = job_lookup.get(entry["job_id"])

        if not worker or not job:
            continue

        # Calculate labor cost
        duration_hours = (entry.get("duration_minutes", 0) or 0) / 60
        hourly_rate = worker.get("hourly_rate", 15.0)
        labor_cost = duration_hours * hourly_rate

        # Convert times to UK timezone for consistent display
        clock_in_uk = utc_to_uk(entry["clock_in"]) if entry.get("clock_in") else None
        clock_out_uk = utc_to_uk(entry["clock_out"]) if entry.get("clock_out") else None

        # Format the time entry data for report
        result.append({
            "id": entry["id"],
            "worker_id": entry["worker_id"],
            "worker_name": worker["name"],
            "job_id": entry["job_id"],
            "job_name": job["name"],
            "job_client": job.get("client", ""),
            "clock_in": clock_in_uk.isoformat() if clock_in_uk else None,
            "clock_out": clock_out_uk.isoformat() if clock_out_uk else None,
            "duration_minutes": entry.get("duration_minutes", 0),
            "hourly_rate": hourly_rate,
            "labor_cost": labor_cost,
            "notes": entry.get("notes", ""),
            "gps_address_in": entry.get("gps_location_in", {}).get("address", "") if entry.get("gps_location_in") else "",
            "gps_address_out": entry.get("gps_location_out", {}).get("address", "") if entry.get("gps_location_out") else "",
            "archived": entry.get("archived", False)
        })

    # Sort by clock_in date (most recent first)
    result.sort(key=lambda x: x["clock_in"] if x["clock_in"] else "1900-01-01", reverse=True)

    return result

@api_router.get("/reports/dashboard")
async def get_dashboard_stats(admin: str = Depends(verify_admin)):
    """Get dashboard statistics (Admin only)"""
    # Get basic counts
    total_workers = await db.workers.count_documents({"active": True, "archived": {"$ne": True}})
    total_jobs = await db.jobs.count_documents({"status": {"$ne": "cancelled"}, "archived": {"$ne": True}})
    active_jobs = await db.jobs.count_documents({"status": "active", "archived": {"$ne": True}})

    # Get total hours this week
    week_start = datetime.utcnow() - timedelta(days=7)
    week_entries = await db.time_entries.find({
        "clock_in": {"$gte": week_start},
        "duration_minutes": {"$exists": True}
    }).to_list(1000)

    total_minutes = sum(entry.get("duration_minutes", 0) or 0 for entry in week_entries)
    total_hours = round(total_minutes / 60, 1)

    # Get total materials cost this month (UK timezone)
    uk_now = get_uk_time()
    month_start_uk = uk_now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    month_start_utc = uk_to_utc(month_start_uk)

    month_materials = await db.materials.find({
        "purchase_date": {"$gte": month_start_utc}
    }).to_list(1000)

    total_materials_cost = sum(mat.get("cost", 0) * mat.get("quantity", 1) for mat in month_materials)

    # Get attendance alerts (workers who haven't logged in before 9am or out after 5pm) - last 7 days
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    nine_am_threshold = timedelta(hours=9)  # 9 AM
    five_pm_threshold = timedelta(hours=17)  # 5 PM

    # Get all non-admin workers
    non_admin_workers = await db.workers.find({
        "role": {"$ne": "admin"},
        "active": True,
        "archived": {"$ne": True}
    }).to_list(1000)

    attendance_alerts = []

    for worker in non_admin_workers:
        # Check each day for the last 7 days
        for i in range(7):
            day_start = (datetime.utcnow() - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            nine_am = day_start + nine_am_threshold
            five_pm = day_start + five_pm_threshold

            # Skip future dates and today if it's before 9 AM
            if day_start.date() == datetime.utcnow().date() and datetime.utcnow().hour < 9:
                continue
            if day_start.date() > datetime.utcnow().date():
                continue

            # Get all time entries for this worker on this day
            day_entries = await db.time_entries.find({
                "worker_id": worker["id"],
                "clock_in": {"$gte": day_start, "$lt": day_end}
            }).to_list(100)

            if not day_entries:
                # No entries for this day (only alert if it's a weekday and not today)
                if day_start.weekday() < 5 and day_start.date() != datetime.utcnow().date():  # Monday=0, Friday=4
                    attendance_alerts.append({
                        "worker_id": worker["id"],
                        "worker_name": worker["name"],
                        "type": "no_clock_in",
                        "date": day_start.date(),
                        "time": None,
                        "message": f"No clock in recorded on {day_start.strftime('%A, %d %B %Y')}"
                    })
            else:
                # Check for late clock-ins
                for entry in day_entries:
                    clock_in_time = entry["clock_in"]

                    # Late clock-in (after 9 AM)
                    if clock_in_time > nine_am:
                        attendance_alerts.append({
                            "worker_id": worker["id"],
                            "worker_name": worker["name"],
                            "type": "late_clock_in",
                            "date": day_start.date(),
                            "time": clock_in_time,
                            "message": f"Clocked in late at {clock_in_time.strftime('%H:%M')} on {day_start.strftime('%A, %d %B %Y')}"
                        })

                    # Late clock-out (after 5 PM) - only check if clocked out
                    if entry.get("clock_out") and entry["clock_out"] > five_pm:
                        attendance_alerts.append({
                            "worker_id": worker["id"],
                            "worker_name": worker["name"],
                            "type": "late_clock_out",
                            "date": day_start.date(),
                            "time": entry["clock_out"],
                            "message": f"Clocked out late at {entry['clock_out'].strftime('%H:%M')} on {day_start.strftime('%A, %d %B %Y')}"
                        })

    # Sort alerts by date (most recent first)
    attendance_alerts.sort(key=lambda x: x["date"] if x["date"] else datetime.min.date(), reverse=True)

    return {
        "total_workers": total_workers,
        "total_jobs": total_jobs,
        "active_jobs": active_jobs,
        "total_hours_this_week": total_hours,
        "total_materials_cost_this_month": total_materials_cost,
        "attendance_alerts": attendance_alerts
    }
@api_router.get("/reports/live-map")
async def get_live_worker_map(admin: str = Depends(verify_admin)):
    """Return currently clocked-in workers with captured GPS locations for the dashboard map."""
    return await get_activity_map(active_only=True, admin=admin)

@api_router.get("/reports/activity-map")
async def get_activity_map(
    active_only: bool = Query(False),
    worker_id: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    date: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Return clock-in/out map markers with date, worker and job filters."""
    filter_dict = {}
    if active_only:
        filter_dict["clock_out"] = None
    if worker_id:
        filter_dict["worker_id"] = worker_id
    if job_id:
        filter_dict["job_id"] = job_id

    if date:
        start_dt = datetime.fromisoformat(date).replace(hour=0, minute=0, second=0, microsecond=0)
        end_dt = start_dt + timedelta(days=1)
        filter_dict["clock_in"] = {"$gte": start_dt, "$lt": end_dt}
    else:
        if start_date:
            start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
            filter_dict["clock_in"] = {"$gte": start_dt}
        if end_date:
            end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
            if "clock_in" in filter_dict:
                filter_dict["clock_in"]["$lte"] = end_dt
            else:
                filter_dict["clock_in"] = {"$lte": end_dt}

    entries = await db.time_entries.find(filter_dict).to_list(1000)
    markers = []

    for entry in entries:
        worker = await db.workers.find_one({"id": entry.get("worker_id")}) or {}
        job = await db.jobs.find_one({"id": entry.get("job_id")}) or {}

        def append_marker(gps, marker_type):
            if not gps:
                return
            latitude = gps.get("latitude")
            longitude = gps.get("longitude")
            if latitude is None or longitude is None:
                return
            markers.append({
                "entry_id": entry.get("id"),
                "worker_id": entry.get("worker_id"),
                "worker_name": worker.get("name", "Unknown worker"),
                "worker_type": worker.get("worker_type", worker.get("role", "worker")),
                "worker_gps_exempt": worker.get("gps_exempt", False),
                "job_id": entry.get("job_id"),
                "job_name": job.get("name", "Unknown job"),
                "job_client": job.get("client", ""),
                "job_location": job.get("location", ""),
                "clock_in": entry.get("clock_in"),
                "clock_out": entry.get("clock_out"),
                "marker_type": marker_type,
                "latitude": latitude,
                "longitude": longitude,
                "accuracy": gps.get("accuracy"),
                "address": gps.get("address", ""),
                "device_id_in": entry.get("device_id_in"),
                "device_id_out": entry.get("device_id_out"),
                "suspicious_flags": entry.get("suspicious_flags", []) or [],
            })

        before_count = len(markers)
        append_marker(entry.get("gps_location_in") or entry.get("gps_location"), "clock_in")
        if not active_only:
            append_marker(entry.get("gps_location_out"), "clock_out")

        # Still return a list record when no GPS exists so the dashboard can show
        # "No GPS recorded" rather than appearing empty. These records are not mapped as pins.
        if len(markers) == before_count:
            markers.append({
                "entry_id": entry.get("id"),
                "worker_id": entry.get("worker_id"),
                "worker_name": worker.get("name", "Unknown worker"),
                "worker_type": worker.get("worker_type", worker.get("role", "worker")),
                "worker_gps_exempt": worker.get("gps_exempt", False),
                "job_id": entry.get("job_id"),
                "job_name": job.get("name", "Unknown job"),
                "job_client": job.get("client", ""),
                "job_location": job.get("location", ""),
                "clock_in": entry.get("clock_in"),
                "clock_out": entry.get("clock_out"),
                "marker_type": "no_gps",
                "latitude": None,
                "longitude": None,
                "accuracy": None,
                "address": "",
                "device_id_in": entry.get("device_id_in"),
                "device_id_out": entry.get("device_id_out"),
                "suspicious_flags": entry.get("suspicious_flags", []) or [],
            })

    markers.sort(key=lambda item: (item.get("worker_name", "").lower(), item.get("marker_type", "")))
    return markers

# Root endpoint for the main app (redirects to API)
@app.get("/")
async def root():
    return {
        "message": "LDA Group Time Tracking System",
        "api_url": "/api/",
        "status": "running",
        "docs": "/docs"
    }

# Handle OPTIONS requests for CORS preflight
@app.options("/{path:path}")
async def options_handler(path: str):
    return {"message": "OK"}

@api_router.get("/")
async def root():
    return {"message": "LDA Group Time Tracking API", "version": "2.0.0"}


# Initialize services
email_service = None
pdf_generator = None

def get_quote_services():
    """Initialize quote services if available"""
    global email_service, pdf_generator
    if EmailService and QuotePDFGenerator:
        if not email_service:
            email_service = EmailService()
        if not pdf_generator:
            pdf_generator = QuotePDFGenerator()
        return email_service, pdf_generator
    return None, None

# ==================== QUOTE SYSTEM ENDPOINTS ====================

@api_router.post("/quotes/{quote_id}/photos")
async def upload_quote_photos(
    quote_id: str,
    files: List[UploadFile] = File(...),
    surveyor_id: str = Depends(get_current_surveyor)
):
    """Upload multiple photos for a quote using Google Drive"""
    if not get_google_drive_service:
        raise HTTPException(status_code=500, detail="Google Drive service not available")

    # Check if quote exists and belongs to surveyor
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    if quote["surveyor_id"] != surveyor_id:
        try:
            verify_admin(None)
        except:
            raise HTTPException(status_code=403, detail="Not authorized to modify this quote")

    # Validate files
    if len(files) > 12:
        raise HTTPException(status_code=400, detail="Maximum 12 photos allowed")

    # Initialize Google Drive service
    drive_service = get_google_drive_service()
    if not drive_service:
        raise HTTPException(status_code=500, detail="Google Drive service initialization failed")

    uploaded_photos = []

    try:
        for file in files:
            # Validate file type
            if not file.content_type or not file.content_type.startswith('image/'):
                raise HTTPException(status_code=400, detail=f"File {file.filename} is not a valid image")

            # Read file content
            file_content = await file.read()

            # Upload to Google Drive
            folder_name = f"Quote_{quote_id}"
            photo_info = drive_service.upload_photo(
                file_content=file_content,
                filename=file.filename,
                folder_name=folder_name
            )

            uploaded_photos.append(photo_info)

    except Exception as e:
        logger.error(f"Error uploading photos: {e}")
        raise HTTPException(status_code=500, detail=f"Error uploading photos: {str(e)}")

    # Update quote with photo information
    await db.quotes.update_one(
        {"id": quote_id},
        {
            "$set": {
                "photos": uploaded_photos,
                "updated_at": datetime.utcnow()
            }
        }
    )

    return {"message": f"Successfully uploaded {len(uploaded_photos)} photos", "photos": uploaded_photos}

@api_router.get("/quotes/{quote_id}/download-pdf")
async def download_quote_pdf(
    quote_id: str,
    surveyor_id: str = Depends(get_current_surveyor)
):
    """Generate and download PDF for a quote"""
    if not QuotePDFGenerator:
        raise HTTPException(status_code=500, detail="PDF generation service not available")

    # Get quote
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    # Check authorization
    if quote["surveyor_id"] != surveyor_id:
        try:
            verify_admin(None)
        except:
            raise HTTPException(status_code=403, detail="Not authorized to access this quote")

    try:
        # Generate PDF
        pdf_generator = QuotePDFGenerator()
        pdf_buffer = pdf_generator.generate_quote_pdf(quote)

        # Return PDF as streaming response
        filename = f"Quote_{quote['quote_number']}.pdf"

        return StreamingResponse(
            io.BytesIO(pdf_buffer.getvalue()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

@api_router.post("/quotes/{quote_id}/mark-sent")
async def mark_quote_sent(
    quote_id: str,
    surveyor_id: str = Depends(get_current_surveyor)
):
    """Mark quote as sent to client (for manual sending workflow)"""
    # Get quote
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    # Check authorization
    if quote["surveyor_id"] != surveyor_id:
        try:
            verify_admin(None)
        except:
            raise HTTPException(status_code=403, detail="Not authorized to modify this quote")

    # Update quote status
    await db.quotes.update_one(
        {"id": quote_id},
        {
            "$set": {
                "status": "sent",
                "sent_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
        }
    )

    return {"message": "Quote marked as sent successfully"}

@api_router.post("/quotes/{quote_id}/client-response")
async def handle_client_response(quote_id: str, response_data: ClientResponse):
    """Handle client response to quote (accept/decline)"""
    # Get quote
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")

    # Update quote with client response
    update_data = {
        "client_response": response_data.response,
        "client_response_date": datetime.utcnow(),
        "client_message": response_data.message,
        "updated_at": datetime.utcnow()
    }

    # Update status based on response
    if response_data.response == "accepted":
        update_data["status"] = "accepted"
    else:
        update_data["status"] = "declined"

    await db.quotes.update_one({"id": quote_id}, {"$set": update_data})

    # Send notification email to info@ldagroup.co.uk
    try:
        email_service, _ = get_quote_services()
        if email_service:
            email_service.send_client_response_notification(quote, response_data)
    except Exception as e:
        logger.warning(f"Failed to send notification email: {e}")

    return {"message": f"Quote {response_data.response} successfully"}

# Startup event
@app.on_event("startup")
async def startup_db_client():
    """Initialize database connection on startup"""
    global db
    try:
        db = await get_database()
        logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize database: {e}")
        raise

# Shutdown event
@app.on_event("shutdown")
async def shutdown_db_client():
    """Close database connection on shutdown"""
    try:
        if db and hasattr(db, 'client'):
            db.client.close()
        logger.info("Database connection closed")
    except Exception as e:
        logger.error(f"Error closing database connection: {e}")

@api_router.get("/reports/materials", response_model=List[Dict[str, Any]])
async def get_materials_report(
    worker_id: Optional[str] = Query(None),
    supplier: Optional[str] = Query(None),
    client: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Get materials report with job/client details and filters (Admin only)"""
    filter_query = {"archived": {"$ne": True}}

    if job_id:
        filter_query["job_id"] = job_id
    if supplier:
        filter_query["supplier"] = {"$regex": supplier, "$options": "i"}
    if worker_id:
        filter_query["worker_id"] = worker_id
    if start_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        filter_query["purchase_date"] = {"$gte": start_dt}
    if end_date:
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        if "purchase_date" in filter_query:
            filter_query["purchase_date"]["$lte"] = end_dt
        else:
            filter_query["purchase_date"] = {"$lte": end_dt}

    materials = await db.materials.find(filter_query).to_list(1000)
    jobs = await db.jobs.find().to_list(1000)
    job_lookup = {job["id"]: job for job in jobs if "id" in job}

    result = []
    for material in materials:
        job = job_lookup.get(material.get("job_id"))
        if not job:
            continue
        if client and client.lower() not in job.get("client", "").lower():
            continue

        purchase_date = utc_to_uk(material.get("purchase_date")) if material.get("purchase_date") else None
        cost = material.get("cost", 0) or 0
        quantity = material.get("quantity", 1) or 1

        result.append({
            "id": material.get("id"),
            "job_id": material.get("job_id"),
            "job_name": job.get("name", "Unknown"),
            "job_client": job.get("client", ""),
            "material_name": material.get("name", ""),
            "name": material.get("name", ""),
            "cost": cost,
            "quantity": quantity,
            "supplier": material.get("supplier", ""),
            "reference": material.get("reference", ""),
            "date": purchase_date.isoformat() if purchase_date else None,
            "purchase_date": purchase_date.isoformat() if purchase_date else None,
            "notes": material.get("notes", ""),
            "total_value": cost * quantity,
            "archived": material.get("archived", False)
        })

    result.sort(key=lambda x: x["date"] if x["date"] else "1900-01-01", reverse=True)
    return result

@api_router.get("/reports/job-costs/{job_id}")
async def get_job_cost_report(job_id: str, admin: str = Depends(verify_admin)):
    """Get detailed job cost report (Admin only)"""
    # Get job details
    job = await db.jobs.find_one({"id": job_id})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Get time entries for this job
    time_entries = await db.time_entries.find({"job_id": job_id}).to_list(1000)
    total_minutes = sum((entry.get("duration_minutes", 0) or 0) for entry in time_entries)
    total_hours = round(total_minutes / 60, 1)

    # Get materials for this job
    materials = await db.materials.find({"job_id": job_id}).to_list(1000)
    total_materials_cost = sum(mat.get("cost", 0) * mat.get("quantity", 1) for mat in materials)

    # Calculate labor cost using worker hourly rates
    labor_cost = 0
    worker_ids = list(set(entry.get("worker_id") for entry in time_entries if entry.get("worker_id")))
    workers = await db.workers.find({"id": {"$in": worker_ids}}).to_list(1000) if worker_ids else []
    worker_rates = {worker["id"]: worker.get("hourly_rate", 15.0) for worker in workers}
    worker_names = {worker["id"]: worker["name"] for worker in workers}

    # Calculate labor cost per entry
    for entry in time_entries:
        if entry.get("duration_minutes"):
            worker_rate = worker_rates.get(entry.get("worker_id"), 15.0)
            entry_hours = entry.get("duration_minutes") / 60
            labor_cost += entry_hours * worker_rate

    total_cost = labor_cost + total_materials_cost
    quoted_cost = job.get("quoted_cost", 0)
    cost_variance = quoted_cost - total_cost

    # Clean data for response
    clean_time_entries = []
    for entry in time_entries:
        clean_entry = {
            "id": entry.get("id"),
            "worker_id": entry.get("worker_id"),
            "worker_name": worker_names.get(entry.get("worker_id"), "Unknown"),
            "clock_in": entry.get("clock_in"),
            "clock_out": entry.get("clock_out"),
            "duration_minutes": entry.get("duration_minutes"),
            "notes": entry.get("notes", ""),
            "hourly_rate": worker_rates.get(entry.get("worker_id"), 15.0)
        }
        clean_time_entries.append(clean_entry)

    clean_materials = []
    for material in materials:
        clean_material = {
            "id": material.get("id"),
            "name": material.get("name"),
            "cost": material.get("cost"),
            "quantity": material.get("quantity"),
            "supplier": material.get("supplier", ""),
            "reference": material.get("reference", ""),
            "purchase_date": material.get("purchase_date"),
            "notes": material.get("notes", "")
        }
        clean_materials.append(clean_material)

    return {
        "job": Job(**job),
        "total_hours": total_hours,
        "labor_cost": labor_cost,
        "materials_cost": total_materials_cost,
        "total_cost": total_cost,
        "quoted_cost": quoted_cost,
        "cost_variance": cost_variance,
        "time_entries": clean_time_entries,
        "materials": clean_materials,
        "time_entries_count": len(time_entries),
        "materials_count": len(materials)
    }

# ==================== CSV EXPORT SYSTEM ====================

@api_router.get("/reports/export/job/{job_id}")
async def export_job_report(job_id: str, admin: str = Depends(verify_admin)):
    """Export comprehensive job report as CSV (Admin only)"""
    # Get job details
    job = await db.jobs.find_one({"id": job_id})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Get time entries and materials
    time_entries = await db.time_entries.find({"job_id": job_id}).to_list(1000)
    materials = await db.materials.find({"job_id": job_id}).to_list(1000)

    # Get worker names
    worker_ids = list(set(entry.get("worker_id") for entry in time_entries))
    workers = await db.workers.find({"id": {"$in": worker_ids}}).to_list(1000)
    worker_names = {worker["id"]: worker["name"] for worker in workers}

    # Calculate totals
    total_minutes = sum((entry.get("duration_minutes", 0) or 0) for entry in time_entries)
    total_hours = round(total_minutes / 60, 1)
    labor_cost = total_hours * 15
    materials_cost = sum(mat.get("cost", 0) * mat.get("quantity", 1) for mat in materials)
    total_cost = labor_cost + materials_cost
    cost_variance = job.get("quoted_cost", 0) - total_cost

    # Create CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Job summary
    writer.writerow(["JOB REPORT - " + job.get("name", "")])
    writer.writerow(["Client", job.get("client", "")])
    writer.writerow(["Location", job.get("location", "")])
    writer.writerow(["Quoted Cost", f"£{job.get('quoted_cost', 0):,.2f}"])
    writer.writerow(["Actual Cost", f"£{total_cost:,.2f}"])
    writer.writerow(["Variance", f"£{cost_variance:,.2f}"])
    writer.writerow([])

    # Time entries section
    writer.writerow(["TIME ENTRIES"])
    writer.writerow(["Worker", "Clock In", "Clock Out", "Duration (hours)", "Labor Cost", "Notes"])

    for entry in time_entries:
        worker_name = worker_names.get(entry.get("worker_id"), "Unknown")
        clock_in = entry.get("clock_in", "")
        clock_out = entry.get("clock_out", "")
        duration_hours = round(entry.get("duration_minutes", 0) / 60, 2) if entry.get("duration_minutes") else 0
        entry_labor_cost = duration_hours * 15

        writer.writerow([
            worker_name,
            format_uk_datetime_for_export(clock_in),
            format_uk_datetime_for_export(clock_out) if clock_out else "Active",
            duration_hours,
            f"£{entry_labor_cost:.2f}",
            entry.get("notes", "")
        ])

    writer.writerow([])
    writer.writerow(["TOTAL LABOR", "", "", total_hours, f"£{labor_cost:.2f}", ""])
    writer.writerow([])

    # Materials section
    writer.writerow(["MATERIALS"])
    writer.writerow(["Material", "Quantity", "Unit Cost", "Total Cost", "Purchase Date", "Notes"])

    for material in materials:
        total_material_cost = material.get("cost", 0) * material.get("quantity", 1)
        purchase_date = material.get("purchase_date", "")

        writer.writerow([
            material.get("name", ""),
            material.get("quantity", 1),
            f"£{material.get('cost', 0):.2f}",
            f"£{total_material_cost:.2f}",
            purchase_date.strftime("%Y-%m-%d") if purchase_date else "",
            material.get("notes", "")
        ])

    writer.writerow([])
    writer.writerow(["TOTAL MATERIALS", "", "", f"£{materials_cost:.2f}", "", ""])

    output.seek(0)
    filename = f"job_report_{job.get('name', 'unknown').replace(' ', '_')}.csv"

    return StreamingResponse(
        io.BytesIO(("\ufeff" + output.getvalue()).encode("utf-8")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/reports/export/time-entries")
async def export_time_entries_csv(
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Export time entries as CSV (Admin only)"""
    # Build filter
    filter_dict = {}
    if job_id:
        filter_dict["job_id"] = job_id

    if start_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        filter_dict["clock_in"] = {"$gte": start_dt}

    if end_date:
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        if "clock_in" in filter_dict:
            filter_dict["clock_in"]["$lte"] = end_dt
        else:
            filter_dict["clock_in"] = {"$lte": end_dt}

    # Get data
    time_entries = await db.time_entries.find(filter_dict).to_list(1000)

    # Get worker and job names with hourly rates
    worker_names = {}
    worker_rates = {}
    job_names = {}

    for entry in time_entries:
        if entry["worker_id"] not in worker_names:
            worker = await db.workers.find_one({"id": entry["worker_id"]})
            if worker:
                worker_names[entry["worker_id"]] = worker.get("name", "Unknown")
                worker_rates[entry["worker_id"]] = worker.get("hourly_rate", 15.0)
            else:
                worker_names[entry["worker_id"]] = "Unknown"
                worker_rates[entry["worker_id"]] = 15.0

        if entry["job_id"] not in job_names:
            job = await db.jobs.find_one({"id": entry["job_id"]})
            job_names[entry["job_id"]] = job.get("name", "Unknown") if job else "Unknown"

    # Create CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Headers
    writer.writerow([
        "Worker Name", "Job Name", "Clock In", "Clock Out",
        "Duration (hours)", "Hourly Rate", "Labor Cost", "Notes",
        "GPS In Lat", "GPS In Lng", "GPS In Accuracy", "GPS In Address",
        "GPS Out Lat", "GPS Out Lng", "GPS Out Accuracy", "GPS Out Address",
        "Device ID In", "Device ID Out", "Suspicious Flags"
    ])

    # Data rows
    for entry in time_entries:
        clock_in = entry.get("clock_in", "")
        clock_out = entry.get("clock_out", "")
        duration_hours = round(entry.get("duration_minutes", 0) / 60, 2) if entry.get("duration_minutes") else ""
        hourly_rate = worker_rates.get(entry["worker_id"], 15.0)
        labor_cost = duration_hours * hourly_rate if duration_hours else 0

        # GPS locations
        gps_in_lat = gps_in_lng = gps_in_accuracy = gps_in_address = ""
        if entry.get("gps_location_in"):
            gps_in_lat = entry["gps_location_in"].get("latitude", "")
            gps_in_lng = entry["gps_location_in"].get("longitude", "")
            gps_in_accuracy = entry["gps_location_in"].get("accuracy", "")
            gps_in_address = entry["gps_location_in"].get("address", "")

        gps_out_lat = gps_out_lng = gps_out_accuracy = gps_out_address = ""
        if entry.get("gps_location_out"):
            gps_out_lat = entry["gps_location_out"].get("latitude", "")
            gps_out_lng = entry["gps_location_out"].get("longitude", "")
            gps_out_accuracy = entry["gps_location_out"].get("accuracy", "")
            gps_out_address = entry["gps_location_out"].get("address", "")

        writer.writerow([
            worker_names.get(entry["worker_id"], "Unknown"),
            job_names.get(entry["job_id"], "Unknown"),
            format_uk_datetime_for_export(clock_in),
            format_uk_datetime_for_export(clock_out),
            duration_hours,
            f"£{hourly_rate:.2f}",
            f"£{labor_cost:.2f}",
            entry.get("notes", ""),
            gps_in_lat, gps_in_lng, gps_in_accuracy, gps_in_address,
            gps_out_lat, gps_out_lng, gps_out_accuracy, gps_out_address,
            entry.get("device_id_in", ""),
            entry.get("device_id_out", ""),
            "; ".join(entry.get("suspicious_flags", []) or [])
        ])

    output.seek(0)

    csv_bytes = ("\ufeff" + output.getvalue()).encode("utf-8")

    return StreamingResponse(
        io.BytesIO(csv_bytes),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=time_entries.csv"}
    )



async def build_time_entries_export_rows(
    worker_id: Optional[str] = None,
    job_id: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    worker_type: Optional[str] = None,
    division: Optional[str] = None,
    trade: Optional[str] = None,
):
    """Shared row builder for PDF timesheet exports."""
    filter_dict = {"archived": {"$ne": True}}
    if worker_id and worker_id != "all":
        filter_dict["worker_id"] = worker_id
    if job_id and job_id != "all":
        filter_dict["job_id"] = job_id
    if start_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        filter_dict["clock_in"] = {"$gte": start_dt}
    if end_date:
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        if "clock_in" in filter_dict:
            filter_dict["clock_in"]["$lte"] = end_dt
        else:
            filter_dict["clock_in"] = {"$lte": end_dt}

    time_entries = await db.time_entries.find(filter_dict).sort("clock_in", 1).to_list(5000)
    worker_ids = list({entry.get("worker_id") for entry in time_entries if entry.get("worker_id")})
    job_ids = list({entry.get("job_id") for entry in time_entries if entry.get("job_id")})

    workers = await db.workers.find({"id": {"$in": worker_ids}}, {"_id": 0}).to_list(1000) if worker_ids else []
    jobs = await db.jobs.find({"id": {"$in": job_ids}}, {"_id": 0}).to_list(1000) if job_ids else []
    worker_lookup = {worker.get("id"): worker for worker in workers}
    job_lookup = {job.get("id"): job for job in jobs}

    rows = []
    for entry in time_entries:
        worker = worker_lookup.get(entry.get("worker_id"), {})
        job = job_lookup.get(entry.get("job_id"), {})

        if worker_type and worker_type != "all" and worker.get("worker_type", "worker") != worker_type:
            continue
        if division and division != "all" and worker.get("division", "") != division:
            continue
        if trade and trade != "all":
            worker_trades = worker.get("trades") or ([worker.get("trade")] if worker.get("trade") else [])
            if trade not in worker_trades:
                continue

        duration_hours = round((entry.get("duration_minutes", 0) or 0) / 60, 2)
        rows.append({
            "worker_name": worker.get("name", "Unknown Worker"),
            "worker_type": worker.get("worker_type", "worker"),
            "division": worker.get("division", ""),
            "job_name": job.get("name", "Unknown Job"),
            "job_client": job.get("client", ""),
            "clock_in": format_uk_datetime_for_export(entry.get("clock_in")),
            "clock_out": format_uk_datetime_for_export(entry.get("clock_out")) if entry.get("clock_out") else "Active",
            "duration_hours": duration_hours,
            "notes": entry.get("notes", ""),
        })
    return rows


@api_router.get("/reports/export/time-entries-pdf")
@api_router.get("/reports/export/time-entries/pdf")
@api_router.get("/reports/export/time-entries.pdf")
async def export_time_entries_pdf(
    worker_id: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    worker_type: Optional[str] = Query(None),
    worker_division: Optional[str] = Query(None),
    job_division: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    group_by: str = Query("worker"),
    admin: str = Depends(verify_admin)
):
    """Export time entries as a PDF timesheet. Supports the frontend route and legacy route aliases."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    except Exception as exc:
        logger.error("PDF export dependency error: %s", exc)
        raise HTTPException(status_code=500, detail="PDF export is not available on this server. Check reportlab is installed.")

    # The frontend currently sends worker_division. The backend row builder expects division.
    effective_division = worker_division or division
    rows = await build_time_entries_export_rows(worker_id, job_id, start_date, end_date, worker_type, effective_division, trade)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=0.8*cm, leftMargin=0.8*cm, topMargin=0.8*cm, bottomMargin=0.8*cm)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("LDA Group - Time Sheet Export", styles["Title"]),
        Paragraph(f"Date range: {start_date or 'All'} to {end_date or 'All'}", styles["Normal"]),
        Paragraph(f"Generated: {get_uk_time().strftime('%d/%m/%Y %H:%M')} UK time", styles["Normal"]),
        Spacer(1, 0.35*cm),
    ]

    if not rows:
        story.append(Paragraph("No time entries found for the selected filters.", styles["Normal"]))
    else:
        if group_by == "job":
            rows.sort(key=lambda row: (row["job_name"].lower(), row["worker_name"].lower(), row["clock_in"]))
        else:
            rows.sort(key=lambda row: (row["worker_name"].lower(), row["job_name"].lower(), row["clock_in"]))

        table_data = [["Worker", "Division", "Job", "Client", "Clock In", "Clock Out", "Hours", "Notes"]]
        total_hours = 0
        for row in rows:
            total_hours += row["duration_hours"] or 0
            table_data.append([
                Paragraph(row["worker_name"], styles["BodyText"]),
                Paragraph(row["division"] or "-", styles["BodyText"]),
                Paragraph(row["job_name"], styles["BodyText"]),
                Paragraph(row["job_client"] or "-", styles["BodyText"]),
                row["clock_in"],
                row["clock_out"],
                f"{row['duration_hours']:.2f}",
                Paragraph(row["notes"] or "", styles["BodyText"]),
            ])
        table_data.append(["TOTAL", "", "", "", "", "", f"{total_hours:.2f}", ""])

        table = Table(table_data, colWidths=[3.2*cm, 2.6*cm, 4.0*cm, 3.0*cm, 3.0*cm, 3.0*cm, 1.5*cm, 6.0*cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d01f2f")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#eeeeee")),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f7f7f7")]),
        ]))
        story.append(table)

    doc.build(story)
    buffer.seek(0)
    filename = f"time_sheet_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

@api_router.get("/reports/export/materials")
async def export_materials_csv(
    worker_id: Optional[str] = Query(None),
    supplier: Optional[str] = Query(None),
    client: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Export materials report as CSV (Admin only)"""
    filter_query = {"archived": {"$ne": True}}

    if start_date and end_date:
        start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
        end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
        filter_query["purchase_date"] = {"$gte": start_dt, "$lte": end_dt}

    if supplier:
        filter_query["supplier"] = {"$regex": supplier, "$options": "i"}

    materials = await db.materials.find(filter_query).to_list(1000)
    jobs = await db.jobs.find().to_list(1000)
    job_lookup = {job["id"]: job for job in jobs if "id" in job}

    # Process materials for export
    export_data = []
    for material in materials:
        job = job_lookup.get(material["job_id"])
        if not job:
            continue

        # Apply filters
        if client and client.lower() not in job.get("client", "").lower():
            continue
        if job_id and job["id"] != job_id:
            continue

        # Convert to UK time for display
        purchase_date = utc_to_uk(material["purchase_date"]) if material.get("purchase_date") else None

        export_data.append({
            "Date": purchase_date.strftime('%Y-%m-%d %H:%M') if purchase_date else "N/A",
            "Job": job["name"],
            "Client": job.get("client", ""),
            "Material": material["name"],
            "Supplier": material.get("supplier", ""),
            "Receipt No": material.get("reference", ""),
            "Quantity": material["quantity"],
            "Unit Cost": f"£{material['cost']:.2f}",
            "Total Value": f"£{material['cost'] * material['quantity']:.2f}",
            "Notes": material.get("notes", "")
        })

    # Sort by date
    export_data.sort(key=lambda x: x["Date"], reverse=True)

    # Create CSV
    output = io.StringIO()
    writer = csv.writer(output)

    # Headers
    writer.writerow([
        "MATERIALS REPORT",
        f"Generated: {get_uk_time().strftime('%Y-%m-%d %H:%M:%S')} UK Time"
    ])
    writer.writerow([])

    if export_data:
        writer.writerow(list(export_data[0].keys()))
        for row in export_data:
            writer.writerow(list(row.values()))
    else:
        writer.writerow(["No materials found for the selected criteria"])

    # Summary
    writer.writerow([])
    writer.writerow(["SUMMARY"])
    writer.writerow([])
    total_value = sum(material["cost"] * material["quantity"] for material in materials)
    writer.writerow(["Total Materials", len(export_data)])
    writer.writerow(["Total Value", f"£{total_value:.2f}"])

    output.seek(0)
    filename = f"materials_report_{datetime.utcnow().strftime('%Y%m%d')}.csv"

    return StreamingResponse(
        io.BytesIO(output.getvalue().encode()),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/reports/export/attendance-alerts")
async def export_attendance_alerts_csv(admin: str = Depends(verify_admin)):
    """Export attendance alerts for the last 7 days as CSV (Admin only)"""
    # Get attendance alerts (same logic as dashboard)
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    nine_am_threshold = timedelta(hours=9)
    five_pm_threshold = timedelta(hours=17)

    non_admin_workers = await db.workers.find({
        "role": {"$ne": "admin"},
        "active": True,
        "archived": {"$ne": True}
    }).to_list(1000)

    attendance_alerts = []

    for worker in non_admin_workers:
        for i in range(7):
            day_start = (datetime.utcnow() - timedelta(days=i)).replace(hour=0, minute=0, second=0, microsecond=0)
            day_end = day_start + timedelta(days=1)
            nine_am = day_start + nine_am_threshold
            five_pm = day_start + five_pm_threshold

            if day_start.date() == datetime.utcnow().date() and datetime.utcnow().hour < 9:
                continue
            if day_start.date() > datetime.utcnow().date():
                continue

            day_entries = await db.time_entries.find({
                "worker_id": worker["id"],
                "clock_in": {"$gte": day_start, "$lt": day_end}
            }).to_list(100)

            if not day_entries:
                if day_start.weekday() < 5 and day_start.date() != datetime.utcnow().date():
                    attendance_alerts.append({
                        "worker_name": worker["name"],
                        "worker_email": worker.get("email", ""),
                        "type": "No Clock In",
                        "date": day_start.strftime('%Y-%m-%d'),
                        "day_of_week": day_start.strftime('%A'),
                        "time": "N/A",
                        "details": f"No clock in recorded on {day_start.strftime('%A, %d %B %Y')}"
                    })
            else:
                for entry in day_entries:
                    clock_in_time = entry["clock_in"]
                    if clock_in_time > nine_am:
                        uk_time = utc_to_uk(clock_in_time)
                        attendance_alerts.append({
                            "worker_name": worker["name"],
                            "worker_email": worker.get("email", ""),
                            "type": "Late Clock In",
                            "date": day_start.strftime('%Y-%m-%d'),
                            "day_of_week": day_start.strftime('%A'),
                            "time": uk_time.strftime('%H:%M') if uk_time else "N/A",
                            "details": f"Clocked in late at {uk_time.strftime('%H:%M') if uk_time else 'Unknown'} on {day_start.strftime('%A, %d %B %Y')}"
                        })

                    if entry.get("clock_out") and entry["clock_out"] > five_pm:
                        uk_time = utc_to_uk(entry["clock_out"])
                        attendance_alerts.append({
                            "worker_name": worker["name"],
                            "worker_email": worker.get("email", ""),
                            "type": "Late Clock Out",
                            "date": day_start.strftime('%Y-%m-%d'),
                            "day_of_week": day_start.strftime('%A'),
                            "time": uk_time.strftime('%H:%M') if uk_time else "N/A",
                            "details": f"Clocked out late at {uk_time.strftime('%H:%M') if uk_time else 'Unknown'} on {day_start.strftime('%A, %d %B %Y')}"
                        })

    # Sort and create CSV
    attendance_alerts.sort(key=lambda x: (x["date"], x["worker_name"]), reverse=True)

    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow([
        "ATTENDANCE ALERTS - LAST 7 DAYS",
        f"Generated: {get_uk_time().strftime('%Y-%m-%d %H:%M:%S')} UK Time"
    ])
    writer.writerow([])
    writer.writerow(["Worker Name", "Worker Email", "Alert Type", "Date", "Day of Week", "Time", "Details"])

    for alert in attendance_alerts:
        writer.writerow([
            alert["worker_name"], alert["worker_email"], alert["type"],
            alert["date"], alert["day_of_week"], alert["time"], alert["details"]
        ])

    # Summary
    writer.writerow([])
    writer.writerow(["SUMMARY"])
    writer.writerow([])

    type_counts = {}
    for alert in attendance_alerts:
        alert_type = alert["type"]
        type_counts[alert_type] = type_counts.get(alert_type, 0) + 1

    for alert_type, count in type_counts.items():
        writer.writerow([alert_type, count])

    writer.writerow([])
    writer.writerow(["Total Alerts", len(attendance_alerts)])

    output.seek(0)
    filename = f"attendance_alerts_{datetime.utcnow().strftime('%Y%m%d')}.csv"

    return StreamingResponse(
        io.BytesIO(output.getvalue().encode()),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ==================== ENHANCED MATERIAL & TIME ENTRY MANAGEMENT ====================

@api_router.put("/materials/{material_id}/archive")
async def archive_material_endpoint(material_id: str, admin: str = Depends(verify_admin)):
    """Archive material (Admin only)"""
    result = await db.materials.update_one(
        {"id": material_id},
        {"$set": {"archived": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Material not found")
    return {"message": "Material archived successfully"}

@api_router.put("/materials/{material_id}/unarchive")
async def unarchive_material_endpoint(material_id: str, admin: str = Depends(verify_admin)):
    """Unarchive material (Admin only)"""
    result = await db.materials.update_one(
        {"id": material_id},
        {"$set": {"archived": False}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Material not found")
    return {"message": "Material unarchived successfully"}

@api_router.delete("/time-entries/{entry_id}")
async def delete_time_entry_endpoint(entry_id: str, admin: str = Depends(verify_admin)):
    """Delete time entry (Admin only)"""
    result = await db.time_entries.delete_one({"id": entry_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Time entry not found")
    return {"message": "Time entry deleted successfully"}

@api_router.put("/time-entries/{entry_id}/archive")
async def archive_time_entry_endpoint(entry_id: str, admin: str = Depends(verify_admin)):
    """Archive time entry (Admin only)"""
    result = await db.time_entries.update_one(
        {"id": entry_id},
        {"$set": {"archived": True}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Time entry not found")
    return {"message": "Time entry archived successfully"}



# ==================== SUBCONTRACTOR SYSTEM ENDPOINTS ====================


def normalise_subcontractor_trades(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        return [item.strip() for item in value.split(",") if item.strip()]
    return []


def normalise_schedule_resource_values(resource_type: Optional[str], resource_id: Optional[str], worker_id: Optional[str]) -> Dict[str, str]:
    final_type = str(resource_type or "worker").strip() or "worker"
    if final_type not in ["worker", "subcontractor_resource"]:
        final_type = "worker"
    final_id = str(resource_id or worker_id or "").strip()
    final_worker_id = final_id if final_type == "worker" else ""
    return {"resource_type": final_type, "resource_id": final_id, "worker_id": final_worker_id}


async def get_subcontractor_resource_display(resource_id: Optional[str]) -> Dict[str, Any]:
    if not resource_id:
        return {}
    resource = await db.subcontractor_resources.find_one({"id": resource_id, "archived": {"$ne": True}}, {"_id": 0})
    if not resource:
        return {}
    company = await db.subcontractors.find_one({"id": resource.get("subcontractor_id"), "archived": {"$ne": True}}, {"_id": 0}) or {}
    company_name = company.get("company_name", "")
    resource_name = resource.get("name", "")
    display_name = f"{company_name} - {resource_name}".strip(" -")
    return {
        "resource": resource,
        "company": company,
        "display_name": display_name or resource_name or "Subcontractor Resource",
    }


@api_router.get("/subcontractors", response_model=List[Dict[str, Any]])
@api_router.get("/subcontractors/", response_model=List[Dict[str, Any]])
async def list_subcontractors(active_only: bool = Query(False), include_archived: bool = Query(False), admin: str = Depends(verify_admin)):
    query: Dict[str, Any] = {}
    if active_only:
        query["active"] = True
    if not include_archived:
        query["archived"] = {"$ne": True}

    companies = await db.subcontractors.find(query, {"_id": 0}).sort("company_name", 1).to_list(5000)
    company_ids = [company.get("id") for company in companies if company.get("id")]
    resource_query: Dict[str, Any] = {"subcontractor_id": {"$in": company_ids}} if company_ids else {"subcontractor_id": {"$in": []}}
    if active_only:
        resource_query["active"] = True
    if not include_archived:
        resource_query["archived"] = {"$ne": True}

    resources = await db.subcontractor_resources.find(resource_query, {"_id": 0}).sort("name", 1).to_list(10000)
    resources_by_company: Dict[str, List[Dict[str, Any]]] = {}
    for resource in resources:
        resources_by_company.setdefault(resource.get("subcontractor_id", ""), []).append(resource)

    for company in companies:
        company["resources"] = resources_by_company.get(company.get("id"), [])

    return companies


@api_router.post("/subcontractors", response_model=SubcontractorCompany)
@api_router.post("/subcontractors/", response_model=SubcontractorCompany)
async def create_subcontractor(subcontractor: SubcontractorCompanyCreate, admin: str = Depends(verify_admin)):
    data = subcontractor.dict()
    data["company_name"] = data.get("company_name", "").strip()
    if not data["company_name"]:
        raise HTTPException(status_code=400, detail="Company name is required")

    duplicate = await db.subcontractors.find_one({
        "company_name": {"$regex": f"^{re.escape(data['company_name'])}$", "$options": "i"},
        "archived": {"$ne": True},
    })
    if duplicate:
        raise HTTPException(status_code=400, detail="A subcontractor with this company name already exists")

    data["trades"] = normalise_subcontractor_trades(data.get("trades"))
    obj = SubcontractorCompany(**data)
    await db.subcontractors.insert_one(obj.dict())
    return obj


@api_router.get("/subcontractors/schedule/resources/all")
async def list_subcontractor_schedule_resources(active_only: bool = Query(True), admin: str = Depends(verify_admin)):
    company_query: Dict[str, Any] = {"archived": {"$ne": True}}
    resource_query: Dict[str, Any] = {"archived": {"$ne": True}}
    if active_only:
        company_query["active"] = True
        resource_query["active"] = True

    companies = await db.subcontractors.find(company_query, {"_id": 0}).to_list(5000)
    company_lookup = {company.get("id"): company for company in companies if company.get("id")}
    resources = await db.subcontractor_resources.find(resource_query, {"_id": 0}).sort("name", 1).to_list(10000)

    output = []
    for resource in resources:
        company = company_lookup.get(resource.get("subcontractor_id"))
        if not company:
            continue
        output.append({
            "resource_type": "subcontractor_resource",
            "resource_id": resource.get("id", ""),
            "display_name": f"{company.get('company_name', '')} - {resource.get('name', '')}".strip(" -"),
            "company_id": company.get("id", ""),
            "company_name": company.get("company_name", ""),
            "resource_name": resource.get("name", ""),
            "trade": resource.get("trade", ""),
            "capacity": resource.get("capacity", 1),
            "active": resource.get("active", True),
        })

    output.sort(key=lambda item: item.get("display_name", "").lower())
    return output


@api_router.get("/subcontractors/{subcontractor_id}", response_model=Dict[str, Any])
async def get_subcontractor(subcontractor_id: str, admin: str = Depends(verify_admin)):
    company = await db.subcontractors.find_one({"id": subcontractor_id, "archived": {"$ne": True}}, {"_id": 0})
    if not company:
        raise HTTPException(status_code=404, detail="Subcontractor not found")
    resources = await db.subcontractor_resources.find({"subcontractor_id": subcontractor_id, "archived": {"$ne": True}}, {"_id": 0}).sort("name", 1).to_list(1000)
    company["resources"] = resources
    return company


@api_router.put("/subcontractors/{subcontractor_id}", response_model=SubcontractorCompany)
async def update_subcontractor(subcontractor_id: str, update: SubcontractorCompanyUpdate, admin: str = Depends(verify_admin)):
    existing = await db.subcontractors.find_one({"id": subcontractor_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Subcontractor not found")

    update_data = {k: v for k, v in update.dict().items() if v is not None}
    if "company_name" in update_data:
        update_data["company_name"] = str(update_data["company_name"]).strip()
        if not update_data["company_name"]:
            raise HTTPException(status_code=400, detail="Company name is required")
        duplicate = await db.subcontractors.find_one({
            "id": {"$ne": subcontractor_id},
            "company_name": {"$regex": f"^{re.escape(update_data['company_name'])}$", "$options": "i"},
            "archived": {"$ne": True},
        })
        if duplicate:
            raise HTTPException(status_code=400, detail="Another subcontractor already has this company name")
    if "trades" in update_data:
        update_data["trades"] = normalise_subcontractor_trades(update_data.get("trades"))
    update_data["updated_at"] = datetime.utcnow()

    await db.subcontractors.update_one({"id": subcontractor_id}, {"$set": update_data})
    updated = await db.subcontractors.find_one({"id": subcontractor_id}, {"_id": 0})
    return SubcontractorCompany(**updated)


@api_router.delete("/subcontractors/{subcontractor_id}")
async def delete_subcontractor(subcontractor_id: str, admin: str = Depends(verify_admin)):
    existing = await db.subcontractors.find_one({"id": subcontractor_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Subcontractor not found")
    now = datetime.utcnow()
    await db.subcontractors.update_one({"id": subcontractor_id}, {"$set": {"active": False, "archived": True, "updated_at": now}})
    await db.subcontractor_resources.update_many({"subcontractor_id": subcontractor_id}, {"$set": {"active": False, "archived": True, "updated_at": now}})
    return {"success": True, "message": "Subcontractor archived"}


@api_router.get("/subcontractors/{subcontractor_id}/resources", response_model=List[SubcontractorResource])
async def list_subcontractor_resources(subcontractor_id: str, active_only: bool = Query(False), include_archived: bool = Query(False), admin: str = Depends(verify_admin)):
    query: Dict[str, Any] = {"subcontractor_id": subcontractor_id}
    if active_only:
        query["active"] = True
    if not include_archived:
        query["archived"] = {"$ne": True}
    resources = await db.subcontractor_resources.find(query, {"_id": 0}).sort("name", 1).to_list(1000)
    return [SubcontractorResource(**resource) for resource in resources]


@api_router.post("/subcontractors/{subcontractor_id}/resources", response_model=SubcontractorResource)
async def create_subcontractor_resource(subcontractor_id: str, resource: SubcontractorResourceCreate, admin: str = Depends(verify_admin)):
    company = await db.subcontractors.find_one({"id": subcontractor_id, "archived": {"$ne": True}})
    if not company:
        raise HTTPException(status_code=404, detail="Subcontractor not found")
    data = resource.dict()
    data["subcontractor_id"] = subcontractor_id
    data["name"] = data.get("name", "").strip()
    if not data["name"]:
        raise HTTPException(status_code=400, detail="Worker/team name is required")
    data["capacity"] = max(1, int(data.get("capacity") or 1))
    obj = SubcontractorResource(**data)
    await db.subcontractor_resources.insert_one(obj.dict())
    return obj


@api_router.put("/subcontractors/resources/{resource_id}", response_model=SubcontractorResource)
async def update_subcontractor_resource(resource_id: str, update: SubcontractorResourceUpdate, admin: str = Depends(verify_admin)):
    existing = await db.subcontractor_resources.find_one({"id": resource_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Subcontractor worker/team not found")
    update_data = {k: v for k, v in update.dict().items() if v is not None}
    if "name" in update_data:
        update_data["name"] = str(update_data["name"]).strip()
        if not update_data["name"]:
            raise HTTPException(status_code=400, detail="Worker/team name is required")
    if "capacity" in update_data:
        update_data["capacity"] = max(1, int(update_data.get("capacity") or 1))
    if "subcontractor_id" in update_data:
        company = await db.subcontractors.find_one({"id": update_data["subcontractor_id"], "archived": {"$ne": True}})
        if not company:
            raise HTTPException(status_code=404, detail="Subcontractor not found")
    update_data["updated_at"] = datetime.utcnow()
    await db.subcontractor_resources.update_one({"id": resource_id}, {"$set": update_data})
    updated = await db.subcontractor_resources.find_one({"id": resource_id}, {"_id": 0})
    return SubcontractorResource(**updated)


@api_router.delete("/subcontractors/resources/{resource_id}")
async def delete_subcontractor_resource(resource_id: str, admin: str = Depends(verify_admin)):
    existing = await db.subcontractor_resources.find_one({"id": resource_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Subcontractor worker/team not found")
    await db.subcontractor_resources.update_one({"id": resource_id}, {"$set": {"active": False, "archived": True, "updated_at": datetime.utcnow()}})
    return {"success": True, "message": "Subcontractor worker/team archived"}

VALID_SCHEDULE_TYPES = {"job", "holiday", "sick", "unavailable"}
ABSENCE_SCHEDULE_TYPES = {"holiday", "sick", "unavailable"}


def normalise_schedule_type(value: Optional[str]) -> str:
    """Return a safe schedule type for older and newer schedule records."""
    schedule_type = str(value or "job").strip().lower()
    return schedule_type if schedule_type in VALID_SCHEDULE_TYPES else "job"


def schedule_type_label(schedule_type: Optional[str]) -> str:
    schedule_type = normalise_schedule_type(schedule_type)
    if schedule_type == "holiday":
        return "Holiday"
    if schedule_type == "sick":
        return "Sick Day"
    if schedule_type == "unavailable":
        return "Unavailable"
    return "Scheduled job"


def is_absence_schedule_type(schedule_type: Optional[str]) -> bool:
    return normalise_schedule_type(schedule_type) in ABSENCE_SCHEDULE_TYPES

# ==================== SCHEDULING ENDPOINTS ====================

@api_router.get("/schedule", response_model=List[Dict[str, Any]])
async def get_schedule_entries(
    start_date: str = Query(...),
    end_date: str = Query(...),
    worker_type: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Get schedule entries between two dates inclusive (Admin only), with optional worker filters."""
    try:
        datetime.fromisoformat(start_date)
        datetime.fromisoformat(end_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="Dates must be in YYYY-MM-DD format")

    filter_query = {
        "scheduled_date": {"$gte": start_date, "$lte": end_date},
        "archived": {"$ne": True}
    }

    if job_id and job_id != "all":
        filter_query["job_id"] = job_id

    worker_filter = {"active": True, "archived": {"$ne": True}}
    if worker_type and worker_type != "all":
        worker_filter["worker_type"] = worker_type
    if division and division != "all":
        worker_filter["division"] = division
    if trade and trade != "all":
        worker_filter["$or"] = [{"trades": trade}, {"trade": trade}]

    worker_filters_applied = any([
        worker_type and worker_type != "all",
        division and division != "all",
        trade and trade != "all"
    ])

    filtered_workers = []
    if worker_filters_applied:
        filtered_workers = await db.workers.find(worker_filter).to_list(1000)
        filtered_worker_ids = [worker["id"] for worker in filtered_workers if "id" in worker]
        filter_query["worker_id"] = {"$in": filtered_worker_ids}

    entries = await db.schedule_entries.find(filter_query).to_list(1000)

    worker_ids = list({entry.get("worker_id") or entry.get("resource_id") for entry in entries if (entry.get("worker_id") or entry.get("resource_id")) and (entry.get("resource_type") or "worker") == "worker"})
    subcontractor_resource_ids = list({entry.get("resource_id") for entry in entries if entry.get("resource_type") == "subcontractor_resource" and entry.get("resource_id")})
    job_ids = list({entry.get("job_id") for entry in entries if entry.get("job_id")})

    workers = await db.workers.find({"id": {"$in": worker_ids}}).to_list(1000) if worker_ids else []
    subcontractor_resources = await db.subcontractor_resources.find({"id": {"$in": subcontractor_resource_ids}}, {"_id": 0}).to_list(1000) if subcontractor_resource_ids else []
    subcontractor_company_ids = list({resource.get("subcontractor_id") for resource in subcontractor_resources if resource.get("subcontractor_id")})
    subcontractor_companies = await db.subcontractors.find({"id": {"$in": subcontractor_company_ids}}, {"_id": 0}).to_list(1000) if subcontractor_company_ids else []
    jobs = await db.jobs.find({"id": {"$in": job_ids}}).to_list(1000) if job_ids else []

    worker_lookup = {worker["id"]: worker for worker in workers if "id" in worker}
    subcontractor_resource_lookup = {resource.get("id"): resource for resource in subcontractor_resources if resource.get("id")}
    subcontractor_company_lookup = {company.get("id"): company for company in subcontractor_companies if company.get("id")}
    job_lookup = {job["id"]: job for job in jobs if "id" in job}

    result = []
    for entry in entries:
        entry.pop("_id", None)
        normalised_resource = normalise_schedule_resource_values(entry.get("resource_type"), entry.get("resource_id"), entry.get("worker_id"))
        entry.update(normalised_resource)
        job = job_lookup.get(entry.get("job_id"), {})
        schedule_type = normalise_schedule_type(entry.get("schedule_type"))
        entry["schedule_type"] = schedule_type
        entry["absence_type"] = entry.get("absence_type") or (schedule_type if is_absence_schedule_type(schedule_type) else None)

        if entry.get("resource_type") == "subcontractor_resource":
            resource = subcontractor_resource_lookup.get(entry.get("resource_id"), {})
            company = subcontractor_company_lookup.get(resource.get("subcontractor_id"), {})
            display_name = entry.get("display_worker_name") or f"{company.get('company_name', '')} - {resource.get('name', '')}".strip(" -") or "Unknown subcontractor"
            entry["worker_name"] = display_name
            entry["display_worker_name"] = display_name
            entry["worker_type"] = "subcontractor"
            entry["worker_division"] = "Subcontractors"
            entry["worker_trades"] = [resource.get("trade", "")] if resource.get("trade") else []
            entry["subcontractor_company_id"] = company.get("id", "")
            entry["subcontractor_company_name"] = company.get("company_name", "")
            entry["subcontractor_resource_name"] = resource.get("name", "")
        else:
            worker = worker_lookup.get(entry.get("worker_id") or entry.get("resource_id"), {})
            entry["worker_name"] = entry.get("display_worker_name") or worker.get("name", "Unknown")
            entry["display_worker_name"] = entry["worker_name"]
            entry["worker_type"] = worker.get("worker_type", "worker")
            entry["worker_division"] = worker.get("division", "")
            entry["worker_trades"] = worker.get("trades", [worker.get("trade", "")] if worker.get("trade") else [])

        if is_absence_schedule_type(schedule_type):
            entry["job_name"] = schedule_type_label(schedule_type)
            entry["job_client"] = ""
            entry["job_location"] = ""
        else:
            entry["job_name"] = job.get("name", "Unknown")
            entry["job_client"] = job.get("client", "")
            entry["job_location"] = job.get("location", "")
        result.append(entry)

    result.sort(key=lambda x: (x.get("scheduled_date", ""), x.get("worker_name", "")))
    return result

@api_router.get("/schedule/by-job", response_model=List[Dict[str, Any]])
async def get_schedule_by_job(
    start_date: str = Query(...),
    end_date: str = Query(...),
    worker_type: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    admin: str = Depends(verify_admin)
):
    """Return schedule entries grouped by job for the admin job-view planner."""
    entries = await get_schedule_entries(
        start_date=start_date,
        end_date=end_date,
        worker_type=worker_type,
        division=division,
        trade=trade,
        job_id=None,
        admin=admin
    )

    grouped = {}
    for entry in entries:
        job_id_value = entry.get("job_id")
        if job_id_value not in grouped:
            grouped[job_id_value] = {
                "job_id": job_id_value,
                "job_name": entry.get("job_name", "Unknown"),
                "job_client": entry.get("job_client", ""),
                "job_location": entry.get("job_location", ""),
                "entries": []
            }
        grouped[job_id_value]["entries"].append(entry)

    result = list(grouped.values())
    result.sort(key=lambda item: item.get("job_name", ""))
    return result

async def _push_gantt_section_to_schedule(request: GanttPushToScheduleRequest):
    """Create schedule entries from a Gantt section. Skips weekends and protects existing allocations unless override is requested."""
    if not request.worker_ids:
        raise HTTPException(status_code=400, detail="Select at least one worker")

    try:
        start = datetime.fromisoformat(request.start_date).date()
        end = datetime.fromisoformat(request.end_date).date()
    except ValueError:
        raise HTTPException(status_code=400, detail="start_date and end_date must be YYYY-MM-DD")

    if end < start:
        raise HTTPException(status_code=400, detail="End date cannot be before start date")

    job = await db.jobs.find_one({"id": request.job_id, "archived": {"$ne": True}})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    workers = await db.workers.find({"id": {"$in": request.worker_ids}, "active": True, "archived": {"$ne": True}}).to_list(1000)
    found_worker_ids = {worker.get("id") for worker in workers}
    missing_worker_ids = [worker_id for worker_id in request.worker_ids if worker_id not in found_worker_ids]
    if missing_worker_ids:
        raise HTTPException(status_code=404, detail=f"Worker(s) not found or inactive: {', '.join(missing_worker_ids)}")

    section_note = request.section_name or "Section"

    dates = []
    current = start
    while current <= end:
        if current.weekday() < 5:
            dates.append(current.isoformat())
        current = current + timedelta(days=1)

    expected_count = len(request.worker_ids) * len(dates)

    if request.replace_existing_for_section and dates:
        # Archive only schedule entries that are linked to this Gantt section.
        # Older records may only have the section name in notes, so keep that as a legacy fallback.
        await db.schedule_entries.update_many(
            {
                "job_id": request.job_id,
                "worker_id": {"$in": request.worker_ids},
                "scheduled_date": {"$in": dates},
                "archived": {"$ne": True},
                "$or": [
                    {"project_section_id": request.section_id},
                    {"linked_gantt_section_id": request.section_id},
                    {"notes": section_note},
                ],
            },
            {"$set": {"archived": True, "updated_date": datetime.utcnow(), "archived_reason": "replaced_by_gantt_section_push"}},
        )

    created = []
    clashes = []
    overridden = []
    worker_lookup = {worker.get("id"): worker for worker in workers}
    job_lookup: Dict[str, Dict[str, Any]] = {request.job_id: job}

    async def get_job_name(job_id: Optional[str]) -> str:
        if not job_id:
            return "Unknown job"
        if job_id not in job_lookup:
            found_job = await db.jobs.find_one({"id": job_id}) or {}
            job_lookup[job_id] = found_job
        return job_lookup.get(job_id, {}).get("name", "Unknown job")

    for worker_id in request.worker_ids:
        worker = worker_lookup.get(worker_id, {})
        for scheduled_date in dates:
            existing = await db.schedule_entries.find_one({
                "resource_type": "worker",
                "resource_id": worker_id,
                "scheduled_date": scheduled_date,
                "archived": {"$ne": True},
            })
            if not existing:
                existing = await db.schedule_entries.find_one({
                    "worker_id": worker_id,
                    "scheduled_date": scheduled_date,
                    "archived": {"$ne": True},
                    "$or": [{"resource_type": {"$exists": False}}, {"resource_type": "worker"}],
                })
            if existing:
                existing_type = normalise_schedule_type(existing.get("schedule_type"))
                existing_name = schedule_type_label(existing_type) if is_absence_schedule_type(existing_type) else await get_job_name(existing.get("job_id"))
                clash_record = {
                    "worker_id": worker_id,
                    "worker_name": worker.get("name", "Unknown"),
                    "scheduled_date": scheduled_date,
                    "existing_entry_id": existing.get("id"),
                    "existing_job_id": existing.get("job_id"),
                    "existing_job_name": existing_name,
                    "existing_schedule_type": existing_type,
                    "existing_notes": existing.get("notes", ""),
                }

                # Holiday/sick/unavailable records are intentional absence blocks. Do not override them
                # from the Gantt, even when normal job clashes are being overridden.
                if request.override_existing_allocations and not is_absence_schedule_type(existing_type):
                    await db.schedule_entries.update_one(
                        {"id": existing.get("id")},
                        {"$set": {"archived": True, "updated_date": datetime.utcnow(), "overridden_by_gantt_section_id": request.section_id}},
                    )
                    overridden.append(clash_record)
                else:
                    clashes.append(clash_record)
                    continue

            entry_obj = ScheduleEntry(
                worker_id=worker_id,
                resource_type="worker",
                resource_id=worker_id,
                display_worker_name=worker.get("name", ""),
                job_id=request.job_id,
                scheduled_date=scheduled_date,
                notes=section_note,
                status="scheduled",
                schedule_type="job",
                project_section_id=request.section_id,
                project_section_name=section_note,
                linked_gantt_section_id=request.section_id,
                schedule_link_mode="linked_to_section",
                hours=8.0,
            )
            await db.schedule_entries.insert_one(entry_obj.dict())
            created.append(entry_obj.dict())

    created_entry_ids = [entry.get("id") for entry in created if entry.get("id")]
    fully_scheduled = expected_count > 0 and len(created) >= expected_count and len(clashes) == 0
    schedule_status = "scheduled" if fully_scheduled else "partial_scheduled" if created else "not_scheduled"
    updated_section = None

    gantt_sections = job.get("gantt_sections") or []
    updated_sections = []
    for section in gantt_sections:
        if section.get("id") == request.section_id:
            existing_ids = section.get("schedule_entry_ids") or []
            merged_ids = list(dict.fromkeys([*existing_ids, *created_entry_ids]))
            section = {
                **section,
                "assigned_worker_ids": request.worker_ids,
                "sent_to_schedule": fully_scheduled,
                "schedule_status": schedule_status,
                "schedule_entry_ids": merged_ids,
                "scheduled_at": datetime.utcnow().isoformat() if created else section.get("scheduled_at", ""),
                "schedule_clashes": clashes,
                "last_schedule_push_at": datetime.utcnow().isoformat(),
            }
            updated_section = section
        updated_sections.append(section)

    await db.jobs.update_one(
        {"id": request.job_id},
        {"$set": {"gantt_sections": updated_sections}}
    )

    return {
        "message": "Gantt section pushed to schedule",
        "created_count": len(created),
        "expected_count": expected_count,
        "fully_scheduled": fully_scheduled,
        "skipped_weekend_count": max(0, ((end - start).days + 1) - len(dates)),
        "clash_count": len(clashes),
        "overridden_count": len(overridden),
        "created": created,
        "created_entry_ids": created_entry_ids,
        "clashes": clashes,
        "overridden": overridden,
        "updated_section": updated_section,
    }

@api_router.post("/gantt/push-to-schedule")
async def push_gantt_to_schedule(request: GanttPushToScheduleRequest, admin: str = Depends(verify_admin)):
    return await _push_gantt_section_to_schedule(request)

@api_router.post("/schedule/from-gantt")
async def schedule_from_gantt(request: GanttPushToScheduleRequest, admin: str = Depends(verify_admin)):
    return await _push_gantt_section_to_schedule(request)



def shift_iso_date(value: Optional[str], delta_days: int) -> Optional[str]:
    """Shift a YYYY-MM-DD date string by delta_days and return YYYY-MM-DD."""
    if not value:
        return value
    try:
        shifted = datetime.fromisoformat(str(value)).date() + timedelta(days=delta_days)
        return shifted.isoformat()
    except Exception:
        return value



def _date_is_weekend(value: Optional[str]) -> bool:
    try:
        return datetime.fromisoformat(str(value)).date().weekday() >= 5
    except Exception:
        return False


def _normalise_section_ids(section: Dict[str, Any]) -> List[str]:
    ids = []
    for entry_id in section.get("schedule_entry_ids") or []:
        if entry_id:
            ids.append(str(entry_id))
    for entry in section.get("schedule_entries") or []:
        if isinstance(entry, dict) and entry.get("id"):
            ids.append(str(entry.get("id")))
        elif entry:
            ids.append(str(entry))
    return list(dict.fromkeys(ids))


async def _find_linked_section_schedule_entries(job: Dict[str, Any], section: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Find schedule entries linked to a Gantt section, including legacy records created before explicit section IDs existed."""
    job_id = job.get("id")
    section_id = section.get("id") or ""
    section_name = section.get("name") or ""
    linked_ids = _normalise_section_ids(section)

    or_terms = []
    if linked_ids:
        or_terms.append({"id": {"$in": linked_ids}})
    if section_id:
        or_terms.append({"project_section_id": section_id})
        or_terms.append({"linked_gantt_section_id": section_id})
    if section_name:
        # Legacy fallback from earlier Gantt pushes where notes carried the section name.
        or_terms.append({"notes": section_name})

    if not or_terms:
        return []

    query = {"job_id": job_id, "archived": {"$ne": True}, "$or": or_terms}
    return await db.schedule_entries.find(query, {"_id": 0}).to_list(5000)




def _normalise_link_value(value: Any) -> str:
    return str(value or "").strip().lower()


def _po_links_to_section(po: Dict[str, Any], section: Dict[str, Any]) -> bool:
    """Return True when a purchase order appears linked to a Gantt section."""
    section_id = _normalise_link_value(section.get("id"))
    section_name = _normalise_link_value(section.get("name"))

    po_section_ids = [
        po.get("project_section_id"),
        po.get("job_section_id"),
        po.get("linked_gantt_section_id"),
    ]
    po_section_names = [
        po.get("project_section_name"),
        po.get("job_section_name"),
        po.get("section_name"),
    ]

    if section_id and any(_normalise_link_value(value) == section_id for value in po_section_ids):
        return True
    if section_name and any(_normalise_link_value(value) == section_name for value in po_section_names):
        return True

    for line in po.get("lines") or []:
        if not isinstance(line, dict):
            continue
        line_section_ids = [
            line.get("project_section_id"),
            line.get("job_section_id"),
            line.get("linked_gantt_section_id"),
        ]
        line_section_names = [
            line.get("project_section_name"),
            line.get("job_section_name"),
            line.get("section_name"),
        ]
        if section_id and any(_normalise_link_value(value) == section_id for value in line_section_ids):
            return True
        if section_name and any(_normalise_link_value(value) == section_name for value in line_section_names):
            return True

    return False


def _po_status_allows_required_date_shift(po: Dict[str, Any]) -> bool:
    status = str(po.get("status") or "").strip().lower().replace(" ", "_")
    blocked = {"cancelled", "canceled", "void", "archived", "rejected", "closed", "paid"}
    return status not in blocked


async def _shift_linked_purchase_orders_for_section(job: Dict[str, Any], section: Dict[str, Any], delta_days: int) -> Dict[str, Any]:
    """Shift required/order/payment planning dates for POs linked to a moved section."""
    if not job or not section or not delta_days:
        return {"shifted_count": 0, "purchase_orders": []}

    candidates = await db.purchase_orders.find({
        "job_id": job.get("id"),
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    shifted = []
    for po in candidates:
        if not _po_status_allows_required_date_shift(po):
            continue
        if not _po_links_to_section(po, section):
            continue

        update = {"updated_at": datetime.utcnow()}
        changed = False
        for date_key in ["required_date", "material_required_date", "delivery_date", "expected_delivery_date", "order_by_date", "expected_payment_date"]:
            if po.get(date_key):
                update[date_key] = shift_iso_date(po.get(date_key), delta_days)
                changed = True

        # Do not shift explicit payment due dates. Cashflow calculates credit terms from required_date.
        # If an explicit payment_due_date is stored, leave it as a deliberate manual override.
        update["last_programme_shift_at"] = datetime.utcnow().isoformat()
        update["last_programme_shift_days"] = delta_days
        update["last_programme_shift_section_id"] = section.get("id", "")
        update["last_programme_shift_section_name"] = section.get("name", "")

        if changed:
            await db.purchase_orders.update_one({"id": po.get("id")}, {"$set": update})
            shifted.append({
                "id": po.get("id"),
                "po_number": po.get("po_number", ""),
                "supplier_name": po.get("supplier_name", ""),
                "old_required_date": po.get("required_date", ""),
                "new_required_date": update.get("required_date", po.get("required_date", "")),
            })

    return {"shifted_count": len(shifted), "purchase_orders": shifted}

def _work_order_links_to_section(work_order: Dict[str, Any], section: Dict[str, Any]) -> bool:
    if not work_order or not section:
        return False
    section_id = _normalise_link_value(section.get("id"))
    section_name = _normalise_link_value(section.get("name"))
    wo_section_ids = [
        work_order.get("section_id"),
        work_order.get("project_section_id"),
        work_order.get("job_section_id"),
    ]
    wo_section_names = [
        work_order.get("section_name"),
        work_order.get("project_section_name"),
        work_order.get("job_section_name"),
    ]
    if section_id and any(_normalise_link_value(value) == section_id for value in wo_section_ids):
        return True
    if section_name and any(_normalise_link_value(value) == section_name for value in wo_section_names):
        return True
    return False


def _work_order_status_allows_date_shift(work_order: Dict[str, Any]) -> bool:
    status = str(work_order.get("status") or "").strip().lower().replace(" ", "_")
    return status not in WORK_ORDER_CANCELLED_STATUSES and status not in {"paid", "closed"}


async def _shift_linked_work_orders_for_section(job: Dict[str, Any], section: Dict[str, Any], delta_days: int) -> Dict[str, Any]:
    """Shift planned contractor WO dates for WOs linked to a moved section."""
    if not job or not section or not delta_days:
        return {"shifted_count": 0, "work_orders": []}

    candidates = await db.work_orders.find({
        "job_id": job.get("id"),
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    shifted = []
    for work_order in candidates:
        if not _work_order_status_allows_date_shift(work_order):
            continue
        if not _work_order_links_to_section(work_order, section):
            continue

        update = {"updated_at": datetime.utcnow()}
        changed = False
        for date_key in ["expected_start_date", "expected_completion_date", "payment_due_date"]:
            if work_order.get(date_key):
                update[date_key] = shift_iso_date(work_order.get(date_key), delta_days)
                changed = True

        update["last_programme_shift_at"] = datetime.utcnow().isoformat()
        update["last_programme_shift_days"] = delta_days
        update["last_programme_shift_section_id"] = section.get("id", "")
        update["last_programme_shift_section_name"] = section.get("name", "")

        if changed:
            await db.work_orders.update_one({"id": work_order.get("id")}, {"$set": update})
            shifted.append({
                "id": work_order.get("id"),
                "wo_number": work_order.get("wo_number", ""),
                "contractor_name": work_order.get("contractor_name", ""),
                "old_start_date": work_order.get("expected_start_date", ""),
                "new_start_date": update.get("expected_start_date", work_order.get("expected_start_date", "")),
            })

    return {"shifted_count": len(shifted), "work_orders": shifted}

async def _build_section_shift_preview(job: Dict[str, Any], section: Dict[str, Any], delta_days: int) -> Dict[str, Any]:
    entries = await _find_linked_section_schedule_entries(job, section)
    worker_ids = list({entry.get("worker_id") for entry in entries if entry.get("worker_id")})
    job_ids = list({entry.get("job_id") for entry in entries if entry.get("job_id")})
    workers = await db.workers.find({"id": {"$in": worker_ids}}, {"_id": 0}).to_list(1000) if worker_ids else []
    jobs = await db.jobs.find({"id": {"$in": job_ids}}, {"_id": 0}).to_list(1000) if job_ids else []
    worker_lookup = {worker.get("id"): worker for worker in workers}
    job_lookup = {item.get("id"): item for item in jobs}

    today_iso = get_uk_time().date().isoformat()
    movable = []
    conflicts = []
    skipped_past_or_today = []
    weekend_warnings = []

    for entry in entries:
        old_date = entry.get("scheduled_date")
        new_date = shift_iso_date(old_date, delta_days)
        worker = worker_lookup.get(entry.get("worker_id"), {})
        if not old_date or not new_date or new_date == old_date:
            continue

        base_record = {
            "entry_id": entry.get("id"),
            "worker_id": entry.get("worker_id"),
            "worker_name": worker.get("name", "Unknown worker"),
            "old_date": old_date,
            "new_date": new_date,
            "hours": entry.get("hours", 8),
            "job_id": entry.get("job_id"),
            "job_name": job_lookup.get(entry.get("job_id"), {}).get("name", job.get("name", "Unknown job")),
            "section_id": section.get("id", ""),
            "section_name": section.get("name", ""),
        }

        if old_date <= today_iso:
            skipped_past_or_today.append({**base_record, "reason": "Past/today schedule entries are not moved automatically"})
            continue

        if _date_is_weekend(new_date):
            weekend_warnings.append({**base_record, "reason": "New date falls on a weekend"})
            continue

        duplicate = await db.schedule_entries.find_one({
            "id": {"$ne": entry.get("id")},
            "worker_id": entry.get("worker_id"),
            "scheduled_date": new_date,
            "archived": {"$ne": True},
        }, {"_id": 0})

        if duplicate:
            duplicate_type = normalise_schedule_type(duplicate.get("schedule_type"))
            duplicate_job = await db.jobs.find_one({"id": duplicate.get("job_id")}, {"_id": 0}) if duplicate.get("job_id") else {}
            conflicts.append({
                **base_record,
                "existing_entry_id": duplicate.get("id"),
                "existing_job_id": duplicate.get("job_id"),
                "existing_job_name": schedule_type_label(duplicate_type) if is_absence_schedule_type(duplicate_type) else (duplicate_job or {}).get("name", "Existing scheduled job"),
                "existing_schedule_type": duplicate_type,
                "existing_notes": duplicate.get("notes", ""),
                "reason": "Worker already has an active schedule entry on the new date",
            })
            continue

        movable.append(base_record)

    blocked_count = len(conflicts) + len(skipped_past_or_today) + len(weekend_warnings)
    return {
        "job_id": job.get("id"),
        "job_name": job.get("name", ""),
        "section_id": section.get("id", ""),
        "section_name": section.get("name", ""),
        "delta_days": delta_days,
        "linked_schedule_count": len(entries),
        "movable_count": len(movable),
        "blocked_count": blocked_count,
        "can_move_all": len(entries) > 0 and blocked_count == 0,
        "movable": movable,
        "conflicts": conflicts,
        "skipped_past_or_today": skipped_past_or_today,
        "weekend_warnings": weekend_warnings,
    }


@api_router.post("/gantt/shift-section-preview")
async def preview_gantt_section_shift(request: GanttShiftSectionRequest, admin: str = Depends(verify_admin)):
    """Preview linked schedule movement before a Gantt section is moved."""
    job = await db.jobs.find_one({"id": request.job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    section = next((dict(item) for item in job.get("gantt_sections") or [] if item.get("id") == request.section_id), None)
    if not section:
        raise HTTPException(status_code=404, detail="Gantt section not found")
    return await _build_section_shift_preview(job, section, request.delta_days)


@api_router.post("/gantt/shift-section")
async def shift_gantt_section(request: GanttShiftSectionRequest, admin: str = Depends(verify_admin)):
    """Move a single Gantt section and safely move only non-conflicting future linked schedule entries."""
    try:
        datetime.fromisoformat(request.start_date)
        datetime.fromisoformat(request.end_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="start_date and end_date must be YYYY-MM-DD")

    job = await db.jobs.find_one({"id": request.job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    sections = [dict(item) for item in (job.get("gantt_sections") or [])]
    section = next((item for item in sections if item.get("id") == request.section_id), None)
    if not section:
        raise HTTPException(status_code=404, detail="Gantt section not found")

    preview = await _build_section_shift_preview(job, section, request.delta_days) if request.shift_schedule else {
        "linked_schedule_count": 0,
        "movable_count": 0,
        "blocked_count": 0,
        "movable": [],
        "conflicts": [],
        "skipped_past_or_today": [],
        "weekend_warnings": [],
    }

    moved_schedule_count = 0
    moved_ids = []
    if request.shift_schedule:
        for item in preview.get("movable") or []:
            await db.schedule_entries.update_one(
                {"id": item.get("entry_id")},
                {"$set": {
                    "scheduled_date": item.get("new_date"),
                    "updated_date": datetime.utcnow(),
                    "shifted_from_date": item.get("old_date"),
                    "shifted_by_gantt_job_id": request.job_id,
                    "shifted_by_gantt_section_id": request.section_id,
                    "schedule_link_mode": "linked_to_section",
                }},
            )
            moved_schedule_count += 1
            moved_ids.append(item.get("entry_id"))

    po_shift_result = {"shifted_count": 0, "purchase_orders": []}
    if request.shift_purchase_orders:
        po_shift_result = await _shift_linked_purchase_orders_for_section(job, section, request.delta_days)

    wo_shift_result = {"shifted_count": 0, "work_orders": []}
    if request.shift_work_orders:
        wo_shift_result = await _shift_linked_work_orders_for_section(job, section, request.delta_days)

    updated_sections = []
    updated_section = None
    has_blocked = bool(preview.get("blocked_count"))
    for item in sections:
        if item.get("id") == request.section_id:
            item = {
                **item,
                "start_date": request.start_date,
                "end_date": request.end_date,
                "last_programme_shift_at": datetime.utcnow().isoformat(),
                "last_programme_shift_days": request.delta_days,
            }
            if request.shift_schedule and preview.get("linked_schedule_count", 0) > 0:
                item["schedule_status"] = "partial_scheduled" if has_blocked else "scheduled"
                item["sent_to_schedule"] = not has_blocked
                item["schedule_move_preview"] = {
                    "moved_schedule_count": moved_schedule_count,
                    "blocked_count": preview.get("blocked_count", 0),
                    "conflict_count": len(preview.get("conflicts") or []),
                    "skipped_past_or_today_count": len(preview.get("skipped_past_or_today") or []),
                    "weekend_warning_count": len(preview.get("weekend_warnings") or []),
                    "moved_at": datetime.utcnow().isoformat(),
                }
            updated_section = item
        updated_sections.append(item)

    await db.jobs.update_one(
        {"id": request.job_id},
        {"$set": {"gantt_sections": updated_sections, "last_programme_shift_at": datetime.utcnow().isoformat()}},
    )
    updated_job = await db.jobs.find_one({"id": request.job_id}, {"_id": 0})
    return {
        "message": "Gantt section shifted",
        "job": updated_job,
        "updated_section": updated_section,
        "moved_schedule_count": moved_schedule_count,
        "moved_schedule_entry_ids": moved_ids,
        "purchase_orders_shifted": po_shift_result.get("shifted_count", 0),
        "shifted_purchase_orders": po_shift_result.get("purchase_orders", []),
        "work_orders_shifted": wo_shift_result.get("shifted_count", 0),
        "shifted_work_orders": wo_shift_result.get("work_orders", []),
        "preview": preview,
    }


@api_router.post("/gantt/shift-project")
async def shift_gantt_project(request: GanttShiftProjectRequest, admin: str = Depends(verify_admin)):
    """Move a project's planned dates, optionally moving its sections and linked schedule entries by the same day offset."""
    try:
        datetime.fromisoformat(request.planned_start_date)
        datetime.fromisoformat(request.planned_end_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="planned_start_date and planned_end_date must be YYYY-MM-DD")

    if request.delta_days == 0:
        raise HTTPException(status_code=400, detail="delta_days must not be zero")

    job = await db.jobs.find_one({"id": request.job_id, "archived": {"$ne": True}})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    original_sections = job.get("gantt_sections") or []
    original_markers = job.get("commercial_markers") or []
    updated_sections = []
    updated_markers = []
    linked_schedule_ids = []

    for section in original_sections:
        section_copy = dict(section)
        if request.shift_sections:
            section_copy["start_date"] = shift_iso_date(section_copy.get("start_date"), request.delta_days)
            section_copy["end_date"] = shift_iso_date(section_copy.get("end_date"), request.delta_days)
            if section_copy.get("schedule_status") == "scheduled":
                section_copy["last_programme_shift_at"] = datetime.utcnow().isoformat()
        for entry_id in section_copy.get("schedule_entry_ids") or []:
            if entry_id:
                linked_schedule_ids.append(entry_id)
        updated_sections.append(section_copy)

    if request.shift_commercial_markers:
        for marker in original_markers:
            marker_copy = dict(marker)
            # Commercial marker/application dates should move with the programme when the project start date shifts.
            for date_key in ["date", "marker_date", "expected_date", "payment_due_date", "retention_due_date"]:
                if marker_copy.get(date_key):
                    marker_copy[date_key] = shift_iso_date(marker_copy.get(date_key), request.delta_days)
            marker_copy["last_programme_shift_at"] = datetime.utcnow().isoformat()
            updated_markers.append(marker_copy)
    else:
        updated_markers = original_markers

    shifted_schedule_count = 0
    clashes = []

    if request.shift_schedule and linked_schedule_ids:
        entries = await db.schedule_entries.find({
            "id": {"$in": list(dict.fromkeys(linked_schedule_ids))},
            "archived": {"$ne": True},
        }).to_list(5000)

        for entry in entries:
            old_date = entry.get("scheduled_date")
            new_date = shift_iso_date(old_date, request.delta_days)
            if not new_date or new_date == old_date:
                continue

            duplicate = await db.schedule_entries.find_one({
                "id": {"$ne": entry.get("id")},
                "worker_id": entry.get("worker_id"),
                "scheduled_date": new_date,
                "archived": {"$ne": True},
            })

            if duplicate:
                duplicate_type = normalise_schedule_type(duplicate.get("schedule_type"))
                clashes.append({
                    "entry_id": entry.get("id"),
                    "worker_id": entry.get("worker_id"),
                    "old_date": old_date,
                    "new_date": new_date,
                    "existing_entry_id": duplicate.get("id"),
                    "existing_job_id": duplicate.get("job_id"),
                    "existing_schedule_type": duplicate_type,
                    "existing_label": schedule_type_label(duplicate_type) if is_absence_schedule_type(duplicate_type) else "Existing scheduled job",
                })
                continue

            await db.schedule_entries.update_one(
                {"id": entry.get("id")},
                {"$set": {
                    "scheduled_date": new_date,
                    "updated_date": datetime.utcnow(),
                    "shifted_from_date": old_date,
                    "shifted_by_gantt_job_id": request.job_id,
                }},
            )
            shifted_schedule_count += 1

    # If some linked schedule entries could not be shifted, mark their sections as partial so the Gantt does not look fully safe.
    clash_entry_ids = {item.get("entry_id") for item in clashes if item.get("entry_id")}
    if clash_entry_ids:
        adjusted_sections = []
        for section in updated_sections:
            section_ids = set(section.get("schedule_entry_ids") or [])
            if section_ids.intersection(clash_entry_ids):
                section = {
                    **section,
                    "sent_to_schedule": False,
                    "schedule_status": "partial_scheduled",
                    "schedule_clashes": clashes,
                }
            adjusted_sections.append(section)
        updated_sections = adjusted_sections

    purchase_orders_shifted = 0
    shifted_purchase_orders = []
    if request.shift_purchase_orders and request.shift_sections:
        for original_section in original_sections:
            result = await _shift_linked_purchase_orders_for_section(job, dict(original_section), request.delta_days)
            purchase_orders_shifted += result.get("shifted_count", 0)
            shifted_purchase_orders.extend(result.get("purchase_orders", []))

    work_orders_shifted = 0
    shifted_work_orders = []
    if request.shift_work_orders and request.shift_sections:
        for original_section in original_sections:
            result = await _shift_linked_work_orders_for_section(job, dict(original_section), request.delta_days)
            work_orders_shifted += result.get("shifted_count", 0)
            shifted_work_orders.extend(result.get("work_orders", []))

    update_doc = {
        "planned_start_date": request.planned_start_date,
        "planned_end_date": request.planned_end_date,
        "gantt_sections": updated_sections if request.shift_sections else original_sections,
        "commercial_markers": updated_markers if request.shift_commercial_markers else original_markers,
        "last_programme_shift_at": datetime.utcnow().isoformat(),
        "last_programme_shift_days": request.delta_days,
    }

    await db.jobs.update_one({"id": request.job_id}, {"$set": update_doc})
    updated_job = await db.jobs.find_one({"id": request.job_id}, {"_id": 0})

    return {
        "message": "Project programme shifted",
        "job": updated_job,
        "delta_days": request.delta_days,
        "sections_shifted": len(updated_sections) if request.shift_sections else 0,
        "commercial_markers_shifted": len(updated_markers) if request.shift_commercial_markers else 0,
        "purchase_orders_shifted": purchase_orders_shifted,
        "shifted_purchase_orders": shifted_purchase_orders,
        "work_orders_shifted": work_orders_shifted,
        "shifted_work_orders": shifted_work_orders,
        "shifted_schedule_count": shifted_schedule_count,
        "clash_count": len(clashes),
        "clashes": clashes,
    }

@api_router.post("/schedule", response_model=ScheduleEntry)
async def create_schedule_entry(schedule_entry: ScheduleEntryCreate, admin: str = Depends(verify_admin)):
    """Allocate a direct worker or subcontractor worker/team to a job, or block a direct worker day as holiday/sick/unavailable."""
    try:
        datetime.fromisoformat(schedule_entry.scheduled_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="scheduled_date must be in YYYY-MM-DD format")

    schedule_type = normalise_schedule_type(schedule_entry.schedule_type)
    if schedule_type == "job" and not schedule_entry.job_id:
        raise HTTPException(status_code=400, detail="Select an active job first")

    resource_values = normalise_schedule_resource_values(schedule_entry.resource_type, schedule_entry.resource_id, schedule_entry.worker_id)
    resource_type = resource_values["resource_type"]
    resource_id = resource_values["resource_id"]
    worker_id = resource_values["worker_id"]

    if not resource_id:
        raise HTTPException(status_code=400, detail="Select a worker or subcontractor team")

    display_worker_name = schedule_entry.display_worker_name or ""

    if resource_type == "subcontractor_resource":
        if schedule_type != "job":
            raise HTTPException(status_code=400, detail="Holiday/sick/unavailable can only be recorded against direct workers")
        resource_display = await get_subcontractor_resource_display(resource_id)
        if not resource_display:
            raise HTTPException(status_code=404, detail="Active subcontractor worker/team not found")
        resource = resource_display.get("resource", {})
        company = resource_display.get("company", {})
        if resource.get("active") is False or company.get("active") is False:
            raise HTTPException(status_code=404, detail="Active subcontractor worker/team not found")
        display_worker_name = display_worker_name or resource_display.get("display_name", "")
    else:
        worker = await db.workers.find_one({"id": resource_id, "active": True, "archived": {"$ne": True}})
        if not worker:
            raise HTTPException(status_code=404, detail="Active worker not found")
        display_worker_name = display_worker_name or worker.get("name", "")

    if schedule_type == "job":
        job = await db.jobs.find_one({"id": schedule_entry.job_id, "status": "active", "archived": {"$ne": True}})
        if not job:
            raise HTTPException(status_code=404, detail="Active job not found")
    else:
        schedule_entry.job_id = None
        schedule_entry.absence_type = schedule_type

    existing = await db.schedule_entries.find_one({
        "resource_type": resource_type,
        "resource_id": resource_id,
        "scheduled_date": schedule_entry.scheduled_date,
        "archived": {"$ne": True},
    })

    # Backwards compatibility for old direct-worker records that only have worker_id.
    if not existing and resource_type == "worker":
        existing = await db.schedule_entries.find_one({
            "worker_id": resource_id,
            "scheduled_date": schedule_entry.scheduled_date,
            "archived": {"$ne": True},
            "$or": [{"resource_type": {"$exists": False}}, {"resource_type": "worker"}],
        })

    if existing:
        raise HTTPException(status_code=400, detail="This worker or subcontractor team already has a schedule entry on this date")

    entry_dict = schedule_entry.dict()
    entry_dict.update(resource_values)
    entry_dict["display_worker_name"] = display_worker_name
    entry_dict["schedule_type"] = schedule_type
    if schedule_type != "job":
        entry_dict["job_id"] = None
        entry_dict["absence_type"] = schedule_type
    entry_obj = ScheduleEntry(**entry_dict)
    await db.schedule_entries.insert_one(entry_obj.dict())
    return entry_obj

@api_router.put("/schedule/{schedule_id}", response_model=ScheduleEntry)
async def update_schedule_entry(schedule_id: str, schedule_update: ScheduleEntryUpdate, admin: str = Depends(verify_admin)):
    """Update a schedule allocation or absence block (Admin only)."""
    existing_entry = await db.schedule_entries.find_one({"id": schedule_id, "archived": {"$ne": True}})
    if not existing_entry:
        raise HTTPException(status_code=404, detail="Schedule entry not found")

    update_dict = {k: v for k, v in schedule_update.dict().items() if v is not None}
    if not update_dict:
        existing_entry.pop("_id", None)
        existing_entry.update(normalise_schedule_resource_values(existing_entry.get("resource_type"), existing_entry.get("resource_id"), existing_entry.get("worker_id")))
        return ScheduleEntry(**existing_entry)

    existing_resource = normalise_schedule_resource_values(existing_entry.get("resource_type"), existing_entry.get("resource_id"), existing_entry.get("worker_id"))
    requested_resource_type = update_dict.get("resource_type", existing_resource["resource_type"])
    requested_resource_id = update_dict.get("resource_id", update_dict.get("worker_id", existing_resource["resource_id"]))
    requested_worker_id = update_dict.get("worker_id", existing_resource["worker_id"])
    resource_values = normalise_schedule_resource_values(requested_resource_type, requested_resource_id, requested_worker_id)
    new_resource_type = resource_values["resource_type"]
    new_resource_id = resource_values["resource_id"]
    new_worker_id = resource_values["worker_id"]
    new_date = update_dict.get("scheduled_date", existing_entry.get("scheduled_date"))
    new_schedule_type = normalise_schedule_type(update_dict.get("schedule_type", existing_entry.get("schedule_type")))
    new_job_id = update_dict.get("job_id", existing_entry.get("job_id"))

    if "scheduled_date" in update_dict:
        try:
            datetime.fromisoformat(update_dict["scheduled_date"])
        except ValueError:
            raise HTTPException(status_code=400, detail="scheduled_date must be in YYYY-MM-DD format")

    if not new_resource_id:
        raise HTTPException(status_code=400, detail="Select a worker or subcontractor team")

    display_worker_name = update_dict.get("display_worker_name") or existing_entry.get("display_worker_name") or ""

    if new_resource_type == "subcontractor_resource":
        if new_schedule_type != "job":
            raise HTTPException(status_code=400, detail="Holiday/sick/unavailable can only be recorded against direct workers")
        resource_display = await get_subcontractor_resource_display(new_resource_id)
        if not resource_display:
            raise HTTPException(status_code=404, detail="Active subcontractor worker/team not found")
        resource = resource_display.get("resource", {})
        company = resource_display.get("company", {})
        if resource.get("active") is False or company.get("active") is False:
            raise HTTPException(status_code=404, detail="Active subcontractor worker/team not found")
        display_worker_name = display_worker_name or resource_display.get("display_name", "")
    else:
        worker = await db.workers.find_one({"id": new_resource_id, "active": True, "archived": {"$ne": True}})
        if not worker:
            raise HTTPException(status_code=404, detail="Active worker not found")
        display_worker_name = display_worker_name or worker.get("name", "")

    if new_schedule_type == "job":
        if not new_job_id:
            raise HTTPException(status_code=400, detail="Select an active job first")
        job = await db.jobs.find_one({"id": new_job_id, "status": "active", "archived": {"$ne": True}})
        if not job:
            raise HTTPException(status_code=404, detail="Active job not found")
        update_dict["job_id"] = new_job_id
        update_dict["absence_type"] = None
    else:
        update_dict["job_id"] = None
        update_dict["absence_type"] = new_schedule_type

    update_dict.update(resource_values)
    update_dict["display_worker_name"] = display_worker_name
    update_dict["schedule_type"] = new_schedule_type

    duplicate = await db.schedule_entries.find_one({
        "id": {"$ne": schedule_id},
        "resource_type": new_resource_type,
        "resource_id": new_resource_id,
        "scheduled_date": new_date,
        "archived": {"$ne": True},
    })

    # Backwards compatibility for old direct-worker records that only have worker_id.
    if not duplicate and new_resource_type == "worker":
        duplicate = await db.schedule_entries.find_one({
            "id": {"$ne": schedule_id},
            "worker_id": new_resource_id,
            "scheduled_date": new_date,
            "archived": {"$ne": True},
            "$or": [{"resource_type": {"$exists": False}}, {"resource_type": "worker"}],
        })

    if duplicate:
        raise HTTPException(status_code=400, detail="This worker or subcontractor team already has a schedule entry on this date")

    update_dict["updated_date"] = datetime.utcnow()
    result = await db.schedule_entries.update_one({"id": schedule_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Schedule entry not found")

    updated_entry = await db.schedule_entries.find_one({"id": schedule_id})
    updated_entry.pop("_id", None)
    updated_entry.update(normalise_schedule_resource_values(updated_entry.get("resource_type"), updated_entry.get("resource_id"), updated_entry.get("worker_id")))
    return ScheduleEntry(**updated_entry)

@api_router.delete("/schedule/{schedule_id}")
async def delete_schedule_entry(schedule_id: str, admin: str = Depends(verify_admin)):
    """Delete a schedule allocation (Admin only)."""
    result = await db.schedule_entries.delete_one({"id": schedule_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Schedule entry not found")
    return {"message": "Schedule entry deleted successfully"}



async def build_schedule_export_data(
    start_date: str,
    end_date: str,
    worker_ids: Optional[str] = None,
    worker_type: Optional[str] = None,
    division: Optional[str] = None,
    trade: Optional[str] = None
):
    """Build schedule rows for CSV/PDF export and worker/job schedule views."""
    try:
        datetime.fromisoformat(start_date)
        datetime.fromisoformat(end_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="Dates must be in YYYY-MM-DD format")

    worker_filter = {"active": True, "archived": {"$ne": True}, "role": {"$ne": "admin"}}
    if worker_type and worker_type != "all":
        worker_filter["worker_type"] = worker_type
    if division and division != "all":
        worker_filter["division"] = division
    if trade and trade != "all":
        worker_filter["$or"] = [{"trades": trade}, {"trade": trade}]

    selected_worker_ids = []
    if worker_ids:
        selected_worker_ids = [worker_id.strip() for worker_id in worker_ids.split(",") if worker_id.strip()]
        if selected_worker_ids:
            worker_filter["id"] = {"$in": selected_worker_ids}

    workers = await db.workers.find(worker_filter).to_list(1000)
    workers.sort(key=lambda worker: worker.get("name", ""))
    allowed_worker_ids = [worker.get("id") for worker in workers if worker.get("id")]

    schedule_filter = {
        "scheduled_date": {"$gte": start_date, "$lte": end_date},
        "archived": {"$ne": True},
    }

    # Important for job exports: when filters or selected workers are applied, only export
    # allocations for the workers that survived those filters.
    if allowed_worker_ids:
        schedule_filter["worker_id"] = {"$in": allowed_worker_ids}
    else:
        schedule_filter["worker_id"] = {"$in": []}

    entries = await db.schedule_entries.find(schedule_filter).to_list(1000)
    job_ids = list({entry.get("job_id") for entry in entries if entry.get("job_id")})
    jobs = await db.jobs.find({"id": {"$in": job_ids}}).to_list(1000) if job_ids else []
    job_lookup = {job["id"]: job for job in jobs if "id" in job}
    worker_lookup = {worker["id"]: worker for worker in workers if "id" in worker}

    entry_lookup = {}
    job_day_lookup = {}
    enriched_entries = []

    for entry in entries:
        entry.pop("_id", None)
        job = job_lookup.get(entry.get("job_id"), {})
        worker_match = worker_lookup.get(entry.get("worker_id"), {})

        schedule_type = normalise_schedule_type(entry.get("schedule_type"))
        entry["schedule_type"] = schedule_type
        entry["absence_type"] = entry.get("absence_type") or (schedule_type if is_absence_schedule_type(schedule_type) else None)
        if is_absence_schedule_type(schedule_type):
            entry["job_name"] = schedule_type_label(schedule_type)
            entry["job_client"] = ""
            entry["job_location"] = ""
        else:
            entry["job_name"] = job.get("name", "Unknown")
            entry["job_client"] = job.get("client", "")
            entry["job_location"] = job.get("location", "")
        entry["worker_name"] = worker_match.get("name", "Unknown")
        entry["worker_type"] = worker_match.get("worker_type", "worker")
        entry["worker_division"] = worker_match.get("division", "")
        entry["worker_trades"] = worker_match.get("trades", [worker_match.get("trade", "")] if worker_match.get("trade") else [])

        enriched_entries.append(entry)
        entry_lookup[(entry.get("worker_id"), entry.get("scheduled_date"))] = entry
        if not is_absence_schedule_type(schedule_type):
            job_day_lookup.setdefault((entry.get("job_id"), entry.get("scheduled_date")), []).append(entry)

    export_jobs = []
    seen_job_ids = set()
    for entry in enriched_entries:
        job_id_value = entry.get("job_id")
        if not job_id_value or job_id_value in seen_job_ids:
            continue
        job = job_lookup.get(job_id_value, {})
        export_jobs.append({
            "id": job_id_value,
            "name": job.get("name", entry.get("job_name", "Unknown")),
            "client": job.get("client", entry.get("job_client", "")),
            "location": job.get("location", entry.get("job_location", "")),
        })
        seen_job_ids.add(job_id_value)

    export_jobs.sort(key=lambda job: job.get("name", ""))

    return {
        "workers": workers,
        "jobs": export_jobs,
        "entries": enriched_entries,
        "entry_lookup": entry_lookup,
        "job_day_lookup": job_day_lookup,
    }

@api_router.get("/schedule/worker/{worker_id}", response_model=List[Dict[str, Any]])
async def get_worker_schedule_entries(
    worker_id: str,
    start_date: str = Query(...),
    end_date: str = Query(...)
):
    """Get one worker's schedule between two dates. Used by the worker frontend."""
    try:
        datetime.fromisoformat(start_date)
        datetime.fromisoformat(end_date)
    except ValueError:
        raise HTTPException(status_code=400, detail="Dates must be in YYYY-MM-DD format")

    worker = await db.workers.find_one({"id": worker_id, "active": True, "archived": {"$ne": True}})
    if not worker:
        raise HTTPException(status_code=404, detail="Worker not found")

    entries = await db.schedule_entries.find({
        "worker_id": worker_id,
        "scheduled_date": {"$gte": start_date, "$lte": end_date},
        "archived": {"$ne": True}
    }).to_list(1000)

    job_ids = list({entry.get("job_id") for entry in entries if entry.get("job_id")})
    jobs = await db.jobs.find({"id": {"$in": job_ids}}).to_list(1000) if job_ids else []
    job_lookup = {job["id"]: job for job in jobs if "id" in job}

    result = []
    for entry in entries:
        entry.pop("_id", None)
        job = job_lookup.get(entry.get("job_id"), {})
        schedule_type = normalise_schedule_type(entry.get("schedule_type"))
        entry["schedule_type"] = schedule_type
        entry["absence_type"] = entry.get("absence_type") or (schedule_type if is_absence_schedule_type(schedule_type) else None)
        entry["worker_name"] = worker.get("name", "")
        if is_absence_schedule_type(schedule_type):
            entry["job_name"] = schedule_type_label(schedule_type)
            entry["job_client"] = ""
            entry["job_location"] = ""
        else:
            entry["job_name"] = job.get("name", "Unknown")
            entry["job_client"] = job.get("client", "")
            entry["job_location"] = job.get("location", "")
        result.append(entry)

    result.sort(key=lambda item: item.get("scheduled_date", ""))
    return result

@api_router.get("/schedule/export")
async def export_schedule(
    start_date: str = Query(...),
    end_date: str = Query(...),
    worker_ids: Optional[str] = Query(None),
    worker_type: Optional[str] = Query(None),
    division: Optional[str] = Query(None),
    trade: Optional[str] = Query(None),
    group_by: str = Query("worker"),
    format: str = Query("csv"),
    admin: str = Depends(verify_admin)
):
    """Export the weekly schedule as CSV or PDF, grouped by worker or by job."""
    group_by = (group_by or "worker").lower().strip()
    if group_by not in ["worker", "job"]:
        group_by = "worker"

    export_data = await build_schedule_export_data(start_date, end_date, worker_ids, worker_type, division, trade)
    workers = export_data["workers"]
    jobs = export_data["jobs"]
    entry_lookup = export_data["entry_lookup"]
    job_day_lookup = export_data["job_day_lookup"]

    start_dt = datetime.fromisoformat(start_date)
    end_dt = datetime.fromisoformat(end_date)
    date_list = []
    cursor = start_dt
    while cursor <= end_dt:
        date_list.append(cursor.strftime("%Y-%m-%d"))
        cursor += timedelta(days=1)

    def worker_cell_text(worker_id: str, day: str) -> str:
        entry = entry_lookup.get((worker_id, day))
        if not entry:
            return "Unallocated"
        lines = [entry.get("job_name", "Scheduled job")]
        if entry.get("job_client"):
            lines.append(entry.get("job_client"))
        if entry.get("job_location"):
            lines.append(entry.get("job_location"))
        if entry.get("notes"):
            lines.append(f"Notes: {entry.get('notes')}")
        return " | ".join(lines)

    def job_cell_text(job_id: str, day: str) -> str:
        entries = job_day_lookup.get((job_id, day), [])
        if not entries:
            return "Unallocated"

        lines = []
        for entry in sorted(entries, key=lambda item: item.get("worker_name", "")):
            worker_line = entry.get("worker_name", "Unknown worker")
            extras = []
            if entry.get("worker_division"):
                extras.append(entry.get("worker_division"))
            if entry.get("notes"):
                extras.append(f"Notes: {entry.get('notes')}")
            if extras:
                worker_line += " (" + "; ".join(extras) + ")"
            lines.append(worker_line)
        return " | ".join(lines)

    export_title = "LDA Group - Weekly Job Schedule" if group_by == "job" else "LDA Group - Weekly Worker Schedule"
    first_column = "Job" if group_by == "job" else "Worker"

    if format.lower() == "pdf":
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        except Exception as e:
            logger.error(f"PDF export dependency error: {e}")
            raise HTTPException(status_code=500, detail="PDF export is not available on this server")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=1*cm, leftMargin=1*cm, topMargin=1*cm, bottomMargin=1*cm)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(export_title, styles["Title"]),
            Paragraph(f"{start_date} to {end_date}", styles["Normal"]),
            Spacer(1, 0.4*cm),
        ]

        headers = [first_column] + [datetime.fromisoformat(day).strftime("%a %d %b") for day in date_list]
        table_data = [headers]

        def worker_pdf_cell_text(worker_id: str, day: str) -> str:
            entry = entry_lookup.get((worker_id, day))
            if not entry:
                return ""
            return entry.get("job_name") or "Scheduled job"

        def job_pdf_cell_text(job_id: str, day: str) -> str:
            entries = job_day_lookup.get((job_id, day), [])
            if not entries:
                return ""
            names = sorted({entry.get("job_name") or "Scheduled job" for entry in entries})
            return "<br/>".join(names)

        if group_by == "job":
            for job in jobs:
                row = [Paragraph(job.get("name", "Unknown"), styles["BodyText"])]
                for day in date_list:
                    row.append(Paragraph(job_pdf_cell_text(job.get("id"), day), styles["BodyText"]))
                table_data.append(row)
        else:
            for worker in workers:
                row = [worker.get("name", "Unknown")]
                for day in date_list:
                    row.append(Paragraph(worker_pdf_cell_text(worker.get("id"), day), styles["BodyText"]))
                table_data.append(row)

        if len(table_data) == 1:
            empty_label = "No scheduled jobs" if group_by == "job" else "No selected workers"
            table_data.append([empty_label] + ["" for _ in date_list])

        col_widths = [3.2*cm] + [3.1*cm for _ in date_list]
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d01f2f")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f7f7")]),
        ]))
        story.append(table)
        doc.build(story)
        buffer.seek(0)
        filename = f"{group_by}_schedule_{start_date}_to_{end_date}.pdf"
        return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([export_title])
    writer.writerow(["Date range", start_date, "to", end_date])
    writer.writerow([])
    writer.writerow([first_column] + [datetime.fromisoformat(day).strftime("%a %d %b %Y") for day in date_list])

    if group_by == "job":
        for job in jobs:
            job_label_parts = [job.get("name", "Unknown")]
            if job.get("client"):
                job_label_parts.append(job.get("client"))
            if job.get("location"):
                job_label_parts.append(job.get("location"))
            job_label = " | ".join(job_label_parts)
            writer.writerow([job_label] + [job_cell_text(job.get("id"), day) for day in date_list])
    else:
        for worker in workers:
            writer.writerow([worker.get("name", "Unknown")] + [worker_cell_text(worker.get("id"), day) for day in date_list])

    output.seek(0)
    filename = f"{group_by}_schedule_{start_date}_to_{end_date}.csv"
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )



# ==================== FINANCE SECTION ENDPOINTS ====================

def finance_to_number(value: Any, fallback: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return fallback
        if isinstance(value, str):
            value = value.replace("£", "").replace(",", "").strip()
            if not value:
                return fallback
        return float(value)
    except (TypeError, ValueError):
        return fallback


def finance_round_money(value: Any) -> float:
    return round(finance_to_number(value, 0.0) + 0.0000001, 2)




def normalise_vat_treatment(value: Optional[str]) -> str:
    treatment = str(value or "standard_20").strip().lower().replace(" ", "_").replace("-", "_")
    aliases = {
        "standard": "standard_20",
        "standard_rate": "standard_20",
        "20": "standard_20",
        "20%": "standard_20",
        "reduced": "reduced_5",
        "reduced_rate": "reduced_5",
        "5": "reduced_5",
        "5%": "reduced_5",
        "zero": "zero_rated",
        "zero_rate": "zero_rated",
        "zero_rated": "zero_rated",
        "exempt": "exempt",
        "vat_exempt": "exempt",
        "outside_scope": "no_vat",
        "no_vat": "no_vat",
        "none": "no_vat",
        "drc": "drc",
        "reverse_charge": "drc",
        "domestic_reverse_charge": "drc",
    }
    return aliases.get(treatment, treatment if treatment in ["standard_20", "reduced_5", "zero_rated", "exempt", "no_vat", "drc"] else "standard_20")


def normalise_job_tax_settings(job: Dict[str, Any]) -> Dict[str, Any]:
    treatment = normalise_vat_treatment(job.get("vat_treatment"))
    drc_enabled = bool(job.get("drc_enabled")) or treatment == "drc"
    if treatment == "drc" or drc_enabled:
        vat_rate = finance_to_number(job.get("vat_rate"), 20.0) or 20.0
        vat_charged_rate = 0.0
        treatment = "drc"
    elif treatment == "reduced_5":
        vat_rate = finance_to_number(job.get("vat_rate"), 5.0) or 5.0
        vat_charged_rate = vat_rate
    elif treatment in ["zero_rated", "exempt", "no_vat"]:
        vat_rate = 0.0
        vat_charged_rate = 0.0
    else:
        vat_rate = finance_to_number(job.get("vat_rate"), 20.0) or 20.0
        vat_charged_rate = vat_rate

    cis_enabled = bool(job.get("cis_enabled"))
    cis_rate = max(0.0, min(100.0, finance_to_number(job.get("cis_rate"), 0.0))) if cis_enabled else 0.0
    basis = str(job.get("cis_deduction_basis") or "labour_only").strip().lower()
    if basis not in ["labour_only", "full_net", "none"]:
        basis = "labour_only"

    return {
        "vat_treatment": treatment,
        "vat_rate": round(vat_rate, 2),
        "vat_charged_rate": round(vat_charged_rate, 2),
        "drc_enabled": drc_enabled,
        "cis_enabled": cis_enabled,
        "cis_rate": round(cis_rate, 2),
        "cis_deduction_basis": basis,
        "tax_notes": job.get("tax_notes", ""),
    }


def calculate_finance_tax_snapshot(job: Dict[str, Any], net_value: float, labour_value: Optional[float] = None, material_value: Optional[float] = None) -> Dict[str, Any]:
    settings = normalise_job_tax_settings(job or {})
    net = round(max(0.0, finance_to_number(net_value)), 2)
    labour = finance_to_number(labour_value, 0.0)
    material = finance_to_number(material_value, 0.0)
    if labour <= 0 and material <= 0:
        # Conservative default: CIS labour-only cannot be calculated accurately without a split,
        # so use the net as the labour basis and flag that it is estimated.
        labour = net
        split_estimated = True
    else:
        split_estimated = False

    vat_value = round(net * (settings["vat_charged_rate"] / 100.0), 2)
    gross_value = round(net + vat_value, 2)

    cis_basis_value = 0.0
    if settings["cis_enabled"] and settings["cis_deduction_basis"] != "none":
        if settings["cis_deduction_basis"] == "full_net":
            cis_basis_value = net
        else:
            cis_basis_value = min(net, max(0.0, labour))
    cis_deduction = round(cis_basis_value * (settings["cis_rate"] / 100.0), 2)
    expected_cash = round(max(0.0, gross_value - cis_deduction), 2)

    return {
        **settings,
        "net_value": net,
        "labour_value": round(max(0.0, labour), 2),
        "material_value": round(max(0.0, material), 2),
        "split_estimated": split_estimated,
        "vat_value": vat_value,
        "gross_value": gross_value,
        "cis_basis_value": round(cis_basis_value, 2),
        "cis_deduction_value": cis_deduction,
        "expected_cash_value": expected_cash,
    }

def parse_iso_date_safe(value: Optional[str]) -> Optional[date]:
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00")).date()
    except Exception:
        return None


def section_finance_values(section: Dict[str, Any]) -> Dict[str, float]:
    labour = finance_to_number(section.get("labour_value"))
    material = finance_to_number(section.get("material_value"))
    subcontractor = finance_to_number(section.get("subcontractor_value"))
    other = finance_to_number(section.get("other_value"))
    calculated_total = labour + material + subcontractor + other
    stored_total = finance_to_number(section.get("section_value") or section.get("total_value") or section.get("value"))
    total = calculated_total if calculated_total > 0 else stored_total
    progress = max(0.0, min(100.0, finance_to_number(section.get("progress_percent"))))
    earned = total * (progress / 100.0)
    return {
        "labour_value": round(labour, 2),
        "material_value": round(material, 2),
        "subcontractor_value": round(subcontractor, 2),
        "other_value": round(other, 2),
        "section_value": round(total, 2),
        "progress_percent": round(progress, 2),
        "earned_value": round(earned, 2),
        "remaining_value": round(max(0.0, total - earned), 2),
    }


def section_planned_value_by_date(section: Dict[str, Any], marker_date: Optional[date]) -> float:
    if not marker_date:
        return 0.0
    values = section_finance_values(section)
    section_value = values["section_value"]
    if section_value <= 0:
        return 0.0
    start = parse_iso_date_safe(section.get("start_date"))
    end = parse_iso_date_safe(section.get("end_date")) or start
    if not start or not end:
        return 0.0
    if end < start:
        start, end = end, start
    if marker_date < start:
        return 0.0
    if marker_date >= end:
        return section_value
    total_days = max(1, (end - start).days + 1)
    days_to_marker = max(0, (marker_date - start).days + 1)
    ratio = max(0.0, min(1.0, days_to_marker / total_days))
    return round(section_value * ratio, 2)


def normalise_finance_marker(marker: Dict[str, Any]) -> Dict[str, Any]:
    marker_type = str(marker.get("type") or marker.get("marker_type") or "application").strip().lower()
    if marker_type in ["final", "final_application", "final_invoice"]:
        marker_type = "final_invoice"
    marker_id = marker.get("id") or str(uuid.uuid4())
    label = marker.get("label") or marker.get("name") or marker_type.replace("_", " ").title()
    raw_deposit_percentage = marker.get("deposit_percentage", marker.get("deposit_percent", marker.get("depositPercentage", "")))
    if raw_deposit_percentage in [None, ""]:
        deposit_percentage = None
    else:
        deposit_percentage = max(0.0, min(100.0, finance_to_number(raw_deposit_percentage)))
    value_mode = marker.get("value_mode") or marker.get("value_type") or marker.get("calculationMode") or marker.get("calculation_mode")
    if not value_mode:
        value_mode = "manual" if marker.get("manual") is True or marker.get("isManual") is True else "auto"
    stored_value = finance_round_money(marker.get("net_value", marker.get("netValue", marker.get("value", marker.get("amount", marker.get("manual_value", 0))))))
    manual_value = finance_round_money(marker.get("manual_value", stored_value if value_mode == "manual" else 0))
    return {
        **marker,
        "id": marker_id,
        "type": marker_type,
        "marker_type": marker_type,
        "label": label,
        "date": marker.get("date") or marker.get("marker_date") or marker.get("expected_date") or marker.get("payment_due_date") or marker.get("retention_due_date") or "",
        "value_mode": "manual" if value_mode == "manual" else "auto",
        "value_type": "manual" if value_mode == "manual" else "auto",
        "calculationMode": "manual" if value_mode == "manual" else "auto",
        "calculation_mode": "manual" if value_mode == "manual" else "auto",
        "manual": value_mode == "manual",
        "isManual": value_mode == "manual",
        "auto": value_mode != "manual",
        "isAuto": value_mode != "manual",
        "manual_value": manual_value,
        "value": stored_value,
        "net_value": stored_value,
        "netValue": stored_value,
        "gross_value": finance_round_money(marker.get("gross_value", stored_value)),
        "deposit_percentage": deposit_percentage,
        "deduct_deposit": marker.get("deduct_deposit") is True or marker.get("deposit_deduction_enabled") is True or str(marker.get("deduct_deposit", "")).lower() == "true",
        "retention_percent": finance_to_number(marker.get("retention_percent", marker.get("retentionPercent", 0))),
        "notes": marker.get("notes") or "",
        "status": marker.get("status") or marker.get("payment_status") or "expected",
    }


def finance_marker_stored_value(marker: Dict[str, Any]) -> float:
    return finance_round_money(marker.get("net_value", marker.get("netValue", marker.get("value", marker.get("amount", marker.get("manual_value", 0))))))


def finance_marker_is_claimed(marker: Dict[str, Any]) -> bool:
    status = str(marker.get("status") or marker.get("payment_status") or "").strip().lower()
    return status in {"claimed", "submitted", "applied", "applied_submitted", "invoiced", "part_received", "part paid", "part_paid", "received", "paid"}


def finance_marker_is_fixed(marker: Dict[str, Any]) -> bool:
    marker_type = marker.get("type")
    return bool(
        marker_type in ["deposit", "final_invoice", "retention", "retention_release"]
        or marker.get("value_mode") == "manual"
        or marker.get("value_type") == "manual"
        or marker.get("calculationMode") == "manual"
        or marker.get("calculation_mode") == "manual"
        or marker.get("manual") is True
        or marker.get("isManual") is True
        or marker.get("locked") is True
        or finance_marker_is_claimed(marker)
    )

def finance_contract_value(job: Dict[str, Any], sections: Optional[List[Dict[str, Any]]] = None) -> float:
    section_source = sections if sections is not None else (job.get("gantt_sections") or [])
    section_total = sum(section_finance_values(section).get("section_value", 0.0) for section in section_source)
    current_contract = finance_to_number(job.get("current_contract_value"))
    if current_contract > 0:
        return round(current_contract, 2)
    base_value = finance_to_number(job.get("original_quoted_cost")) or finance_to_number(job.get("quoted_cost"))
    approved_variations = finance_to_number(job.get("approved_variations_total"))
    if base_value > 0 or approved_variations > 0:
        return round(base_value + approved_variations, 2)
    return round(section_total, 2)


def finance_deposit_marker_value(job: Dict[str, Any], marker: Dict[str, Any], contract_value: Optional[float] = None) -> float:
    contract_value = finance_contract_value(job) if contract_value is None else contract_value
    deposit_percentage = marker.get("deposit_percentage")
    if deposit_percentage is not None:
        return round(contract_value * (finance_to_number(deposit_percentage) / 100.0), 2)
    return round(finance_to_number(marker.get("manual_value") or marker.get("value") or marker.get("amount")), 2)


def build_finance_payment_schedule(job: Dict[str, Any], sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    markers = [normalise_finance_marker(marker) for marker in (job.get("commercial_markers") or [])]
    markers = sorted([marker for marker in markers if marker.get("date")], key=lambda item: item.get("date") or "")
    contract_value = finance_contract_value(job, sections)
    earned_to_date = sum(section_finance_values(section).get("earned_value", 0.0) for section in sections)

    fixed_total = 0.0
    auto_indexes: List[int] = []
    for index, marker in enumerate(markers):
        marker_type = marker.get("type")
        if marker_type == "deposit":
            fixed_total += finance_deposit_marker_value(job, marker, contract_value)
        elif marker_type in ["application", "interim"] and not finance_marker_is_fixed(marker):
            auto_indexes.append(index)
        else:
            fixed_total += finance_marker_stored_value(marker)

    if auto_indexes and contract_value > 0:
        remaining = max(0.0, finance_round_money(contract_value - fixed_total))
        base = math.floor((remaining / len(auto_indexes)) * 100) / 100
        allocated = 0.0
        for i, marker_index in enumerate(auto_indexes):
            value = finance_round_money(remaining - allocated) if i == len(auto_indexes) - 1 else finance_round_money(base)
            allocated = finance_round_money(allocated + value)
            markers[marker_index].update({
                "value": value,
                "net_value": value,
                "netValue": value,
                "gross_value": value,
                "manual_value": 0.0,
                "value_mode": "auto",
                "value_type": "auto",
                "calculationMode": "auto",
                "calculation_mode": "auto",
                "manual": False,
                "isManual": False,
                "auto": True,
                "isAuto": True,
            })

    previous_application_earned_cumulative = 0.0
    previous_net_payments = 0.0
    rows = []

    for marker in markers:
        marker_type = marker.get("type")
        marker_date = parse_iso_date_safe(marker.get("date"))
        manual_value = finance_to_number(marker.get("manual_value")) or finance_marker_stored_value(marker)
        value_mode = marker.get("value_mode") or "auto"
        cumulative_planned = sum(section_planned_value_by_date(section, marker_date) for section in sections)
        cumulative_earned = sum(min(section_finance_values(section).get("earned_value", 0.0), section_planned_value_by_date(section, marker_date)) for section in sections)
        gross_value = 0.0
        earned_period = 0.0
        deposit_deduction = 0.0

        if marker_type == "deposit":
            gross_value = finance_deposit_marker_value(job, marker, contract_value)
            earned_period = gross_value
        elif marker_type in ["application", "interim"]:
            capped_planned = max(0.0, min(contract_value or cumulative_planned, cumulative_planned))
            capped_earned = max(0.0, min(contract_value or cumulative_earned, cumulative_earned))
            stored_value = finance_marker_stored_value(marker)
            if stored_value > 0:
                gross_value = stored_value
            elif value_mode == "manual" and manual_value > 0:
                gross_value = manual_value
            else:
                gross_value = max(0.0, capped_planned - previous_net_payments)
            earned_period = max(0.0, capped_earned - previous_application_earned_cumulative)
            previous_application_earned_cumulative = max(previous_application_earned_cumulative, capped_earned)
        elif marker_type == "final_invoice":
            stored_value = finance_marker_stored_value(marker)
            gross_value = stored_value if stored_value > 0 else (manual_value if value_mode == "manual" and manual_value > 0 else max(0.0, contract_value - previous_net_payments))
            earned_period = max(0.0, earned_to_date - previous_application_earned_cumulative)
            previous_application_earned_cumulative = max(previous_application_earned_cumulative, earned_to_date)
        elif marker_type == "retention":
            retention_percent = finance_to_number(marker.get("retention_percent"))
            gross_value = manual_value if manual_value > 0 else (contract_value * (retention_percent / 100.0) if retention_percent > 0 else 0.0)
            earned_period = gross_value
        else:
            stored_value = finance_marker_stored_value(marker)
            gross_value = manual_value if value_mode == "manual" and manual_value > 0 else (stored_value or cumulative_planned)
            earned_period = min(gross_value, cumulative_earned)

        net_value = max(0.0, gross_value - deposit_deduction)
        net_earned = net_value if marker_type == "deposit" else max(0.0, earned_period - deposit_deduction)
        risk_value = max(0.0, net_value - net_earned)
        previous_net_payments += net_value

        rows.append({
            **marker,
            "gross_value": finance_round_money(gross_value),
            "deposit_deduction": finance_round_money(deposit_deduction),
            "net_value": finance_round_money(net_value),
            "netValue": finance_round_money(net_value),
            "value": finance_round_money(net_value),
            "earned_value": finance_round_money(net_earned),
            "risk_value": finance_round_money(risk_value),
            "contract_value": finance_round_money(contract_value),
        })

    return rows


async def build_finance_project_summary(job_id: str) -> Dict[str, Any]:
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    sections = []
    for section in job.get("gantt_sections") or []:
        values = section_finance_values(section)
        sections.append({
            "id": section.get("id"),
            "name": section.get("name", "Section"),
            "start_date": section.get("start_date", ""),
            "end_date": section.get("end_date", ""),
            "status": section.get("status", ""),
            **values,
        })

    records = await db.finance_records.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(1000)
    total_value = round(sum(item["section_value"] for item in sections), 2)
    earned_value = round(sum(item["earned_value"] for item in sections), 2)
    remaining_value = round(max(0.0, total_value - earned_value), 2)
    submitted_to_date = round(sum(finance_to_number(item.get("submitted_value")) for item in records), 2)
    certified_to_date = round(sum(finance_to_number(item.get("certified_value")) for item in records), 2)
    invoiced_to_date = round(sum(finance_to_number(item.get("invoice_value")) for item in records), 2)
    paid_to_date = round(sum(finance_to_number(item.get("paid_value")) for item in records), 2)
    retention_held = round(sum(finance_to_number(item.get("retention_value")) for item in records if not item.get("retention_paid_date")), 2)
    outstanding = round(max(0.0, invoiced_to_date - paid_to_date), 2)

    payment_schedule = build_finance_payment_schedule(job, job.get("gantt_sections") or [])
    application_forecast = []
    for marker in payment_schedule:
        if marker.get("type") not in ["application", "interim", "final_invoice"]:
            continue
        linked_records = [record for record in records if record.get("application_marker_id") == marker.get("id")]
        submitted = round(sum(finance_to_number(item.get("submitted_value")) for item in linked_records), 2)
        certified = round(sum(finance_to_number(item.get("certified_value")) for item in linked_records), 2)
        invoiced = round(sum(finance_to_number(item.get("invoice_value")) for item in linked_records), 2)
        paid = round(sum(finance_to_number(item.get("paid_value")) for item in linked_records), 2)
        forecast_value = finance_to_number(marker.get("net_value"))
        earned_marker_value = finance_to_number(marker.get("earned_value"))
        shortfall = round(max(0.0, forecast_value - earned_marker_value), 2)
        application_forecast.append({
            "marker_id": marker.get("id"),
            "label": marker.get("label"),
            "date": marker.get("date"),
            "type": marker.get("type"),
            "value_mode": marker.get("value_mode") or "auto",
            "gross_value": finance_to_number(marker.get("gross_value")),
            "deposit_deduction": finance_to_number(marker.get("deposit_deduction")),
            "planned_value": finance_to_number(marker.get("gross_value")),
            "forecast_value": round(forecast_value, 2),
            "net_forecast_value": round(forecast_value, 2),
            "earned_value": round(earned_marker_value, 2),
            "shortfall": shortfall,
            "suggested_claim": round(min(earned_marker_value, forecast_value), 2),
            "submitted_value": submitted,
            "certified_value": certified,
            "invoice_value": invoiced,
            "paid_value": paid,
            "status": "at_risk" if shortfall > 0 else "on_track",
        })
    application_forecast.sort(key=lambda item: item.get("date") or "")

    next_application = None
    today = datetime.utcnow().date()
    future_apps = [item for item in application_forecast if parse_iso_date_safe(item.get("date")) and parse_iso_date_safe(item.get("date")) >= today]
    if future_apps:
        next_application = future_apps[0]
    elif application_forecast:
        next_application = application_forecast[-1]

    return {
        "job": {
            "id": job.get("id"),
            "name": job.get("name"),
            "client": job.get("client", ""),
            "location": job.get("location", ""),
            "quoted_cost": job.get("quoted_cost", 0),
            "planned_start_date": job.get("planned_start_date"),
            "planned_end_date": job.get("planned_end_date"),
            "status": job.get("status", "active"),
        },
        "sections": sections,
        "markers": markers,
        "application_forecast": application_forecast,
        "next_application": next_application,
        "records": records,
        "summary": {
            "contract_value": total_value,
            "earned_value": earned_value,
            "remaining_value": remaining_value,
            "submitted_to_date": submitted_to_date,
            "certified_to_date": certified_to_date,
            "invoiced_to_date": invoiced_to_date,
            "paid_to_date": paid_to_date,
            "retention_held": retention_held,
            "outstanding": outstanding,
            "application_risk": next_application.get("shortfall", 0) if next_application else 0,
        },
    }


@api_router.get("/finance/projects")
async def get_finance_projects(admin: str = Depends(verify_admin)):
    """Return active projects that can be reviewed in the Finance section."""
    jobs = await db.jobs.find({"archived": {"$ne": True}}, {"_id": 0}).sort("name", 1).to_list(1000)
    result = []
    for job in jobs:
        sections = job.get("gantt_sections") or []
        section_value = round(sum(section_finance_values(section)["section_value"] for section in sections), 2)
        earned_value = round(sum(section_finance_values(section)["earned_value"] for section in sections), 2)
        result.append({
            "id": job.get("id"),
            "name": job.get("name"),
            "client": job.get("client", ""),
            "location": job.get("location", ""),
            "status": job.get("status", "active"),
            "include_in_gantt": job.get("include_in_gantt", False),
            "section_count": len(sections),
            "section_value": section_value,
            "earned_value": earned_value,
            "commercial_marker_count": len(job.get("commercial_markers") or []),
        })
    return result


@api_router.get("/finance/project/{job_id}/summary")
async def get_finance_project_summary(job_id: str, admin: str = Depends(verify_admin)):
    return await build_finance_project_summary(job_id)


@api_router.get("/finance/project/{job_id}/records")
async def get_finance_records(job_id: str, admin: str = Depends(verify_admin)):
    records = await db.finance_records.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).sort("created_date", -1).to_list(1000)
    return records


@api_router.post("/finance/records", response_model=FinanceRecord)
async def create_finance_record(record: FinanceRecordCreate, admin: str = Depends(verify_admin)):
    job = await db.jobs.find_one({"id": record.job_id, "archived": {"$ne": True}})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    record_dict = record.dict()
    if finance_to_number(record_dict.get("retention_value")) <= 0 and finance_to_number(record_dict.get("retention_percent")) > 0:
        base_value = finance_to_number(record_dict.get("invoice_value")) or finance_to_number(record_dict.get("certified_value")) or finance_to_number(record_dict.get("submitted_value"))
        record_dict["retention_value"] = round(base_value * (finance_to_number(record_dict.get("retention_percent")) / 100.0), 2)
    record_obj = FinanceRecord(**record_dict)
    await db.finance_records.insert_one(record_obj.dict())
    return record_obj


@api_router.put("/finance/records/{record_id}", response_model=FinanceRecord)
async def update_finance_record(record_id: str, record_update: FinanceRecordUpdate, admin: str = Depends(verify_admin)):
    existing = await db.finance_records.find_one({"id": record_id, "archived": {"$ne": True}})
    if not existing:
        raise HTTPException(status_code=404, detail="Finance record not found")
    update_dict = {k: v for k, v in record_update.dict().items() if v is not None}
    if update_dict:
        base_value = finance_to_number(update_dict.get("invoice_value", existing.get("invoice_value"))) or finance_to_number(update_dict.get("certified_value", existing.get("certified_value"))) or finance_to_number(update_dict.get("submitted_value", existing.get("submitted_value")))
        retention_percent = finance_to_number(update_dict.get("retention_percent", existing.get("retention_percent")))
        if "retention_value" not in update_dict and retention_percent > 0 and base_value > 0:
            update_dict["retention_value"] = round(base_value * (retention_percent / 100.0), 2)
        update_dict["updated_date"] = datetime.utcnow()
        result = await db.finance_records.update_one({"id": record_id}, {"$set": update_dict})
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Finance record not found")
    updated = await db.finance_records.find_one({"id": record_id}, {"_id": 0})
    return FinanceRecord(**updated)


@api_router.delete("/finance/records/{record_id}")
async def delete_finance_record(record_id: str, admin: str = Depends(verify_admin)):
    result = await db.finance_records.update_one({"id": record_id}, {"$set": {"archived": True, "updated_date": datetime.utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Finance record not found")
    return {"message": "Finance record archived successfully"}


@api_router.get("/finance/project/{job_id}/export.csv")
async def export_finance_project_csv(job_id: str, admin: str = Depends(verify_admin)):
    data = await build_finance_project_summary(job_id)
    output = io.StringIO()
    writer = csv.writer(output)
    job = data["job"]
    summary = data["summary"]
    writer.writerow(["LDA Group - Finance Summary"])
    writer.writerow(["Project", job.get("name", "")])
    writer.writerow(["Client", job.get("client", "")])
    writer.writerow([])
    writer.writerow(["Contract Value", summary["contract_value"]])
    writer.writerow(["Earned Value", summary["earned_value"]])
    writer.writerow(["Remaining Value", summary["remaining_value"]])
    writer.writerow(["Submitted To Date", summary["submitted_to_date"]])
    writer.writerow(["Certified To Date", summary["certified_to_date"]])
    writer.writerow(["Invoiced To Date", summary["invoiced_to_date"]])
    writer.writerow(["Paid To Date", summary["paid_to_date"]])
    writer.writerow(["Outstanding", summary["outstanding"]])
    writer.writerow(["Retention Held", summary["retention_held"]])
    writer.writerow([])
    writer.writerow(["Section", "Start", "End", "Labour", "Materials", "Subcontractor", "Other", "Total", "Progress %", "Earned", "Remaining"])
    for section in data["sections"]:
        writer.writerow([section.get("name"), section.get("start_date"), section.get("end_date"), section.get("labour_value"), section.get("material_value"), section.get("subcontractor_value"), section.get("other_value"), section.get("section_value"), section.get("progress_percent"), section.get("earned_value"), section.get("remaining_value")])
    writer.writerow([])
    writer.writerow(["Application", "Date", "Planned", "Earned", "Shortfall", "Suggested Claim", "Submitted", "Certified", "Invoiced", "Paid", "Status"])
    for app in data["application_forecast"]:
        writer.writerow([app.get("label"), app.get("date"), app.get("forecast_value"), app.get("earned_value"), app.get("shortfall"), app.get("suggested_claim"), app.get("submitted_value"), app.get("certified_value"), app.get("invoice_value"), app.get("paid_value"), app.get("status")])
    output.seek(0)
    filename = f"finance_{str(job.get('name', 'project')).replace(' ', '_')}_{datetime.utcnow().strftime('%Y%m%d')}.csv"
    return StreamingResponse(io.BytesIO(output.getvalue().encode("utf-8-sig")), media_type="text/csv", headers={"Content-Disposition": f"attachment; filename={filename}"})



# ==================== COMPANY FINANCE DASHBOARD ENDPOINTS ====================
# These routes power the new company-level finance dashboard.
# They intentionally use a separate MongoDB collection from the existing
# project application finance records so the older project finance workflow is preserved.

class DashboardFinanceRecordBase(BaseModel):
    project_id: Optional[str] = None
    project_name: Optional[str] = None

    # deposit, interim, final, retention, variation, other
    type: str = "interim"

    description: Optional[str] = None

    expected_date: Optional[str] = None
    expected_amount: float = 0.0

    anticipated_date: Optional[str] = None
    anticipated_amount: Optional[float] = None

    # expected, anticipated, at_risk, received, overdue
    status: str = "expected"

    received_date: Optional[str] = None
    received_amount: Optional[float] = None

    notes: Optional[str] = None

    # Links dashboard tracking records back to Gantt commercial markers so
    # the Finance page can supersede the original planned marker once tracked.
    linked_marker_id: Optional[str] = None
    source_marker_id: Optional[str] = None

    # Lightweight audit fields. The frontend sends the logged-in user details
    # where available, and the backend preserves them.
    created_by: Optional[str] = None
    created_by_email: Optional[str] = None
    created_by_role: Optional[str] = None
    updated_by: Optional[str] = None
    updated_by_email: Optional[str] = None
    updated_by_role: Optional[str] = None


class DashboardFinanceRecordCreate(DashboardFinanceRecordBase):
    pass


class DashboardFinanceRecordUpdate(BaseModel):
    project_id: Optional[str] = None
    project_name: Optional[str] = None
    type: Optional[str] = None
    description: Optional[str] = None

    expected_date: Optional[str] = None
    expected_amount: Optional[float] = None

    anticipated_date: Optional[str] = None
    anticipated_amount: Optional[float] = None

    status: Optional[str] = None

    received_date: Optional[str] = None
    received_amount: Optional[float] = None

    notes: Optional[str] = None
    linked_marker_id: Optional[str] = None
    source_marker_id: Optional[str] = None
    created_by: Optional[str] = None
    created_by_email: Optional[str] = None
    created_by_role: Optional[str] = None
    updated_by: Optional[str] = None
    updated_by_email: Optional[str] = None
    updated_by_role: Optional[str] = None
    archived: Optional[bool] = None


def dashboard_finance_to_float(value: Any, fallback: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return fallback
        return float(value)
    except (TypeError, ValueError):
        return fallback


def serialize_dashboard_finance_record(record: Dict[str, Any]) -> Dict[str, Any]:
    if not record:
        return {}

    return {
        "id": record.get("id"),
        "project_id": record.get("project_id"),
        "project_name": record.get("project_name"),
        "type": record.get("type", "interim"),
        "description": record.get("description"),
        "expected_date": record.get("expected_date"),
        "expected_amount": dashboard_finance_to_float(record.get("expected_amount")),
        "anticipated_date": record.get("anticipated_date"),
        "anticipated_amount": (
            dashboard_finance_to_float(record.get("anticipated_amount"))
            if record.get("anticipated_amount") is not None
            else None
        ),
        "status": record.get("status", "expected"),
        "received_date": record.get("received_date"),
        "received_amount": (
            dashboard_finance_to_float(record.get("received_amount"))
            if record.get("received_amount") is not None
            else None
        ),
        "notes": record.get("notes"),
        "linked_marker_id": record.get("linked_marker_id"),
        "source_marker_id": record.get("source_marker_id"),
        "created_by": record.get("created_by"),
        "created_by_email": record.get("created_by_email"),
        "created_by_role": record.get("created_by_role"),
        "updated_by": record.get("updated_by"),
        "updated_by_email": record.get("updated_by_email"),
        "updated_by_role": record.get("updated_by_role"),
        "created_at": record.get("created_at"),
        "updated_at": record.get("updated_at"),
        "archived": record.get("archived", False),
    }


def calculate_dashboard_finance_status(record: Dict[str, Any]) -> str:
    status = record.get("status") or "expected"

    if status == "received":
        return "received"

    expected_date = parse_iso_date_safe(record.get("expected_date"))
    anticipated_date = parse_iso_date_safe(record.get("anticipated_date"))
    due_date = anticipated_date or expected_date
    today = datetime.utcnow().date()

    if due_date and due_date < today:
        return "overdue"

    return status


async def backfill_dashboard_project_name(data: Dict[str, Any]) -> Dict[str, Any]:
    if data.get("project_id") and not data.get("project_name"):
        job = await db.jobs.find_one({"id": data["project_id"]}, {"_id": 0})
        if job:
            data["project_name"] = (
                job.get("name")
                or job.get("job_name")
                or job.get("address")
                or job.get("location")
                or "Unnamed project"
            )
    return data


FINANCE_AUDIT_FIELD_LABELS = {
    "project_name": "project name",
    "type": "finance type",
    "description": "description",
    "expected_date": "expected date",
    "expected_amount": "expected amount",
    "anticipated_date": "anticipated / due date",
    "anticipated_amount": "anticipated amount",
    "status": "status",
    "received_date": "received date",
    "received_amount": "received amount",
    "notes": "notes",
    "linked_marker_id": "linked marker",
    "source_marker_id": "source marker",
}

FINANCE_AUDIT_TRACKED_FIELDS = list(FINANCE_AUDIT_FIELD_LABELS.keys())
FINANCE_MONEY_FIELDS = {"expected_amount", "anticipated_amount", "received_amount"}

def audit_actor_from_data(data: Dict[str, Any], existing: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
    existing = existing or {}
    name = (data.get("updated_by") or data.get("created_by") or existing.get("updated_by") or existing.get("created_by") or "Unknown user")
    email = (data.get("updated_by_email") or data.get("created_by_email") or existing.get("updated_by_email") or existing.get("created_by_email") or "")
    role = (data.get("updated_by_role") or data.get("created_by_role") or existing.get("updated_by_role") or existing.get("created_by_role") or "unknown")
    return {
        "name": str(name).strip() or "Unknown user",
        "email": str(email).strip(),
        "role": str(role).strip() or "unknown",
    }

def normalise_audit_compare_value(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, float):
        return round(value, 2)
    if isinstance(value, int):
        return value
    return str(value).strip()

def format_audit_value(field: str, value: Any) -> str:
    if value is None or value == "":
        return "blank"
    if field in FINANCE_MONEY_FIELDS:
        try:
            return f"£{float(value):,.2f}"
        except Exception:
            return str(value)
    if field == "status":
        return str(value).replace("_", " ").title()
    return str(value)

def audit_record_label(record: Dict[str, Any]) -> str:
    project = record.get("project_name") or "Unallocated project"
    description = record.get("description") or record.get("type") or "finance record"
    return f"{project} – {description}"

def build_audit_description(action: str, field: Optional[str], old_value: Any, new_value: Any, record: Dict[str, Any], actor: Dict[str, str]) -> str:
    actor_name = actor.get("name") or "Unknown user"
    label = audit_record_label(record)
    if action == "create":
        return f"{actor_name} created {label}."
    if action == "delete":
        return f"{actor_name} archived {label}."
    field_label = FINANCE_AUDIT_FIELD_LABELS.get(field or "", field or "field")
    return f"{actor_name} changed {label} {field_label} from {format_audit_value(field or '', old_value)} to {format_audit_value(field or '', new_value)}."

async def insert_finance_audit_log(action: str, record_id: str, record: Dict[str, Any], actor: Dict[str, str], field: Optional[str] = None, old_value: Any = None, new_value: Any = None) -> Dict[str, Any]:
    now = datetime.utcnow().isoformat()
    log = {
        "id": str(uuid.uuid4()),
        "record_type": "finance_dashboard_record",
        "record_id": record_id,
        "project_id": record.get("project_id"),
        "project_name": record.get("project_name"),
        "action": action,
        "field": field,
        "field_label": FINANCE_AUDIT_FIELD_LABELS.get(field or "", field),
        "old_value": old_value,
        "new_value": new_value,
        "changed_by": actor.get("name"),
        "changed_by_email": actor.get("email"),
        "changed_by_role": actor.get("role"),
        "changed_at": now,
        "description": build_audit_description(action, field, old_value, new_value, record, actor),
    }
    await db.audit_logs.insert_one(log)
    log.pop("_id", None)
    return log

async def insert_finance_update_audit_logs(existing: Dict[str, Any], updated: Dict[str, Any], update_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    actor = audit_actor_from_data(update_data, existing)
    logs = []
    for field in FINANCE_AUDIT_TRACKED_FIELDS:
        if field not in update_data:
            continue
        old_value = existing.get(field)
        new_value = updated.get(field)
        if normalise_audit_compare_value(old_value) == normalise_audit_compare_value(new_value):
            continue
        logs.append(await insert_finance_audit_log("update", updated.get("id"), updated, actor, field, old_value, new_value))
    return logs

def serialize_audit_log(log: Dict[str, Any]) -> Dict[str, Any]:
    log = dict(log or {})
    log.pop("_id", None)
    return log


@api_router.get("/finance-records")
async def get_dashboard_finance_records(
    project_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    type: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
):
    query: Dict[str, Any] = {"archived": {"$ne": True}}

    if project_id:
        query["project_id"] = project_id

    if status:
        query["status"] = status

    if type:
        query["type"] = type

    if start_date or end_date:
        date_query: Dict[str, Any] = {}

        if start_date:
            date_query["$gte"] = start_date

        if end_date:
            date_query["$lte"] = end_date

        query["expected_date"] = date_query

    records = await db.finance_dashboard_records.find(query, {"_id": 0}).sort("expected_date", 1).to_list(2000)

    output = []
    for record in records:
        serialized = serialize_dashboard_finance_record(record)
        serialized["status"] = calculate_dashboard_finance_status(serialized)
        output.append(serialized)

    return output


@api_router.post("/finance-records")
async def create_dashboard_finance_record(record: DashboardFinanceRecordCreate):
    now = datetime.utcnow().isoformat()

    data = record.dict()
    data["id"] = str(uuid.uuid4())
    data["created_at"] = now
    data["updated_at"] = now
    data["archived"] = False

    created_by = (data.get("created_by") or "").strip()
    created_by_email = (data.get("created_by_email") or "").strip()
    created_by_role = (data.get("created_by_role") or "").strip()
    data["created_by"] = created_by or "Unknown user"
    data["created_by_email"] = created_by_email
    data["created_by_role"] = created_by_role or "unknown"
    data["updated_by"] = (data.get("updated_by") or data["created_by"]).strip() or data["created_by"]
    data["updated_by_email"] = (data.get("updated_by_email") or data["created_by_email"]).strip()
    data["updated_by_role"] = (data.get("updated_by_role") or data["created_by_role"]).strip() or data["created_by_role"]

    if data.get("anticipated_amount") is None:
        data["anticipated_amount"] = data.get("expected_amount", 0)

    if not data.get("anticipated_date"):
        data["anticipated_date"] = data.get("expected_date")

    data = await backfill_dashboard_project_name(data)

    await db.finance_dashboard_records.insert_one(data)
    await insert_finance_audit_log("create", data["id"], data, audit_actor_from_data(data))

    serialized = serialize_dashboard_finance_record(data)
    serialized["status"] = calculate_dashboard_finance_status(serialized)
    return serialized


@api_router.put("/finance-records/{record_id}")
async def update_dashboard_finance_record(record_id: str, update: DashboardFinanceRecordUpdate):
    existing = await db.finance_dashboard_records.find_one({"id": record_id, "archived": {"$ne": True}})

    if not existing:
        raise HTTPException(status_code=404, detail="Finance record not found")

    update_data = {
        key: value
        for key, value in update.dict().items()
        if value is not None
    }

    update_data["updated_at"] = datetime.utcnow().isoformat()

    # Do not accidentally blank out audit fields when the frontend has no user context.
    for audit_key in ["created_by", "created_by_email", "created_by_role", "updated_by", "updated_by_email", "updated_by_role"]:
        if audit_key in update_data and (update_data[audit_key] is None or str(update_data[audit_key]).strip() == ""):
            update_data.pop(audit_key, None)

    if not update_data.get("updated_by"):
        update_data["updated_by"] = existing.get("updated_by") or existing.get("created_by") or "Unknown user"
    if not update_data.get("updated_by_email"):
        update_data["updated_by_email"] = existing.get("updated_by_email") or existing.get("created_by_email") or ""
    if not update_data.get("updated_by_role"):
        update_data["updated_by_role"] = existing.get("updated_by_role") or existing.get("created_by_role") or "unknown"

    if update_data.get("project_id") and not update_data.get("project_name"):
        update_data = await backfill_dashboard_project_name(update_data)

    await db.finance_dashboard_records.update_one(
        {"id": record_id},
        {"$set": update_data}
    )

    updated = await db.finance_dashboard_records.find_one({"id": record_id}, {"_id": 0})
    await insert_finance_update_audit_logs(existing, updated, update_data)
    serialized = serialize_dashboard_finance_record(updated)
    serialized["status"] = calculate_dashboard_finance_status(serialized)

    return serialized


@api_router.delete("/finance-records/{record_id}")
async def delete_dashboard_finance_record(record_id: str):
    existing = await db.finance_dashboard_records.find_one({"id": record_id, "archived": {"$ne": True}}, {"_id": 0})
    result = await db.finance_dashboard_records.update_one(
        {"id": record_id, "archived": {"$ne": True}},
        {"$set": {"archived": True, "updated_at": datetime.utcnow().isoformat()}},
    )

    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Finance record not found")

    if existing:
        await insert_finance_audit_log("delete", record_id, existing, audit_actor_from_data({}, existing))

    return {"success": True, "message": "Finance record deleted"}


@api_router.get("/finance-records/{record_id}/audit-logs")
async def get_dashboard_finance_record_audit_logs(record_id: str):
    logs = await db.audit_logs.find({
        "record_type": "finance_dashboard_record",
        "record_id": record_id,
    }, {"_id": 0}).sort("changed_at", -1).to_list(500)
    return [serialize_audit_log(log) for log in logs]


@api_router.get("/audit-logs")
async def get_audit_logs(
    record_type: Optional[str] = Query(None),
    record_id: Optional[str] = Query(None),
    project_id: Optional[str] = Query(None),
    changed_by_email: Optional[str] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    limit: int = Query(200),
):
    query: Dict[str, Any] = {}
    if record_type:
        query["record_type"] = record_type
    if record_id:
        query["record_id"] = record_id
    if project_id:
        query["project_id"] = project_id
    if changed_by_email:
        query["changed_by_email"] = changed_by_email
    if start_date or end_date:
        date_query: Dict[str, Any] = {}
        if start_date:
            date_query["$gte"] = start_date
        if end_date:
            date_query["$lte"] = end_date
        query["changed_at"] = date_query

    safe_limit = max(1, min(1000, limit))
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("changed_at", -1).to_list(safe_limit)
    return [serialize_audit_log(log) for log in logs]


@api_router.get("/finance-dashboard")
async def get_company_finance_dashboard():
    today = datetime.utcnow().date()
    lookahead_end = today + timedelta(days=28)

    # Include past unpaid records so overdue money is visible in the summary,
    # while the forecast grid itself remains a 4-week lookahead.
    records = await db.finance_dashboard_records.find({
        "archived": {"$ne": True},
        "expected_date": {"$lte": lookahead_end.isoformat()},
    }, {"_id": 0}).sort("expected_date", 1).to_list(2000)

    all_records = []

    for record in records:
        serialized = serialize_dashboard_finance_record(record)
        serialized["status"] = calculate_dashboard_finance_status(serialized)
        all_records.append(serialized)

    # Contractor Work Orders are outgoing committed costs. They are kept separate
    # from income/application records so FinancePage can show contractor spend
    # without mixing it into client cash-in totals.
    work_orders = await db.work_orders.find({
        "archived": {"$ne": True},
        "status": {"$nin": list(WORK_ORDER_CANCELLED_STATUSES)},
    }, {"_id": 0}).sort("payment_due_date", 1).to_list(5000)

    contractor_cost_records = []
    for work_order in work_orders:
        status_key = str(work_order.get("status") or "").strip().lower()
        is_committed = status_key in WORK_ORDER_COMMITTED_STATUSES
        due_date = work_order.get("payment_due_date") or work_order.get("expected_completion_date") or work_order.get("expected_start_date")
        due_date_obj = parse_iso_date_safe(due_date)
        in_next_4_weeks = bool(due_date_obj and today <= due_date_obj <= lookahead_end)
        is_overdue = bool(due_date_obj and due_date_obj < today and status_key not in {"paid", "complete"})

        contractor_cost_records.append({
            "id": work_order.get("id", ""),
            "wo_number": work_order.get("wo_number", ""),
            "job_id": work_order.get("job_id", ""),
            "job_name": work_order.get("job_name", ""),
            "section_id": work_order.get("section_id", ""),
            "section_name": work_order.get("section_name", ""),
            "contractor_id": work_order.get("contractor_id", ""),
            "contractor_name": work_order.get("contractor_name", ""),
            "trade": work_order.get("trade", ""),
            "description": work_order.get("description", ""),
            "status": work_order.get("status", ""),
            "is_committed": is_committed,
            "net_amount": round(finance_to_number(work_order.get("net_amount")), 2),
            "vat_amount": round(finance_to_number(work_order.get("vat_amount")), 2),
            "gross_amount": round(finance_to_number(work_order.get("gross_amount")), 2),
            "cis_deduction": round(finance_to_number(work_order.get("cis_deduction")), 2),
            "payment_due_date": due_date,
            "expected_start_date": work_order.get("expected_start_date"),
            "expected_completion_date": work_order.get("expected_completion_date"),
            "in_next_4_weeks": in_next_4_weeks,
            "is_overdue": is_overdue,
        })

    weeks = []

    for i in range(4):
        week_start = today + timedelta(days=i * 7)
        week_end = week_start + timedelta(days=6)

        week_records = []

        for record in all_records:
            record_date = parse_iso_date_safe(
                record.get("anticipated_date") or record.get("expected_date")
            )

            if record_date and week_start <= record_date <= week_end:
                week_records.append(record)

        expected_total = sum(
            dashboard_finance_to_float(r.get("expected_amount"))
            for r in week_records
        )

        anticipated_total = sum(
            dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount")))
            for r in week_records
        )

        received_total = sum(
            dashboard_finance_to_float(r.get("received_amount"))
            for r in week_records
            if r.get("status") == "received"
        )

        at_risk_total = sum(
            dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount")))
            for r in week_records
            if r.get("status") == "at_risk"
        )

        overdue_total = sum(
            dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount")))
            for r in week_records
            if r.get("status") == "overdue"
        )

        weeks.append({
            "week_index": i + 1,
            "start_date": week_start.isoformat(),
            "end_date": week_end.isoformat(),
            "expected_total": round(expected_total, 2),
            "anticipated_total": round(anticipated_total, 2),
            "received_total": round(received_total, 2),
            "at_risk_total": round(at_risk_total, 2),
            "overdue_total": round(overdue_total, 2),
            "records": week_records,
        })

    overdue_records = [record for record in all_records if record.get("status") == "overdue"]
    received_records = [record for record in all_records if record.get("status") == "received"]
    at_risk_records = [record for record in all_records if record.get("status") == "at_risk"]

    future_records = []
    for record in all_records:
        record_date = parse_iso_date_safe(record.get("anticipated_date") or record.get("expected_date"))
        if record_date and today <= record_date <= lookahead_end:
            future_records.append(record)

    contractor_committed = [record for record in contractor_cost_records if record.get("is_committed")]
    contractor_next_4_weeks = [record for record in contractor_committed if record.get("in_next_4_weeks")]
    contractor_overdue = [record for record in contractor_committed if record.get("is_overdue")]

    summary = {
        "expected_next_4_weeks": round(sum(dashboard_finance_to_float(r.get("expected_amount")) for r in future_records), 2),
        "anticipated_next_4_weeks": round(sum(dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount"))) for r in future_records), 2),
        "received_next_4_weeks": round(sum(dashboard_finance_to_float(r.get("received_amount")) for r in received_records), 2),
        "at_risk_next_4_weeks": round(sum(dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount"))) for r in at_risk_records), 2),
        "overdue_next_4_weeks": round(sum(dashboard_finance_to_float(r.get("anticipated_amount"), dashboard_finance_to_float(r.get("expected_amount"))) for r in overdue_records), 2),
        "contractor_committed_net_total": round(sum(record.get("net_amount", 0) for record in contractor_committed), 2),
        "contractor_committed_gross_total": round(sum(record.get("gross_amount", 0) for record in contractor_committed), 2),
        "contractor_due_next_4_weeks_net": round(sum(record.get("net_amount", 0) for record in contractor_next_4_weeks), 2),
        "contractor_due_next_4_weeks_gross": round(sum(record.get("gross_amount", 0) for record in contractor_next_4_weeks), 2),
        "contractor_overdue_net": round(sum(record.get("net_amount", 0) for record in contractor_overdue), 2),
        "contractor_overdue_gross": round(sum(record.get("gross_amount", 0) for record in contractor_overdue), 2),
        "work_order_count": len(contractor_cost_records),
        "committed_work_order_count": len(contractor_committed),
        "record_count": len(all_records),
    }

    return {
        "summary": summary,
        "weeks": weeks,
        "records": all_records,
        "work_orders": contractor_cost_records,
        "contractor_costs": {
            "records": contractor_cost_records,
            "committed_records": contractor_committed,
            "next_4_weeks_records": contractor_next_4_weeks,
            "overdue_records": contractor_overdue,
            "committed_net_total": summary["contractor_committed_net_total"],
            "committed_gross_total": summary["contractor_committed_gross_total"],
            "due_next_4_weeks_net": summary["contractor_due_next_4_weeks_net"],
            "due_next_4_weeks_gross": summary["contractor_due_next_4_weeks_gross"],
            "overdue_net": summary["contractor_overdue_net"],
            "overdue_gross": summary["contractor_overdue_gross"],
        },
    }

# ==================== PURCHASE ORDER SYSTEM ====================

def clean_mongo_doc(doc: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not doc:
        return None
    doc.pop("_id", None)
    return doc


def calculate_po_line_totals(line: Dict[str, Any]) -> Dict[str, Any]:
    quantity = float(line.get("quantity") or 0)
    unit_cost = float(line.get("unit_cost") or 0)
    vat_rate = float(line.get("vat_rate") or 0)

    # Quote imports can provide supplier gross amounts that already include VAT.
    # Preserve those exact totals so VAT is not added twice and pennies do not drift.
    if line.get("prices_include_vat") and line.get("source_line_gross_total") is not None:
        gross_total = round(float(line.get("source_line_gross_total") or 0), 2)
        net_total = round(float(line.get("source_line_net_total") or (gross_total / (1 + vat_rate / 100) if vat_rate else gross_total)), 2)
        vat_total = round(float(line.get("source_line_vat_total") or (gross_total - net_total)), 2)
        unit_cost = round(net_total / quantity, 4) if quantity else 0.0
    else:
        net_total = round(quantity * unit_cost, 2)
        vat_total = round(net_total * (vat_rate / 100), 2)
        gross_total = round(net_total + vat_total, 2)

    line["quantity"] = quantity
    line["unit_cost"] = unit_cost
    line["vat_rate"] = vat_rate
    line["net_total"] = net_total
    line["vat_total"] = vat_total
    line["gross_total"] = gross_total
    line.setdefault("id", str(uuid.uuid4()))
    line.setdefault("received_quantity", 0.0)
    line.setdefault("material_status", "committed")
    return line


def calculate_po_totals(po_dict: Dict[str, Any]) -> Dict[str, Any]:
    lines = [calculate_po_line_totals(dict(line)) for line in po_dict.get("lines", [])]
    po_dict["lines"] = lines
    po_dict["net_total"] = round(sum(line.get("net_total", 0) for line in lines), 2)
    po_dict["vat_total"] = round(sum(line.get("vat_total", 0) for line in lines), 2)
    po_dict["gross_total"] = round(sum(line.get("gross_total", 0) for line in lines), 2)
    return po_dict


async def next_po_number() -> str:
    year = datetime.utcnow().year
    prefix = f"PO-{year}-"
    latest = await db.purchase_orders.find_one({"po_number": {"$regex": f"^{prefix}"}}, sort=[("po_number", -1)])
    if latest and latest.get("po_number"):
        try:
            next_number = int(latest["po_number"].split("-")[-1]) + 1
        except Exception:
            next_number = 1
    else:
        next_number = 1
    return f"{prefix}{str(next_number).zfill(4)}"


def normalise_quote_text(text: str) -> str:
    """Clean extracted quote text while preserving useful line breaks."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\u00a0]+", " ", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(line for line in lines if line)


def parse_money_value(value: Any) -> Optional[float]:
    """Parse UK money strings like £1,234.56 into floats."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    cleaned = str(value).strip()
    cleaned = cleaned.replace("£", "").replace(",", "").replace(" ", "")
    cleaned = cleaned.replace("GBP", "").replace("gbp", "")
    cleaned = re.sub(r"[^0-9.\-]", "", cleaned)
    if cleaned in {"", ".", "-"}:
        return None
    try:
        return round(float(cleaned), 2)
    except Exception:
        return None


def extract_quote_number(text: str) -> str:
    cleaned_text = normalise_quote_text(text)
    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]

    # Prefer exact label-next-line layouts such as:
    # Quote Number\nQU-1757
    labels = {"quote number", "quotenumber", "quotation number", "quotationnumber", "quote no", "quoteno", "estimate number", "estimatenumber"}
    for idx, line in enumerate(lines):
        compact = re.sub(r"[^a-z0-9]", "", line.lower())
        if compact in {re.sub(r"[^a-z0-9]", "", label) for label in labels}:
            for candidate in lines[idx + 1: idx + 4]:
                value = candidate.strip().strip(".,;:")
                if re.search(r"[A-Z]", value, re.IGNORECASE) and re.search(r"\d", value) and len(value) <= 35:
                    return value

    patterns = [
        r"(?:quote\s*number|quotation\s*number|quote\s*no\.?|quotation\s*no\.?)\s*[:#-]?\s*([A-Z0-9][A-Z0-9\-/]{2,})",
        r"(?:quote|quotation)\s*(?:ref|reference)\s*[:#-]?\s*([A-Z0-9][A-Z0-9\-/]{2,})",
        r"(?:ref|reference)\s*[:#-]?\s*([A-Z]{1,5}-?\d{2,})",
    ]
    for pattern in patterns:
        match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if match:
            value = match.group(1).strip().strip(".,;:")
            if re.search(r"\d", value) and not re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}$", value):
                return value
    return ""


def extract_first_date_for_patterns(text: str, labels: List[str]) -> str:
    cleaned_text = normalise_quote_text(text)
    date_pattern = r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2}|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4})"
    for label in labels:
        pattern = rf"{label}\s*[:#-]?\s*{date_pattern}"
        match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return ""


def extract_email_from_text(text: str) -> str:
    match = re.search(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", text or "", re.IGNORECASE)
    return match.group(0).strip() if match else ""


def extract_supplier_name_from_text(text: str) -> str:
    """Best-effort supplier name extraction from supplier quote text.

    Patch 2.1 deliberately avoids returning the customer name, company numbers,
    VAT numbers, phone numbers or address fragments as the supplier.
    """
    cleaned_text = normalise_quote_text(text)
    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]

    # Xero-style estimates often show the supplier block immediately after the VAT number.
    for idx, line in enumerate(lines):
        if re.fullmatch(r"(?:vat\s*)?number", line, flags=re.IGNORECASE) or re.sub(r"[^a-z]", "", line.lower()) == "vatnumber" or line.lower().startswith("vat number"):
            block = []
            for candidate in lines[idx + 1: idx + 10]:
                c = candidate.strip(" -|,. ")
                lower = c.lower()
                if not c:
                    continue
                if re.fullmatch(r"[0-9 ]{6,}", c):
                    continue
                if any(term in lower for term in ["unit", "road", "street", "park", "upon tyne", "postcode", "tel", "phone", "019", "email", "to supply", "description"]):
                    break
                block.append(c)
                joined = " ".join(block)
                if any(term in joined.lower() for term in ["ltd", "limited", "plc", "llp", "t/a", "trading as"]) and len(joined) >= 8:
                    # Allow one or two following trading-name fragments where useful.
                    continue
            joined = " ".join(block).strip()
            if joined and any(term in joined.lower() for term in ["ltd", "limited", "plc", "llp", "t/a", "trading as"]):
                return joined[:140]

    # Look for company-like runs in the first part of the document.
    stop_words = ["description", "quantity", "qty", "unit price", "amount", "total", "terms", "payment"]
    search_lines = []
    for line in lines[:35]:
        if any(stop in line.lower() for stop in stop_words):
            break
        search_lines.append(line)

    for window_size in [3, 2, 1]:
        for idx in range(0, max(0, len(search_lines) - window_size + 1)):
            joined = " ".join(search_lines[idx: idx + window_size]).strip(" -|,. ")
            lower = joined.lower()
            if any(skip in lower for skip in ["company registration", "registered office", "lda group", "ldagroup", "quote number", "estimate", "invoice", "date", "expiry", "reference"]):
                continue
            if "@" in joined or re.search(r"£|\d{3,}", joined):
                continue
            if re.search(r"\b(ltd|limited|plc|llp|t/a|trading as)\b", joined, re.IGNORECASE) and 5 <= len(joined) <= 140:
                return joined

    # Final fallback: first non-customer, non-address useful line.
    skip_terms = [
        "quote", "quotation", "estimate", "invoice", "date", "page", "tel", "phone",
        "email", "vat", "company reg", "registered office", "account", "delivery", "address", "customer",
        "subtotal", "total", "amount", "qty", "quantity", "unit", "price", "lda group", "ldagroup",
    ]
    for line in lines[:18]:
        lower = line.lower()
        if any(term in lower for term in skip_terms):
            continue
        if "@" in line or re.search(r"£|\d+\.\d{2}", line):
            continue
        if len(line) < 3 or len(line) > 90:
            continue
        return line.strip(" -|,.")
    return ""


def quote_appears_vat_inclusive(text: str) -> bool:
    """Return True when quote wording suggests line/totals already include VAT."""
    lower = normalise_quote_text(text).lower()
    inclusive_patterns = [
        "includes vat",
        "include vat",
        "including vat",
        "inc vat",
        "inc. vat",
        "vat included",
        "prices include vat",
        "amount gbp",
    ]
    return any(pattern in lower for pattern in inclusive_patterns)


def extract_quote_totals(text: str) -> Dict[str, Optional[float]]:
    cleaned_text = normalise_quote_text(text)
    result = {"net_total": None, "vat_total": None, "gross_total": None, "vat_inclusive": quote_appears_vat_inclusive(cleaned_text)}
    money = r"£?\s*([0-9]+(?:,[0-9]{3})*(?:\.\d{2})?)"

    # Specific Xero-style: INCLUDES VAT 20% 577.16 / TOTAL GBP 3,463.00
    include_vat_matches = re.findall(r"includes?\s+vat\s*(?:\d+(?:\.\d+)?%\s*)?" + money, cleaned_text, flags=re.IGNORECASE)
    if include_vat_matches:
        result["vat_total"] = parse_money_value(include_vat_matches[-1])
        result["vat_inclusive"] = True

    total_gbp_matches = re.findall(r"(?:total\s*gbp|total\s+£|grand\s+total|total\s+due|amount\s+due|gross\s+total)\s*[:\-]?\s*" + money, cleaned_text, flags=re.IGNORECASE)
    if total_gbp_matches:
        result["gross_total"] = parse_money_value(total_gbp_matches[-1])

    patterns = {
        "net_total": [
            rf"(?:sub\s*total|subtotal|net\s*total|goods\s*total|total\s*net|net\s*amount)\s*[:\-]?\s*{money}",
        ],
        "vat_total": [
            rf"(?:vat|v\.a\.t\.|tax)\s*(?:total|amount)?\s*[:\-]?\s*{money}",
        ],
        "gross_total": [
            rf"(?:grand\s*total|total\s*due|amount\s*due|gross\s*total|total\s*inc\.?\s*vat|total\s*including\s*vat)\s*[:\-]?\s*{money}",
            rf"(?:^|\n)\s*total\s*[:\-]?\s*{money}",
        ],
    }
    for key, key_patterns in patterns.items():
        if result.get(key) is not None:
            continue
        for pattern in key_patterns:
            matches = re.findall(pattern, cleaned_text, flags=re.IGNORECASE | re.MULTILINE)
            if matches:
                value = parse_money_value(matches[-1] if isinstance(matches[-1], str) else matches[-1][0])
                if value is not None:
                    result[key] = value
                    break

    if result["gross_total"] is not None and result["vat_total"] is not None and result["net_total"] is None:
        result["net_total"] = round(result["gross_total"] - result["vat_total"], 2)
    if result["gross_total"] is None and result["net_total"] is not None and result["vat_total"] is not None:
        result["gross_total"] = round(result["net_total"] + result["vat_total"], 2)
    if result["vat_total"] is None and result["gross_total"] is not None and result["net_total"] is not None:
        result["vat_total"] = round(result["gross_total"] - result["net_total"], 2)
    if result["net_total"] is None and result["gross_total"] is not None and result.get("vat_inclusive"):
        # Fallback to 20% extraction if VAT amount is not explicitly shown.
        result["net_total"] = round(result["gross_total"] / 1.2, 2)
        result["vat_total"] = round(result["gross_total"] - result["net_total"], 2)
    return result


def extract_quote_table_lines(text: str) -> List[str]:
    """Return only the likely line-item section, not headers/footers/terms."""
    lines = [line.strip() for line in normalise_quote_text(text).splitlines() if line.strip()]
    start_idx = None
    header_terms = ["description", "qty", "quantity", "unit", "price", "vat", "amount", "total", "gbp"]
    for idx, line in enumerate(lines):
        lower = line.lower()
        score = sum(1 for term in header_terms if term in lower)
        if score >= 3 and ("description" in lower or "item" in lower or "product" in lower or "service" in lower):
            start_idx = idx + 1
            break
    if start_idx is None:
        # Fallback: start after wording such as 'To Supply' or 'Please see estimate'.
        for idx, line in enumerate(lines):
            lower = line.lower()
            if any(marker in lower for marker in ["to supply", "initial estimate", "please see the following"]):
                start_idx = idx + 1
                break
    if start_idx is None:
        start_idx = 0

    table_lines = []
    stop_markers = [
        "includes vat", "include vat", "subtotal", "sub total", "total gbp", "grand total",
        "total due", "terms", "payment", "bank details", "all goods remain", "warranty", "returns",
        "company registration no", "registered office",
    ]
    for line in lines[start_idx:]:
        lower = line.lower()
        if any(marker in lower for marker in stop_markers):
            break
        table_lines.append(line)
    return table_lines


def line_looks_like_numeric_item_row(line: str) -> bool:
    """Detect lines that contain qty, VAT/discount and final amount columns."""
    stripped = line.strip()
    if not stripped:
        return False
    moneyish = re.findall(r"(?:£\s*)?\d+(?:,\d{3})*(?:\.\d{2})?%?", stripped)
    if len(moneyish) < 3:
        return False
    # Numeric rows often start with quantity, but can also end a description line.
    return bool(re.search(r"\b\d+(?:\.\d+)?\s+\d", stripped) or re.search(r"\d+(?:\.\d+)?%", stripped))


def parse_numeric_item_row(line: str) -> Optional[Dict[str, float]]:
    """Extract quantity, VAT rate and final line amount from a flexible numeric row.

    Works from the right-hand side of the row so product names/dimensions such as
    "Bench 45 Low" or "840mm x 840mm" are not mistaken for quantities.
    """
    tokens = re.findall(r"(?:£\s*)?\d+(?:,\d{3})*(?:\.\d{2})?%?", line)
    if len(tokens) < 2:
        return None
    numbers = [parse_money_value(token) for token in tokens]
    if any(num is None for num in numbers):
        numbers = [num for num in numbers if num is not None]
    if len(numbers) < 2:
        return None

    amount = numbers[-1]
    percent_positions = [idx for idx, token in enumerate(tokens) if "%" in token]
    vat_rate = 20.0
    if percent_positions:
        rate = parse_money_value(tokens[percent_positions[-1]])
        if rate is not None and 0 <= rate <= 100:
            vat_rate = rate

    # Expected right-hand patterns:
    # qty unit discount% vat% amount  -> quantity is five tokens from the end
    # qty unit vat% amount            -> quantity is four tokens from the end
    if len(percent_positions) >= 2 and len(numbers) >= 5:
        quantity = numbers[-5]
    elif len(percent_positions) >= 1 and len(numbers) >= 4:
        quantity = numbers[-4]
    elif len(numbers) >= 3:
        quantity = numbers[-3]
    else:
        quantity = 1

    if quantity is None or quantity <= 0 or quantity > 10000:
        quantity = 1
    if amount is None or amount <= 0:
        return None
    return {"quantity": float(quantity), "amount": float(amount), "vat_rate": float(vat_rate)}


def parse_quote_lines_from_text(text: str) -> List[Dict[str, Any]]:
    """Best-effort quote line extraction. Always requires review before PO creation.

    Patch 2.1 reads table-like areas only and uses the supplier's final amount column,
    rather than every number in the document.
    """
    table_lines = extract_quote_table_lines(text)
    vat_inclusive = quote_appears_vat_inclusive(text)
    parsed_lines: List[Dict[str, Any]] = []
    description_buffer: List[str] = []
    skip_description_terms = ["description", "quantity", "unit price", "discount", "amount gbp", "vat"]

    def flush_line(description: str, numeric: Dict[str, float]):
        desc = " ".join(description.split()).strip(" -|,. ")
        if not desc or len(desc) < 3:
            return
        if any(term == desc.lower() for term in skip_description_terms):
            return
        quantity = float(numeric.get("quantity") or 1)
        amount = float(numeric.get("amount") or 0)
        vat_rate = float(numeric.get("vat_rate") or 20)
        if amount <= 0 or quantity <= 0:
            return
        # The PO model stores unit_cost as NET. If supplier amount is VAT-inclusive,
        # strip VAT here so the PO does not add VAT twice.
        if vat_inclusive and vat_rate > 0:
            gross_line_total = round(amount, 2)
            net_line_total = round(gross_line_total / (1 + vat_rate / 100), 2)
            vat_line_total = round(gross_line_total - net_line_total, 2)
        else:
            net_line_total = round(amount, 2)
            vat_line_total = round(net_line_total * (vat_rate / 100), 2)
            gross_line_total = round(net_line_total + vat_line_total, 2)
        unit_cost = round(net_line_total / quantity, 4)
        parsed_lines.append({
            "description": desc[:220],
            "quantity": quantity,
            "unit_cost": unit_cost,
            "vat_rate": vat_rate,
            "net_total": net_line_total,
            "vat_total": vat_line_total,
            "gross_total": gross_line_total,
            "prices_include_vat": bool(vat_inclusive),
            "source_line_net_total": net_line_total,
            "source_line_vat_total": vat_line_total,
            "source_line_gross_total": gross_line_total,
            "cost_category": "Materials",
        })

    for raw_line in table_lines:
        line = " ".join(raw_line.split())
        lower = line.lower()
        if not line:
            continue
        if any(term in lower for term in ["subtotal", "total", "terms", "payment", "company registration", "registered office"]):
            break
        if any(term == lower for term in skip_description_terms):
            continue

        numeric = parse_numeric_item_row(line) if line_looks_like_numeric_item_row(line) else None
        if numeric:
            # Remove the numeric tokens from the line; anything left at the front is description.
            description_part = re.sub(r"(?:£\s*)?\d+(?:,\d{3})*(?:\.\d{2})?%?", " ", line)
            description_part = re.sub(r"\s{2,}", " ", description_part).strip(" -|,. ")
            description = " ".join(description_buffer + ([description_part] if description_part else []))
            flush_line(description, numeric)
            description_buffer = []
        else:
            # Ignore footnote/options that do not have their own price row.
            if re.search(r"\*\*|\bmore\b|optional|option", line, re.IGNORECASE):
                continue
            description_buffer.append(line)
            # Avoid a runaway buffer on messy PDFs.
            if len(description_buffer) > 8:
                description_buffer = description_buffer[-8:]

    # De-duplicate identical rows caused by PDF extraction artefacts.
    unique_lines = []
    seen = set()
    for line in parsed_lines:
        key = (line.get("description", "").lower(), float(line.get("quantity") or 0), float(line.get("unit_cost") or 0), float(line.get("vat_rate") or 0))
        if key in seen:
            continue
        seen.add(key)
        unique_lines.append(line)
    return unique_lines[:80]


def extract_phone_from_text(text: str) -> str:
    """Best-effort UK phone extraction for quick supplier creation."""
    cleaned_text = normalise_quote_text(text)
    patterns = [
        r"(?:tel|telephone|phone|t)\s*[:#-]?\s*((?:\+44\s?|0)\d[\d\s().-]{8,})",
        r"\b((?:\+44\s?|0)\d{2,5}[\s.-]?\d{3,4}[\s.-]?\d{3,4})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, cleaned_text, re.IGNORECASE)
        if match:
            value = re.sub(r"\s+", " ", match.group(1)).strip(" .,-")
            if len(re.sub(r"\D", "", value)) >= 10:
                return value[:40]
    return ""


def extract_supplier_vat_number_from_text(text: str) -> str:
    cleaned_text = normalise_quote_text(text)
    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        lower = line.lower()
        if "vat" in lower and "number" in lower:
            same_line = re.search(r"(?:vat\s*number|vat\s*no\.?)\s*[:#-]?\s*([A-Z0-9][A-Z0-9 \-]{4,25})", line, re.IGNORECASE)
            if same_line:
                return same_line.group(1).strip(" .,-")
            for candidate in lines[idx + 1: idx + 3]:
                digits = re.sub(r"\D", "", candidate)
                if 7 <= len(digits) <= 15:
                    return candidate.strip(" .,-")[:30]
    return ""


def extract_supplier_address_from_text(text: str, supplier_name: str = "") -> str:
    """Extract a short supplier address block from the header area for inline supplier creation."""
    cleaned_text = normalise_quote_text(text)
    lines = [line.strip() for line in cleaned_text.splitlines() if line.strip()]
    if not lines:
        return ""

    start_idx = None
    supplier_norm = normalise_supplier_match_name(supplier_name) if supplier_name else ""
    if supplier_norm:
        supplier_tokens = set(supplier_match_tokens(supplier_norm))
        for idx, line in enumerate(lines[:45]):
            line_tokens = set(supplier_match_tokens(line))
            if supplier_tokens and len(supplier_tokens & line_tokens) >= min(2, len(supplier_tokens)):
                start_idx = idx + 1
                break

    if start_idx is None:
        # Fallback to the block after VAT Number in Xero-style supplier headers.
        for idx, line in enumerate(lines[:35]):
            if "vat" in line.lower() and "number" in line.lower():
                start_idx = idx + 2
                break

    if start_idx is None:
        return ""

    address_lines = []
    stop_terms = [
        "to supply", "description", "quantity", "quote number", "reference", "date", "expiry",
        "total", "includes vat", "terms", "payment", "company registration", "registered office",
    ]
    for candidate in lines[start_idx:start_idx + 12]:
        lower = candidate.lower()
        if any(term in lower for term in stop_terms):
            break
        if "@" in candidate:
            continue
        if re.search(r"(?:\+44|0)\d", candidate):
            continue
        if re.fullmatch(r"[0-9 ]{6,}", candidate):
            continue
        if len(candidate) > 120:
            break
        address_lines.append(candidate.strip(" ,"))

    # Remove accidental supplier-name repeats from the top of the address block.
    while address_lines and supplier_norm and supplier_token_overlap(address_lines[0], supplier_name) >= 0.5:
        address_lines.pop(0)

    return "\n".join(address_lines[:7]).strip()


def ocr_space_configured() -> bool:
    return bool(os.environ.get("OCR_SPACE_API_KEY", "").strip())


def extract_text_with_ocr_space(filename: str, content_type: str, content: bytes) -> str:
    """OCR fallback for scanned PDFs/images using OCR.space over HTTPS.

    This avoids requiring Tesseract/Poppler binaries on Render. Set OCR_SPACE_API_KEY
    in Render to enable it. The free OCR.space endpoint is fine for testing, but a
    paid/keyed plan is more reliable for production volume.
    """
    api_key = os.environ.get("OCR_SPACE_API_KEY", "").strip()
    if not api_key:
        return ""

    endpoint = os.environ.get("OCR_SPACE_API_URL", "https://api.ocr.space/parse/image").strip()
    language = os.environ.get("OCR_SPACE_LANGUAGE", "eng").strip() or "eng"
    max_bytes = int(os.environ.get("OCR_SPACE_MAX_BYTES", str(10 * 1024 * 1024)))
    if len(content) > max_bytes:
        logger.warning("OCR skipped for %s because file is larger than OCR_SPACE_MAX_BYTES", filename)
        return ""

    data = {
        "apikey": api_key,
        "language": language,
        "isOverlayRequired": "false",
        "scale": "true",
        "OCREngine": os.environ.get("OCR_SPACE_ENGINE", "2"),
        "detectOrientation": "true",
    }
    files = {"file": (filename or "quote_upload", content, content_type or "application/octet-stream")}
    try:
        response = requests.post(endpoint, data=data, files=files, timeout=90)
        response.raise_for_status()
        payload = response.json()
    except Exception as exc:
        logger.warning("OCR.space extraction failed for %s: %s", filename, exc)
        return ""

    if payload.get("IsErroredOnProcessing"):
        logger.warning("OCR.space processing error for %s: %s", filename, payload.get("ErrorMessage"))
        return ""

    parsed_results = payload.get("ParsedResults") or []
    text_parts = []
    for item in parsed_results:
        parsed_text = item.get("ParsedText") or ""
        if parsed_text.strip():
            text_parts.append(parsed_text)
    return normalise_quote_text("\n".join(text_parts))


async def extract_text_from_upload(file: UploadFile, content: bytes) -> Dict[str, Any]:
    """Extract text from an uploaded quote.

    Returns metadata so the frontend can tell the user whether OCR was used or
    whether OCR needs configuring.
    """
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()
    result = {
        "text": "",
        "method": "none",
        "ocr_attempted": False,
        "ocr_configured": ocr_space_configured(),
        "ocr_used": False,
    }

    if filename.endswith(".pdf") or "pdf" in content_type:
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = normalise_quote_text("\n".join(page.extract_text() or "" for page in reader.pages))
            # Digital PDFs should produce real text. If they do not, try OCR.
            if len(text.strip()) >= 40:
                result.update({"text": text, "method": "pdf_text"})
                return result
        except Exception as exc:
            logger.warning("PDF text extraction failed: %s", exc)

        result["ocr_attempted"] = True
        ocr_text = extract_text_with_ocr_space(file.filename or "quote.pdf", content_type or "application/pdf", content)
        if ocr_text:
            result.update({"text": ocr_text, "method": "ocr_space", "ocr_used": True})
        return result

    if filename.endswith(".docx") or "wordprocessingml" in content_type:
        try:
            from docx import Document
            document = Document(io.BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            result.update({"text": normalise_quote_text("\n".join(parts)), "method": "docx_text"})
            return result
        except Exception as exc:
            logger.warning("DOCX text extraction failed: %s", exc)
            return result

    if filename.endswith((".txt", ".csv")) or "text" in content_type or "csv" in content_type:
        try:
            result.update({"text": normalise_quote_text(content.decode("utf-8", errors="ignore")), "method": "plain_text"})
            return result
        except Exception:
            return result

    if filename.endswith((".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff")) or content_type.startswith("image/"):
        result["ocr_attempted"] = True
        ocr_text = extract_text_with_ocr_space(file.filename or "quote_image", content_type or "image/jpeg", content)
        if ocr_text:
            result.update({"text": ocr_text, "method": "ocr_space", "ocr_used": True})
        return result

    return result


def normalise_supplier_match_name(value: str) -> str:
    """Normalise supplier names for safer matching without over-matching short words like Test."""
    value = (value or "").lower()
    value = value.replace("&", " and ")
    value = re.sub(r"\bt\s*/\s*a\b", " trading as ", value)
    value = re.sub(r"[^a-z0-9\s]", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def supplier_match_tokens(value: str) -> List[str]:
    stop = {
        "ltd", "limited", "plc", "llp", "company", "co", "the", "and", "trading", "as", "ta",
        "fireplaces", "fireplace", "services", "service", "group", "uk", "gb", "estimate", "quote", "quotation",
    }
    tokens = []
    for token in normalise_supplier_match_name(value).split():
        if len(token) < 3:
            continue
        if token in stop:
            continue
        if token.isdigit():
            continue
        tokens.append(token)
    return tokens


def supplier_token_overlap(a: str, b: str) -> float:
    a_tokens = set(supplier_match_tokens(a))
    b_tokens = set(supplier_match_tokens(b))
    if not a_tokens or not b_tokens:
        return 0.0
    shared = len(a_tokens & b_tokens)
    return max(shared / len(a_tokens), shared / len(b_tokens))


async def match_supplier_from_quote(supplier_name: str, supplier_email: str) -> Dict[str, Any]:
    """Safely match an extracted supplier to an existing supplier.

    Patch 2.2 deliberately avoids weak substring matches. A short supplier such as
    "Test" must not match a real supplier quote unless the email is an exact match.
    """
    from difflib import SequenceMatcher

    suppliers = await db.suppliers.find({"archived": {"$ne": True}}, {"_id": 0}).to_list(1000)
    supplier_name_norm = normalise_supplier_match_name(supplier_name)
    supplier_email_norm = (supplier_email or "").strip().lower()
    supplier_domain = supplier_email_norm.split("@")[-1] if "@" in supplier_email_norm else ""

    generic_domains = {
        "gmail.com", "outlook.com", "hotmail.com", "live.com", "icloud.com", "yahoo.com", "aol.com",
    }

    best_match = None
    best_score = 0
    for supplier in suppliers:
        saved_name_raw = supplier.get("name") or ""
        saved_name_norm = normalise_supplier_match_name(saved_name_raw)
        saved_emails = [
            (supplier.get("orders_email") or "").strip().lower(),
            (supplier.get("accounts_email") or "").strip().lower(),
        ]
        saved_domains = [email.split("@")[-1] for email in saved_emails if "@" in email]

        score = 0

        # Exact mailbox match is safe.
        if supplier_email_norm and supplier_email_norm in saved_emails:
            score = max(score, 100)

        # Domain match is only safe for business domains, or where the name is also similar.
        overlap = supplier_token_overlap(supplier_name_norm, saved_name_norm)
        ratio = SequenceMatcher(None, supplier_name_norm, saved_name_norm).ratio() if supplier_name_norm and saved_name_norm else 0

        if supplier_domain and supplier_domain in saved_domains and supplier_domain not in generic_domains:
            if overlap >= 0.35 or ratio >= 0.70:
                score = max(score, 90)
            else:
                score = max(score, 65)

        # Name matching must be strong. No raw substring matching.
        if supplier_name_norm and saved_name_norm:
            if supplier_name_norm == saved_name_norm:
                score = max(score, 95)
            elif ratio >= 0.86:
                score = max(score, 88)
            elif overlap >= 0.67 and len(set(supplier_match_tokens(supplier_name_norm)) & set(supplier_match_tokens(saved_name_norm))) >= 2:
                score = max(score, 82)
            elif overlap >= 0.50 and ratio >= 0.70:
                score = max(score, 74)

        if score > best_score:
            best_score = score
            best_match = supplier

    # Require a strong match. Anything lower should be a user review/create-supplier step.
    if best_match and best_score >= 70:
        return {"matched_supplier_id": best_match.get("id"), "matched_supplier_name": best_match.get("name"), "match_score": best_score}
    return {"matched_supplier_id": None, "matched_supplier_name": "", "match_score": best_score if best_match else 0}


def po_company_details() -> Dict[str, Any]:
    """Company details shown on PO PDFs. Override these with Render env vars if needed."""
    address_lines = os.environ.get(
        "PO_COMPANY_ADDRESS_LINES",
        "LDA Group Building Services Ltd|Newcastle upon Tyne|United Kingdom",
    )
    return {
        "name": os.environ.get("PO_COMPANY_NAME", "LDA Group Building Services Ltd"),
        "address_lines": [line.strip() for line in address_lines.split("|") if line.strip()],
        "phone": os.environ.get("PO_COMPANY_PHONE", ""),
        "email": os.environ.get("PO_COMPANY_EMAIL", "info@ldagroup.co.uk"),
        "website": os.environ.get("PO_COMPANY_WEBSITE", "www.ldagroup.co.uk"),
        "vat_number": os.environ.get("PO_COMPANY_VAT_NUMBER", ""),
        "company_number": os.environ.get("PO_COMPANY_NUMBER", ""),
        "logo_url": os.environ.get("PO_COMPANY_LOGO_URL", "https://ldagroup.co.uk/wp-content/uploads/2022/01/lda-group-200x200.png"),
    }


def po_pdf_escape(value: Any) -> str:
    from xml.sax.saxutils import escape
    return escape(str(value or ""))


def get_po_logo_flowable(width_cm: float = 2.2):
    """Return a ReportLab image flowable for the LDA logo, or None if unavailable."""
    try:
        from reportlab.lib.units import cm
        from reportlab.platypus import Image
        details = po_company_details()
        logo_url = details.get("logo_url")
        if not logo_url:
            return None
        response = requests.get(logo_url, timeout=8)
        response.raise_for_status()
        image = Image(io.BytesIO(response.content))
        image.drawWidth = width_cm * cm
        image.drawHeight = width_cm * cm
        return image
    except Exception as exc:
        logger.warning("Could not load PO logo for PDF: %s", exc)
        return None


def generate_purchase_order_pdf_bytes(po: Dict[str, Any]) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    except Exception as exc:
        logger.error("PO PDF dependency error: %s", exc)
        raise HTTPException(status_code=500, detail="PDF generation is not available on this server. Check reportlab is installed.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=1.2*cm, leftMargin=1.2*cm, topMargin=1.0*cm, bottomMargin=1.0*cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SmallMuted", parent=styles["BodyText"], fontSize=8, leading=10, textColor=colors.HexColor("#475569")))
    styles.add(ParagraphStyle(name="RightTitle", parent=styles["Title"], alignment=2, fontSize=18, leading=22, textColor=colors.HexColor("#111827")))
    story = []

    company = po_company_details()
    company_lines = [f"<b>{po_pdf_escape(company.get('name'))}</b>"]
    company_lines.extend(po_pdf_escape(line) for line in company.get("address_lines", []))
    contact_bits = [bit for bit in [company.get("phone"), company.get("email"), company.get("website")] if bit]
    if contact_bits:
        company_lines.append(po_pdf_escape(" | ".join(contact_bits)))
    reg_bits = []
    if company.get("company_number"):
        reg_bits.append(f"Company No: {company.get('company_number')}")
    if company.get("vat_number"):
        reg_bits.append(f"VAT No: {company.get('vat_number')}")
    if reg_bits:
        company_lines.append(po_pdf_escape(" | ".join(reg_bits)))

    logo = get_po_logo_flowable()
    header_left = logo if logo else Paragraph(f"<b>{po_pdf_escape(company.get('name'))}</b>", styles["Heading2"])
    header_right = Paragraph(
        f"<b>PURCHASE ORDER</b><br/><font size='10'>PO Number: {po_pdf_escape(po.get('po_number', ''))}</font><br/>"
        f"<font size='9'>Status: {po_pdf_escape(str(po.get('status', 'draft')).replace('_', ' ').title())}</font><br/>"
        f"<font size='9'>Date: {get_uk_time().strftime('%d/%m/%Y')}</font>",
        styles["RightTitle"],
    )
    header_table = Table([[header_left, header_right]], colWidths=[6.5*cm, 11.5*cm])
    header_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(header_table)
    story.append(Paragraph("<br/>".join(company_lines), styles["SmallMuted"]))
    story.append(Spacer(1, 0.35*cm))

    supplier_lines = [po_pdf_escape(po.get("supplier_name", ""))]
    if po.get("supplier_email"):
        supplier_lines.append(po_pdf_escape(po.get("supplier_email")))
    if po.get("supplier_address"):
        supplier_lines.append(po_pdf_escape(po.get("supplier_address")))

    job_lines = [po_pdf_escape(po.get("job_name", ""))]
    if po.get("delivery_address"):
        job_lines.append(po_pdf_escape(po.get("delivery_address")))

    detail_data = [
        [Paragraph("<b>Supplier</b>", styles["BodyText"]), Paragraph("<b>Job / Delivery</b>", styles["BodyText"])],
        [
            Paragraph("<br/>".join(supplier_lines) or "-", styles["BodyText"]),
            Paragraph("<br/>".join(job_lines) or "-", styles["BodyText"]),
        ],
        [
            Paragraph(f"<b>Supplier Quote Ref:</b> {po_pdf_escape(po.get('supplier_quote_number') or '-')}", styles["BodyText"]),
            Paragraph(f"<b>Required Date:</b> {po_pdf_escape(format_uk_date_only(po.get('required_date')))}", styles["BodyText"]),
        ],
    ]
    detail_table = Table(detail_data, colWidths=[9*cm, 9*cm])
    detail_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(detail_table)
    story.append(Spacer(1, 0.45*cm))

    table_data = [["Description", "Qty", "Unit", "VAT %", "Net", "VAT", "Gross"]]
    for line in po.get("lines", []):
        qty = finance_material_to_float(line.get("quantity"), 0.0)
        unit_cost = finance_material_to_float(line.get("unit_cost"), 0.0)
        vat_rate = finance_material_to_float(line.get("vat_rate"), 0.0)
        net_total = finance_material_to_float(line.get("net_total"), qty * unit_cost)
        vat_total = finance_material_to_float(line.get("vat_total"), net_total * vat_rate / 100)
        gross_total = finance_material_to_float(line.get("gross_total"), net_total + vat_total)
        table_data.append([
            Paragraph(po_pdf_escape(line.get("description", "")), styles["BodyText"]),
            f"{qty:g}",
            f"£{unit_cost:,.2f}",
            f"{vat_rate:g}%",
            f"£{net_total:,.2f}",
            f"£{vat_total:,.2f}",
            f"£{gross_total:,.2f}",
        ])
    table_data.extend([
        ["", "", "", "", "Net", "", f"£{finance_material_to_float(po.get('net_total')):,.2f}"],
        ["", "", "", "", "VAT", "", f"£{finance_material_to_float(po.get('vat_total')):,.2f}"],
        ["", "", "", "", "Gross", "", f"£{finance_material_to_float(po.get('gross_total')):,.2f}"],
    ])
    line_table = Table(table_data, colWidths=[7*cm, 1.3*cm, 2*cm, 1.5*cm, 2*cm, 2*cm, 2.2*cm], repeatRows=1)
    line_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d01f2f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("BACKGROUND", (4, -3), (-1, -1), colors.HexColor("#f3f4f6")),
        ("FONTNAME", (4, -3), (-1, -1), "Helvetica-Bold"),
    ]))
    story.append(line_table)

    if po.get("notes"):
        story.append(Spacer(1, 0.45*cm))
        story.append(Paragraph("<b>Notes</b>", styles["Heading3"]))
        story.append(Paragraph(po_pdf_escape(po.get("notes", "")).replace("\n", "<br/>"), styles["BodyText"]))

    story.append(Spacer(1, 0.45*cm))
    story.append(Paragraph("Please confirm receipt and advise expected delivery date.", styles["Normal"]))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_purchase_orders_export_pdf_bytes(purchase_orders: List[Dict[str, Any]], title: str = "Purchase Orders Export") -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    except Exception as exc:
        logger.error("PO export PDF dependency error: %s", exc)
        raise HTTPException(status_code=500, detail="PDF generation is not available on this server. Check reportlab is installed.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=1.0*cm, leftMargin=1.0*cm, topMargin=1.0*cm, bottomMargin=1.0*cm)
    styles = getSampleStyleSheet()
    story = []
    company = po_company_details()
    logo = get_po_logo_flowable(width_cm=1.6)
    heading = Paragraph(f"<b>{po_pdf_escape(title)}</b><br/><font size='9'>{po_pdf_escape(company.get('name'))} | Generated {get_uk_time().strftime('%d/%m/%Y %H:%M')}</font>", styles["Title"])
    header = Table([[logo or "", heading]], colWidths=[2.4*cm, 24*cm])
    header.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "TOP")]))
    story.append(header)
    story.append(Spacer(1, 0.3*cm))

    totals = {
        "net": sum(finance_material_to_float(po.get("net_total")) for po in purchase_orders),
        "vat": sum(finance_material_to_float(po.get("vat_total")) for po in purchase_orders),
        "gross": sum(finance_material_to_float(po.get("gross_total")) for po in purchase_orders),
    }
    story.append(Paragraph(f"<b>POs:</b> {len(purchase_orders)} &nbsp;&nbsp; <b>Net:</b> £{totals['net']:,.2f} &nbsp;&nbsp; <b>VAT:</b> £{totals['vat']:,.2f} &nbsp;&nbsp; <b>Gross:</b> £{totals['gross']:,.2f}", styles["Normal"]))
    story.append(Spacer(1, 0.25*cm))

    table_data = [["PO", "Date", "Supplier", "Job", "Required", "Status", "Net", "VAT", "Gross"]]
    for po in purchase_orders:
        created = finance_material_iso(po.get("created_at")) or ""
        table_data.append([
            po_pdf_escape(po.get("po_number")),
            created,
            Paragraph(po_pdf_escape(po.get("supplier_name") or "-"), styles["BodyText"]),
            Paragraph(po_pdf_escape(po.get("job_name") or "-"), styles["BodyText"]),
            format_uk_date_only(po.get("required_date")),
            po_pdf_escape(str(po.get("status") or "draft").replace("_", " ").title()),
            f"£{finance_material_to_float(po.get('net_total')):,.2f}",
            f"£{finance_material_to_float(po.get('vat_total')):,.2f}",
            f"£{finance_material_to_float(po.get('gross_total')):,.2f}",
        ])
    table = Table(table_data, colWidths=[3*cm, 2.3*cm, 4.2*cm, 5.2*cm, 2.3*cm, 3.0*cm, 2.2*cm, 2.2*cm, 2.3*cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#111827")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#cbd5e1")),
        ("FONTSIZE", (0,0), (-1,-1), 7),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("ALIGN", (6,1), (-1,-1), "RIGHT"),
    ]))
    story.append(table)
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

@api_router.get("/suppliers")
async def get_suppliers(include_archived: bool = Query(False), admin: str = Depends(verify_admin)):
    filter_dict = {} if include_archived else {"archived": {"$ne": True}}
    suppliers = await db.suppliers.find(filter_dict, {"_id": 0}).sort("name", 1).to_list(1000)
    return suppliers




def normalise_supplier_import_value(value: Any) -> str:
    return str(value or "").strip()


def normalise_supplier_match_key(value: str) -> str:
    cleaned = re.sub(r"[^a-z0-9]+", " ", str(value or "").lower())
    noise = {"ltd", "limited", "plc", "uk", "the", "and", "t", "a", "ta", "t/a", "company", "co"}
    tokens = [token for token in cleaned.split() if token and token not in noise]
    return " ".join(tokens)


def get_csv_value(row: Dict[str, Any], aliases: List[str]) -> str:
    lowered = {str(key or "").strip().lower().replace(" ", "_"): value for key, value in row.items()}
    for alias in aliases:
        key = alias.strip().lower().replace(" ", "_")
        if key in lowered and lowered[key] is not None:
            return normalise_supplier_import_value(lowered[key])
    return ""


def parse_bool_csv(value: Any, default: bool = True) -> bool:
    text = str(value or "").strip().lower()
    if not text:
        return default
    if text in {"true", "yes", "y", "1", "active"}:
        return True
    if text in {"false", "no", "n", "0", "inactive", "archived"}:
        return False
    return default


@api_router.get("/suppliers/import-template")
async def download_supplier_import_template(admin: str = Depends(verify_admin)):
    """Download a CSV template for supplier bulk imports."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "supplier_name",
        "contact_name",
        "orders_email",
        "accounts_email",
        "phone",
        "address",
        "vat_number",
        "payment_terms",
        "payment_terms_days",
        "payment_requirement",
        "lead_time_days",
        "notes",
        "active",
    ])
    writer.writerow([
        "Example Supplier Ltd",
        "Sales Team",
        "orders@example-supplier.co.uk",
        "accounts@example-supplier.co.uk",
        "0191 000 0000",
        "Example address, Newcastle",
        "123456789",
        "30 days",
        "30",
        "credit_terms",
        "0",
        "Main supplier notes",
        "true",
    ])
    output.seek(0)
    return StreamingResponse(
        io.BytesIO(("\ufeff" + output.getvalue()).encode("utf-8")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=supplier_import_template.csv"},
    )


def parse_supplier_int(value: Any, default: int = 0) -> int:
    try:
        if value is None or value == "":
            return default
        return max(0, int(round(float(str(value).strip()))))
    except Exception:
        return default


def normalise_supplier_payment_requirement(value: Any) -> str:
    text = str(value or "").strip().lower().replace(" ", "_").replace("-", "_")
    if text in ["proforma", "pro_forma", "pay_before_order", "advance", "upfront", "prepayment"]:
        return "proforma"
    if text in ["immediate", "due_on_receipt", "cod", "cash", "cash_on_delivery", "0_days"]:
        return "immediate"
    return "credit_terms"


@api_router.post("/suppliers/import-csv")
async def import_suppliers_csv(
    file: UploadFile = File(...),
    update_existing: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    """Bulk create/update suppliers from a CSV file."""
    if not file.filename or not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Please upload a CSV file")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The uploaded CSV file is empty")
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="CSV file is too large. Maximum size is 5MB")

    try:
        decoded = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        decoded = content.decode("latin-1")

    reader = csv.DictReader(io.StringIO(decoded))
    if not reader.fieldnames:
        raise HTTPException(status_code=400, detail="CSV file has no header row")

    existing_suppliers = await db.suppliers.find({"archived": {"$ne": True}}, {"_id": 0}).to_list(10000)
    by_email = {}
    by_name = {}
    for supplier in existing_suppliers:
        for email_field in ["orders_email", "accounts_email"]:
            email = str(supplier.get(email_field) or "").strip().lower()
            if email:
                by_email[email] = supplier
        name_key = normalise_supplier_match_key(supplier.get("name", ""))
        if name_key:
            by_name[name_key] = supplier

    created = 0
    updated = 0
    skipped = 0
    errors = []
    processed = []

    for row_number, row in enumerate(reader, start=2):
        try:
            name = get_csv_value(row, ["supplier_name", "name", "supplier", "company", "company_name"])
            orders_email = get_csv_value(row, ["orders_email", "order_email", "email", "supplier_email", "sales_email"])
            accounts_email = get_csv_value(row, ["accounts_email", "account_email", "accounts", "accounts_contact_email"])
            contact_name = get_csv_value(row, ["contact_name", "contact", "main_contact"])
            phone = get_csv_value(row, ["phone", "telephone", "tel", "mobile"])
            address = get_csv_value(row, ["address", "supplier_address", "registered_address"])
            vat_number = get_csv_value(row, ["vat_number", "vat", "vat_no", "vat_registration", "vat_reg"])
            payment_terms = get_csv_value(row, ["payment_terms", "terms"])
            payment_terms_days = parse_supplier_int(get_csv_value(row, ["payment_terms_days", "terms_days", "days_to_pay", "days_to_pay_after_invoice", "credit_days"]), 30)
            payment_requirement = normalise_supplier_payment_requirement(get_csv_value(row, ["payment_requirement", "payment_type", "terms_type", "payment_method", "supplier_payment_requirement"]))
            lead_time_days = parse_supplier_int(get_csv_value(row, ["lead_time_days", "material_lead_time_days", "ordering_lead_time_days", "lead_time"]), 0)
            notes = get_csv_value(row, ["notes", "note", "comments"])
            active = parse_bool_csv(get_csv_value(row, ["active", "status"]), default=True)

            if not name:
                skipped += 1
                errors.append({"row": row_number, "message": "Missing supplier_name/name"})
                continue

            match = None
            for email in [orders_email, accounts_email]:
                email_key = email.strip().lower()
                if email_key and email_key in by_email:
                    match = by_email[email_key]
                    break

            if not match:
                name_key = normalise_supplier_match_key(name)
                if name_key and name_key in by_name:
                    match = by_name[name_key]

            supplier_doc = {
                "name": name,
                "contact_name": contact_name,
                "orders_email": orders_email,
                "accounts_email": accounts_email,
                "phone": phone,
                "address": address,
                "vat_number": vat_number,
                "payment_terms": payment_terms or f"{payment_terms_days} days",
                "payment_terms_days": payment_terms_days,
                "payment_requirement": payment_requirement,
                "lead_time_days": lead_time_days,
                "notes": notes,
                "active": active,
                "archived": False,
            }

            if match:
                if not update_existing:
                    skipped += 1
                    processed.append({"row": row_number, "supplier_name": name, "action": "skipped_existing"})
                    continue
                update_doc = {k: v for k, v in supplier_doc.items() if v not in [None, ""] or k in ["active", "archived"]}
                update_doc["updated_at"] = datetime.utcnow()
                await db.suppliers.update_one({"id": match["id"]}, {"$set": update_doc})
                updated += 1
                processed.append({"row": row_number, "supplier_name": name, "action": "updated"})
                # Keep lookup current for subsequent rows.
                match.update(update_doc)
            else:
                supplier_obj = Supplier(**supplier_doc)
                await db.suppliers.insert_one(supplier_obj.dict())
                created += 1
                processed.append({"row": row_number, "supplier_name": name, "action": "created"})
                supplier_lookup = supplier_obj.dict()
                for email in [supplier_obj.orders_email, supplier_obj.accounts_email]:
                    email_key = str(email or "").strip().lower()
                    if email_key:
                        by_email[email_key] = supplier_lookup
                name_key = normalise_supplier_match_key(supplier_obj.name)
                if name_key:
                    by_name[name_key] = supplier_lookup
        except Exception as exc:
            skipped += 1
            errors.append({"row": row_number, "message": str(exc)})

    return {
        "message": "Supplier CSV import complete",
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "errors": errors,
        "processed": processed[:100],
    }

@api_router.post("/suppliers", response_model=Supplier)
async def create_supplier(supplier: SupplierCreate, admin: str = Depends(verify_admin)):
    existing = await db.suppliers.find_one({"name": {"$regex": f"^{re.escape(supplier.name)}$", "$options": "i"}, "archived": {"$ne": True}})
    if existing:
        raise HTTPException(status_code=400, detail="A supplier with this name already exists")
    supplier_obj = Supplier(**supplier.dict())
    await db.suppliers.insert_one(supplier_obj.dict())
    return supplier_obj


@api_router.put("/suppliers/{supplier_id}", response_model=Supplier)
async def update_supplier(supplier_id: str, supplier_update: SupplierUpdate, admin: str = Depends(verify_admin)):
    update_dict = {k: v for k, v in supplier_update.dict().items() if v is not None}
    update_dict["updated_at"] = datetime.utcnow()
    result = await db.suppliers.update_one({"id": supplier_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Supplier not found")
    updated = await db.suppliers.find_one({"id": supplier_id})
    return Supplier(**updated)


@api_router.delete("/suppliers/{supplier_id}")
async def archive_supplier(supplier_id: str, admin: str = Depends(verify_admin)):
    result = await db.suppliers.update_one({"id": supplier_id}, {"$set": {"archived": True, "active": False, "updated_at": datetime.utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Supplier not found")
    return {"message": "Supplier archived successfully"}



def normalise_work_order_status(value: Optional[str]) -> str:
    status = str(value or "price_received").strip().lower().replace(" ", "_").replace("-", "_")
    aliases = {
        "price requested": "price_requested",
        "awaiting_price": "price_requested",
        "price received": "price_received",
        "quote_received": "price_received",
        "quoted": "price_received",
        "in progress": "in_progress",
        "awaiting sign off": "awaiting_sign_off",
        "awaiting_signoff": "awaiting_sign_off",
        "sign_off": "awaiting_sign_off",
        "invoice received": "invoice_received",
        "part paid": "part_paid",
        "retention held": "retention_held",
        "complete": "complete",
        "completed": "complete",
        "closed": "closed",
        "cancelled": "cancelled",
        "canceled": "cancelled",
    }
    status = aliases.get(status, status)
    return status if status in WORK_ORDER_STATUSES else "price_received"




def work_order_date_to_iso(value: Any) -> str:
    raw = str(value or "").strip()[:10]
    if not raw:
        return ""
    try:
        return datetime.fromisoformat(raw).date().isoformat()
    except Exception:
        return raw


def work_order_get_friday_on_or_after(value: Any) -> str:
    iso = work_order_date_to_iso(value) or datetime.utcnow().date().isoformat()
    start = datetime.fromisoformat(iso)
    days_until_friday = (4 - start.weekday()) % 7
    return (start + timedelta(days=days_until_friday)).date().isoformat()


def work_order_add_months(date_value: datetime, months: int) -> datetime:
    month = date_value.month - 1 + months
    year = date_value.year + month // 12
    month = month % 12 + 1
    day = min(date_value.day, 28)
    return date_value.replace(year=year, month=month, day=day)


def build_default_work_order_payment_schedule(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Build default WO payment lines server-side so finance still works if the UI only sends a mode."""
    mode = str(data.get("payment_schedule_mode") or "single").strip().lower().replace(" ", "_").replace("-", "_")
    if mode not in {"single", "weekly", "monthly", "manual", "milestone"}:
        mode = "single"

    net_total = round(max(0.0, finance_to_number(data.get("net_amount"), 0.0)), 2)
    if net_total <= 0:
        return []

    vat_rate = max(0.0, finance_to_number(data.get("vat_rate"), 20.0))
    cis_applicable = bool(data.get("cis_applicable"))
    cis_rate = max(0.0, min(100.0, finance_to_number(data.get("cis_rate"), 20.0))) if cis_applicable else 0.0
    retention_percent = max(0.0, min(100.0, finance_to_number(data.get("retention_percent"), 0.0)))

    start_iso = work_order_date_to_iso(data.get("expected_start_date") or data.get("payment_due_date") or datetime.utcnow().date().isoformat())
    completion_iso = work_order_date_to_iso(data.get("expected_completion_date") or data.get("payment_due_date") or start_iso)
    due_fallback = work_order_date_to_iso(data.get("payment_due_date") or completion_iso or start_iso)
    release_date = work_order_date_to_iso(data.get("retention_release_date"))

    due_dates: List[str] = []
    try:
        start = datetime.fromisoformat(start_iso)
        completion = datetime.fromisoformat(completion_iso)
    except Exception:
        start = datetime.utcnow()
        completion = start

    if mode == "weekly":
        cursor = datetime.fromisoformat(work_order_get_friday_on_or_after(start_iso))
        while cursor <= completion or not due_dates:
            due_dates.append(cursor.date().isoformat())
            cursor += timedelta(days=7)
            if len(due_dates) > 104:
                break
    elif mode == "monthly":
        cursor = start.replace(day=min(28, start.day))
        while cursor <= completion or not due_dates:
            due_dates.append(cursor.date().isoformat())
            cursor = work_order_add_months(cursor, 1)
            if len(due_dates) > 36:
                break
    elif mode == "milestone":
        due_dates = [item for item in [start_iso, completion_iso] if item]
    else:
        due_dates = [due_fallback]

    unique_dates: List[str] = []
    for item in due_dates:
        if item and item not in unique_dates:
            unique_dates.append(item)
    if not unique_dates:
        unique_dates = [due_fallback or datetime.utcnow().date().isoformat()]

    split_count = max(1, len(unique_dates))
    remaining_net = net_total
    remaining_retention = round(net_total * (retention_percent / 100), 2)
    lines: List[Dict[str, Any]] = []

    for index, due_date in enumerate(unique_dates, start=1):
        is_last = index == len(unique_dates)
        net = round(remaining_net, 2) if is_last else round(net_total / split_count, 2)
        remaining_net = round(remaining_net - net, 2)
        vat = round(net * (vat_rate / 100), 2)
        gross = round(net + vat, 2)
        cis = round(net * (cis_rate / 100), 2)
        retention = round(remaining_retention, 2) if is_last else round(net * (retention_percent / 100), 2)
        remaining_retention = round(remaining_retention - retention, 2)
        if mode == "weekly":
            label = f"Week {index} invoice"
        elif mode == "monthly":
            label = f"Month {index} invoice"
        elif mode == "milestone":
            label = "Start milestone" if index == 1 else "Completion milestone"
        else:
            label = "Main payment"
        lines.append({
            "id": f"line-{index}",
            "label": label,
            "due_date": due_date,
            "percentage": round((net / max(1.0, net_total)) * 100, 2),
            "net_amount": net,
            "vat_amount": vat,
            "gross_amount": gross,
            "cis_deduction": cis,
            "retention_amount": retention,
            "cash_amount": round(gross - cis - retention, 2),
            "status": "invoice_expected",
            "is_retention_release": False,
            "notes": "",
        })

    retention_total = round(net_total * (retention_percent / 100), 2)
    if retention_total > 0:
        lines.append({
            "id": "retention-release",
            "label": "Retention release",
            "due_date": release_date,
            "percentage": 0,
            "net_amount": 0,
            "vat_amount": 0,
            "gross_amount": 0,
            "cis_deduction": 0,
            "retention_amount": -retention_total,
            "cash_amount": retention_total,
            "status": "retention_release_due" if release_date else "held",
            "is_retention_release": True,
            "notes": str(data.get("retention_release_rule") or "Release after sign-off"),
        })

    return lines

def normalise_work_order_payment_schedule(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    raw_lines = data.get("payment_schedule")
    if not isinstance(raw_lines, list):
        raw_lines = []
    # If the UI sends only a payment mode (for example weekly) but no lines, build
    # the schedule here so the saved WO and Finance cashflow both have staged payments.
    if not raw_lines:
        raw_lines = build_default_work_order_payment_schedule(data)

    vat_rate = max(0.0, finance_to_number(data.get("vat_rate"), 20.0))
    cis_applicable = bool(data.get("cis_applicable"))
    cis_rate = max(0.0, min(100.0, finance_to_number(data.get("cis_rate"), 20.0))) if cis_applicable else 0.0
    cleaned: List[Dict[str, Any]] = []

    for index, line in enumerate(raw_lines, start=1):
        if not isinstance(line, dict):
            continue
        net_amount = round(max(0.0, finance_to_number(line.get("net_amount"), 0.0)), 2)
        vat_amount = finance_to_number(line.get("vat_amount"), 0.0)
        if vat_amount <= 0 and net_amount > 0 and vat_rate > 0:
            vat_amount = net_amount * (vat_rate / 100)
        gross_amount = finance_to_number(line.get("gross_amount"), 0.0)
        if gross_amount <= 0 and net_amount > 0:
            gross_amount = net_amount + vat_amount
        cis_deduction = finance_to_number(line.get("cis_deduction"), 0.0)
        if cis_applicable and cis_deduction <= 0 and net_amount > 0:
            cis_deduction = net_amount * (cis_rate / 100)
        retention_amount = finance_to_number(line.get("retention_amount"), 0.0)
        cash_amount = finance_to_number(line.get("cash_amount"), 0.0)
        if cash_amount <= 0 and (gross_amount > 0 or retention_amount != 0):
            cash_amount = gross_amount - cis_deduction - retention_amount

        status = str(line.get("status") or "not_due").strip().lower().replace(" ", "_").replace("-", "_")
        if status not in {"not_due", "invoice_expected", "invoice_received", "approved", "part_paid", "paid", "held", "retention_release_due"}:
            status = "not_due"

        cleaned.append({
            "id": str(line.get("id") or f"line-{index}"),
            "label": str(line.get("label") or f"Payment {index}"),
            "due_date": str(line.get("due_date") or "")[:10],
            "percentage": finance_to_number(line.get("percentage"), 0.0),
            "net_amount": round(net_amount, 2),
            "vat_amount": round(vat_amount, 2),
            "gross_amount": round(gross_amount, 2),
            "cis_deduction": round(cis_deduction, 2),
            "retention_amount": round(retention_amount, 2),
            "cash_amount": round(cash_amount, 2),
            "status": status,
            "is_retention_release": bool(line.get("is_retention_release")),
            "notes": str(line.get("notes") or ""),
        })

    return cleaned

def calculate_work_order_totals(data: Dict[str, Any]) -> Dict[str, Any]:
    pricing_type = str(data.get("pricing_type") or "fixed_price").strip().lower().replace(" ", "_")
    if pricing_type not in {"fixed_price", "day_rate", "hourly", "item_rate"}:
        pricing_type = "fixed_price"

    quantity = max(0.0, finance_to_number(data.get("quantity"), 1.0))
    rate = max(0.0, finance_to_number(data.get("rate"), 0.0))
    net_amount = finance_to_number(data.get("net_amount"), 0.0)
    if net_amount <= 0 and pricing_type in {"day_rate", "hourly", "item_rate"}:
        net_amount = quantity * rate
    net_amount = max(0.0, net_amount)

    vat_rate = max(0.0, finance_to_number(data.get("vat_rate"), 20.0))
    vat_amount = finance_to_number(data.get("vat_amount"), 0.0)
    if vat_amount <= 0 and vat_rate > 0:
        vat_amount = net_amount * (vat_rate / 100)
    gross_amount = finance_to_number(data.get("gross_amount"), 0.0)
    if gross_amount <= 0:
        gross_amount = net_amount + vat_amount

    cis_applicable = bool(data.get("cis_applicable"))
    cis_rate = max(0.0, min(100.0, finance_to_number(data.get("cis_rate"), 20.0))) if cis_applicable else 0.0
    cis_deduction = finance_to_number(data.get("cis_deduction"), 0.0)
    if cis_applicable and cis_deduction <= 0:
        cis_deduction = net_amount * (cis_rate / 100)

    payment_terms_days = int(finance_to_number(data.get("payment_terms_days"), 30))
    payment_requirement = str(data.get("payment_requirement") or "credit_terms").strip().lower()
    if payment_requirement not in {"credit_terms", "proforma", "immediate"}:
        payment_requirement = "credit_terms"

    if not data.get("payment_due_date") and data.get("expected_completion_date"):
        try:
            completion = datetime.fromisoformat(str(data.get("expected_completion_date"))[:10])
            if payment_requirement in {"proforma", "immediate"}:
                data["payment_due_date"] = completion.date().isoformat()
            else:
                data["payment_due_date"] = (completion + timedelta(days=payment_terms_days)).date().isoformat()
        except Exception:
            pass

    data["pricing_type"] = pricing_type
    data["quantity"] = quantity
    data["rate"] = rate
    data["net_amount"] = round(net_amount, 2)
    data["vat_rate"] = vat_rate
    data["vat_amount"] = round(vat_amount, 2)
    data["gross_amount"] = round(gross_amount, 2)
    data["cis_applicable"] = cis_applicable
    data["cis_rate"] = cis_rate
    data["cis_deduction"] = round(cis_deduction, 2)
    data["payment_requirement"] = payment_requirement
    data["payment_terms_days"] = payment_terms_days
    payment_schedule_mode = str(data.get("payment_schedule_mode") or "single").strip().lower().replace(" ", "_").replace("-", "_")
    if payment_schedule_mode not in {"single", "weekly", "monthly", "manual", "milestone"}:
        payment_schedule_mode = "single"
    data["payment_schedule_mode"] = payment_schedule_mode
    data["retention_percent"] = max(0.0, min(100.0, finance_to_number(data.get("retention_percent"), 0.0)))
    retention_release_rule = str(data.get("retention_release_rule") or "manual").strip().lower().replace(" ", "_").replace("-", "_")
    if retention_release_rule not in {"manual", "supervisor_sign_off", "defects_period", "client_payment_received"}:
        retention_release_rule = "manual"
    data["retention_release_rule"] = retention_release_rule
    data["retention_release_date"] = str(data.get("retention_release_date") or "")[:10] or None
    data["payment_schedule"] = normalise_work_order_payment_schedule(data)
    data["status"] = normalise_work_order_status(data.get("status"))
    data["is_committed_cost"] = data["status"] in WORK_ORDER_COMMITTED_STATUSES
    return data


async def get_next_work_order_number() -> str:
    today_prefix = f"WO-{datetime.utcnow().strftime('%Y%m%d')}-"
    latest = await db.work_orders.find_one({"wo_number": {"$regex": f"^{today_prefix}"}}, sort=[("wo_number", -1)])
    if latest and latest.get("wo_number"):
        try:
            next_number = int(str(latest["wo_number"]).split("-")[-1]) + 1
        except Exception:
            next_number = 1
    else:
        next_number = 1
    return f"{today_prefix}{next_number:03d}"


async def enrich_work_order_job_and_contractor(data: Dict[str, Any]) -> Dict[str, Any]:
    job_id = str(data.get("job_id") or "").strip()
    if not job_id:
        raise HTTPException(status_code=400, detail="job_id is required")

    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    data["job_id"] = job_id
    data["job_name"] = data.get("job_name") or job.get("display_name") or job.get("name", "")
    data["job_number"] = data.get("job_number") if data.get("job_number") is not None else job.get("job_number")
    data["division"] = data.get("division") or job.get("division", "")

    section_id = str(data.get("section_id") or "").strip()
    if section_id:
        section = next((item for item in job.get("gantt_sections") or [] if str(item.get("id")) == section_id), None)
        if section:
            data["section_name"] = data.get("section_name") or section.get("name", "")
            data["expected_start_date"] = data.get("expected_start_date") or section.get("start_date")
            data["expected_completion_date"] = data.get("expected_completion_date") or section.get("end_date")

    contractor_id = str(data.get("contractor_id") or "").strip()
    if contractor_id:
        contractor = await db.workers.find_one({"id": contractor_id, "archived": {"$ne": True}}, {"_id": 0})
        if contractor:
            data["contractor_name"] = data.get("contractor_name") or contractor.get("name", "")
            data["contractor_email"] = data.get("contractor_email") or contractor.get("email", "")
            if not data.get("trade"):
                trades = contractor.get("trades") or []
                data["trade"] = trades[0] if trades else contractor.get("trade", "")

    return data


@api_router.get("/work-orders")
async def get_work_orders(
    status: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    contractor_id: Optional[str] = Query(None),
    section_id: Optional[str] = Query(None),
    committed_only: bool = Query(False),
    include_cancelled: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    filter_dict: Dict[str, Any] = {"archived": {"$ne": True}}
    if status:
        filter_dict["status"] = normalise_work_order_status(status)
    elif not include_cancelled:
        filter_dict["status"] = {"$nin": list(WORK_ORDER_CANCELLED_STATUSES)}
    if committed_only:
        filter_dict["status"] = {"$in": list(WORK_ORDER_COMMITTED_STATUSES)}
    if job_id:
        filter_dict["job_id"] = job_id
    if contractor_id:
        filter_dict["contractor_id"] = contractor_id
    if section_id:
        filter_dict["section_id"] = section_id
    return await db.work_orders.find(filter_dict, {"_id": 0}).sort("created_at", -1).to_list(5000)


@api_router.get("/jobs/{job_id}/work-orders")
async def get_job_work_orders(job_id: str, include_cancelled: bool = Query(False), admin: str = Depends(verify_admin)):
    filter_dict: Dict[str, Any] = {"job_id": job_id, "archived": {"$ne": True}}
    if not include_cancelled:
        filter_dict["status"] = {"$nin": list(WORK_ORDER_CANCELLED_STATUSES)}
    return await db.work_orders.find(filter_dict, {"_id": 0}).sort("created_at", -1).to_list(5000)


@api_router.get("/work-orders/{work_order_id}", response_model=WorkOrder)
async def get_work_order(work_order_id: str, admin: str = Depends(verify_admin)):
    work_order = await db.work_orders.find_one({"id": work_order_id, "archived": {"$ne": True}}, {"_id": 0})
    if not work_order:
        raise HTTPException(status_code=404, detail="Work order not found")
    return WorkOrder(**work_order)


@api_router.post("/work-orders", response_model=WorkOrder)
async def create_work_order(work_order: WorkOrderCreate, admin: str = Depends(verify_admin)):
    data = work_order.dict()
    data = await enrich_work_order_job_and_contractor(data)
    data = calculate_work_order_totals(data)
    if not data.get("wo_number"):
        data["wo_number"] = await get_next_work_order_number()
    obj = WorkOrder(**data)
    await db.work_orders.insert_one(obj.dict())
    return obj


@api_router.put("/work-orders/{work_order_id}", response_model=WorkOrder)
async def update_work_order(work_order_id: str, update: WorkOrderUpdate, admin: str = Depends(verify_admin)):
    existing = await db.work_orders.find_one({"id": work_order_id, "archived": {"$ne": True}}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Work order not found")

    update_data = {k: v for k, v in update.dict().items() if v is not None}
    if not update_data:
        return WorkOrder(**existing)

    merged = {**existing, **update_data}
    merged = await enrich_work_order_job_and_contractor(merged)
    merged = calculate_work_order_totals(merged)
    merged["updated_at"] = datetime.utcnow()

    await db.work_orders.update_one({"id": work_order_id}, {"$set": merged})
    updated = await db.work_orders.find_one({"id": work_order_id}, {"_id": 0})
    return WorkOrder(**updated)


@api_router.put("/work-orders/{work_order_id}/archive")
async def archive_work_order(work_order_id: str, admin: str = Depends(verify_admin)):
    result = await db.work_orders.update_one(
        {"id": work_order_id},
        {"$set": {"archived": True, "status": "cancelled", "updated_at": datetime.utcnow()}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Work order not found")
    return {"success": True, "message": "Work order archived"}


@api_router.delete("/work-orders/{work_order_id}")
async def delete_work_order(work_order_id: str, admin: str = Depends(verify_admin)):
    result = await db.work_orders.update_one(
        {"id": work_order_id},
        {"$set": {"archived": True, "status": "cancelled", "updated_at": datetime.utcnow()}},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Work order not found")
    return {"success": True, "message": "Work order deleted"}


@api_router.get("/purchase-orders")
async def get_purchase_orders(
    status: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    supplier_id: Optional[str] = Query(None),
    include_cancelled: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    filter_dict = {}
    if status:
        filter_dict["status"] = status
    elif not include_cancelled:
        filter_dict["status"] = {"$ne": "cancelled"}
    if job_id:
        filter_dict["job_id"] = job_id
    if supplier_id:
        filter_dict["supplier_id"] = supplier_id
    purchase_orders = await db.purchase_orders.find(filter_dict, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return purchase_orders


def parse_po_id_list(ids: Optional[str] = None, po_ids: Optional[List[str]] = None) -> List[str]:
    values: List[str] = []
    if ids:
        values.extend([item.strip() for item in str(ids).split(",") if item.strip()])
    if po_ids:
        values.extend([str(item).strip() for item in po_ids if str(item).strip()])
    # Preserve order but remove duplicates.
    seen = set()
    cleaned = []
    for value in values:
        if value not in seen:
            seen.add(value)
            cleaned.append(value)
    return cleaned


async def find_purchase_orders_for_export(
    status: Optional[str] = None,
    job_id: Optional[str] = None,
    supplier_id: Optional[str] = None,
    ids: Optional[str] = None,
    include_cancelled: bool = True,
) -> List[Dict[str, Any]]:
    filter_dict: Dict[str, Any] = {}
    selected_ids = parse_po_id_list(ids)
    if selected_ids:
        filter_dict["id"] = {"$in": selected_ids}
    else:
        if status:
            filter_dict["status"] = status
        elif not include_cancelled:
            filter_dict["status"] = {"$ne": "cancelled"}
        if job_id:
            filter_dict["job_id"] = job_id
        if supplier_id:
            filter_dict["supplier_id"] = supplier_id

    purchase_orders = await db.purchase_orders.find(filter_dict, {"_id": 0}).sort("created_at", -1).to_list(10000)
    if selected_ids:
        order = {po_id: index for index, po_id in enumerate(selected_ids)}
        purchase_orders.sort(key=lambda po: order.get(po.get("id"), 999999))
    return purchase_orders


@api_router.get("/purchase-orders/export.csv")
async def export_purchase_orders_csv(
    status: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    supplier_id: Optional[str] = Query(None),
    ids: Optional[str] = Query(None),
    include_cancelled: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    purchase_orders = await find_purchase_orders_for_export(status, job_id, supplier_id, ids, include_cancelled)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["LDA Group - Purchase Orders Export"])
    writer.writerow(["Generated", get_uk_time().strftime("%d/%m/%Y %H:%M")])
    writer.writerow([])
    writer.writerow([
        "PO Number", "Created", "Required Date", "Supplier", "Supplier Email", "Job", "Job Number", "Division",
        "Status", "Quote Ref", "Line Description", "Quantity", "Unit Cost", "VAT Rate", "Line Net", "Line VAT", "Line Gross",
        "PO Net", "PO VAT", "PO Gross", "Requested By", "Approved By", "Sent At", "Notes",
    ])
    for po in purchase_orders:
        lines = po.get("lines") or [{}]
        for line in lines:
            writer.writerow([
                po.get("po_number", ""),
                format_uk_datetime_for_export(po.get("created_at")),
                format_uk_date_only(po.get("required_date")),
                po.get("supplier_name", ""),
                po.get("supplier_email", ""),
                po.get("job_name", ""),
                po.get("job_number", ""),
                po.get("division", ""),
                po.get("status", ""),
                po.get("supplier_quote_number", ""),
                line.get("description", ""),
                line.get("quantity", ""),
                line.get("unit_cost", ""),
                line.get("vat_rate", ""),
                line.get("net_total", ""),
                line.get("vat_total", ""),
                line.get("gross_total", ""),
                po.get("net_total", ""),
                po.get("vat_total", ""),
                po.get("gross_total", ""),
                po.get("requested_by_name", ""),
                po.get("approved_by_name", ""),
                format_uk_datetime_for_export(po.get("sent_at")),
                po.get("notes", ""),
            ])
    output.seek(0)
    filename = f"purchase_orders_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.csv"
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@api_router.get("/purchase-orders/export.pdf")
async def export_purchase_orders_pdf(
    status: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    supplier_id: Optional[str] = Query(None),
    ids: Optional[str] = Query(None),
    include_cancelled: bool = Query(True),
    admin: str = Depends(verify_admin),
):
    purchase_orders = await find_purchase_orders_for_export(status, job_id, supplier_id, ids, include_cancelled)
    pdf_bytes = generate_purchase_orders_export_pdf_bytes(purchase_orders)
    filename = f"purchase_orders_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.pdf"
    return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})


def build_po_approval_notification_payload(po: Dict[str, Any], requested_by: str = "") -> Dict[str, Any]:
    """Build a compact approval notification payload for Power Automate/SMTP."""
    po_number = po.get("po_number") or "Purchase Order"
    required_date_display = format_uk_date_only(po.get("required_date"))
    if required_date_display == "-":
        required_date_display = "TBC"

    subject = f"PO Approval Required - {po_number} - {po.get('job_name', '')}"
    lines = po.get("lines") or []
    line_summary = "\n".join(
        f"- {line.get('description', '')} | Qty: {line.get('quantity', '')} | Net: £{float(line.get('net_total') or 0):,.2f}"
        for line in lines[:10]
    )
    if len(lines) > 10:
        line_summary += f"\n- plus {len(lines) - 10} more line(s)"

    body_text = f"""A purchase order has been created and requires approval.

PO Number: {po_number}
Supplier: {po.get('supplier_name', '')}
Job: {po.get('job_name', '')}
Required Date: {required_date_display}
Requested By: {po.get('requested_by_name') or requested_by or 'Admin'}

Net Total: £{float(po.get('net_total') or 0):,.2f}
VAT Total: £{float(po.get('vat_total') or 0):,.2f}
Gross Total: £{float(po.get('gross_total') or 0):,.2f}

Line Summary:
{line_summary or '-'}

Notes:
{po.get('notes') or '-'}
"""

    line_rows = "".join(
        f"""
        <tr>
          <td style=\"padding:6px 8px; border-bottom:1px solid #e5e7eb;\">{line.get('description', '')}</td>
          <td style=\"padding:6px 8px; border-bottom:1px solid #e5e7eb; text-align:right;\">{line.get('quantity', '')}</td>
          <td style=\"padding:6px 8px; border-bottom:1px solid #e5e7eb; text-align:right;\">£{float(line.get('net_total') or 0):,.2f}</td>
        </tr>
        """
        for line in lines[:10]
    ) or "<tr><td colspan=\"3\" style=\"padding:6px 8px;\">No line items</td></tr>"

    body_html = f"""
    <div style=\"font-family:Arial, sans-serif; color:#111827;\">
      <h2 style=\"margin:0 0 12px;\">Purchase Order Approval Required</h2>
      <p>A purchase order has been created and requires approval.</p>
      <table style=\"border-collapse:collapse; font-size:14px; margin-bottom:14px;\">
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>PO Number:</strong></td><td>{po_number}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Supplier:</strong></td><td>{po.get('supplier_name', '')}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Job:</strong></td><td>{po.get('job_name', '')}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Required Date:</strong></td><td>{required_date_display}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Requested By:</strong></td><td>{po.get('requested_by_name') or requested_by or 'Admin'}</td></tr>
      </table>
      <table style=\"border-collapse:collapse; font-size:14px; margin-bottom:14px; min-width:360px;\">
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Net Total:</strong></td><td>£{float(po.get('net_total') or 0):,.2f}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>VAT Total:</strong></td><td>£{float(po.get('vat_total') or 0):,.2f}</td></tr>
        <tr><td style=\"padding:4px 14px 4px 0;\"><strong>Gross Total:</strong></td><td><strong>£{float(po.get('gross_total') or 0):,.2f}</strong></td></tr>
      </table>
      <h3 style=\"font-size:15px; margin:14px 0 6px;\">Line Summary</h3>
      <table style=\"border-collapse:collapse; font-size:13px; width:100%; max-width:720px;\">
        <thead>
          <tr style=\"background:#f3f4f6;\">
            <th style=\"padding:7px 8px; text-align:left;\">Description</th>
            <th style=\"padding:7px 8px; text-align:right;\">Qty</th>
            <th style=\"padding:7px 8px; text-align:right;\">Net</th>
          </tr>
        </thead>
        <tbody>{line_rows}</tbody>
      </table>
      <p style=\"margin-top:14px;\"><strong>Notes:</strong><br>{po.get('notes') or '-'}</p>
    </div>
    """.strip()

    return {
        "subject": subject,
        "body_text": body_text,
        "body_html": body_html,
        "required_date_display": required_date_display,
    }


async def send_po_approval_notification(po: Dict[str, Any], requested_by: str = "") -> Dict[str, Any]:
    """Send an internal PO approval notification without blocking PO creation on failure.

    Preferred configuration:
    - POWER_AUTOMATE_PO_APPROVAL_URL
    - POWER_AUTOMATE_PO_APPROVAL_SECRET (optional)

    Fallback configuration:
    - SMTP_HOST / SMTP_PORT / SMTP_USERNAME / SMTP_PASSWORD / SMTP_FROM_EMAIL

    Recipient defaults to info@ldagroup.co.uk and can be overridden with PO_APPROVAL_NOTIFY_EMAIL.
    """
    approval_email = os.environ.get("PO_APPROVAL_NOTIFY_EMAIL", "info@ldagroup.co.uk").strip() or "info@ldagroup.co.uk"
    notification = build_po_approval_notification_payload(po, requested_by=requested_by)
    po_number = po.get("po_number", "purchase_order")

    try:
        pdf_bytes = generate_purchase_order_pdf_bytes(po)
    except Exception as exc:
        logger.warning("Could not generate PO approval PDF attachment for %s: %s", po_number, exc)
        pdf_bytes = b""

    power_automate_url = os.environ.get("POWER_AUTOMATE_PO_APPROVAL_URL", "").strip()
    power_automate_secret = os.environ.get("POWER_AUTOMATE_PO_APPROVAL_SECRET", "").strip()

    if power_automate_url:
        public_backend_url = (
            os.environ.get("PUBLIC_BACKEND_URL")
            or os.environ.get("BACKEND_PUBLIC_URL")
            or os.environ.get("REACT_APP_BACKEND_URL")
            or ""
        ).strip().rstrip("/")
        approval_callback_secret = os.environ.get("PO_APPROVAL_CALLBACK_SECRET", power_automate_secret).strip()
        approval_callback_url = f"{public_backend_url}/api/purchase-orders/{po.get('id')}/approval-response" if public_backend_url and po.get("id") else ""

        payload = {
            "secret": power_automate_secret,
            "callback_secret": approval_callback_secret,
            "approval_callback_url": approval_callback_url,
            "notification_type": "purchase_order_approval_required",
            "to": approval_email,
            "po_id": po.get("id"),
            "po_number": po_number,
            "supplier_id": po.get("supplier_id", ""),
            "supplier_name": po.get("supplier_name", ""),
            "supplier_email": po.get("supplier_email", ""),
            "job_id": po.get("job_id", ""),
            "job_name": po.get("job_name", ""),
            "job_number": po.get("job_number", ""),
            "division": po.get("division", ""),
            "required_date": notification["required_date_display"],
            "delivery_address": po.get("delivery_address", ""),
            "status": po.get("status", "draft"),
            "requested_by": po.get("requested_by_name") or requested_by or "Admin",
            "requested_by_email": po.get("requested_by_email", ""),
            "requested_by_role": po.get("requested_by_role", ""),
            "net_total": po.get("net_total", 0),
            "vat_total": po.get("vat_total", 0),
            "gross_total": po.get("gross_total", 0),
            "subject": notification["subject"],
            "body_text": notification["body_text"],
            "body_html": notification["body_html"],
            "pdf_filename": f"{po_number}.pdf",
            "pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8") if pdf_bytes else "",
        }

        response = requests.post(power_automate_url, json=payload, timeout=20)
        if response.status_code < 200 or response.status_code >= 300:
            raise RuntimeError(f"Power Automate approval notification failed: {response.status_code} - {response.text[:500]}")
        return {"sent": True, "method": "power_automate", "to": approval_email}

    smtp_host = os.environ.get("SMTP_HOST", "").strip()
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_username = os.environ.get("SMTP_USERNAME", "").strip()
    smtp_password = os.environ.get("SMTP_PASSWORD", "").strip()
    smtp_from = os.environ.get("SMTP_FROM_EMAIL", smtp_username).strip()
    smtp_from_name = os.environ.get("SMTP_FROM_NAME", "LDA Group").strip()
    smtp_reply_to = os.environ.get("SMTP_REPLY_TO", os.environ.get("PO_REPLY_TO_EMAIL", "")).strip()
    smtp_use_tls = os.environ.get("SMTP_USE_TLS", "true").lower() != "false"

    if not smtp_host or not smtp_from:
        return {"sent": False, "method": "not_configured", "to": approval_email}

    message = EmailMessage()
    message["From"] = f"{smtp_from_name} <{smtp_from}>"
    message["To"] = approval_email
    message["Subject"] = notification["subject"]
    if smtp_reply_to:
        message["Reply-To"] = smtp_reply_to
    message.set_content(notification["body_text"])
    message.add_alternative(notification["body_html"], subtype="html")
    if pdf_bytes:
        message.add_attachment(pdf_bytes, maintype="application", subtype="pdf", filename=f"{po_number}.pdf")

    with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as smtp:
        if smtp_use_tls:
            smtp.starttls()
        if smtp_username and smtp_password:
            smtp.login(smtp_username, smtp_password)
        smtp.send_message(message)

    return {"sent": True, "method": "smtp", "to": approval_email}



def build_po_requester_options_payload(po: Dict[str, Any]) -> Dict[str, Any]:
    """Build the email body sent to the person who raised an approved PO."""
    po_number = po.get("po_number", "purchase_order")
    subject = f"PO {po_number} approved - choose next action"
    required_date_display = format_uk_date_only(po.get("required_date"))
    if required_date_display == "-":
        required_date_display = "TBC"

    body_text = f"""Hi {po.get('requested_by_name') or ''},

Purchase order {po_number} has been approved.

Supplier: {po.get('supplier_name', '')}
Job: {po.get('job_name', '')}
Required date: {required_date_display}
Gross total: £{float(po.get('gross_total') or 0):,.2f}

Please choose whether to send the PO to the supplier and assign the materials to the job, or assign the materials to the job only.

LDA Group
"""

    body_html = f"""
    <div style="font-family:Arial, sans-serif; font-size:14px; color:#111827;">
      <p>Hi {po.get('requested_by_name') or ''},</p>
      <p>Purchase order <strong>{po_number}</strong> has been approved.</p>
      <table style="border-collapse:collapse; font-family:Arial, sans-serif; font-size:14px;">
        <tr><td style="padding:4px 14px 4px 0;"><strong>Supplier:</strong></td><td>{po.get('supplier_name', '')}</td></tr>
        <tr><td style="padding:4px 14px 4px 0;"><strong>Job:</strong></td><td>{po.get('job_name', '')}</td></tr>
        <tr><td style="padding:4px 14px 4px 0;"><strong>Required date:</strong></td><td>{required_date_display}</td></tr>
        <tr><td style="padding:4px 14px 4px 0;"><strong>Gross total:</strong></td><td><strong>£{float(po.get('gross_total') or 0):,.2f}</strong></td></tr>
      </table>
      <p>Please choose the next action:</p>
      <ul>
        <li><strong>Send email to supplier and assign materials</strong> - sends the approved PO to the supplier and commits the PO line items to the job materials.</li>
        <li><strong>Assign materials to job only</strong> - commits the PO line items to the job materials without emailing the supplier.</li>
      </ul>
    </div>
    """.strip()

    return {
        "subject": subject,
        "body_text": body_text,
        "body_html": body_html,
        "required_date_display": required_date_display,
    }


async def send_po_requester_options_notification(po: Dict[str, Any]) -> Dict[str, Any]:
    """After a PO is approved, ask the requester what should happen next."""
    requester_email = (po.get("requested_by_email") or "").strip()
    if not requester_email:
        return {"sent": False, "method": "not_configured", "to": "", "error": "No requester email is saved against this PO"}

    power_automate_url = os.environ.get("POWER_AUTOMATE_PO_REQUESTER_OPTIONS_URL", "").strip()
    power_automate_secret = os.environ.get("POWER_AUTOMATE_PO_REQUESTER_OPTIONS_SECRET", os.environ.get("PO_APPROVAL_CALLBACK_SECRET", "")).strip()
    if not power_automate_url:
        return {"sent": False, "method": "not_configured", "to": requester_email, "error": "POWER_AUTOMATE_PO_REQUESTER_OPTIONS_URL is not configured"}

    notification = build_po_requester_options_payload(po)
    po_number = po.get("po_number", "purchase_order")
    try:
        pdf_bytes = generate_purchase_order_pdf_bytes(po)
    except Exception as exc:
        logger.warning("Could not generate requester options PDF attachment for %s: %s", po_number, exc)
        pdf_bytes = b""

    public_backend_url = (
        os.environ.get("PUBLIC_BACKEND_URL")
        or os.environ.get("BACKEND_PUBLIC_URL")
        or os.environ.get("REACT_APP_BACKEND_URL")
        or ""
    ).strip().rstrip("/")
    callback_url = f"{public_backend_url}/api/purchase-orders/{po.get('id')}/requester-action-response" if public_backend_url and po.get("id") else ""

    payload = {
        "secret": power_automate_secret,
        "callback_secret": power_automate_secret,
        "requester_action_callback_url": callback_url,
        "notification_type": "purchase_order_requester_action_required",
        "to": requester_email,
        "po_id": po.get("id"),
        "po_number": po_number,
        "supplier_id": po.get("supplier_id", ""),
        "supplier_name": po.get("supplier_name", ""),
        "supplier_email": po.get("supplier_email", ""),
        "job_id": po.get("job_id", ""),
        "job_name": po.get("job_name", ""),
        "required_date": notification["required_date_display"],
        "gross_total": po.get("gross_total", 0),
        "requested_by_name": po.get("requested_by_name", ""),
        "requested_by_email": requester_email,
        "subject": notification["subject"],
        "body_text": notification["body_text"],
        "body_html": notification["body_html"],
        "pdf_filename": f"{po_number}.pdf",
        "pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8") if pdf_bytes else "",
        "options": "Send email to supplier and assign materials, Assign materials to job only",
    }

    try:
        response = requests.post(power_automate_url, json=payload, timeout=30)
    except Exception as exc:
        logger.exception("Failed to call requester options Power Automate flow: %s", exc)
        raise RuntimeError(f"Failed to call requester options flow: {exc}")

    if response.status_code < 200 or response.status_code >= 300:
        raise RuntimeError(f"Requester options Power Automate flow failed: {response.status_code} - {response.text[:500]}")

    return {"sent": True, "method": "power_automate", "to": requester_email}


@api_router.delete("/purchase-orders/bulk-delete")
async def bulk_delete_purchase_orders(request: PurchaseOrderBulkDeleteRequest, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    selected_ids = parse_po_id_list(po_ids=request.po_ids)
    if not selected_ids:
        raise HTTPException(status_code=400, detail="No purchase orders selected")
    result = await db.purchase_orders.delete_many({"id": {"$in": selected_ids}})
    return {
        "message": f"Deleted {result.deleted_count} purchase order(s)",
        "deleted_count": result.deleted_count,
        "requested_count": len(selected_ids),
        "deleted_by": super_admin.get("name") or super_admin.get("email") or "Super Admin",
    }


@api_router.get("/purchase-orders/{po_id}")
async def get_purchase_order(po_id: str, admin: str = Depends(verify_admin)):
    po = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    return po


@api_router.post("/purchase-orders", response_model=PurchaseOrder)
async def create_purchase_order(po_data: PurchaseOrderCreate, admin: str = Depends(verify_admin)):
    po_dict = po_data.dict()
    supplier = await db.suppliers.find_one({"id": po_data.supplier_id}) or {}
    job = await db.jobs.find_one({"id": po_data.job_id}) or {}
    if not supplier:
        raise HTTPException(status_code=404, detail="Supplier not found")
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    requested_po_number = str(po_dict.get("po_number") or "").strip()
    if requested_po_number:
        duplicate_po = await db.purchase_orders.find_one({"po_number": requested_po_number})
        if duplicate_po:
            raise HTTPException(status_code=400, detail="A purchase order with this PO number already exists")
        po_dict["po_number"] = requested_po_number
    else:
        po_dict["po_number"] = await next_po_number()
    po_dict["supplier_name"] = po_dict.get("supplier_name") or supplier.get("name", "")
    po_dict["supplier_email"] = po_dict.get("supplier_email") or supplier.get("orders_email") or supplier.get("accounts_email") or ""
    po_dict["job_name"] = po_dict.get("job_name") or job.get("display_name") or job.get("name", "")
    po_dict["job_number"] = po_dict.get("job_number") or job.get("job_number")
    po_dict["division"] = po_dict.get("division") or job.get("division", "")
    po_dict["delivery_address"] = po_dict.get("delivery_address") or job.get("location", "")
    po_dict["requested_by_user_id"] = po_dict.get("requested_by_user_id") or admin
    po_dict["requested_by_name"] = po_dict.get("requested_by_name") or admin
    po_dict["requested_by_email"] = po_dict.get("requested_by_email") or (admin if "@" in str(admin) else "")
    po_dict["requested_by_role"] = po_dict.get("requested_by_role") or ""
    po_dict.update(_normalise_po_planning_fields(po_dict, supplier))
    po_dict["status"] = po_dict.get("status") or "pending_approval"
    po_dict = calculate_po_totals(po_dict)

    po_obj = PurchaseOrder(**po_dict)
    po_doc = po_obj.dict()
    await db.purchase_orders.insert_one(po_doc)

    try:
        notification_result = await send_po_approval_notification(po_doc, requested_by=admin)
        await db.purchase_orders.update_one(
            {"id": po_obj.id},
            {"$set": {
                "approval_notification_sent": bool(notification_result.get("sent")),
                "approval_notification_method": notification_result.get("method", "unknown"),
                "approval_notification_to": notification_result.get("to", "info@ldagroup.co.uk"),
                "approval_notification_at": datetime.utcnow() if notification_result.get("sent") else None,
                "approval_notification_error": "" if notification_result.get("sent") else "Approval notification email is not configured",
            }}
        )
    except Exception as exc:
        logger.exception("PO approval notification failed for %s: %s", po_obj.po_number, exc)
        await db.purchase_orders.update_one(
            {"id": po_obj.id},
            {"$set": {
                "approval_notification_sent": False,
                "approval_notification_method": "failed",
                "approval_notification_to": os.environ.get("PO_APPROVAL_NOTIFY_EMAIL", "info@ldagroup.co.uk"),
                "approval_notification_error": str(exc)[:1000],
            }}
        )

    return po_obj


@api_router.put("/purchase-orders/{po_id}", response_model=PurchaseOrder)
async def update_purchase_order(po_id: str, po_update: PurchaseOrderUpdate, admin: str = Depends(verify_admin)):
    existing = await db.purchase_orders.find_one({"id": po_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    update_dict = {k: v for k, v in po_update.dict().items() if v is not None}
    if "lines" in update_dict:
        temp = calculate_po_totals({"lines": update_dict["lines"]})
        update_dict["lines"] = temp["lines"]
        update_dict["net_total"] = temp["net_total"]
        update_dict["vat_total"] = temp["vat_total"]
        update_dict["gross_total"] = temp["gross_total"]

    planning_keys = {"supplier_id", "required_date", "payment_requirement", "payment_terms_days", "lead_time_days", "payment_due_date", "order_by_date", "expected_payment_date"}
    if planning_keys.intersection(update_dict.keys()):
        supplier_id = update_dict.get("supplier_id") or existing.get("supplier_id")
        supplier = await db.suppliers.find_one({"id": supplier_id}) if supplier_id else {}
        merged_po = {**existing, **update_dict}
        update_dict.update(_normalise_po_planning_fields(merged_po, supplier or {}))

    update_dict["updated_at"] = datetime.utcnow()
    result = await db.purchase_orders.update_one({"id": po_id}, {"$set": update_dict})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    updated = await db.purchase_orders.find_one({"id": po_id})
    return PurchaseOrder(**updated)


@api_router.delete("/purchase-orders/{po_id}")
async def delete_purchase_order(po_id: str, admin: str = Depends(verify_admin)):
    result = await db.purchase_orders.delete_one({"id": po_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    return {"message": "Purchase order deleted successfully"}


@api_router.post("/purchase-orders/{po_id}/approval-response")
async def update_purchase_order_from_approval_response(po_id: str, response: PurchaseOrderApprovalResponseRequest):
    """Update a PO after Power Automate Send email with options returns Approve/Reject.

    This endpoint is intentionally protected by a shared secret instead of Basic auth,
    because it is called server-to-server by Power Automate after the approver clicks an option.
    """
    expected_secret = (
        os.environ.get("PO_APPROVAL_CALLBACK_SECRET")
        or os.environ.get("POWER_AUTOMATE_PO_APPROVAL_SECRET")
        or ""
    ).strip()

    if expected_secret and not secrets.compare_digest(str(response.secret or ""), expected_secret):
        raise HTTPException(status_code=403, detail="Invalid approval callback secret")

    decision_raw = (response.decision or response.selected_option or "").strip().lower()
    decision_raw = decision_raw.replace(" ", "_").replace("-", "_")

    if decision_raw in ["approve", "approved", "yes", "accept", "accepted"]:
        new_status = "approved"
    elif decision_raw in ["reject", "rejected", "no", "decline", "declined"]:
        new_status = "rejected"
    else:
        raise HTTPException(status_code=400, detail="Approval decision must be Approve or Reject")

    po = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")

    now = datetime.utcnow()
    responder_name = response.responder_name or response.responder_email or "Power Automate approval"
    comments = response.comments or response.reason or ""

    update = {
        "status": new_status,
        "approval_response": "approved" if new_status == "approved" else "rejected",
        "approval_response_at": now,
        "approval_response_by": responder_name,
        "approval_response_email": response.responder_email or "",
        "approval_response_comments": comments,
        "updated_at": now,
    }

    if new_status == "approved":
        update.update({
            "approved_by_user_id": response.responder_email or "power_automate",
            "approved_by_name": responder_name,
            "approved_at": now,
        })
    else:
        update.update({
            "rejected_by_user_id": response.responder_email or "power_automate",
            "rejected_by_name": responder_name,
            "rejected_at": now,
            "rejection_reason": comments,
        })

    await db.purchase_orders.update_one({"id": po_id}, {"$set": update})
    updated = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})

    requester_notification_result = {"sent": False, "method": "not_required", "to": ""}
    if new_status == "approved":
        try:
            requester_notification_result = await send_po_requester_options_notification(updated)
            await db.purchase_orders.update_one(
                {"id": po_id},
                {"$set": {
                    "requester_options_notification_sent": bool(requester_notification_result.get("sent")),
                    "requester_options_notification_method": requester_notification_result.get("method", "unknown"),
                    "requester_options_notification_to": requester_notification_result.get("to", ""),
                    "requester_options_notification_at": datetime.utcnow() if requester_notification_result.get("sent") else None,
                    "requester_options_notification_error": requester_notification_result.get("error", "") if not requester_notification_result.get("sent") else "",
                }}
            )
        except Exception as exc:
            logger.exception("Requester options notification failed for PO %s: %s", po_id, exc)
            await db.purchase_orders.update_one(
                {"id": po_id},
                {"$set": {
                    "requester_options_notification_sent": False,
                    "requester_options_notification_method": "failed",
                    "requester_options_notification_to": updated.get("requested_by_email", ""),
                    "requester_options_notification_error": str(exc)[:1000],
                }}
            )
            requester_notification_result = {"sent": False, "method": "failed", "to": updated.get("requested_by_email", ""), "error": str(exc)}

    return {
        "success": True,
        "message": f"Purchase order {updated.get('po_number', po_id)} marked as {new_status}",
        "po_id": po_id,
        "po_number": updated.get("po_number"),
        "status": updated.get("status"),
        "requester_options_notification": requester_notification_result,
    }


@api_router.post("/purchase-orders/{po_id}/approve")
async def approve_purchase_order(po_id: str, admin: str = Depends(verify_admin)):
    update = {
        "status": "approved",
        "approved_by_user_id": admin,
        "approved_by_name": admin,
        "approved_at": datetime.utcnow(),
        "updated_at": datetime.utcnow(),
    }
    result = await db.purchase_orders.update_one({"id": po_id}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Purchase order not found")

    updated = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    requester_notification_result = {"sent": False, "method": "not_configured", "to": ""}
    try:
        requester_notification_result = await send_po_requester_options_notification(updated)
        await db.purchase_orders.update_one(
            {"id": po_id},
            {"$set": {
                "requester_options_notification_sent": bool(requester_notification_result.get("sent")),
                "requester_options_notification_method": requester_notification_result.get("method", "unknown"),
                "requester_options_notification_to": requester_notification_result.get("to", ""),
                "requester_options_notification_at": datetime.utcnow() if requester_notification_result.get("sent") else None,
                "requester_options_notification_error": requester_notification_result.get("error", "") if not requester_notification_result.get("sent") else "",
            }}
        )
    except Exception as exc:
        logger.exception("Requester options notification failed for PO %s: %s", po_id, exc)
        await db.purchase_orders.update_one(
            {"id": po_id},
            {"$set": {
                "requester_options_notification_sent": False,
                "requester_options_notification_method": "failed",
                "requester_options_notification_to": (updated or {}).get("requested_by_email", ""),
                "requester_options_notification_error": str(exc)[:1000],
            }}
        )

    return {"message": "Purchase order approved", "requester_options_notification": requester_notification_result}


@api_router.post("/purchase-orders/{po_id}/mark-sent")
async def mark_purchase_order_sent(po_id: str, admin: str = Depends(verify_admin)):
    result = await db.purchase_orders.update_one({"id": po_id}, {"$set": {"status": "sent", "sent_at": datetime.utcnow(), "sent_by_user_id": admin, "sent_by_name": admin, "updated_at": datetime.utcnow()}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    return {"message": "Purchase order marked as sent"}


@api_router.get("/purchase-orders/{po_id}/pdf")
async def download_purchase_order_pdf(po_id: str, admin: str = Depends(verify_admin)):
    po = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    pdf_bytes = generate_purchase_order_pdf_bytes(po)
    filename = f"{po.get('po_number', 'purchase_order')}.pdf"
    return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})



@api_router.post("/purchase-orders/{po_id}/requester-action-response")
async def update_purchase_order_from_requester_action_response(po_id: str, response: PurchaseOrderRequesterActionResponseRequest):
    """Handle the post-approval requester choice from Power Automate.

    Options:
    - Send email to supplier and assign materials
    - Assign materials to job only
    """
    expected_secret = (
        os.environ.get("POWER_AUTOMATE_PO_REQUESTER_OPTIONS_SECRET")
        or os.environ.get("PO_APPROVAL_CALLBACK_SECRET")
        or ""
    ).strip()

    if expected_secret and not secrets.compare_digest(str(response.secret or ""), expected_secret):
        raise HTTPException(status_code=403, detail="Invalid requester action callback secret")

    decision_raw = (response.decision or response.selected_option or "").strip().lower()
    decision_key = decision_raw.replace("&", "and").replace("+", "and")
    decision_key = re.sub(r"[^a-z0-9]+", "_", decision_key).strip("_")

    send_supplier = any(token in decision_key for token in [
        "send_email_to_supplier_and_assign_materials",
        "send_to_supplier_and_assign_materials",
        "send_email_supplier_assign_materials",
        "send_supplier",
    ])
    assign_only = any(token in decision_key for token in [
        "assign_materials_to_job_only",
        "assign_materials_only",
        "assign_only",
    ])

    if not send_supplier and not assign_only:
        raise HTTPException(status_code=400, detail="Requester action must be supplier email + materials, or assign materials only")

    po = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    if po.get("status") == "rejected":
        raise HTTPException(status_code=400, detail="Rejected purchase orders cannot be actioned")
    if po.get("status") not in ["approved", "sent", "materials_assigned", "received", "invoiced", "closed"]:
        raise HTTPException(status_code=400, detail="Purchase order must be approved before requester actions can be completed")

    responder_name = response.responder_name or response.responder_email or po.get("requested_by_name") or "Requester"
    responder_email = response.responder_email or po.get("requested_by_email") or ""

    email_result = None
    material_result = None

    if send_supplier:
        email_result = await send_purchase_order_email(po_id, admin=responder_email or responder_name)

    try:
        material_result = await assign_purchase_order_materials(po_id, admin=responder_email or responder_name)
    except HTTPException as exc:
        if exc.status_code == 400 and "already been assigned" in str(exc.detail).lower():
            material_result = {"message": "PO materials were already assigned to the job", "materials_created": 0, "already_assigned": True}
        else:
            raise

    await db.purchase_orders.update_one(
        {"id": po_id},
        {"$set": {
            "requester_action_response": "send_supplier_and_assign_materials" if send_supplier else "assign_materials_only",
            "requester_action_response_at": datetime.utcnow(),
            "requester_action_response_by": responder_name,
            "requester_action_response_email": responder_email,
            "requester_action_response_comments": response.comments or "",
            "updated_at": datetime.utcnow(),
        }}
    )

    updated = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    return {
        "success": True,
        "po_id": po_id,
        "po_number": updated.get("po_number"),
        "status": updated.get("status"),
        "decision": response.decision or response.selected_option,
        "email_result": email_result,
        "material_result": material_result,
    }



@api_router.post("/purchase-orders/{po_id}/send-email")
async def send_purchase_order_email(po_id: str, admin: str = Depends(verify_admin)):
    po = await db.purchase_orders.find_one({"id": po_id}, {"_id": 0})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")

    if po.get("status") in ["draft", "pending_approval", "rejected", "cancelled"]:
        raise HTTPException(status_code=400, detail="Purchase order must be approved before it can be emailed to the supplier")

    supplier_email = (po.get("supplier_email") or "").strip()
    if not supplier_email:
        raise HTTPException(status_code=400, detail="No supplier email address is saved against this purchase order")

    pdf_bytes = generate_purchase_order_pdf_bytes(po)
    po_number = po.get("po_number", "purchase_order")
    subject = f"Purchase Order {po_number} - LDA Group"
    filename = f"{po_number}.pdf"

    required_date_display = format_uk_date_only(po.get("required_date"))
    if required_date_display == "-":
        required_date_display = "TBC"

    body_text = f"""Hi,

Please find attached purchase order {po_number} for the following job:

Job: {po.get('job_name', '')}
Required date: {required_date_display}
Delivery address: {po.get('delivery_address') or 'TBC'}

Please confirm receipt and advise expected delivery date.

Kind regards,
LDA Group
"""

    body_html = f"""
    <p>Hi,</p>
    <p>Please find attached purchase order <strong>{po_number}</strong> for the following job:</p>
    <table style="border-collapse:collapse; font-family:Arial, sans-serif; font-size:14px;">
      <tr><td style="padding:4px 12px 4px 0;"><strong>Job:</strong></td><td>{po.get('job_name', '')}</td></tr>
      <tr><td style="padding:4px 12px 4px 0;"><strong>Required date:</strong></td><td>{required_date_display}</td></tr>
      <tr><td style="padding:4px 12px 4px 0;"><strong>Delivery address:</strong></td><td>{po.get('delivery_address') or 'TBC'}</td></tr>
    </table>
    <p>Please confirm receipt and advise expected delivery date.</p>
    <p>Kind regards,<br>LDA Group</p>
    """.strip()

    # Preferred route: Power Automate webhook over HTTPS.
    # This avoids Render free-tier SMTP port restrictions on 25/465/587.
    power_automate_url = os.environ.get("POWER_AUTOMATE_PO_EMAIL_URL", "").strip()
    power_automate_secret = os.environ.get("POWER_AUTOMATE_PO_EMAIL_SECRET", "").strip()

    if power_automate_url:
        payload = {
            "secret": power_automate_secret,
            "po_id": po.get("id"),
            "po_number": po_number,
            "supplier_name": po.get("supplier_name", ""),
            "supplier_email": supplier_email,
            "reply_to": os.environ.get("PO_REPLY_TO_EMAIL", os.environ.get("SMTP_FROM_EMAIL", "info@ldagroup.co.uk")).strip(),
            "subject": subject,
            "body_text": body_text,
            "body_html": body_html,
            "pdf_filename": filename,
            "pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8"),
            "job_id": po.get("job_id"),
            "job_name": po.get("job_name", ""),
            "required_date": required_date_display,
            "delivery_address": po.get("delivery_address") or "",
            "net_total": po.get("net_total", 0),
            "vat_total": po.get("vat_total", 0),
            "gross_total": po.get("gross_total", 0),
        }

        try:
            response = requests.post(power_automate_url, json=payload, timeout=60)
        except Exception as exc:
            logger.exception("Failed to call Power Automate PO email flow: %s", exc)
            raise HTTPException(status_code=500, detail=f"Failed to call Power Automate PO email flow: {exc}")

        if response.status_code < 200 or response.status_code >= 300:
            response_text = response.text[:1000] if response.text else "No response body"
            logger.error("Power Automate PO email flow failed: %s %s", response.status_code, response_text)
            raise HTTPException(status_code=500, detail=f"Power Automate PO email flow failed: {response.status_code} - {response_text}")

        await db.purchase_orders.update_one(
            {"id": po_id},
            {"$set": {
                "status": "sent",
                "sent_at": datetime.utcnow(),
                "sent_by_user_id": admin,
                "sent_by_name": admin,
                "sent_to": supplier_email,
                "email_subject": subject,
                "email_method": "power_automate",
                "updated_at": datetime.utcnow(),
            }}
        )
        return {"message": "Purchase order email sent via Power Automate", "sent_to": supplier_email, "method": "power_automate"}

    # Fallback route: SMTP. This is retained for paid Render instances or other hosts.
    smtp_host = os.environ.get("SMTP_HOST", "").strip()
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_username = os.environ.get("SMTP_USERNAME", "").strip()
    smtp_password = os.environ.get("SMTP_PASSWORD", "").strip()
    smtp_from = os.environ.get("SMTP_FROM_EMAIL", smtp_username).strip()
    smtp_from_name = os.environ.get("SMTP_FROM_NAME", "LDA Group").strip()
    smtp_reply_to = os.environ.get("SMTP_REPLY_TO", os.environ.get("PO_REPLY_TO_EMAIL", "")).strip()
    smtp_use_tls = os.environ.get("SMTP_USE_TLS", "true").lower() != "false"
    if not smtp_host or not smtp_from:
        raise HTTPException(status_code=500, detail="Email is not configured. Add POWER_AUTOMATE_PO_EMAIL_URL and POWER_AUTOMATE_PO_EMAIL_SECRET in Render, or configure SMTP settings.")

    message = EmailMessage()
    message["From"] = f"{smtp_from_name} <{smtp_from}>"
    message["To"] = supplier_email
    message["Subject"] = subject
    if smtp_reply_to:
        message["Reply-To"] = smtp_reply_to
    message.set_content(body_text)
    message.add_attachment(pdf_bytes, maintype="application", subtype="pdf", filename=filename)

    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as smtp:
            if smtp_use_tls:
                smtp.starttls()
            if smtp_username and smtp_password:
                smtp.login(smtp_username, smtp_password)
            smtp.send_message(message)
    except Exception as exc:
        logger.exception("Failed to send PO email: %s", exc)
        raise HTTPException(status_code=500, detail=f"Failed to send PO email: {exc}")

    await db.purchase_orders.update_one(
        {"id": po_id},
        {"$set": {
            "status": "sent",
            "sent_at": datetime.utcnow(),
            "sent_by_user_id": admin,
            "sent_by_name": admin,
            "sent_to": supplier_email,
            "email_subject": subject,
            "email_method": "smtp",
            "updated_at": datetime.utcnow(),
        }}
    )
    return {"message": "Purchase order email sent", "sent_to": supplier_email, "method": "smtp"}


@api_router.post("/purchase-orders/{po_id}/assign-materials")
async def assign_purchase_order_materials(po_id: str, admin: str = Depends(verify_admin)):
    po = await db.purchase_orders.find_one({"id": po_id})
    if not po:
        raise HTTPException(status_code=404, detail="Purchase order not found")
    if po.get("materials_assigned"):
        raise HTTPException(status_code=400, detail="Materials have already been assigned for this PO")
    materials_to_insert = []
    updated_lines = []
    for line in po.get("lines", []):
        material_id = str(uuid.uuid4())
        material = {
            "id": material_id,
            "job_id": po.get("job_id"),
            "name": line.get("description", "PO material"),
            "cost": float(line.get("unit_cost") or 0),
            "quantity": int(float(line.get("quantity") or 0)) if float(line.get("quantity") or 0).is_integer() else float(line.get("quantity") or 0),
            "supplier": po.get("supplier_name", ""),
            "reference": po.get("po_number", ""),
            "purchase_date": datetime.utcnow(),
            "notes": f"Assigned from purchase order {po.get('po_number', '')}",
            "created_date": datetime.utcnow(),
            "source_type": "purchase_order",
            "purchase_order_id": po_id,
            "purchase_order_line_id": line.get("id"),
            "job_section_id": line.get("job_section_id") or po.get("project_section_id") or "",
            "job_section_name": line.get("job_section_name") or po.get("project_section_name") or "",
            "status": "committed",
        }
        materials_to_insert.append(material)
        line["material_id"] = material_id
        line["material_status"] = "committed"
        updated_lines.append(line)
    if materials_to_insert:
        await db.materials.insert_many(materials_to_insert)
    await db.purchase_orders.update_one({"id": po_id}, {"$set": {"lines": updated_lines, "materials_assigned": True, "materials_assigned_at": datetime.utcnow(), "status": "materials_assigned", "updated_at": datetime.utcnow()}})
    return {"message": "PO materials assigned to job", "materials_created": len(materials_to_insert)}


@api_router.post("/purchase-orders/import-quote")
async def import_purchase_order_quote(
    file: UploadFile = File(...),
    job_id: Optional[str] = Query(None),
    admin: str = Depends(verify_admin),
):
    """Upload a supplier quote and return structured data for PO review.

    Patch 2 improves digital PDF/DOCX/TXT extraction, supplier matching, totals detection,
    and line-item parsing. Scanned image OCR is still deliberately left for a later patch.
    """
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Quote file is too large. Maximum size is 10MB.")

    filename = file.filename or "uploaded_quote"
    content_type = file.content_type or "application/octet-stream"
    extraction = await extract_text_from_upload(file, content)
    extracted_text = extraction.get("text", "")
    lines = parse_quote_lines_from_text(extracted_text) if extracted_text else []
    quote_number = extract_quote_number(extracted_text) if extracted_text else ""
    supplier_name = extract_supplier_name_from_text(extracted_text) if extracted_text else ""
    supplier_email = extract_email_from_text(extracted_text) if extracted_text else ""
    supplier_phone = extract_phone_from_text(extracted_text) if extracted_text else ""
    supplier_vat_number = extract_supplier_vat_number_from_text(extracted_text) if extracted_text else ""
    supplier_address = extract_supplier_address_from_text(extracted_text, supplier_name) if extracted_text else ""
    quote_date = extract_first_date_for_patterns(extracted_text, [r"quote\s*date", r"quotation\s*date", r"date"]) if extracted_text else ""
    expiry_date = extract_first_date_for_patterns(extracted_text, [r"valid\s*until", r"expiry\s*date", r"expires", r"quote\s*valid\s*until"]) if extracted_text else ""
    totals = extract_quote_totals(extracted_text) if extracted_text else {"net_total": None, "vat_total": None, "gross_total": None, "vat_inclusive": False}
    vat_inclusive = bool(totals.get("vat_inclusive"))
    supplier_match = await match_supplier_from_quote(supplier_name, supplier_email) if extracted_text else {"matched_supplier_id": None, "matched_supplier_name": "", "match_score": 0}

    warnings = []
    if not extracted_text:
        if extraction.get("ocr_attempted") and not extraction.get("ocr_configured"):
            warnings.append("No readable text could be extracted. This looks like a scanned PDF/image. Add OCR_SPACE_API_KEY in Render to enable OCR reading.")
        elif extraction.get("ocr_attempted"):
            warnings.append("OCR was attempted, but no readable text could be extracted. Please enter the PO details manually or try a clearer scan.")
        else:
            warnings.append("No readable text could be extracted. Please enter the PO details manually.")
    elif extraction.get("ocr_used"):
        warnings.append("OCR was used to read this quote. Please carefully review supplier details, quantities and prices before creating the PO.")
    if extracted_text and not lines:
        warnings.append("Text was extracted, but line items could not be confidently detected. Please enter the PO lines manually.")
    if extracted_text and not supplier_match.get("matched_supplier_id"):
        if supplier_name:
            warnings.append("Supplier was detected but not matched to an existing supplier. Check or create the supplier before making the PO.")
        else:
            warnings.append("Supplier name could not be confidently detected.")
    if extracted_text and not quote_number:
        warnings.append("Quote reference could not be confidently detected.")
    if extracted_text and vat_inclusive:
        warnings.append("Quote appears to include VAT already. The importer has converted line amounts back to net values so VAT is not added twice.")

    line_net_total = round(sum(float(line.get("net_total") if line.get("net_total") is not None else (float(line.get("quantity") or 0) * float(line.get("unit_cost") or 0))) for line in lines), 2)
    line_gross_total = round(sum(float(line.get("gross_total") if line.get("gross_total") is not None else ((float(line.get("quantity") or 0) * float(line.get("unit_cost") or 0)) * (1 + (float(line.get("vat_rate") or 0) / 100)))) for line in lines), 2)
    extracted_net = totals.get("net_total")
    extracted_gross = totals.get("gross_total")
    totals_match = None
    if lines and extracted_gross is not None and vat_inclusive:
        totals_match = abs(line_gross_total - extracted_gross) <= max(1.0, extracted_gross * 0.03)
        if not totals_match:
            warnings.append("Extracted line gross total does not match the quote gross total. Review quantities and prices before creating the PO.")
    elif lines and extracted_net is not None:
        totals_match = abs(line_net_total - extracted_net) <= max(1.0, extracted_net * 0.03)
        if not totals_match:
            warnings.append("Extracted line total does not match the quote net total. Review quantities and prices before creating the PO.")

    confidence_score = 0
    if extracted_text:
        confidence_score += 20
    if lines:
        confidence_score += 30
    if supplier_match.get("matched_supplier_id"):
        confidence_score += 20
    elif supplier_name:
        confidence_score += 10
    if quote_number:
        confidence_score += 10
    if extracted_net is not None or extracted_gross is not None:
        confidence_score += 15
    if totals_match is True:
        confidence_score += 5

    if confidence_score >= 75:
        confidence = "high"
    elif confidence_score >= 45:
        confidence = "medium"
    else:
        confidence = "low"

    upload_id = str(uuid.uuid4())
    upload_doc = {
        "id": upload_id,
        "filename": filename,
        "content_type": content_type,
        "size_bytes": len(content),
        "job_id": job_id,
        "uploaded_by": admin,
        "uploaded_at": datetime.utcnow(),
        "extracted_text": extracted_text[:40000],
        "quote_number": quote_number,
        "quote_date": quote_date,
        "expiry_date": expiry_date,
        "supplier_name": supplier_name,
        "supplier_email": supplier_email,
        "supplier_phone": supplier_phone,
        "supplier_vat_number": supplier_vat_number,
        "supplier_address": supplier_address,
        "extraction_method": extraction.get("method", "none"),
        "ocr_attempted": extraction.get("ocr_attempted", False),
        "ocr_used": extraction.get("ocr_used", False),
        "ocr_configured": extraction.get("ocr_configured", False),
        "totals": totals,
        "vat_inclusive": vat_inclusive,
        "lines": lines,
        "confidence": confidence,
        "confidence_score": confidence_score,
        "warnings": warnings,
    }
    # Store content only for small files to avoid hitting MongoDB document limits.
    if len(content) <= 2 * 1024 * 1024:
        upload_doc["content_base64"] = base64.b64encode(content).decode("utf-8")
    await db.purchase_order_quote_uploads.insert_one(upload_doc)

    return {
        "upload_id": upload_id,
        "filename": filename,
        "content_type": content_type,
        "quote_number": quote_number,
        "quote_date": quote_date,
        "expiry_date": expiry_date,
        "supplier_name": supplier_name,
        "supplier_email": supplier_email,
        "supplier_phone": supplier_phone,
        "supplier_vat_number": supplier_vat_number,
        "supplier_address": supplier_address,
        "extraction_method": extraction.get("method", "none"),
        "ocr_attempted": extraction.get("ocr_attempted", False),
        "ocr_used": extraction.get("ocr_used", False),
        "ocr_configured": extraction.get("ocr_configured", False),
        "matched_supplier_id": supplier_match.get("matched_supplier_id"),
        "matched_supplier_name": supplier_match.get("matched_supplier_name"),
        "supplier_match_score": supplier_match.get("match_score", 0),
        "lines": lines,
        "line_net_total": line_net_total,
        "quote_net_total": totals.get("net_total"),
        "quote_vat_total": totals.get("vat_total"),
        "quote_gross_total": totals.get("gross_total"),
        "vat_inclusive": vat_inclusive,
        "line_gross_total": line_gross_total,
        "totals_match": totals_match,
        "confidence": confidence,
        "confidence_score": confidence_score,
        "warnings": warnings,
        "warning": " ".join(warnings) if warnings else "Quote imported. Please review extracted details before creating the PO.",
        "extracted_text_preview": extracted_text[:1800] if extracted_text else "",
        "extracted_text_length": len(extracted_text or ""),
    }


@api_router.get("/jobs/{job_id}/purchase-orders")
async def get_job_purchase_orders(job_id: str, admin: str = Depends(verify_admin)):
    purchase_orders = await db.purchase_orders.find({"job_id": job_id}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    committed_statuses = {"approved", "sent", "materials_assigned", "part_received"}
    committed_value = sum(po.get("net_total", 0) for po in purchase_orders if po.get("status") in committed_statuses)
    actual_value = sum(po.get("net_total", 0) for po in purchase_orders if po.get("status") in {"received", "invoiced", "closed"})
    return {"purchase_orders": purchase_orders, "committed_value": committed_value, "actual_value": actual_value}



# ==================== FINANCE MATERIAL SPEND DASHBOARD ENDPOINTS ====================
# These routes feed Finance > Material Spend. They combine actual material entries,
# purchase order commitments and forecast material allowances from Gantt sections.

MATERIAL_PO_CASH_OUT_STATUSES = {
    "draft",
    "raised",
    "sent",
    "sent_to_supplier",
    "awaiting_delivery",
    "part_delivered",
    "delivered",
    "awaiting_invoice",
    "invoiced",
    "part_paid",
    "disputed",
}

MATERIAL_PO_PAID_STATUSES = {"paid"}
MATERIAL_PO_CANCELLED_STATUSES = {"cancelled", "void", "archived"}


def finance_material_parse_date(value: Any) -> Optional[date]:
    if not value:
        return None
    if isinstance(value, datetime):
        return value.date()
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00")[:10]).date()
    except Exception:
        return None


def finance_material_iso(value: Any) -> str:
    parsed = finance_material_parse_date(value)
    return parsed.isoformat() if parsed else ""


def finance_material_date_in_range(value: Any, start_date: Optional[str], end_date: Optional[str]) -> bool:
    parsed = finance_material_parse_date(value)
    if not parsed:
        return False
    start = finance_material_parse_date(start_date) if start_date else None
    end = finance_material_parse_date(end_date) if end_date else None
    if start and parsed < start:
        return False
    if end and parsed > end:
        return False
    return True


def finance_material_to_float(value: Any, fallback: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return fallback
        return float(value)
    except (TypeError, ValueError):
        return fallback


def finance_material_money(value: Any) -> float:
    return round(finance_material_to_float(value), 2)


def normalise_po_status(value: Any) -> str:
    return str(value or "draft").strip().lower().replace(" ", "_").replace("-", "_")


def material_receipt_present(material: Dict[str, Any]) -> bool:
    receipt_keys = [
        "receipt_url",
        "receipt_file_url",
        "receipt_file_id",
        "receipt_image",
        "receipt_photo",
        "receipt_attachment",
        "attachment_url",
        "file_url",
    ]
    if any(material.get(key) for key in receipt_keys):
        return True
    reference = str(material.get("reference") or material.get("receipt_number") or "").strip()
    return bool(reference)


def get_material_allowance_for_job(job: Dict[str, Any]) -> float:
    sections = job.get("gantt_sections") or []
    material_total = 0.0
    for section in sections:
        material_total += finance_material_to_float(section.get("material_value"))
    if material_total <= 0:
        material_total = finance_material_to_float(job.get("material_allowance") or job.get("materials_allowance"))
    return round(material_total, 2)


def get_section_forecast_material_value(section: Dict[str, Any]) -> float:
    material_value = finance_material_to_float(section.get("material_value"))
    if material_value > 0:
        return round(material_value, 2)
    # Fallback for older sections where only a total section value exists.
    # We deliberately keep this fallback conservative so it does not inflate material forecast.
    return 0.0


async def build_material_spend_dashboard_data(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    job_id: Optional[str] = None,
    supplier: Optional[str] = None,
    status: Optional[str] = None,
    spend_type: Optional[str] = None,
) -> Dict[str, Any]:
    """Build Finance > Material Spend without double-counting forecast allowances.

    Forecast material values from Gantt sections are allowances. When a PO or an
    actual material entry is linked to the same Gantt section, that value should
    reduce the remaining forecast rather than sit on top of it.

    Example:
      section forecast material = £567.00
      linked PO net value        = £34.81
      remaining forecast         = £532.19
      total exposure             = £567.00, not £601.81
    """
    today = datetime.utcnow().date()
    if not start_date:
        start_date = today.isoformat()
    if not end_date:
        end_date = (today + timedelta(days=27)).isoformat()

    def safe_text(value: Any) -> str:
        return str(value or "").strip()

    def section_key(job_value: Any, section_id: Any = "", section_name: Any = "") -> Optional[str]:
        job_value = safe_text(job_value)
        section_id = safe_text(section_id)
        section_name = safe_text(section_name).lower()
        if not job_value:
            return None
        if section_id:
            return f"{job_value}::id::{section_id}"
        if section_name:
            return f"{job_value}::name::{section_name}"
        return None

    def section_keys_for_section(job_value: Any, section: Dict[str, Any]) -> List[str]:
        keys = []
        by_id = section_key(job_value, section.get("id"), "")
        by_name = section_key(job_value, "", section.get("name"))
        if by_id:
            keys.append(by_id)
        if by_name and by_name not in keys:
            keys.append(by_name)
        return keys

    def add_deduction(target: Dict[str, float], keys: List[str], amount: float) -> None:
        amount = round(max(0.0, finance_material_to_float(amount)), 2)
        if amount <= 0:
            return
        for key in keys:
            if key:
                target[key] = round(target.get(key, 0.0) + amount, 2)

    def get_line_net(line: Dict[str, Any]) -> float:
        if line.get("source_line_net_total") not in [None, ""]:
            return finance_material_money(line.get("source_line_net_total"))
        if line.get("net_total") not in [None, ""] and finance_material_to_float(line.get("net_total")) > 0:
            return finance_material_money(line.get("net_total"))
        return finance_material_money(finance_material_to_float(line.get("quantity"), 1.0) * finance_material_to_float(line.get("unit_cost")))

    def get_line_vat(line: Dict[str, Any]) -> float:
        if line.get("source_line_vat_total") not in [None, ""]:
            return finance_material_money(line.get("source_line_vat_total"))
        if line.get("vat_total") not in [None, ""] and finance_material_to_float(line.get("vat_total")) > 0:
            return finance_material_money(line.get("vat_total"))
        return finance_material_money(get_line_net(line) * (finance_material_to_float(line.get("vat_rate")) / 100.0))

    def get_line_gross(line: Dict[str, Any]) -> float:
        if line.get("source_line_gross_total") not in [None, ""]:
            return finance_material_money(line.get("source_line_gross_total"))
        if line.get("gross_total") not in [None, ""] and finance_material_to_float(line.get("gross_total")) > 0:
            return finance_material_money(line.get("gross_total"))
        return finance_material_money(get_line_net(line) + get_line_vat(line))

    jobs = await db.jobs.find({"archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    job_lookup = {job.get("id"): job for job in jobs if job.get("id")}

    rows: List[Dict[str, Any]] = []
    actual_net_by_section: Dict[str, float] = {}
    po_net_by_section: Dict[str, float] = {}

    # 1) Actual material entries.
    # Materials created from a PO assignment are skipped here because the PO row
    # already represents that commitment. This avoids PO + assigned material double-counting.
    material_query: Dict[str, Any] = {"archived": {"$ne": True}}
    if job_id:
        material_query["job_id"] = job_id
    if supplier:
        material_query["supplier"] = {"$regex": supplier, "$options": "i"}

    materials = await db.materials.find(material_query, {"_id": 0}).to_list(10000)
    for material in materials:
        if material.get("purchase_order_id") or str(material.get("source_type") or "").lower() == "purchase_order":
            continue

        job = job_lookup.get(material.get("job_id"), {})
        material_date = finance_material_iso(material.get("purchase_date") or material.get("date") or material.get("created_date"))
        net = finance_material_money(material.get("net_total"))
        if net <= 0:
            net = finance_material_money(finance_material_to_float(material.get("cost")) * finance_material_to_float(material.get("quantity"), 1.0))
        vat = finance_material_money(material.get("vat_total") or material.get("vat"))
        gross = finance_material_money(material.get("gross_total") or (net + vat))
        row_status = str(material.get("status") or "approved").strip().lower().replace(" ", "_")
        receipt_ok = material_receipt_present(material)

        section_keys = []
        material_section_id = material.get("job_section_id") or material.get("project_section_id") or material.get("section_id")
        material_section_name = material.get("job_section_name") or material.get("project_section_name") or material.get("section_name")
        key = section_key(material.get("job_id"), material_section_id, material_section_name)
        if key:
            section_keys.append(key)
            # Add a name fallback too where available, so older rows still match if an ID changed.
            name_key = section_key(material.get("job_id"), "", material_section_name)
            if name_key and name_key not in section_keys:
                section_keys.append(name_key)
            add_deduction(actual_net_by_section, section_keys, net)

        rows.append({
            "id": material.get("id") or str(uuid.uuid4()),
            "date": material_date,
            "type": "actual",
            "type_label": "Actual",
            "job_id": material.get("job_id", ""),
            "job_name": job.get("name") or job.get("display_name") or material.get("job_name") or "Unknown job",
            "job_client": job.get("client", ""),
            "supplier": material.get("supplier") or "Unknown supplier",
            "description": material.get("name") or material.get("description") or "Material entry",
            "net": net,
            "vat": vat,
            "gross": gross,
            "status": row_status,
            "due_date": "",
            "paid_date": material_date,
            "source": "Material Entry",
            "source_id": material.get("id"),
            "job_section_id": safe_text(material_section_id),
            "job_section_name": safe_text(material_section_name),
            "receipt": receipt_ok,
            "receipt_reference": material.get("reference") or material.get("receipt_number") or "",
            "notes": material.get("notes", ""),
            "exposure_net": net,
            "cash_out_gross": gross,
        })

    # 2) Purchase orders / committed spend.
    po_query: Dict[str, Any] = {}
    if job_id:
        po_query["job_id"] = job_id
    if supplier:
        po_query["supplier_name"] = {"$regex": supplier, "$options": "i"}
    purchase_orders = await db.purchase_orders.find(po_query, {"_id": 0}).to_list(10000)

    for po in purchase_orders:
        po_status = normalise_po_status(po.get("status"))
        if po_status in MATERIAL_PO_CANCELLED_STATUSES:
            continue
        po_date = finance_material_iso(po.get("created_at") or po.get("approved_at") or po.get("sent_at"))
        due_date = finance_material_iso(_calculate_po_expected_payment_date(po) or po.get("expected_payment_date") or po.get("required_date") or po.get("created_at"))
        paid_date = finance_material_iso(po.get("paid_date") or po.get("payment_date"))
        net = finance_material_money(po.get("net_total"))
        vat = finance_material_money(po.get("vat_total"))
        gross = finance_material_money(po.get("gross_total") or (net + vat))
        row_type = "po_paid" if po_status in MATERIAL_PO_PAID_STATUSES else "po"
        source_label = po.get("po_number") or "Purchase Order"

        # Deduct linked PO line values from their Gantt section forecast.
        # This is the important double-counting fix.
        linked_line_total = 0.0
        lines = po.get("lines") or []
        for line in lines:
            line_net = get_line_net(line)
            if line_net <= 0:
                continue
            line_section_id = line.get("job_section_id") or line.get("project_section_id") or po.get("project_section_id")
            line_section_name = line.get("job_section_name") or line.get("project_section_name") or po.get("project_section_name")
            keys = []
            by_id = section_key(po.get("job_id"), line_section_id, "")
            by_name = section_key(po.get("job_id"), "", line_section_name)
            if by_id:
                keys.append(by_id)
            if by_name and by_name not in keys:
                keys.append(by_name)
            if keys:
                add_deduction(po_net_by_section, keys, line_net)
                linked_line_total += line_net

        # If the PO has a top-level project section but no line-level links, deduct the whole PO.
        if linked_line_total <= 0 and (po.get("project_section_id") or po.get("project_section_name")):
            keys = []
            by_id = section_key(po.get("job_id"), po.get("project_section_id"), "")
            by_name = section_key(po.get("job_id"), "", po.get("project_section_name"))
            if by_id:
                keys.append(by_id)
            if by_name and by_name not in keys:
                keys.append(by_name)
            add_deduction(po_net_by_section, keys, net)

        rows.append({
            "id": po.get("id") or str(uuid.uuid4()),
            "date": po_date,
            "type": row_type,
            "type_label": "PO Paid" if row_type == "po_paid" else "PO",
            "job_id": po.get("job_id", ""),
            "job_name": po.get("job_name") or (job_lookup.get(po.get("job_id"), {}) or {}).get("name") or "Unknown job",
            "job_client": (job_lookup.get(po.get("job_id"), {}) or {}).get("client", ""),
            "supplier": po.get("supplier_name") or "Unknown supplier",
            "description": po.get("supplier_quote_number") or po.get("notes") or source_label,
            "net": net,
            "vat": vat,
            "gross": gross,
            "status": po_status,
            "due_date": due_date,
            "paid_date": paid_date,
            "source": source_label,
            "source_id": po.get("id"),
            "job_section_id": safe_text(po.get("project_section_id")),
            "job_section_name": safe_text(po.get("project_section_name")),
            "receipt": bool(po.get("supplier_quote_number") or po.get("source_file_name")),
            "receipt_reference": po.get("supplier_quote_number") or po.get("source_file_name") or "",
            "notes": po.get("notes", ""),
            "exposure_net": net,
            "cash_out_gross": gross,
        })

    # 3) Remaining forecast material spend from Gantt section material allowances.
    # This is now a REMAINING forecast, not the original forecast plus POs.
    for job in jobs:
        if job_id and job.get("id") != job_id:
            continue
        for section in job.get("gantt_sections") or []:
            section_material = get_section_forecast_material_value(section)
            if section_material <= 0:
                continue
            forecast_date = finance_material_iso(section.get("start_date") or section.get("end_date") or job.get("planned_start_date"))
            if not forecast_date:
                continue

            keys = section_keys_for_section(job.get("id"), section)
            actual_deducted = round(max(actual_net_by_section.get(key, 0.0) for key in keys) if keys else 0.0, 2)
            po_deducted = round(max(po_net_by_section.get(key, 0.0) for key in keys) if keys else 0.0, 2)
            remaining_forecast = round(max(0.0, section_material - actual_deducted - po_deducted), 2)
            covered_value = round(min(section_material, actual_deducted + po_deducted), 2)

            if remaining_forecast <= 0:
                continue

            forecast_notes = [
                f"Original material allowance: £{section_material:,.2f}",
            ]
            if po_deducted > 0:
                forecast_notes.append(f"Less linked POs: £{po_deducted:,.2f}")
            if actual_deducted > 0:
                forecast_notes.append(f"Less linked actual spend: £{actual_deducted:,.2f}")

            rows.append({
                "id": f"forecast-{job.get('id')}-{section.get('id') or section.get('name')}",
                "date": forecast_date,
                "type": "forecast",
                "type_label": "Remaining Forecast",
                "job_id": job.get("id", ""),
                "job_name": job.get("name") or job.get("display_name") or "Unknown job",
                "job_client": job.get("client", ""),
                "supplier": "Various",
                "description": f"{section.get('name') or 'Gantt section'} remaining material forecast",
                "net": remaining_forecast,
                "vat": 0.0,
                "gross": remaining_forecast,
                "status": "forecast",
                "due_date": forecast_date,
                "paid_date": "",
                "source": "Forecast",
                "source_id": section.get("id"),
                "job_section_id": safe_text(section.get("id")),
                "job_section_name": safe_text(section.get("name")),
                "receipt": None,
                "receipt_reference": "",
                "notes": " | ".join(forecast_notes),
                "forecast_original_net": section_material,
                "forecast_deducted_po_net": po_deducted,
                "forecast_deducted_actual_net": actual_deducted,
                "forecast_covered_net": covered_value,
                "exposure_net": remaining_forecast,
                "cash_out_gross": remaining_forecast,
            })

    # Apply display filters after all sources are normalised.
    filtered_rows = []
    for row in rows:
        row_date_for_range = row.get("due_date") or row.get("date") or row.get("paid_date")
        if not finance_material_date_in_range(row_date_for_range, start_date, end_date):
            continue
        if spend_type and spend_type != "all":
            if spend_type == "po" and row.get("type") not in ["po", "po_paid"]:
                continue
            if spend_type != "po" and row.get("type") != spend_type:
                continue
        if status and status != "all" and row.get("status") != status:
            continue
        if supplier and supplier.lower() not in str(row.get("supplier", "")).lower():
            continue
        filtered_rows.append(row)

    filtered_rows.sort(key=lambda item: (item.get("due_date") or item.get("date") or "", item.get("job_name") or ""))

    # Summaries use sensible time windows independent of the active filter where needed.
    uk_now = get_uk_time()
    month_start = uk_now.replace(day=1).date()
    month_end = (month_start.replace(day=28) + timedelta(days=4)).replace(day=1) - timedelta(days=1)
    month_start_iso = month_start.isoformat()
    month_end_iso = month_end.isoformat()

    actual_this_month = sum(row.get("exposure_net", row.get("net", 0.0)) for row in rows if row.get("type") == "actual" and finance_material_date_in_range(row.get("date"), month_start_iso, month_end_iso))
    po_this_month = sum(row.get("exposure_net", row.get("net", 0.0)) for row in rows if row.get("type") in ["po", "po_paid"] and finance_material_date_in_range(row.get("date"), month_start_iso, month_end_iso))
    due_next_4_weeks = sum(row.get("cash_out_gross", row.get("gross", 0.0)) for row in rows if row.get("type") == "po" and finance_material_date_in_range(row.get("due_date"), today.isoformat(), (today + timedelta(days=27)).isoformat()))
    unreceipted_spend = sum(row.get("exposure_net", row.get("net", 0.0)) for row in rows if row.get("type") == "actual" and row.get("receipt") is False)
    forecast_next_4_weeks = sum(row.get("exposure_net", row.get("net", 0.0)) for row in rows if row.get("type") == "forecast" and finance_material_date_in_range(row.get("date"), today.isoformat(), (today + timedelta(days=27)).isoformat()))

    job_summary: Dict[str, Dict[str, Any]] = {}
    for job in jobs:
        if job_id and job.get("id") != job_id:
            continue
        allowance = get_material_allowance_for_job(job)
        job_summary[job.get("id")] = {
            "job_id": job.get("id"),
            "job_name": job.get("name") or job.get("display_name") or "Unknown job",
            "client": job.get("client", ""),
            "allowance": allowance,
            "actual_spend": 0.0,
            "po_commitments": 0.0,
            "forecast_spend": 0.0,
            "total_committed": 0.0,
            "variance": allowance,
            "percent_used": 0.0,
        }

    for row in rows:
        jid = row.get("job_id")
        if not jid or jid not in job_summary:
            continue
        exposure_value = row.get("exposure_net", row.get("net", 0.0))
        if row.get("type") == "actual":
            job_summary[jid]["actual_spend"] += exposure_value
        elif row.get("type") == "po":
            job_summary[jid]["po_commitments"] += exposure_value
        elif row.get("type") == "forecast":
            job_summary[jid]["forecast_spend"] += exposure_value

    jobs_over_allowance = 0
    for item in job_summary.values():
        item["actual_spend"] = round(item["actual_spend"], 2)
        item["po_commitments"] = round(item["po_commitments"], 2)
        item["forecast_spend"] = round(item["forecast_spend"], 2)
        item["total_committed"] = round(item["actual_spend"] + item["po_commitments"], 2)
        item["total_exposure"] = round(item["actual_spend"] + item["po_commitments"] + item["forecast_spend"], 2)
        item["variance"] = round(item["allowance"] - item["total_committed"], 2)
        item["forecast_variance"] = round(item["allowance"] - item["total_exposure"], 2)
        item["percent_used"] = round((item["total_committed"] / item["allowance"] * 100.0), 1) if item["allowance"] > 0 else 0.0
        item["percent_exposed"] = round((item["total_exposure"] / item["allowance"] * 100.0), 1) if item["allowance"] > 0 else 0.0
        if item["allowance"] > 0 and item["total_committed"] > item["allowance"]:
            jobs_over_allowance += 1

    supplier_summary: Dict[str, float] = {}
    for row in rows:
        if row.get("type") in ["actual", "po", "po_paid"]:
            supplier_summary[row.get("supplier") or "Unknown supplier"] = supplier_summary.get(row.get("supplier") or "Unknown supplier", 0.0) + row.get("cash_out_gross", row.get("gross", 0.0))
    top_suppliers = [
        {"supplier": supplier_name, "gross": round(total, 2)}
        for supplier_name, total in sorted(supplier_summary.items(), key=lambda item: item[1], reverse=True)[:8]
    ]

    forecast_income = 0.0
    finance_records = await db.finance_dashboard_records.find({"archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    for record in finance_records:
        record_date = record.get("anticipated_date") or record.get("expected_date")
        if finance_material_date_in_range(record_date, start_date, end_date):
            forecast_income += finance_material_to_float(record.get("anticipated_amount") or record.get("expected_amount"))

    forecast_material_spend = sum(row.get("exposure_net", row.get("net", 0.0)) for row in filtered_rows if row.get("type") in ["forecast", "po", "actual"])
    supplier_payments_due = sum(row.get("cash_out_gross", row.get("gross", 0.0)) for row in filtered_rows if row.get("type") == "po")

    attention = [
        {
            "type": "missing_receipts",
            "title": "Transactions missing receipts",
            "count": len([row for row in rows if row.get("type") == "actual" and row.get("receipt") is False]),
            "value": round(unreceipted_spend, 2),
            "detail": "Upload receipts for actual material entries.",
        },
        {
            "type": "pos_due",
            "title": "POs due or awaiting delivery",
            "count": len([row for row in rows if row.get("type") == "po" and normalise_po_status(row.get("status")) not in MATERIAL_PO_PAID_STATUSES]),
            "value": round(due_next_4_weeks, 2),
            "detail": "Check supplier delivery and invoice status.",
        },
        {
            "type": "over_allowance",
            "title": "Jobs over material allowance",
            "count": jobs_over_allowance,
            "value": round(sum(abs(item["variance"]) for item in job_summary.values() if item["variance"] < 0), 2),
            "detail": "Review job material allowance variances.",
        },
    ]

    return {
        "start_date": start_date,
        "end_date": end_date,
        "summary": {
            "actual_spend_this_month": round(actual_this_month, 2),
            "purchase_orders_this_month": round(po_this_month, 2),
            "supplier_payments_due": round(due_next_4_weeks, 2),
            "unreceipted_spend": round(unreceipted_spend, 2),
            "forecast_spend_next_4_weeks": round(forecast_next_4_weeks, 2),
            "jobs_over_allowance": jobs_over_allowance,
            "forecast_income": round(forecast_income, 2),
            "forecast_material_spend": round(forecast_material_spend, 2),
            "supplier_payments_due_filtered": round(supplier_payments_due, 2),
            "net_forecast_position": round(forecast_income - forecast_material_spend, 2),
        },
        "rows": filtered_rows,
        "top_suppliers": top_suppliers,
        "job_summary": list(job_summary.values()),
        "attention": attention,
    }


@api_router.get("/finance/material-spend")
@api_router.get("/finance/material-spend-dashboard")
async def get_finance_material_spend_dashboard(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    supplier: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    type: Optional[str] = Query(None),
):
    return await build_material_spend_dashboard_data(start_date, end_date, job_id, supplier, status, type)


@api_router.get("/finance/material-spend/export.csv")
@api_router.get("/finance/material-spend-dashboard/export.csv")
async def export_finance_material_spend_dashboard_csv(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    job_id: Optional[str] = Query(None),
    supplier: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    type: Optional[str] = Query(None),
):
    data = await build_material_spend_dashboard_data(start_date, end_date, job_id, supplier, status, type)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["LDA Group - Finance Material Spend Export"])
    writer.writerow(["Date range", data.get("start_date"), "to", data.get("end_date")])
    writer.writerow([])
    writer.writerow(["Date", "Type", "Job", "Supplier", "Description", "Net", "VAT", "Gross", "Status", "Due Date", "Paid Date", "Source", "Receipt"])
    for row in data.get("rows", []):
        writer.writerow([
            row.get("date", ""),
            row.get("type_label", row.get("type", "")),
            row.get("job_name", ""),
            row.get("supplier", ""),
            row.get("description", ""),
            row.get("net", 0),
            row.get("vat", 0),
            row.get("gross", 0),
            row.get("status", ""),
            row.get("due_date", ""),
            row.get("paid_date", ""),
            row.get("source", ""),
            "yes" if row.get("receipt") is True else "no" if row.get("receipt") is False else "n/a",
        ])
    output.seek(0)
    filename = f"finance_material_spend_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.csv"
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8-sig")),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ==================== SYSTEM STATUS & API INFO ====================

@api_router.get("/system/status")
async def get_system_status_endpoint():
    """Get system and service status"""
    # Check database connection
    db_status = False
    try:
        if db:
            await db.command("ping")
            db_status = True
    except Exception as e:
        logger.error(f"Database connection check failed: {e}")

    # Check quote services
    services_status = {
        "email_service": EmailService is not None,
        "pdf_generator": QuotePDFGenerator is not None,
        "google_drive": get_google_drive_service is not None and callable(get_google_drive_service)
    }

    return {
        "system": "operational",
        "database": db_status,
        "services": services_status,
        "timestamp": datetime.utcnow().isoformat(),
        "uk_time": get_uk_time().strftime('%Y-%m-%d %H:%M:%S %Z')
    }

@api_router.get("/api-info")
async def api_info_endpoint():
    """Get comprehensive API information"""
    return {
        "message": "LDA Group Time Tracking API",
        "version": "2.0.0",
        "features": [
            "Worker Management & Authentication",
            "Time Tracking with GPS",
            "Material Management",
            "Job Management & Costing",
            "Advanced Reporting & Analytics",
            "Comprehensive CSV Exports",
            "Quote System with PDF Generation",
            "Surveyor Authentication",
            "UK Timezone Support",
            "Real-time Dashboard Statistics",
            "Worker Scheduling Board"
        ],
        "main_endpoints": {
            "authentication": ["/api/admin/login", "/api/surveyors/login", "/api/surveyors/register"],
            "workers": ["/api/workers", "/api/workers/{id}", "/api/workers/{id}/archive"],
            "jobs": ["/api/jobs", "/api/jobs/{id}", "/api/jobs/{id}/archive", "/api/jobs/{id}/post-work-photos"],
            "time_tracking": ["/api/time-entries", "/api/time-entries/clock-in", "/api/time-entries/{id}/clock-out"],
            "materials": ["/api/materials", "/api/materials/{id}", "/api/materials/{id}/archive"],
            "quotes": ["/api/quotes", "/api/quotes/{id}", "/api/quotes/{id}/photos", "/api/quotes/{id}/download-pdf"],
            "reports": ["/api/reports/dashboard", "/api/reports/time-entries", "/api/reports/materials", "/api/reports/job-costs/{id}"],
            "schedule": ["/api/schedule", "/api/schedule/{id}", "/api/schedule/worker/{worker_id}", "/api/schedule/export"],
            "exports": ["/api/reports/export/job/{id}", "/api/reports/export/time-entries", "/api/reports/export/materials", "/api/reports/export/attendance-alerts"],
            "system": ["/api/system/status", "/api/api-info"]
        },
        "total_endpoints": 50
    }


# ==================== VARIATION SYSTEM HELPERS & ENDPOINTS ====================

VARIATION_ACTIVE_STATUSES = {"draft", "pending_client_approval", "approved", "rejected", "cancelled"}
VARIATION_PENDING_STATUSES = {"draft", "pending_client_approval"}


def variation_public_base_url() -> str:
    return (
        os.environ.get("PUBLIC_FRONTEND_URL")
        or os.environ.get("FRONTEND_URL")
        or ""
    ).strip()


def variation_backend_base_url() -> str:
    return (
        os.environ.get("PUBLIC_BACKEND_URL")
        or os.environ.get("BACKEND_URL")
        or ""
    ).strip().rstrip("/")


def calculate_variation_totals(data: Dict[str, Any]) -> Dict[str, float]:
    material = finance_to_number(data.get("material_value"))
    labour = finance_to_number(data.get("labour_value"))
    subcontractor = finance_to_number(data.get("subcontractor_value"))
    other = finance_to_number(data.get("other_value"))
    net = round(material + labour + subcontractor + other, 2)
    vat_rate = finance_to_number(data.get("vat_rate"), 20.0)
    vat = round(net * (vat_rate / 100.0), 2)
    return {
        "material_value": round(material, 2),
        "labour_value": round(labour, 2),
        "subcontractor_value": round(subcontractor, 2),
        "other_value": round(other, 2),
        "net_total": net,
        "vat_rate": vat_rate,
        "vat_total": vat,
        "gross_total": round(net + vat, 2),
    }


async def next_variation_number(job_id: str) -> str:
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0}) or {}
    prefix = job.get("job_number") or job.get("display_name") or job.get("name") or "VAR"
    prefix = str(prefix).replace(":", "").replace("/", "-")
    count = await db.variations.count_documents({"job_id": job_id})
    return f"VO-{prefix}-{str(count + 1).zfill(3)}"


def build_variation_approval_link(token: str) -> str:
    """Build the client-facing variation approval link.

    The React frontend uses HashRouter, so approval links must be:
    https://frontend.example/#/variation-approval/TOKEN

    Keep PUBLIC_FRONTEND_URL as the clean site root in Render, for example:
    https://lda-group.vercel.app
    This helper strips any accidental /#/ or # from the env var to avoid broken
    links such as https://site/#/#/variation-approval/TOKEN.
    """
    clean_token = str(token or "").strip()
    frontend = variation_public_base_url()

    if frontend:
        clean_frontend = frontend.strip().rstrip("/")

        if "/#/" in clean_frontend:
            clean_frontend = clean_frontend.split("/#/")[0].rstrip("/")
        elif clean_frontend.endswith("/#"):
            clean_frontend = clean_frontend[:-2].rstrip("/")
        elif clean_frontend.endswith("#"):
            clean_frontend = clean_frontend[:-1].rstrip("/")

        return f"{clean_frontend}/#/variation-approval/{clean_token}"

    backend = variation_backend_base_url()
    if backend:
        return f"{backend}/api/variations/public/{clean_token}"

    return f"/api/variations/public/{clean_token}"


async def recalculate_job_variation_totals(job_id: str) -> Dict[str, Any]:
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        return {}
    variations = await db.variations.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).to_list(5000)
    approved_total = round(sum(finance_to_number(v.get("net_total")) for v in variations if v.get("status") == "approved"), 2)
    pending_total = round(sum(finance_to_number(v.get("net_total")) for v in variations if v.get("status") in VARIATION_PENDING_STATUSES), 2)
    rejected_total = round(sum(finance_to_number(v.get("net_total")) for v in variations if v.get("status") == "rejected"), 2)
    original_value = finance_to_number(job.get("original_quoted_cost")) or finance_to_number(job.get("quoted_cost"))
    current_value = round(original_value + approved_total, 2)
    patch = {
        # Keep the tender/original value, but also update quoted_cost so existing
        # Jobs/Gantt/Finance screens that still read quoted_cost show the live value.
        "original_quoted_cost": original_value,
        "quoted_cost": current_value,
        "approved_variations_total": approved_total,
        "pending_variations_total": pending_total,
        "rejected_variations_total": rejected_total,
        "current_contract_value": current_value,
        "updated_at": datetime.utcnow(),
    }
    await db.jobs.update_one({"id": job_id}, {"$set": patch})
    return {**job, **patch}


async def create_or_update_variation_finance_record(variation: Dict[str, Any]) -> None:
    try:
        if variation.get("status") != "approved":
            return
        job = await db.jobs.find_one({"id": variation.get("job_id")}, {"_id": 0}) or {}
        expected_date = variation.get("required_date") or datetime.utcnow().date().isoformat()
        record_id = f"variation-{variation.get('id')}"
        variation_net = finance_to_number(variation.get("net_total"))
        standalone_invoice = bool(variation.get("standalone_variation_invoice"))
        tax_snapshot = calculate_finance_tax_snapshot(
            job,
            variation_net,
            labour_value=variation.get("labour_value"),
            material_value=variation.get("material_value"),
        )
        record = {
            "id": record_id,
            "project_id": variation.get("job_id"),
            "project_name": variation.get("job_name") or job.get("name") or "",
            "type": "variation",
            "description": variation.get("title") or "Approved variation",
            "expected_date": expected_date,
            # Approved variations are included in the contract/application valuation by default.
            # They remain visible here as references but are excluded from totals to prevent double-counting.
            "expected_amount": variation_net if standalone_invoice else 0.0,
            "anticipated_date": expected_date,
            "anticipated_amount": variation_net if standalone_invoice else 0.0,
            "display_amount": variation_net,
            "reference_amount": variation_net,
            "status": "expected" if standalone_invoice else "reference",
            "source": "variation",
            "source_id": variation.get("id"),
            "included_in_contract_value": not standalone_invoice,
            "exclude_from_finance_totals": not standalone_invoice,
            "standalone_variation_invoice": standalone_invoice,
            "tax_treatment_snapshot": tax_snapshot,
            "expected_cash_value": tax_snapshot.get("expected_cash_value") if standalone_invoice else 0.0,
            "cis_deduction_value": tax_snapshot.get("cis_deduction_value"),
            "vat_value": tax_snapshot.get("vat_value"),
            "gross_value": tax_snapshot.get("gross_value"),
            "notes": " | ".join([item for item in [variation.get("scope_of_works") or variation.get("description") or "", "Included in contract/application forecast" if not standalone_invoice else "Standalone variation invoice"] if item]),
            "created_by": variation.get("created_by_name", ""),
            "created_by_email": variation.get("created_by_email", ""),
            "created_by_role": variation.get("created_by_role", ""),
            "updated_by": variation.get("approved_by_name") or variation.get("created_by_name", ""),
            "updated_by_email": variation.get("approved_by_email") or variation.get("created_by_email", ""),
            "updated_by_role": "client" if variation.get("approved_by_email") else variation.get("created_by_role", ""),
            "updated_at": datetime.utcnow(),
            "archived": False,
        }
        existing = await db.finance_dashboard_records.find_one({"id": record_id})
        if existing:
            await db.finance_dashboard_records.update_one({"id": record_id}, {"$set": record})
        else:
            record["created_at"] = datetime.utcnow()
            await db.finance_dashboard_records.insert_one(record)
    except Exception as exc:
        logger.warning("Could not create variation finance record: %s", exc)


async def add_variation_to_gantt(variation_id: str) -> Dict[str, Any]:
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    job = await db.jobs.find_one({"id": variation.get("job_id")}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if variation.get("added_to_gantt") and variation.get("gantt_section_id"):
        return variation
    section_id = variation.get("gantt_section_id") or f"variation-section-{variation_id}"
    sections = job.get("gantt_sections") or []
    if not any(section.get("id") == section_id for section in sections):
        start_date = variation.get("required_date") or job.get("planned_end_date") or job.get("planned_start_date") or ""
        new_section = {
            "id": section_id,
            "name": f"Variation - {variation.get('title') or variation.get('variation_number')}",
            "start_date": start_date,
            "end_date": start_date,
            "status": "planned",
            "notes": "\n".join([variation.get("client_instruction_summary", ""), variation.get("scope_of_works", "")]).strip(),
            "required_trades": [],
            "assigned_worker_ids": [],
            "sent_to_schedule": False,
            "schedule_entry_ids": [],
            "schedule_status": "not_scheduled",
            "labour_value": finance_to_number(variation.get("labour_value")),
            "material_value": finance_to_number(variation.get("material_value")),
            "subcontractor_value": finance_to_number(variation.get("subcontractor_value")),
            "other_value": finance_to_number(variation.get("other_value")),
            "section_value": finance_to_number(variation.get("net_total")),
            "progress_percent": 0,
            "variation_id": variation_id,
            "variation_number": variation.get("variation_number", ""),
        }
        sections.append(new_section)
        await db.jobs.update_one({"id": job.get("id")}, {"$set": {"gantt_sections": sections, "include_in_gantt": True, "updated_at": datetime.utcnow()}})
    patch = {"added_to_gantt": True, "gantt_section_id": section_id, "updated_at": datetime.utcnow()}
    await db.variations.update_one({"id": variation_id}, {"$set": patch})
    return await db.variations.find_one({"id": variation_id}, {"_id": 0})


async def approve_variation(variation_id: str, approver: Dict[str, Any], signature: str = "", comments: str = "") -> Dict[str, Any]:
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    patch = {
        "status": "approved",
        "approved_by_name": approver.get("name") or approver.get("email") or "Approved",
        "approved_by_email": approver.get("email", ""),
        "approved_at": datetime.utcnow(),
        "client_signature": signature or variation.get("client_signature", ""),
        "client_comments": comments or variation.get("client_comments", ""),
        "updated_at": datetime.utcnow(),
    }
    await db.variations.update_one({"id": variation_id}, {"$set": patch})
    updated = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if updated.get("add_to_gantt_on_approval"):
        updated = await add_variation_to_gantt(variation_id)
    await recalculate_job_variation_totals(updated.get("job_id"))
    await create_or_update_variation_finance_record(updated)

    # Once the client has approved, send the final approved PDF record back to
    # the person who originally raised the variation. This gives the PM/admin
    # an immediate audit trail without needing to download it from the app.
    originator_notification = await send_variation_originator_approval_notification(updated)
    await db.variations.update_one({"id": variation_id}, {"$set": {
        "originator_approval_notification_status": "sent" if originator_notification.get("sent") else originator_notification.get("method", "not_configured"),
        "originator_approval_notification_result": originator_notification,
        "originator_approval_sent_at": datetime.utcnow() if originator_notification.get("sent") else updated.get("originator_approval_sent_at"),
        "updated_at": datetime.utcnow(),
    }})

    return await db.variations.find_one({"id": variation_id}, {"_id": 0})


def variation_originator_email(variation: Dict[str, Any]) -> str:
    """Best-effort email for the person who raised/requested the variation."""
    for key in [
        "created_by_email",
        "requested_by_email",
        "raised_by_email",
        "submitted_by_email",
        "updated_by_email",
    ]:
        value = str(variation.get(key) or "").strip()
        if value and "@" in value:
            return value
    return ""


def variation_originator_name(variation: Dict[str, Any]) -> str:
    for key in [
        "created_by_name",
        "requested_by_name",
        "raised_by_name",
        "submitted_by_name",
        "updated_by_name",
    ]:
        value = str(variation.get(key) or "").strip()
        if value:
            return value
    email = variation_originator_email(variation)
    return email or "Originator"


def build_variation_originator_approval_payload(variation: Dict[str, Any]) -> Dict[str, Any]:
    recipient = variation_originator_email(variation)
    recipient_name = variation_originator_name(variation)
    backend = variation_backend_base_url()
    pdf_link = f"{backend}/api/variations/{variation.get('id')}/pdf" if backend and variation.get("id") else ""
    approval_link = variation.get("approval_link") or build_variation_approval_link(variation.get("approval_token", ""))
    approved_by = variation.get("approved_by_name") or variation.get("approved_by_email") or "the client"
    approved_at = variation_display_date(variation.get("approved_at"))

    email_subject = f"Variation approved - {variation.get('variation_number')} - {variation.get('job_name')}"
    email_html = f"""
    <div style="font-family:Arial,sans-serif;color:#0f172a;line-height:1.45;max-width:720px">
      <div style="border-bottom:4px solid #d01f2f;padding-bottom:12px;margin-bottom:18px">
        <h2 style="margin:0;color:#0f172a">Variation Approved</h2>
        <p style="margin:4px 0 0;color:#64748b">{variation.get('variation_number') or ''}</p>
      </div>
      <p>Hi {recipient_name},</p>
      <p>The following variation has been approved by the client. A PDF approval record is attached.</p>
      <table style="border-collapse:collapse;width:100%;font-size:14px">
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Project</td><td style="padding:8px;border:1px solid #cbd5e1">{variation.get('job_name') or ''}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Variation</td><td style="padding:8px;border:1px solid #cbd5e1">{variation.get('title') or ''}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Approved by</td><td style="padding:8px;border:1px solid #cbd5e1">{approved_by}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Approved date</td><td style="padding:8px;border:1px solid #cbd5e1">{approved_at}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Net</td><td style="padding:8px;border:1px solid #cbd5e1">{variation_money(variation.get('net_total'))}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">VAT</td><td style="padding:8px;border:1px solid #cbd5e1">{variation_money(variation.get('vat_total'))}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Gross</td><td style="padding:8px;border:1px solid #cbd5e1;color:#b91c1c;font-weight:bold">{variation_money(variation.get('gross_total'))}</td></tr>
      </table>
      <p style="font-size:12px;color:#64748b;margin-top:16px">Client approval page:<br />{approval_link}</p>
    </div>
    """

    payload = {
        "notification_type": "variation_approved_originator",
        "variation_id": variation.get("id"),
        "variation_number": variation.get("variation_number"),
        "job_id": variation.get("job_id"),
        "job_name": variation.get("job_name"),
        "client_name": recipient_name,
        "client_email": recipient,
        "to_email": recipient,
        "to_name": recipient_name,
        "originator_email": recipient,
        "originator_name": recipient_name,
        "title": variation.get("title"),
        "description": variation.get("description"),
        "client_instruction_summary": variation.get("client_instruction_summary"),
        "scope_of_works": variation.get("scope_of_works"),
        "net_total": variation.get("net_total"),
        "vat_total": variation.get("vat_total"),
        "gross_total": variation.get("gross_total"),
        "approval_link": approval_link,
        "pdf_link": pdf_link,
        "email_subject": email_subject,
        "email_html": email_html,
    }
    payload.update(build_variation_pdf_attachment_payload(variation))
    return payload


async def send_variation_originator_approval_notification(variation: Dict[str, Any]) -> Dict[str, Any]:
    recipient = variation_originator_email(variation)
    if not recipient:
        return {"sent": False, "method": "not_configured", "error": "No originator email saved against this variation"}

    # You can use a dedicated approval-notification flow, or leave it blank to
    # reuse the same Power Automate HTTP trigger as the client approval email.
    power_automate_url = (
        os.environ.get("POWER_AUTOMATE_VARIATION_APPROVED_URL", "").strip()
        or os.environ.get("POWER_AUTOMATE_VARIATION_APPROVAL_URL", "").strip()
    )
    power_automate_secret = (
        os.environ.get("POWER_AUTOMATE_VARIATION_APPROVED_SECRET", "").strip()
        or os.environ.get("POWER_AUTOMATE_VARIATION_APPROVAL_SECRET", "").strip()
    )

    payload = build_variation_originator_approval_payload(variation)

    if power_automate_url:
        headers = {"Content-Type": "application/json"}
        if power_automate_secret:
            headers["x-lda-secret"] = power_automate_secret
        try:
            response = requests.post(power_automate_url, json=payload, headers=headers, timeout=20)
            response.raise_for_status()
            return {"sent": True, "method": "power_automate", "to": recipient}
        except Exception as exc:
            logger.exception("Failed to send variation approved Power Automate notification: %s", exc)
            return {"sent": False, "method": "power_automate", "to": recipient, "error": str(exc)}

    smtp_host = os.environ.get("SMTP_HOST", "").strip()
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_username = os.environ.get("SMTP_USERNAME", "").strip()
    smtp_password = os.environ.get("SMTP_PASSWORD", "").strip()
    smtp_from = os.environ.get("SMTP_FROM_EMAIL", smtp_username).strip()
    smtp_from_name = os.environ.get("SMTP_FROM_NAME", "LDA Group").strip()

    if smtp_host and smtp_username and smtp_password and smtp_from:
        msg = EmailMessage()
        msg["Subject"] = payload.get("email_subject", "Variation approved")
        msg["From"] = f"{smtp_from_name} <{smtp_from}>"
        msg["To"] = recipient
        msg.set_content(f"Variation approved: {variation.get('variation_number')} - {variation.get('job_name')}\n")
        msg.add_alternative(payload.get("email_html", ""), subtype="html")
        if payload.get("pdf_base64"):
            try:
                msg.add_attachment(
                    base64.b64decode(payload["pdf_base64"]),
                    maintype="application",
                    subtype="pdf",
                    filename=payload.get("pdf_filename") or "variation.pdf",
                )
            except Exception as attach_exc:
                logger.warning("Could not attach approved variation PDF to SMTP email: %s", attach_exc)
        try:
            with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as smtp:
                if os.environ.get("SMTP_USE_TLS", "true").lower() != "false":
                    smtp.starttls()
                smtp.login(smtp_username, smtp_password)
                smtp.send_message(msg)
            return {"sent": True, "method": "smtp", "to": recipient}
        except Exception as exc:
            logger.exception("Failed to send variation approved SMTP email: %s", exc)
            return {"sent": False, "method": "smtp", "to": recipient, "error": str(exc)}

    return {"sent": False, "method": "not_configured", "to": recipient, "error": "No Power Automate or SMTP settings configured"}


async def send_variation_client_notification(variation: Dict[str, Any]) -> Dict[str, Any]:
    recipient = str(variation.get("client_email") or "").strip()
    if not recipient:
        return {"sent": False, "method": "not_configured", "error": "No client email saved against this variation"}
    token = variation.get("approval_token") or ""
    approval_link = build_variation_approval_link(token) if token else variation.get("approval_link", "")

    backend = variation_backend_base_url()
    pdf_link = f"{backend}/api/variations/{variation.get('id')}/pdf" if backend and variation.get("id") else ""
    email_subject = f"Variation approval required - {variation.get('variation_number')} - {variation.get('job_name')}"
    email_html = f"""
    <div style="font-family:Arial,sans-serif;color:#0f172a;line-height:1.45;max-width:720px">
      <div style="border-bottom:4px solid #d01f2f;padding-bottom:12px;margin-bottom:18px">
        <h2 style="margin:0;color:#0f172a">LDA Group - Variation Approval Required</h2>
        <p style="margin:4px 0 0;color:#64748b">{variation.get('variation_number') or ''}</p>
      </div>
      <p>Please review the following variation request:</p>
      <table style="border-collapse:collapse;width:100%;font-size:14px">
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Project</td><td style="padding:8px;border:1px solid #cbd5e1">{variation.get('job_name') or ''}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Variation</td><td style="padding:8px;border:1px solid #cbd5e1">{variation.get('title') or ''}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Net</td><td style="padding:8px;border:1px solid #cbd5e1">{variation_money(variation.get('net_total'))}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">VAT</td><td style="padding:8px;border:1px solid #cbd5e1">{variation_money(variation.get('vat_total'))}</td></tr>
        <tr><td style="padding:8px;border:1px solid #cbd5e1;background:#f8fafc;font-weight:bold">Gross</td><td style="padding:8px;border:1px solid #cbd5e1;color:#b91c1c;font-weight:bold">{variation_money(variation.get('gross_total'))}</td></tr>
      </table>
      <p style="margin-top:18px"><a href="{approval_link}" style="background:#d01f2f;color:white;text-decoration:none;padding:12px 18px;border-radius:8px;display:inline-block;font-weight:bold">Review, Query or Approve Variation</a></p>
      <p style="font-size:12px;color:#64748b">If the button does not work, copy and paste this link into your browser:<br />{approval_link}</p>
    </div>
    """
    payload = {
        "variation_id": variation.get("id"),
        "variation_number": variation.get("variation_number"),
        "job_id": variation.get("job_id"),
        "job_name": variation.get("job_name"),
        "client_name": variation.get("client_name"),
        "client_email": recipient,
        "to_email": recipient,
        "to_name": variation.get("client_name"),
        "originator_email": variation_originator_email(variation),
        "originator_name": variation_originator_name(variation),
        "reply_to_email": variation_originator_email(variation),
        "reply_to_name": variation_originator_name(variation),
        "title": variation.get("title"),
        "description": variation.get("description"),
        "client_instruction_summary": variation.get("client_instruction_summary"),
        "scope_of_works": variation.get("scope_of_works"),
        "net_total": variation.get("net_total"),
        "vat_total": variation.get("vat_total"),
        "gross_total": variation.get("gross_total"),
        "approval_link": approval_link,
        "pdf_link": pdf_link,
        "email_subject": email_subject,
        "email_html": email_html,
    }
    payload.update(build_variation_pdf_attachment_payload(variation))
    power_automate_url = os.environ.get("POWER_AUTOMATE_VARIATION_APPROVAL_URL", "").strip()
    power_automate_secret = os.environ.get("POWER_AUTOMATE_VARIATION_APPROVAL_SECRET", "").strip()
    if power_automate_url:
        headers = {"Content-Type": "application/json"}
        if power_automate_secret:
            headers["x-lda-secret"] = power_automate_secret
        try:
            response = requests.post(power_automate_url, json=payload, headers=headers, timeout=20)
            response.raise_for_status()
            return {"sent": True, "method": "power_automate", "to": recipient}
        except Exception as exc:
            logger.exception("Failed to send variation approval Power Automate notification: %s", exc)
            return {"sent": False, "method": "power_automate", "to": recipient, "error": str(exc)}
    smtp_host = os.environ.get("SMTP_HOST", "").strip()
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    smtp_username = os.environ.get("SMTP_USERNAME", "").strip()
    smtp_password = os.environ.get("SMTP_PASSWORD", "").strip()
    smtp_from = os.environ.get("SMTP_FROM_EMAIL", smtp_username).strip()
    smtp_from_name = os.environ.get("SMTP_FROM_NAME", "LDA Group").strip()
    smtp_reply_to = os.environ.get("SMTP_REPLY_TO", "").strip()
    if smtp_host and smtp_username and smtp_password and smtp_from:
        msg = EmailMessage()
        msg["Subject"] = email_subject
        msg["From"] = f"{smtp_from_name} <{smtp_from}>"
        msg["To"] = recipient
        if smtp_reply_to:
            msg["Reply-To"] = smtp_reply_to
        msg.set_content(
            f"Please review the following variation request:\n\n"
            f"Project: {variation.get('job_name')}\n"
            f"Variation: {variation.get('title')}\n"
            f"Net: {variation_money(variation.get('net_total'))}\n"
            f"VAT: {variation_money(variation.get('vat_total'))}\n"
            f"Gross: {variation_money(variation.get('gross_total'))}\n\n"
            f"Approval link: {approval_link}\n"
        )
        msg.add_alternative(email_html, subtype="html")
        attachment = build_variation_pdf_attachment_payload(variation)
        if attachment.get("pdf_base64"):
            try:
                msg.add_attachment(
                    base64.b64decode(attachment["pdf_base64"]),
                    maintype="application",
                    subtype="pdf",
                    filename=attachment.get("pdf_filename") or "variation.pdf",
                )
            except Exception as attach_exc:
                logger.warning("Could not attach variation PDF to SMTP email: %s", attach_exc)
        try:
            with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as smtp:
                if os.environ.get("SMTP_USE_TLS", "true").lower() != "false":
                    smtp.starttls()
                smtp.login(smtp_username, smtp_password)
                smtp.send_message(msg)
            return {"sent": True, "method": "smtp", "to": recipient}
        except Exception as exc:
            logger.exception("Failed to send variation approval SMTP email: %s", exc)
            return {"sent": False, "method": "smtp", "to": recipient, "error": str(exc)}
    return {"sent": False, "method": "not_configured", "to": recipient, "error": "No Power Automate or SMTP variation approval settings configured"}




def variation_money(value: Any) -> str:
    return f"£{finance_to_number(value):,.2f}"


def variation_display_date(value: Any) -> str:
    if not value:
        return "-"
    try:
        value_str = str(value)
        if len(value_str) >= 10 and value_str[4] == "-" and value_str[7] == "-":
            parsed = datetime.fromisoformat(value_str[:10])
            return parsed.strftime("%d %b %Y")
        if isinstance(value, datetime):
            return value.strftime("%d %b %Y")
        return value_str
    except Exception:
        return str(value)

def safe_variation_pdf_filename(variation: Dict[str, Any]) -> str:
    number = str(variation.get("variation_number") or "variation").strip() or "variation"
    title = str(variation.get("title") or variation.get("job_name") or "approval").strip() or "approval"
    filename = f"{number} - {title}.pdf"
    filename = re.sub(r'[\\/:*?"<>|]+', "-", filename)
    filename = re.sub(r"\s+", " ", filename).strip()
    return filename[:150] if filename else "variation.pdf"


def build_variation_pdf_attachment_payload(variation: Dict[str, Any]) -> Dict[str, str]:
    try:
        pdf_bytes = generate_variation_pdf_bytes(variation)
        return {
            "pdf_filename": safe_variation_pdf_filename(variation),
            "pdf_base64": base64.b64encode(pdf_bytes).decode("utf-8"),
        }
    except Exception as exc:
        logger.exception("Could not generate variation PDF attachment: %s", exc)
        return {"pdf_filename": "", "pdf_base64": ""}


def generate_variation_pdf_bytes(variation: Dict[str, Any]) -> bytes:
    """Create a formal PDF record for a variation approval/instruction."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    except Exception:
        raise HTTPException(status_code=500, detail="PDF generation is not available on this server. Check reportlab is installed.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.3 * cm,
        leftMargin=1.3 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LdaTitle", parent=styles["Heading1"], fontSize=16, leading=20, spaceAfter=6, textColor=colors.HexColor("#0f172a")))
    styles.add(ParagraphStyle(name="LdaHeading", parent=styles["Heading2"], fontSize=10.5, leading=13, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor("#0f172a")))
    styles.add(ParagraphStyle(name="LdaBody", parent=styles["BodyText"], fontSize=8.5, leading=11, textColor=colors.HexColor("#334155")))
    styles.add(ParagraphStyle(name="LdaSmall", parent=styles["BodyText"], fontSize=7.5, leading=9.5, textColor=colors.HexColor("#64748b")))

    def p(value, style="LdaBody"):
        safe = str(value or "-").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        return Paragraph(safe, styles[style])

    story = []

    logo_url = os.environ.get("LDA_LOGO_URL", "https://ldagroup.co.uk/wp-content/uploads/2022/01/lda-group-200x200.png").strip()
    logo_cell = ""
    if logo_url:
        try:
            logo_response = requests.get(logo_url, timeout=8)
            logo_response.raise_for_status()
            logo_image = Image(io.BytesIO(logo_response.content), width=1.55 * cm, height=1.55 * cm)
            logo_cell = logo_image
        except Exception:
            logo_cell = ""

    header = Table(
        [[logo_cell, [Paragraph("LDA Group", styles["LdaTitle"]), Paragraph("Variation Approval Record", styles["LdaSmall"])] , Paragraph(str(variation.get("variation_number") or "Variation"), styles["LdaTitle"])]],
        colWidths=[2.0 * cm, 9.4 * cm, 5.2 * cm],
    )
    header.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (2, 0), (2, 0), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LINEBELOW", (0, 0), (-1, -1), 1, colors.HexColor("#d01f2f")),
    ]))
    story.append(header)
    story.append(Spacer(1, 0.25 * cm))

    status_text = str(variation.get("status") or "").replace("_", " ").title()
    summary_rows = [
        [p("Project", "LdaSmall"), p(variation.get("job_name") or variation.get("job_id")), p("Status", "LdaSmall"), p(status_text)],
        [p("Variation title", "LdaSmall"), p(variation.get("title")), p("Required date", "LdaSmall"), p(variation_display_date(variation.get("required_date")))],
        [p("Client", "LdaSmall"), p(variation.get("client_name") or "-"), p("Client email", "LdaSmall"), p(variation.get("client_email") or "-")],
        [p("Raised by", "LdaSmall"), p(variation.get("created_by_name") or "-"), p("Raised date", "LdaSmall"), p(variation_display_date(variation.get("created_at")))],
    ]
    summary = Table(summary_rows, colWidths=[3.0 * cm, 6.2 * cm, 3.0 * cm, 4.4 * cm])
    summary.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f8fafc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(summary)

    story.append(Paragraph("Client instruction / conversation summary", styles["LdaHeading"]))
    story.append(p(variation.get("client_instruction_summary") or "No client instruction summary recorded."))
    story.append(Paragraph("Scope / description", styles["LdaHeading"]))
    story.append(p(variation.get("scope_of_works") or variation.get("description") or "No scope recorded."))

    cost_rows = [
        ["Materials", variation_money(variation.get("material_value"))],
        ["Labour", variation_money(variation.get("labour_value"))],
        ["Subcontractor", variation_money(variation.get("subcontractor_value"))],
        ["Other", variation_money(variation.get("other_value"))],
        ["Net variation", variation_money(variation.get("net_total"))],
        [f"VAT @ {finance_to_number(variation.get('vat_rate'), 20):g}%", variation_money(variation.get("vat_total"))],
        ["Gross total", variation_money(variation.get("gross_total"))],
    ]
    cost_table = Table(cost_rows, colWidths=[11.6 * cm, 5.0 * cm])
    cost_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (-1, 3), colors.white),
        ("BACKGROUND", (0, 4), (-1, 4), colors.HexColor("#f8fafc")),
        ("BACKGROUND", (0, 6), (-1, 6), colors.HexColor("#fee2e2")),
        ("TEXTCOLOR", (0, 6), (-1, 6), colors.HexColor("#991b1b")),
        ("FONTNAME", (0, 4), (-1, 6), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(Paragraph("Cost breakdown", styles["LdaHeading"]))
    story.append(cost_table)

    story.append(Paragraph("Approval / response", styles["LdaHeading"]))
    approval_rows = [
        [p("Approved by", "LdaSmall"), p(variation.get("approved_by_name") or variation.get("rejected_by_name") or "-")],
        [p("Email", "LdaSmall"), p(variation.get("approved_by_email") or variation.get("rejected_by_email") or "-")],
        [p("Date", "LdaSmall"), p(variation_display_date(variation.get("approved_at") or variation.get("rejected_at")))],
        [p("Signature", "LdaSmall"), p(variation.get("client_signature") or "-")],
        [p("Comments", "LdaSmall"), p(variation.get("client_comments") or variation.get("rejection_reason") or "-")],
    ]
    approval_table = Table(approval_rows, colWidths=[3.0 * cm, 13.6 * cm])
    approval_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(approval_table)

    approval_link = variation.get("approval_link") or build_variation_approval_link(variation.get("approval_token", ""))
    if approval_link:
        story.append(Spacer(1, 0.35 * cm))
        story.append(Paragraph("Client review link", styles["LdaHeading"]))
        button_text = f'<link href="{approval_link}"><font color="white"><b>Review, Query or Approve Variation</b></font></link>'
        button_table = Table([[Paragraph(button_text, styles["LdaBody"])]], colWidths=[16.6 * cm])
        button_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#d01f2f")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING", (0, 0), (-1, -1), 10),
            ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#d01f2f")),
        ]))
        story.append(button_table)
        story.append(Spacer(1, 0.12 * cm))
        story.append(Paragraph(f"If the button does not work, copy and paste this link into your browser:<br/>{approval_link}", styles["LdaSmall"]))

    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("This document records the variation request and client response held within the LDA Work App.", styles["LdaSmall"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


@api_router.get("/variations")
async def get_variations(job_id: Optional[str] = Query(None), status: Optional[str] = Query(None), include_archived: bool = Query(False), admin: str = Depends(verify_admin)):
    filters: Dict[str, Any] = {}
    if job_id:
        filters["job_id"] = job_id
    if status and status != "all":
        filters["status"] = status
    if not include_archived:
        filters["archived"] = {"$ne": True}
    return await db.variations.find(filters, {"_id": 0}).sort("created_at", -1).to_list(5000)


@api_router.get("/jobs/{job_id}/variations")
async def get_job_variations(job_id: str, admin: str = Depends(verify_admin)):
    return await db.variations.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).sort("created_at", -1).to_list(1000)


@api_router.get("/jobs/{job_id}/variation-summary")
async def get_job_variation_summary(job_id: str, admin: str = Depends(verify_admin)):
    job = await recalculate_job_variation_totals(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    variations = await db.variations.find({"job_id": job_id, "archived": {"$ne": True}}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return {
        "job_id": job_id,
        "job_name": job.get("display_name") or job.get("name") or "",
        "original_quoted_cost": finance_to_number(job.get("original_quoted_cost") or job.get("quoted_cost")),
        "approved_variations_total": finance_to_number(job.get("approved_variations_total")),
        "pending_variations_total": finance_to_number(job.get("pending_variations_total")),
        "rejected_variations_total": finance_to_number(job.get("rejected_variations_total")),
        "current_contract_value": finance_to_number(job.get("current_contract_value") or job.get("quoted_cost")),
        "variation_count": len(variations),
        "approved_count": len([v for v in variations if v.get("status") == "approved"]),
        "pending_count": len([v for v in variations if v.get("status") in VARIATION_PENDING_STATUSES]),
        "rejected_count": len([v for v in variations if v.get("status") == "rejected"]),
        "variations": variations,
    }


@api_router.post("/variations", response_model=Variation)
async def create_variation(variation: VariationCreate, admin: str = Depends(verify_admin)):
    job = await db.jobs.find_one({"id": variation.job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    data = variation.dict()
    totals = calculate_variation_totals(data)
    token = secrets.token_urlsafe(32)
    doc = {
        **data,
        **totals,
        "id": str(uuid.uuid4()),
        "variation_number": await next_variation_number(variation.job_id),
        "job_name": job.get("display_name") or job.get("name") or "",
        "job_number": job.get("job_number"),
        "client_name": job.get("client") or "",
        "status": variation.status if variation.status in VARIATION_ACTIVE_STATUSES else "pending_client_approval",
        "approval_token": token,
        "approval_link": build_variation_approval_link(token),
        "approval_notification_status": "not_sent",
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow(),
        "archived": False,
    }
    await db.variations.insert_one(doc)
    await recalculate_job_variation_totals(variation.job_id)
    if doc.get("status") == "pending_client_approval" and doc.get("client_email"):
        notification_result = await send_variation_client_notification(doc)
        await db.variations.update_one({"id": doc["id"]}, {"$set": {
            "approval_sent_at": datetime.utcnow() if notification_result.get("sent") else None,
            "approval_notification_status": "sent" if notification_result.get("sent") else notification_result.get("method", "not_configured"),
            "approval_notification_result": notification_result,
        }})
        doc = await db.variations.find_one({"id": doc["id"]}, {"_id": 0})
    return Variation(**doc)


@api_router.get("/variations/{variation_id}")
async def get_variation(variation_id: str, admin: str = Depends(verify_admin)):
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    return variation


@api_router.get("/variations/{variation_id}/pdf")
async def download_variation_pdf(variation_id: str, admin: str = Depends(verify_admin)):
    variation = await db.variations.find_one({"id": variation_id, "archived": {"$ne": True}}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    pdf_bytes = generate_variation_pdf_bytes(variation)
    safe_ref = re.sub(r"[^A-Za-z0-9_-]+", "_", str(variation.get("variation_number") or variation_id)).strip("_") or "variation"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={safe_ref}.pdf"},
    )


@api_router.put("/variations/{variation_id}")
async def update_variation(variation_id: str, update: VariationUpdate, admin: str = Depends(verify_admin)):
    existing = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Variation not found")
    if existing.get("status") == "approved":
        raise HTTPException(status_code=400, detail="Approved variations cannot be edited. Create a further variation instead.")
    update_data = {k: v for k, v in update.dict().items() if v is not None}
    update_data.update(calculate_variation_totals({**existing, **update_data}))
    update_data["updated_at"] = datetime.utcnow()
    await db.variations.update_one({"id": variation_id}, {"$set": update_data})
    updated = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    await recalculate_job_variation_totals(updated.get("job_id"))
    return updated


@api_router.delete("/variations/{variation_id}")
async def archive_variation(variation_id: str, super_admin: Dict[str, Any] = Depends(get_super_admin_user)):
    """Soft-delete a variation. Super Admin only.

    If the variation had already been approved, this reverses its commercial impact by
    archiving the variation and recalculating the job variation totals. Any finance
    record created from the variation is archived, and the linked Gantt section is
    removed so the job value/programme stay in sync.
    """
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")

    now = datetime.utcnow()
    actor_name = super_admin.get("name") or super_admin.get("email") or "Super Admin"

    await db.variations.update_one(
        {"id": variation_id},
        {"$set": {
            "archived": True,
            "archived_at": now,
            "archived_by": actor_name,
            "archived_by_email": super_admin.get("email", ""),
            "status_before_archive": variation.get("status", ""),
            "updated_at": now,
        }},
    )

    # Hide any finance/dashboard record created for this variation.
    await db.finance_dashboard_records.update_one(
        {"id": f"variation-{variation_id}"},
        {"$set": {"archived": True, "updated_at": now}},
    )

    # Remove the linked Gantt section, if one was created from this variation.
    job_id = variation.get("job_id")
    gantt_section_id = variation.get("gantt_section_id")
    if job_id:
        job = await db.jobs.find_one({"id": job_id}, {"_id": 0}) or {}
        sections = job.get("gantt_sections") or []
        filtered_sections = [
            section for section in sections
            if section.get("variation_id") != variation_id and section.get("id") != gantt_section_id
        ]
        if len(filtered_sections) != len(sections):
            await db.jobs.update_one(
                {"id": job_id},
                {"$set": {"gantt_sections": filtered_sections, "updated_at": now}},
            )

        await recalculate_job_variation_totals(job_id)

    return {"success": True, "message": "Variation deleted", "variation_id": variation_id}


@api_router.post("/variations/{variation_id}/submit-for-approval")
async def submit_variation_for_approval(variation_id: str, admin: str = Depends(verify_admin)):
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    token = variation.get("approval_token") or secrets.token_urlsafe(32)
    patch = {
        "status": "pending_client_approval",
        "approval_token": token,
        "approval_link": build_variation_approval_link(token),
        "updated_at": datetime.utcnow(),
    }
    await db.variations.update_one({"id": variation_id}, {"$set": patch})
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    notification_result = await send_variation_client_notification(variation)
    await db.variations.update_one({"id": variation_id}, {"$set": {
        "approval_sent_at": datetime.utcnow() if notification_result.get("sent") else variation.get("approval_sent_at"),
        "approval_notification_status": "sent" if notification_result.get("sent") else notification_result.get("method", "not_configured"),
        "approval_notification_result": notification_result,
        "updated_at": datetime.utcnow(),
    }})
    await recalculate_job_variation_totals(variation.get("job_id"))
    return await db.variations.find_one({"id": variation_id}, {"_id": 0})


@api_router.post("/variations/{variation_id}/approve")
async def admin_approve_variation(variation_id: str, admin: str = Depends(verify_admin)):
    return await approve_variation(variation_id, {"name": admin, "email": admin})


@api_router.post("/variations/{variation_id}/reject")
async def admin_reject_variation(variation_id: str, response: VariationApprovalResponse, admin: str = Depends(verify_admin)):
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    patch = {
        "status": "rejected",
        "rejected_by_name": response.name or admin,
        "rejected_by_email": response.email or admin,
        "rejected_at": datetime.utcnow(),
        "rejection_reason": response.reason or response.comments,
        "client_comments": response.comments,
        "updated_at": datetime.utcnow(),
    }
    await db.variations.update_one({"id": variation_id}, {"$set": patch})
    updated = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    await recalculate_job_variation_totals(updated.get("job_id"))
    return updated


@api_router.post("/variations/{variation_id}/add-to-gantt")
async def variation_add_to_gantt(variation_id: str, admin: str = Depends(verify_admin)):
    variation = await db.variations.find_one({"id": variation_id}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation not found")
    if variation.get("status") != "approved":
        raise HTTPException(status_code=400, detail="Only approved variations can be added to the Gantt")
    updated = await add_variation_to_gantt(variation_id)
    await recalculate_job_variation_totals(updated.get("job_id"))
    return updated


@api_router.get("/variations/public/{token}")
async def public_get_variation(token: str):
    variation = await db.variations.find_one({"approval_token": token, "archived": {"$ne": True}}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation approval link not found")
    return variation


@api_router.post("/variations/public/{token}/response")
async def public_variation_response(token: str, response: VariationApprovalResponse):
    variation = await db.variations.find_one({"approval_token": token, "archived": {"$ne": True}}, {"_id": 0})
    if not variation:
        raise HTTPException(status_code=404, detail="Variation approval link not found")
    decision = str(response.decision or "").strip().lower()
    if decision in ["approve", "approved", "accept", "accepted"]:
        return await approve_variation(variation.get("id"), {"name": response.name or "Client", "email": response.email or variation.get("client_email", "")}, signature=response.signature, comments=response.comments)
    if decision in ["reject", "rejected", "decline", "declined"]:
        patch = {
            "status": "rejected",
            "rejected_by_name": response.name or "Client",
            "rejected_by_email": response.email or variation.get("client_email", ""),
            "rejected_at": datetime.utcnow(),
            "rejection_reason": response.reason or response.comments,
            "client_comments": response.comments,
            "updated_at": datetime.utcnow(),
        }
        await db.variations.update_one({"id": variation.get("id")}, {"$set": patch})
        updated = await db.variations.find_one({"id": variation.get("id")}, {"_id": 0})
        await recalculate_job_variation_totals(updated.get("job_id"))
        return updated
    raise HTTPException(status_code=400, detail="Decision must be approve or reject")


# ==================== PRICE BUILDER SYSTEM ====================

class PriceBuilderRateUpdate(BaseModel):
    sor_library: Optional[str] = None
    granular_unit_net_rate: Optional[float] = None
    est_labour_value_per_uom: Optional[float] = None
    est_materials_other_value_per_uom: Optional[float] = None
    est_labour_minutes_per_uom: Optional[float] = None
    est_labour_hours_per_uom: Optional[float] = None
    labour_rate_used: Optional[float] = None
    labour_allocation_percent: Optional[float] = None
    split_confidence: Optional[str] = None
    split_note: Optional[str] = None
    pricing_split_note: Optional[str] = None
    line_type: Optional[str] = None
    change_reason: str = ""

class PriceBuilderSORCreate(BaseModel):
    sor_library: str = "Building Services"
    granular_sor_code: str
    parent_nhf_code: str = "LDA"
    job_type: str = "Manual / Specialist"
    trade_required: str = "Specialist"
    description: str
    uom: str = "item"
    granular_unit_net_rate: float = 0.0
    vat_rate: float = 0.2
    line_type: str = "Priceable LDA item"
    labour_rate_used: float = 37.0
    est_labour_value_per_uom: float = 0.0
    est_materials_other_value_per_uom: float = 0.0
    est_labour_hours_per_uom: Optional[float] = None
    est_labour_minutes_per_uom: Optional[float] = None
    split_confidence: str = "Medium"
    split_note: str = "Manual LDA SOR item created from Price Builder."
    pricing_split_note: str = "Manual LDA item. Rate to be reviewed against supplier/subcontractor quote where applicable."
    change_reason: str = "New SOR created from Price Builder"

class BOQMatchRow(BaseModel):
    description: str = ""
    quantity: float = 1.0
    uom: str = "item"

class PriceBuildPayload(BaseModel):
    build_name: str = "New Price Build"
    quote_reference: str = ""
    project_id: str = ""
    status: str = "Draft"
    lines: List[Dict[str, Any]] = []


def pb_now_iso() -> str:
    return datetime.utcnow().isoformat()


def pb_clean(value: Any) -> str:
    if value is None:
        return ""
    try:
        if isinstance(value, float) and math.isnan(value):
            return ""
    except Exception:
        pass
    return str(value).strip()


def pb_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        if isinstance(value, str):
            value = value.replace("£", "").replace(",", "").replace("%", "").strip()
            if value == "":
                return default
        number_value = float(value)
        if math.isnan(number_value):
            return default
        return number_value
    except Exception:
        return default


def pb_money(value: Any, default: float = 0.0) -> float:
    return round(pb_float(value, default), 2)


def pb_code(value: Any) -> str:
    text = pb_clean(value)
    if text.endswith(".0") and text.replace(".0", "").isdigit():
        text = text[:-2]
    return text.replace(" ", "")


def pb_doc(doc: Dict[str, Any]) -> Dict[str, Any]:
    if not doc:
        return {}
    output = {}
    for key, value in doc.items():
        if key == "_id":
            output["id"] = str(value)
        elif isinstance(value, datetime):
            output[key] = value.isoformat()
        else:
            output[key] = value
    return output


def pb_line_totals(line: Dict[str, Any]) -> Dict[str, Any]:
    quantity = pb_float(line.get("quantity"), 0)
    library_rate = pb_money(line.get("library_unit_rate"), 0)
    override = line.get("override_unit_rate")
    effective_rate = pb_money(override, library_rate) if override not in [None, ""] else library_rate
    vat_rate = pb_float(line.get("vat_rate"), 0.2)
    net = round(quantity * effective_rate, 2)
    vat = round(net * vat_rate, 2)
    gross = round(net + vat, 2)

    labour_each = pb_money(line.get("est_labour_value_per_uom"), 0)
    materials_each = pb_money(line.get("est_materials_other_value_per_uom"), 0)
    imported_split_total = labour_each + materials_each
    scale = effective_rate / imported_split_total if imported_split_total > 0 else 1

    line["effective_unit_rate"] = effective_rate
    line["net_amount"] = net
    line["vat_amount"] = vat
    line["gross_amount"] = gross
    line["est_labour_total"] = round(quantity * labour_each * scale, 2)
    line["est_materials_other_total"] = round(quantity * materials_each * scale, 2)
    line["est_labour_hours_total"] = round(quantity * pb_float(line.get("est_labour_hours_per_uom"), 0), 2)
    return line


def pb_totals(lines: List[Dict[str, Any]]) -> Dict[str, Any]:
    return {
        "net_total": round(sum(pb_money(line.get("net_amount")) for line in lines), 2),
        "vat_total": round(sum(pb_money(line.get("vat_amount")) for line in lines), 2),
        "gross_total": round(sum(pb_money(line.get("gross_amount")) for line in lines), 2),
        "est_labour_total": round(sum(pb_money(line.get("est_labour_total")) for line in lines), 2),
        "est_materials_other_total": round(sum(pb_money(line.get("est_materials_other_total")) for line in lines), 2),
        "est_labour_hours_total": round(sum(pb_float(line.get("est_labour_hours_total")) for line in lines), 2),
        "line_count": len(lines),
    }


SOR_LIBRARIES = ["NHF Rates", "Building Services", "Construction"]
DEFAULT_SOR_LIBRARY = "NHF Rates"

def normalise_sor_library(value: Optional[str]) -> str:
    requested = str(value or DEFAULT_SOR_LIBRARY).strip()
    if requested.lower() in ["all", "all libraries", "all sor libraries"]:
        return "All Libraries"
    for library in SOR_LIBRARIES:
        if requested.lower() == library.lower():
            return library
    return DEFAULT_SOR_LIBRARY

def sor_library_query(value: Optional[str]) -> Dict[str, Any]:
    library = normalise_sor_library(value)
    if library == "All Libraries":
        return {}
    return {"sor_library": library}

async def backfill_sor_library(default_library: str = DEFAULT_SOR_LIBRARY):
    try:
        await db.sor_rates.update_many(
            {"$or": [{"sor_library": {"$exists": False}}, {"sor_library": ""}, {"sor_library": None}]},
            {"$set": {"sor_library": default_library}},
        )
    except Exception as exc:
        logger.warning("Could not backfill SOR library values: %s", exc)

async def ensure_price_builder_indexes():
    try:
        await backfill_sor_library()
        # Earlier versions used a single unique index on granular_sor_code.  We now allow
        # the same SOR code to exist in different libraries, so drop that legacy index if present.
        try:
            await db.sor_rates.drop_index("granular_sor_code_1")
        except Exception:
            pass
        await db.sor_rates.create_index([("sor_library", 1), ("granular_sor_code", 1)], unique=True)
        await db.sor_rates.create_index("sor_library")
        await db.sor_rates.create_index("parent_nhf_code")
        await db.sor_rates.create_index("job_type")
        await db.sor_rates.create_index("trade_required")
        await db.sor_rates.create_index("line_type")
        await db.price_builds.create_index("id", unique=True)
        await db.price_builds.create_index("project_id")
        await db.price_builds.create_index("created_at")
    except Exception as exc:
        logger.warning("Could not create price builder indexes: %s", exc)


@api_router.post("/price-builder/import-sor-library")
async def import_price_builder_sor_library(
    file: UploadFile = File(...),
    replace_existing: bool = Query(False),
    sor_library: str = Query(DEFAULT_SOR_LIBRARY),
    admin: str = Depends(verify_admin),
):
    """Import an SOR library workbook/CSV into MongoDB.

    Supports both:
    - NHF/LDA granular workbook column names, e.g. Granular SOR Code / Granular Task Description
    - Construction/Building Services SOR exports, e.g. SOR Code / Description / Category / Section

    If replace_existing=true, only the selected sor_library is replaced.
    """
    filename = (file.filename or "").lower()
    if not filename.endswith((".xlsx", ".xls", ".csv")):
        raise HTTPException(status_code=400, detail="Upload an Excel workbook (.xlsx/.xls) or CSV file.")

    try:
        import pandas as pd
        from io import BytesIO
        content = await file.read()

        if filename.endswith(".csv"):
            frame = pd.read_csv(BytesIO(content))
        else:
            excel_file = pd.ExcelFile(BytesIO(content))
            # Prefer SOR Rates, but fall back to the first sheet so Construction/BS exports import cleanly.
            sheet_name = "SOR Rates" if "SOR Rates" in excel_file.sheet_names else excel_file.sheet_names[0]
            frame = pd.read_excel(excel_file, sheet_name=sheet_name)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read uploaded SOR file: {exc}")

    def first_value(row, *names):
        for name in names:
            if name in row.index:
                value = row.get(name)
                if value is not None and str(value).strip().lower() not in ["", "nan", "none"]:
                    return value
        return ""

    def first_money(row, *names):
        return pb_money(first_value(row, *names))

    def first_float(row, *names, default=0):
        value = first_value(row, *names)
        if value == "":
            return default
        return pb_float(value, default)

    await ensure_price_builder_indexes()
    selected_library = normalise_sor_library(sor_library)
    if selected_library == "All Libraries":
        selected_library = DEFAULT_SOR_LIBRARY

    if replace_existing:
        await db.sor_rates.delete_many({"sor_library": selected_library})

    imported = 0
    skipped = 0
    now = pb_now_iso()

    for _, row in frame.iterrows():
        sor_code = pb_code(first_value(row, "Granular SOR Code", "SOR Code", "Code", "Item Code", "Rate Code"))
        description = pb_clean(first_value(row, "Granular Task Description", "Description", "Item Description", "BOQ Description", "Line Item Description"))

        if not sor_code or not description:
            skipped += 1
            continue

        row_library = pb_clean(first_value(row, "SOR Library", "Library"))
        row_library = normalise_sor_library(row_library) if row_library else selected_library
        if row_library == "All Libraries":
            row_library = selected_library

        unit_rate = first_money(row, "Granular Unit Net Rate", "Unit Net Rate", "Rate", "Unit Rate", "Net Rate")
        labour_value = first_money(row, "Est. Labour Value / UOM", "Labour Value / UOM", "Labour Value", "Labour")
        materials_value = first_money(row, "Est. Materials/Other Value / UOM", "Materials/Other Value / UOM", "Materials / Other Value", "Materials Value", "Materials")
        labour_hours = first_float(row, "Est. Labour Hours / UOM", "Labour Hours / UOM", "Estimated Labour Hours", "Labour Hours")
        labour_minutes = first_float(row, "Est. Labour Minutes / UOM", "Labour Minutes / UOM", "Estimated Labour Minutes", "Labour Minutes")
        labour_rate = first_money(row, "Labour Rate Used (£/hr)", "Labour Rate", "Trade Rate")
        labour_percent = first_float(row, "Labour Allocation %", "Labour %", "Labour Percent")

        if labour_value == 0 and materials_value == 0 and unit_rate > 0:
            labour_percent_for_calc = labour_percent if labour_percent > 0 else 0.60
            if labour_percent_for_calc > 1:
                labour_percent_for_calc = labour_percent_for_calc / 100
            labour_value = round(unit_rate * labour_percent_for_calc, 2)
            materials_value = round(unit_rate - labour_value, 2)

        if labour_rate == 0:
            labour_rate = 37.0

        if labour_hours == 0 and labour_value > 0 and labour_rate > 0:
            labour_hours = round(labour_value / labour_rate, 4)

        if labour_minutes == 0 and labour_hours > 0:
            labour_minutes = round(labour_hours * 60, 1)

        uom = pb_clean(first_value(row, "UOM", "Unit", "Unit of Measure")) or "item"
        line_type = pb_clean(first_value(row, "Line Type", "Rate Mode", "Pricing Mode")) or "Rate"
        category = pb_clean(first_value(row, "Job Type", "Category / Section", "Category", "Section"))
        subcategory = pb_clean(first_value(row, "Subcategory", "Sub Category"))
        trade_required = pb_clean(first_value(row, "Trade Required", "Trade", "Default Trade")) or "Construction Operative"
        source_file = pb_clean(first_value(row, "Source File", "Source Workbook"))
        source_sheet = pb_clean(first_value(row, "Source Sheet", "Worksheet"))
        source_row = pb_clean(first_value(row, "Source Row", "Source Ref", "Row Ref"))

        if source_file or source_sheet or source_row:
            source_row_text = " > ".join([part for part in [source_file, source_sheet, source_row] if part])
        else:
            source_row_text = pb_clean(first_value(row, "Source Row"))

        doc = {
            "sor_library": row_library,
            "parent_nhf_code": pb_code(first_value(row, "Parent NHF Code", "Parent Code", "Parent SOR Code")),
            "granular_sor_code": sor_code,
            "job_type": category,
            "subcategory": subcategory,
            "trade_required": trade_required,
            "description": description,
            "uom": uom,
            "source_quantity_basis": first_float(row, "Source Quantity Basis", "Source Quantity", "Quantity", "Qty"),
            "nhf_baseline_rate": first_money(row, "NHF Baseline Rate", "Baseline Rate"),
            "split_percent": first_float(row, "Split %", "Split Percent"),
            "granular_unit_net_rate": unit_rate,
            "vat_rate": first_float(row, "VAT Rate", "VAT", default=0.2) or 0.2,
            "line_type": line_type,
            "parent_price_item": pb_clean(first_value(row, "Parent Price Item")),
            "original_nhf_description": pb_clean(first_value(row, "Original NHF Description")),
            "pricing_split_note": pb_clean(first_value(row, "Pricing / Split Note", "Pricing Note", "Notes", "Build-Up Note")),
            "source_row": source_row_text,
            "source_file": source_file,
            "source_sheet": source_sheet,
            "labour_rate_used": labour_rate,
            "labour_allocation_percent": labour_percent,
            "est_labour_value_per_uom": labour_value,
            "est_materials_other_value_per_uom": materials_value,
            "est_labour_hours_per_uom": labour_hours,
            "est_labour_minutes_per_uom": labour_minutes,
            "split_confidence": pb_clean(first_value(row, "Split Confidence", "Confidence", "Pricing Confidence")),
            "split_note": pb_clean(first_value(row, "Split Note", "Estimation Note", "Rate Note")),
            "selection_label": f"{description} — {uom} — {line_type}",
            "archived": bool(first_float(row, "Archived", default=0)),
            "updated_at": now,
        }

        await db.sor_rates.update_one(
            {"sor_library": row_library, "granular_sor_code": sor_code},
            {"$set": doc, "$setOnInsert": {"created_at": now}},
            upsert=True,
        )
        imported += 1

    return {
        "success": True,
        "sor_library": selected_library,
        "imported_or_updated": imported,
        "skipped": skipped,
        "replace_existing": replace_existing,
        "columns_seen": list(frame.columns),
    }


@api_router.get("/price-builder/filters")
async def get_price_builder_filters(sor_library: str = Query("All Libraries"), admin: str = Depends(verify_admin)):
    await ensure_price_builder_indexes()
    query = sor_library_query(sor_library)
    existing_libraries = sorted([value for value in await db.sor_rates.distinct("sor_library") if value])
    libraries = []
    for library in SOR_LIBRARIES + existing_libraries:
        if library and library not in libraries:
            libraries.append(library)
    return {
        "libraries": libraries,
        "job_types": sorted([value for value in await db.sor_rates.distinct("job_type", query) if value]),
        "trades": sorted([value for value in await db.sor_rates.distinct("trade_required", query) if value]),
        "line_types": sorted([value for value in await db.sor_rates.distinct("line_type", query) if value]),
        "uoms": sorted([value for value in await db.sor_rates.distinct("uom", query) if value]),
    }


@api_router.get("/price-builder/rates")
async def search_price_builder_rates(
    search: str = "",
    job_type: str = "",
    trade: str = "",
    line_type: str = "",
    sor_library: str = Query("All Libraries"),
    include_non_priceable: bool = False,
    include_archived: bool = False,
    limit: int = Query(50, ge=1, le=200),
    skip: int = Query(0, ge=0),
    admin: str = Depends(verify_admin),
):
    await ensure_price_builder_indexes()
    query: Dict[str, Any] = sor_library_query(sor_library)
    if not include_archived:
        query["archived"] = {"$ne": True}
    search_text = search.strip()
    exact_code_search = bool(re.fullmatch(r"[A-Za-z0-9 ._-]+", search_text)) and bool(re.search(r"\d", search_text))

    if job_type:
        query["job_type"] = job_type
    if trade:
        query["trade_required"] = trade
    if line_type:
        query["line_type"] = line_type
    elif not include_non_priceable and not exact_code_search:
        query["line_type"] = {"$regex": "priceable|single|parent", "$options": "i"}

    if search_text:
        escaped = re.escape(search_text.replace(" ", ""))
        loose = re.escape(search_text)
        query["$or"] = [
            {"granular_sor_code": {"$regex": escaped, "$options": "i"}},
            {"parent_nhf_code": {"$regex": escaped, "$options": "i"}},
            {"description": {"$regex": loose, "$options": "i"}},
            {"original_nhf_description": {"$regex": loose, "$options": "i"}},
            {"parent_price_item": {"$regex": loose, "$options": "i"}},
        ]

    rates = [pb_doc(doc) async for doc in db.sor_rates.find(query).sort([("job_type", 1), ("description", 1)]).skip(skip).limit(limit)]
    total = await db.sor_rates.count_documents(query)
    return {"rates": rates, "total": total, "skip": skip, "limit": limit}


@api_router.post("/price-builder/rates")
async def create_price_builder_sor_rate(payload: PriceBuilderSORCreate, admin_user: Dict[str, Any] = Depends(get_super_admin_user)):
    """Create a new reusable LDA SOR item. Super Admin only."""
    await ensure_price_builder_indexes()
    selected_library = normalise_sor_library(payload.sor_library)
    if selected_library == "All Libraries":
        selected_library = "Building Services"
    sor_code = pb_code(payload.granular_sor_code)
    if not sor_code:
        raise HTTPException(status_code=400, detail="SOR code is required")
    if not pb_clean(payload.description):
        raise HTTPException(status_code=400, detail="Description is required")

    existing = await db.sor_rates.find_one({"sor_library": selected_library, "granular_sor_code": sor_code})
    if existing:
        raise HTTPException(status_code=400, detail=f"A SOR item with this code already exists in {selected_library}")

    unit_rate = pb_money(payload.granular_unit_net_rate)
    labour_value = pb_money(payload.est_labour_value_per_uom)
    materials_value = pb_money(payload.est_materials_other_value_per_uom)
    labour_rate = pb_money(payload.labour_rate_used, 37.0) or 37.0

    labour_minutes = pb_float(payload.est_labour_minutes_per_uom)
    labour_hours = pb_float(payload.est_labour_hours_per_uom)
    if labour_hours == 0 and labour_minutes > 0:
        labour_hours = round(labour_minutes / 60, 4)
    if labour_minutes == 0 and labour_hours > 0:
        labour_minutes = round(labour_hours * 60, 1)
    if labour_hours == 0 and labour_value > 0 and labour_rate > 0:
        labour_hours = round(labour_value / labour_rate, 4)
        labour_minutes = round(labour_hours * 60, 1)

    if unit_rate > 0 and labour_value == 0 and materials_value == 0:
        labour_value = round(unit_rate * 0.45, 2)
        materials_value = round(unit_rate - labour_value, 2)

    now = pb_now_iso()
    admin = admin_user.get("email") or admin_user.get("name") or "Super Admin"
    doc = {
        "sor_library": selected_library,
        "parent_nhf_code": pb_code(payload.parent_nhf_code) or "LDA",
        "granular_sor_code": sor_code,
        "job_type": pb_clean(payload.job_type) or "Manual / Specialist",
        "trade_required": pb_clean(payload.trade_required) or "Specialist",
        "description": pb_clean(payload.description),
        "uom": pb_clean(payload.uom) or "item",
        "source_quantity_basis": 1,
        "nhf_baseline_rate": unit_rate,
        "split_percent": 1,
        "granular_unit_net_rate": unit_rate,
        "vat_rate": pb_float(payload.vat_rate, 0.2) or 0.2,
        "line_type": pb_clean(payload.line_type) or "Priceable LDA item",
        "parent_price_item": "Manual LDA SOR",
        "original_nhf_description": "Manual LDA SOR",
        "pricing_split_note": pb_clean(payload.pricing_split_note),
        "source_row": "Price Builder manual SOR",
        "labour_rate_used": labour_rate,
        "labour_allocation_percent": round((labour_value / unit_rate) * 100, 2) if unit_rate > 0 else 0,
        "est_labour_value_per_uom": labour_value,
        "est_materials_other_value_per_uom": materials_value,
        "est_labour_hours_per_uom": labour_hours,
        "est_labour_minutes_per_uom": labour_minutes,
        "split_confidence": pb_clean(payload.split_confidence) or "Medium",
        "split_note": pb_clean(payload.split_note),
        "selection_label": f"{pb_clean(payload.description)} — {pb_clean(payload.uom) or 'item'} — {pb_clean(payload.line_type) or 'Rate'}",
        "created_at": now,
        "updated_at": now,
        "created_by": admin,
        "updated_by": admin,
        "archived": False,
        "rate_history": [{
            "id": str(uuid.uuid4()),
            "changed_at": now,
            "changed_by": admin,
            "reason": payload.change_reason or "New SOR created from Price Builder",
            "old": {},
            "new": {"granular_sor_code": sor_code, "granular_unit_net_rate": unit_rate},
        }],
    }
    await db.sor_rates.insert_one(doc)
    return {"success": True, "rate": pb_doc(doc)}


def pb_simple_tokens(text: str) -> List[str]:
    words = re.findall(r"[a-zA-Z0-9]+", (text or "").lower())
    stop = {"and", "or", "the", "to", "of", "for", "with", "in", "on", "inc", "including", "supply", "install", "new", "renew", "replace", "item", "nr", "no"}
    return [w for w in words if len(w) > 2 and w not in stop]


def pb_match_confidence(description: str, rate_doc: Dict[str, Any]) -> int:
    boq_tokens = set(pb_simple_tokens(description))
    rate_tokens = set(pb_simple_tokens(" ".join([
        rate_doc.get("description", ""),
        rate_doc.get("job_type", ""),
        rate_doc.get("trade_required", ""),
        rate_doc.get("original_nhf_description", ""),
    ])))
    if not boq_tokens or not rate_tokens:
        return 0
    overlap = boq_tokens.intersection(rate_tokens)
    score = int(round((len(overlap) / max(len(boq_tokens), 1)) * 100))
    if description and rate_doc.get("description") and description.lower() in rate_doc.get("description", "").lower():
        score = max(score, 90)
    return min(score, 99)


async def pb_match_one_boq_line(description: str, uom: str = "", limit: int = 5, sor_library: str = "All Libraries") -> List[Dict[str, Any]]:
    tokens = pb_simple_tokens(description)
    query: Dict[str, Any] = {"archived": {"$ne": True}, **sor_library_query(sor_library)}
    if tokens:
        important = tokens[:8]
        query["$or"] = []
        for token in important:
            query["$or"].extend([
                {"description": {"$regex": re.escape(token), "$options": "i"}},
                {"original_nhf_description": {"$regex": re.escape(token), "$options": "i"}},
                {"job_type": {"$regex": re.escape(token), "$options": "i"}},
                {"trade_required": {"$regex": re.escape(token), "$options": "i"}},
            ])
    if uom:
        # do not force UOM as BOQs vary between nr/each/item, but it helps ordering later
        pass
    docs = await db.sor_rates.find(query, {"_id": 0}).limit(80).to_list(80)
    scored = []
    for doc in docs:
        score = pb_match_confidence(description, doc)
        if uom and str(doc.get("uom", "")).lower() == str(uom).lower():
            score = min(score + 8, 99)
        if score > 0:
            out = pb_doc(doc)
            out["match_confidence"] = score
            scored.append(out)
    scored.sort(key=lambda item: (item.get("match_confidence", 0), item.get("granular_unit_net_rate", 0)), reverse=True)
    return scored[:limit]


@api_router.post("/price-builder/match-boq")
async def match_price_builder_boq(file: UploadFile = File(...), sor_library: str = Query("All Libraries"), admin: str = Depends(verify_admin)):
    """Upload a CSV/XLSX BOQ and return likely SOR matches for each row."""
    filename = (file.filename or "").lower()
    content = await file.read()
    rows: List[Dict[str, Any]] = []

    try:
        if filename.endswith((".xlsx", ".xls")):
            import pandas as pd
            from io import BytesIO
            frame = pd.read_excel(BytesIO(content))
            raw_rows = frame.to_dict("records")
        elif filename.endswith(".csv"):
            text = content.decode("utf-8-sig", errors="ignore")
            raw_rows = list(csv.DictReader(io.StringIO(text)))
        else:
            raise HTTPException(status_code=400, detail="Upload a BOQ as .xlsx, .xls or .csv")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read BOQ file: {exc}")

    def pick(row: Dict[str, Any], names: List[str]) -> Any:
        lookup = {str(k).strip().lower(): v for k, v in row.items()}
        for name in names:
            if name.lower() in lookup:
                return lookup[name.lower()]
        for key, value in lookup.items():
            if any(name.lower() in key for name in names):
                return value
        return ""

    for idx, row in enumerate(raw_rows[:500]):
        description = pb_clean(pick(row, ["description", "item", "works", "scope", "boq description", "task", "name"]))
        if not description:
            values = [pb_clean(v) for v in row.values() if pb_clean(v)]
            description = values[0] if values else ""
        if not description:
            continue
        quantity = pb_float(pick(row, ["qty", "quantity", "quant", "amount"]), 1) or 1
        uom = pb_clean(pick(row, ["uom", "unit", "units", "measure"])) or "item"
        matches = await pb_match_one_boq_line(description, uom, limit=5, sor_library=sor_library)
        rows.append({
            "id": str(uuid.uuid4()),
            "source_row": idx + 1,
            "description": description,
            "quantity": quantity,
            "uom": uom,
            "matches": matches,
        })

    return {"success": True, "rows": rows, "row_count": len(rows)}


@api_router.get("/price-builder/rates/{sor_code}")
async def get_price_builder_rate(sor_code: str, sor_library: str = Query("All Libraries"), admin: str = Depends(verify_admin)):
    await ensure_price_builder_indexes()
    query = {"granular_sor_code": sor_code, **sor_library_query(sor_library)}
    rate = await db.sor_rates.find_one(query)
    if not rate:
        raise HTTPException(status_code=404, detail="SOR rate not found")
    siblings = []
    parent_code = rate.get("parent_nhf_code")
    if parent_code:
        sibling_query = {"parent_nhf_code": parent_code}
        if rate.get("sor_library"):
            sibling_query["sor_library"] = rate.get("sor_library")
        siblings = [pb_doc(doc) async for doc in db.sor_rates.find(sibling_query).sort("granular_sor_code", 1)]
    return {"rate": pb_doc(rate), "siblings": siblings}


@api_router.put("/price-builder/rates/{sor_code}")
async def update_price_builder_master_rate(sor_code: str, payload: PriceBuilderRateUpdate, sor_library: str = Query("All Libraries"), admin_user: Dict[str, Any] = Depends(get_super_admin_user)):
    await ensure_price_builder_indexes()
    selected_library = normalise_sor_library(payload.sor_library or sor_library)
    query = {"granular_sor_code": sor_code, **sor_library_query(selected_library)}
    existing = await db.sor_rates.find_one(query)
    if not existing:
        raise HTTPException(status_code=404, detail="SOR rate not found")

    update_data: Dict[str, Any] = {}
    numeric_fields = [
        "granular_unit_net_rate",
        "est_labour_value_per_uom",
        "est_materials_other_value_per_uom",
        "est_labour_minutes_per_uom",
        "est_labour_hours_per_uom",
        "labour_rate_used",
        "labour_allocation_percent",
    ]
    payload_dict = payload.dict()
    for field in numeric_fields:
        if payload_dict.get(field) is not None:
            update_data[field] = pb_float(payload_dict.get(field))

    if "est_labour_minutes_per_uom" in update_data and "est_labour_hours_per_uom" not in update_data:
        update_data["est_labour_hours_per_uom"] = round(update_data["est_labour_minutes_per_uom"] / 60, 4)
    if "est_labour_hours_per_uom" in update_data and "est_labour_minutes_per_uom" not in update_data:
        update_data["est_labour_minutes_per_uom"] = round(update_data["est_labour_hours_per_uom"] * 60, 1)

    for field in ["split_confidence", "split_note", "pricing_split_note", "line_type"]:
        if payload_dict.get(field) is not None:
            update_data[field] = pb_clean(payload_dict.get(field))

    if "granular_unit_net_rate" in update_data and update_data["granular_unit_net_rate"] > 0:
        labour_value = update_data.get("est_labour_value_per_uom", existing.get("est_labour_value_per_uom", 0))
        update_data["labour_allocation_percent"] = round((pb_float(labour_value) / update_data["granular_unit_net_rate"]) * 100, 2)

    update_data["updated_at"] = pb_now_iso()
    admin = admin_user.get("email") or admin_user.get("name") or "Super Admin"
    update_data["updated_by"] = admin

    audit_entry = {
        "id": str(uuid.uuid4()),
        "changed_at": pb_now_iso(),
        "changed_by": admin,
        "reason": payload.change_reason or "Master rate updated from Price Builder",
        "old": {key: existing.get(key) for key in update_data.keys() if key not in ["updated_at", "updated_by"]},
        "new": {key: update_data.get(key) for key in update_data.keys() if key not in ["updated_at", "updated_by"]},
    }

    await db.sor_rates.update_one(
        query,
        {"$set": update_data, "$push": {"rate_history": audit_entry}},
    )
    updated = await db.sor_rates.find_one(query)
    return {"success": True, "rate": pb_doc(updated)}


@api_router.post("/price-builds")
async def create_price_build(payload: PriceBuildPayload, admin: str = Depends(verify_admin)):
    await ensure_price_builder_indexes()
    lines = [pb_line_totals(dict(line)) for line in payload.lines]
    doc = payload.dict()
    doc["id"] = str(uuid.uuid4())
    doc["lines"] = lines
    doc["totals"] = pb_totals(lines)
    doc["created_by"] = admin
    doc["created_at"] = pb_now_iso()
    doc["updated_at"] = pb_now_iso()
    await db.price_builds.insert_one(doc)
    return pb_doc(doc)


@api_router.get("/price-builds")
async def list_price_builds(project_id: str = "", limit: int = Query(50, ge=1, le=200), admin: str = Depends(verify_admin)):
    await ensure_price_builder_indexes()
    query = {"project_id": project_id} if project_id else {}
    builds = [pb_doc(doc) async for doc in db.price_builds.find(query).sort("updated_at", -1).limit(limit)]
    return {"builds": builds}


@api_router.get("/price-builds/{build_id}")
async def get_price_build(build_id: str, admin: str = Depends(verify_admin)):
    doc = await db.price_builds.find_one({"id": build_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Price build not found")
    return pb_doc(doc)


@api_router.put("/price-builds/{build_id}")
async def update_price_build(build_id: str, payload: PriceBuildPayload, admin: str = Depends(verify_admin)):
    lines = [pb_line_totals(dict(line)) for line in payload.lines]
    update_data = payload.dict()
    update_data["lines"] = lines
    update_data["totals"] = pb_totals(lines)
    update_data["updated_by"] = admin
    update_data["updated_at"] = pb_now_iso()
    result = await db.price_builds.update_one({"id": build_id}, {"$set": update_data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Price build not found")
    doc = await db.price_builds.find_one({"id": build_id})
    return pb_doc(doc)


@api_router.delete("/price-builds/{build_id}")
async def delete_price_build(build_id: str, admin: str = Depends(verify_admin)):
    result = await db.price_builds.delete_one({"id": build_id})
    return {"success": result.deleted_count == 1}



# ==================== CROSS-PROCESS SYNC HEALTH CHECKS ====================

def _date_in_range(value: Optional[str], start_value: Optional[str], end_value: Optional[str]) -> bool:
    value_date = parse_iso_date_safe(value)
    start_date = parse_iso_date_safe(start_value)
    end_date = parse_iso_date_safe(end_value)
    if not value_date or not start_date or not end_date:
        return True
    if end_date < start_date:
        start_date, end_date = end_date, start_date
    return start_date <= value_date <= end_date


def _date_after(value: Optional[str], compare_value: Optional[str]) -> bool:
    value_date = parse_iso_date_safe(value)
    compare_date = parse_iso_date_safe(compare_value)
    return bool(value_date and compare_date and value_date > compare_date)


def _date_before(value: Optional[str], compare_value: Optional[str]) -> bool:
    value_date = parse_iso_date_safe(value)
    compare_date = parse_iso_date_safe(compare_value)
    return bool(value_date and compare_date and value_date < compare_date)


def _safe_int(value: Any, fallback: int = 0) -> int:
    try:
        if value is None or value == "":
            return fallback
        return int(round(float(value)))
    except Exception:
        return fallback


def _iso_add_days(value: Optional[str], days: int) -> Optional[str]:
    if not value:
        return None
    parsed = parse_iso_date_safe(value)
    if not parsed:
        return None
    return (parsed + timedelta(days=days)).isoformat()


def _po_required_date(po: Dict[str, Any]) -> Optional[str]:
    for key in ["required_date", "material_required_date", "delivery_date", "expected_delivery_date"]:
        value = po.get(key)
        if parse_iso_date_safe(value):
            return str(value)[:10]
    return None


def _calculate_po_order_by_date(po: Dict[str, Any]) -> Optional[str]:
    required_date = _po_required_date(po)
    if not required_date:
        return po.get("order_by_date")
    lead_time_days = max(0, _safe_int(po.get("lead_time_days"), 0))
    return _iso_add_days(required_date, -lead_time_days) if lead_time_days else (po.get("order_by_date") or required_date)


def _calculate_po_expected_payment_date(po: Dict[str, Any]) -> Optional[str]:
    """Return the cashflow date a PO should use when no actual paid date exists.

    Explicit payment_due_date is treated as a manual override. Otherwise:
    - proforma = order-by date / created date
    - immediate = approved/sent/created date
    - credit terms = required date + payment_terms_days
    """
    manual_due = po.get("payment_due_date") or po.get("invoice_due_date")
    if parse_iso_date_safe(manual_due):
        return str(manual_due)[:10]

    requirement = normalise_supplier_payment_requirement(po.get("payment_requirement") or "credit_terms")
    required_date = _po_required_date(po)

    if requirement == "proforma":
        return _calculate_po_order_by_date(po) or finance_material_iso(po.get("created_at")) or required_date

    if requirement == "immediate":
        return finance_material_iso(po.get("sent_at") or po.get("approved_at") or po.get("created_at")) or required_date

    if required_date:
        terms_days = max(0, _safe_int(po.get("payment_terms_days"), 30))
        return _iso_add_days(required_date, terms_days)

    return finance_material_iso(po.get("expected_payment_date") or po.get("created_at"))


def _normalise_po_planning_fields(po: Dict[str, Any], supplier: Optional[Dict[str, Any]] = None, preserve_manual_due: bool = True) -> Dict[str, Any]:
    """Return a patch that makes PO planning/cashflow fields consistent.

    This does not mark the PO as sent/approved/paid; it only fills planning fields used by
    Project Management and the Finance cashflow forecast.
    """
    supplier = supplier or {}
    patch: Dict[str, Any] = {}

    requirement = normalise_supplier_payment_requirement(po.get("payment_requirement") or supplier.get("payment_requirement") or "credit_terms")
    terms_days = max(0, _safe_int(po.get("payment_terms_days"), _safe_int(supplier.get("payment_terms_days"), 30)))
    lead_time_days = max(0, _safe_int(po.get("lead_time_days"), _safe_int(supplier.get("lead_time_days"), 0)))

    patch["payment_requirement"] = requirement
    patch["payment_terms_days"] = terms_days
    patch["lead_time_days"] = lead_time_days

    working = {**po, **patch}
    order_by_date = _calculate_po_order_by_date(working)
    if order_by_date:
        patch["order_by_date"] = order_by_date
        working["order_by_date"] = order_by_date

    expected_payment_date = _calculate_po_expected_payment_date(working)
    if expected_payment_date:
        patch["expected_payment_date"] = expected_payment_date

    # Preserve explicit manual payment_due_date. If none exists, leave the manual field blank and
    # store the calculated cashflow date in expected_payment_date.
    if preserve_manual_due and po.get("payment_due_date"):
        patch["payment_due_date"] = po.get("payment_due_date")

    return patch


def _check_severity_rank(severity: str) -> int:
    return {"critical": 3, "warning": 2, "info": 1, "ok": 0}.get(str(severity or "").lower(), 0)


def _sync_check(severity: str, category: str, title: str, detail: str = "", action: str = "", record_type: str = "", record_id: str = "") -> Dict[str, Any]:
    return {
        "severity": severity,
        "category": category,
        "title": title,
        "detail": detail,
        "action": action,
        "record_type": record_type,
        "record_id": record_id,
    }


@api_router.get("/process-sync/job/{job_id}")
async def get_job_process_sync_health(job_id: str, admin: str = Depends(verify_admin)):
    """Review whether linked project-management, schedule, PO and cashflow data is still aligned for a job.

    This is intentionally a read-only diagnostic endpoint. It does not change programme, schedule,
    purchase order or finance records. The Project Management page uses it before/after programme
    changes to spot downstream records that need attention.
    """
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    sections = [dict(section) for section in (job.get("gantt_sections") or [])]
    section_by_id = {str(section.get("id") or ""): section for section in sections if section.get("id")}
    section_by_name = {_normalise_link_value(section.get("name")): section for section in sections if section.get("name")}

    checks: List[Dict[str, Any]] = []

    schedule_entries = await db.schedule_entries.find({
        "job_id": job_id,
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    purchase_orders = await db.purchase_orders.find({
        "job_id": job_id,
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    finance_records = await db.finance_records.find({
        "job_id": job_id,
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    active_schedule_ids = {entry.get("id") for entry in schedule_entries if entry.get("id")}
    schedule_by_section: Dict[str, List[Dict[str, Any]]] = {}
    manual_schedule_count = 0

    for entry in schedule_entries:
        section_id = str(entry.get("project_section_id") or entry.get("linked_gantt_section_id") or "")
        section_name_key = _normalise_link_value(entry.get("project_section_name") or entry.get("notes"))
        section = section_by_id.get(section_id) or section_by_name.get(section_name_key)

        if section:
            schedule_by_section.setdefault(str(section.get("id") or section.get("name")), []).append(entry)
            if not _date_in_range(entry.get("scheduled_date"), section.get("start_date"), section.get("end_date")):
                checks.append(_sync_check(
                    "warning",
                    "Schedule",
                    "Schedule entry sits outside its Gantt section dates",
                    f"{entry.get('worker_name') or entry.get('worker_id') or 'Worker'} is scheduled on {entry.get('scheduled_date')} but section '{section.get('name', 'Section')}' runs {section.get('start_date') or '-'} to {section.get('end_date') or '-' }.",
                    "Review the section dates or move this schedule entry.",
                    "schedule_entry",
                    entry.get("id", ""),
                ))
        else:
            manual_schedule_count += 1
            if entry.get("schedule_link_mode") == "linked_to_section" or entry.get("project_section_id") or entry.get("linked_gantt_section_id"):
                checks.append(_sync_check(
                    "warning",
                    "Schedule",
                    "Schedule entry has a broken Gantt section link",
                    f"A schedule entry on {entry.get('scheduled_date')} points at a section that could not be found on this job.",
                    "Re-send the section to schedule or detach this entry as manual.",
                    "schedule_entry",
                    entry.get("id", ""),
                ))

    for section in sections:
        section_key = str(section.get("id") or section.get("name") or "")
        linked_entries = schedule_by_section.get(section_key, [])
        section_ids = set(_normalise_section_ids(section))
        missing_ids = [entry_id for entry_id in section_ids if entry_id not in active_schedule_ids]
        if section.get("sent_to_schedule") is True and not linked_entries and not section_ids:
            checks.append(_sync_check(
                "warning",
                "Gantt",
                "Section is marked scheduled but has no active linked schedule entries",
                f"'{section.get('name', 'Section')}' is green/scheduled on the Gantt but no active schedule entries were found.",
                "Re-send this section to schedule or reset its schedule status.",
                "gantt_section",
                section.get("id", ""),
            ))
        if missing_ids:
            checks.append(_sync_check(
                "info",
                "Gantt",
                "Section contains old schedule entry references",
                f"'{section.get('name', 'Section')}' still references {len(missing_ids)} schedule entry id(s) that are no longer active.",
                "This is usually harmless, but re-sending the section will refresh the links.",
                "gantt_section",
                section.get("id", ""),
            ))

    if manual_schedule_count and sections:
        checks.append(_sync_check(
            "info",
            "Schedule",
            "Some job schedule entries are manual rather than Gantt-linked",
            f"{manual_schedule_count} active schedule entr{'y is' if manual_schedule_count == 1 else 'ies are'} not linked to a Gantt section. These will still affect labour cashflow, but will not move automatically when a section moves.",
            "Leave as manual if intentional, or schedule from Project Management if they should follow section moves.",
            "schedule_entry",
            "",
        ))

    linked_po_count = 0
    unlinked_po_count = 0
    for po in purchase_orders:
        status = str(po.get("status") or "draft").lower().replace(" ", "_")
        if status in {"cancelled", "canceled", "void", "archived", "rejected"}:
            continue

        linked_section = next((section for section in sections if _po_links_to_section(po, section)), None)
        if linked_section:
            linked_po_count += 1
            required_date = po.get("required_date") or po.get("material_required_date") or po.get("delivery_date") or po.get("expected_delivery_date")
            if required_date and linked_section.get("end_date") and _date_after(required_date, linked_section.get("end_date")):
                checks.append(_sync_check(
                    "warning",
                    "Purchase Orders",
                    "PO required date is after the linked section ends",
                    f"{po.get('po_number') or 'PO'} / {po.get('supplier_name') or 'Supplier'} is required on {required_date}, after section '{linked_section.get('name', 'Section')}' ends on {linked_section.get('end_date')}.",
                    "Check whether the PO date should move with the section or whether the PO is linked to the correct section.",
                    "purchase_order",
                    po.get("id", ""),
                ))
            if po.get("order_by_date") and required_date and _date_after(po.get("order_by_date"), required_date):
                checks.append(_sync_check(
                    "critical",
                    "Purchase Orders",
                    "PO order-by date is after the required date",
                    f"{po.get('po_number') or 'PO'} has order-by date {po.get('order_by_date')} but required date {required_date}.",
                    "Set the order-by date before the required date.",
                    "purchase_order",
                    po.get("id", ""),
                ))
            if str(po.get("payment_requirement") or "credit_terms") == "credit_terms" and not po.get("payment_due_date") and not required_date:
                checks.append(_sync_check(
                    "warning",
                    "Cashflow",
                    "Credit terms PO has no required date for cashflow timing",
                    f"{po.get('po_number') or 'PO'} has credit terms but no required/delivery date, so finance may fall back to created/sent date.",
                    "Add a required date so the cashflow can calculate required date + payment terms.",
                    "purchase_order",
                    po.get("id", ""),
                ))
            if po.get("payment_due_date") and required_date and _date_before(po.get("payment_due_date"), required_date):
                checks.append(_sync_check(
                    "warning",
                    "Cashflow",
                    "PO payment due date is before materials are required",
                    f"{po.get('po_number') or 'PO'} has payment due date {po.get('payment_due_date')} before required date {required_date}.",
                    "Check whether this is a proforma/immediate payment or an incorrect manual override.",
                    "purchase_order",
                    po.get("id", ""),
                ))
        else:
            unlinked_po_count += 1

    if unlinked_po_count and sections:
        checks.append(_sync_check(
            "info",
            "Purchase Orders",
            "Some POs are not linked to a Gantt section",
            f"{unlinked_po_count} active PO(s) on this job are not linked to a section. They can still appear in finance, but section moves will not automatically update their planning dates.",
            "Link future PO requests from the section in Project Management where possible.",
            "purchase_order",
            "",
        ))

    project_start = job.get("planned_start_date") or job.get("start_date")
    project_end = job.get("planned_end_date") or job.get("end_date")
    markers = [normalise_finance_marker(marker) for marker in (job.get("commercial_markers") or [])]
    for marker in markers:
        marker_date = marker.get("date")
        if marker_date and project_start and project_end and not _date_in_range(marker_date, project_start, project_end):
            checks.append(_sync_check(
                "info",
                "Commercial Markers",
                "Commercial marker is outside the planned project dates",
                f"'{marker.get('label') or marker.get('type')}' is dated {marker_date}; project dates are {project_start} to {project_end}.",
                "Leave if intentional, otherwise move the marker or project dates.",
                "commercial_marker",
                marker.get("id", ""),
            ))

    marker_ids = {marker.get("id") for marker in markers if marker.get("id")}
    for record in finance_records:
        marker_id = record.get("application_marker_id")
        if marker_id and marker_id not in marker_ids:
            checks.append(_sync_check(
                "info",
                "Finance",
                "Finance record points at a missing commercial marker",
                f"Finance record '{record.get('label') or record.get('invoice_number') or record.get('type')}' references an application marker that is no longer on the Gantt.",
                "Leave if this is historical, otherwise re-link or archive the finance record.",
                "finance_record",
                record.get("id", ""),
            ))

    if not checks:
        checks.append(_sync_check(
            "ok",
            "Process Sync",
            "No cross-process issues found",
            "Gantt sections, linked schedules, purchase orders and commercial markers appear aligned for this job.",
            "No action needed.",
            "job",
            job_id,
        ))

    counts = {"critical": 0, "warning": 0, "info": 0, "ok": 0}
    for check in checks:
        severity = str(check.get("severity") or "info").lower()
        counts[severity] = counts.get(severity, 0) + 1

    checks.sort(key=lambda item: (-_check_severity_rank(item.get("severity", "")), item.get("category", ""), item.get("title", "")))

    return {
        "job": {
            "id": job.get("id"),
            "name": job.get("name"),
            "client": job.get("client", ""),
            "planned_start_date": project_start,
            "planned_end_date": project_end,
        },
        "summary": {
            "checks_total": len(checks),
            "critical": counts.get("critical", 0),
            "warning": counts.get("warning", 0),
            "info": counts.get("info", 0),
            "ok": counts.get("ok", 0),
            "sections": len(sections),
            "schedule_entries": len(schedule_entries),
            "linked_purchase_orders": linked_po_count,
            "unlinked_purchase_orders": unlinked_po_count,
            "finance_records": len(finance_records),
            "commercial_markers": len(markers),
        },
        "checks": checks,
    }



class ProcessSyncRepairRequest(BaseModel):
    repair_schedule_links: bool = True
    repair_section_schedule_status: bool = True
    repair_po_section_links: bool = True
    repair_po_cashflow_dates: bool = True


@api_router.post("/process-sync/job/{job_id}/repair")
async def repair_job_process_sync(job_id: str, request: ProcessSyncRepairRequest, admin: str = Depends(verify_admin)):
    """Apply safe cross-process repairs for one job.

    Safe repairs are deliberately limited to data-linking and calculated planning fields:
    - refresh Gantt section schedule_entry_ids / schedule_status from active schedule entries
    - backfill missing section IDs on legacy schedule entries that match by section name
    - backfill missing PO section IDs where the PO/line already names the section
    - backfill PO payment requirement, terms, order-by and expected payment dates

    It does not move historical records, overwrite manual payment_due_date values, approve/send POs,
    or change commercial marker dates.
    """
    job = await db.jobs.find_one({"id": job_id, "archived": {"$ne": True}}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    sections = [dict(section) for section in (job.get("gantt_sections") or [])]
    section_by_id = {str(section.get("id") or ""): section for section in sections if section.get("id")}
    section_by_name = {_normalise_link_value(section.get("name")): section for section in sections if section.get("name")}
    now = datetime.utcnow()
    repairs: List[Dict[str, Any]] = []

    schedule_entries = await db.schedule_entries.find({
        "job_id": job_id,
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    # 1) Backfill legacy schedule links and build a fresh section -> schedule-entry map.
    schedule_by_section_id: Dict[str, List[Dict[str, Any]]] = {str(section.get("id")): [] for section in sections if section.get("id")}
    if request.repair_schedule_links or request.repair_section_schedule_status:
        for entry in schedule_entries:
            section_id = str(entry.get("project_section_id") or entry.get("linked_gantt_section_id") or "")
            section_name_key = _normalise_link_value(entry.get("project_section_name") or entry.get("notes"))
            section = section_by_id.get(section_id) or section_by_name.get(section_name_key)
            if not section:
                continue

            real_section_id = str(section.get("id") or "")
            if real_section_id:
                schedule_by_section_id.setdefault(real_section_id, []).append(entry)

            needs_link_update = (
                request.repair_schedule_links
                and real_section_id
                and (
                    entry.get("project_section_id") != real_section_id
                    or entry.get("linked_gantt_section_id") != real_section_id
                    or entry.get("project_section_name") != section.get("name", "")
                    or entry.get("schedule_link_mode") not in ["linked_to_section", "manually_adjusted"]
                )
            )
            if needs_link_update:
                await db.schedule_entries.update_one(
                    {"id": entry.get("id")},
                    {"$set": {
                        "project_section_id": real_section_id,
                        "linked_gantt_section_id": real_section_id,
                        "project_section_name": section.get("name", ""),
                        "schedule_link_mode": "linked_to_section",
                        "updated_date": now,
                    }},
                )
                repairs.append({
                    "category": "Schedule",
                    "title": "Linked schedule entry to Gantt section",
                    "detail": f"Schedule entry {entry.get('id')} linked to section '{section.get('name', 'Section')}'.",
                    "record_type": "schedule_entry",
                    "record_id": entry.get("id", ""),
                })

    # 2) Refresh section schedule IDs and RAG/schedule status from actual active entries.
    if request.repair_section_schedule_status and sections:
        refreshed_sections = []
        sections_changed = 0
        for section in sections:
            section_id = str(section.get("id") or "")
            linked_entries = schedule_by_section_id.get(section_id, [])
            linked_ids = [entry.get("id") for entry in linked_entries if entry.get("id")]
            old_ids = section.get("schedule_entry_ids") or []
            old_status = section.get("schedule_status") or "not_scheduled"
            old_sent = bool(section.get("sent_to_schedule"))

            if linked_ids:
                new_status = "scheduled"
                new_sent = True
            else:
                new_status = "not_scheduled"
                new_sent = False

            if linked_ids != old_ids or new_status != old_status or new_sent != old_sent:
                section = {
                    **section,
                    "schedule_entry_ids": linked_ids,
                    "sent_to_schedule": new_sent,
                    "schedule_status": new_status,
                    "last_process_sync_repair_at": now.isoformat(),
                }
                sections_changed += 1
                repairs.append({
                    "category": "Gantt",
                    "title": "Refreshed section schedule status",
                    "detail": f"Section '{section.get('name', 'Section')}' now references {len(linked_ids)} active schedule entr{'y' if len(linked_ids) == 1 else 'ies'}.",
                    "record_type": "gantt_section",
                    "record_id": section_id,
                })
            refreshed_sections.append(section)

        if sections_changed:
            sections = refreshed_sections
            await db.jobs.update_one(
                {"id": job_id},
                {"$set": {"gantt_sections": refreshed_sections, "last_process_sync_repair_at": now.isoformat()}},
            )

    # 3) Backfill PO section links and payment/cashflow planning dates.
    purchase_orders = await db.purchase_orders.find({
        "job_id": job_id,
        "archived": {"$ne": True},
    }, {"_id": 0}).to_list(5000)

    supplier_ids = list({po.get("supplier_id") for po in purchase_orders if po.get("supplier_id")})
    suppliers = await db.suppliers.find({"id": {"$in": supplier_ids}}, {"_id": 0}).to_list(1000) if supplier_ids else []
    supplier_lookup = {supplier.get("id"): supplier for supplier in suppliers}

    for po in purchase_orders:
        status = str(po.get("status") or "draft").lower().replace(" ", "_")
        if status in {"cancelled", "canceled", "void", "archived", "rejected"}:
            continue

        po_update: Dict[str, Any] = {}

        if request.repair_po_section_links:
            linked_section = next((section for section in sections if _po_links_to_section(po, section)), None)
            if linked_section:
                section_id = linked_section.get("id", "")
                section_name = linked_section.get("name", "")
                if section_id and po.get("project_section_id") != section_id:
                    po_update["project_section_id"] = section_id
                if section_name and po.get("project_section_name") != section_name:
                    po_update["project_section_name"] = section_name

                changed_lines = []
                line_changed = False
                for line in po.get("lines") or []:
                    if not isinstance(line, dict):
                        changed_lines.append(line)
                        continue
                    new_line = dict(line)
                    line_names_match = _normalise_link_value(line.get("job_section_name") or line.get("project_section_name") or line.get("section_name")) == _normalise_link_value(section_name)
                    line_ids_missing = not (line.get("job_section_id") or line.get("project_section_id"))
                    if line_names_match or line_ids_missing:
                        if section_id and new_line.get("job_section_id") != section_id:
                            new_line["job_section_id"] = section_id
                            line_changed = True
                        if section_name and new_line.get("job_section_name") != section_name:
                            new_line["job_section_name"] = section_name
                            line_changed = True
                    changed_lines.append(new_line)
                if line_changed:
                    po_update["lines"] = changed_lines

        if request.repair_po_cashflow_dates:
            planning_patch = _normalise_po_planning_fields({**po, **po_update}, supplier_lookup.get(po.get("supplier_id"), {}))
            for key, value in planning_patch.items():
                if value is not None and str(po.get(key, "")) != str(value):
                    po_update[key] = value

        if po_update:
            po_update["updated_at"] = now
            po_update["last_process_sync_repair_at"] = now.isoformat()
            await db.purchase_orders.update_one({"id": po.get("id")}, {"$set": po_update})
            repairs.append({
                "category": "Purchase Orders",
                "title": "Repaired PO planning/link fields",
                "detail": f"{po.get('po_number') or 'PO'} updated fields: {', '.join(sorted(k for k in po_update.keys() if k not in ['updated_at', 'last_process_sync_repair_at']))}.",
                "record_type": "purchase_order",
                "record_id": po.get("id", ""),
            })

    health = await get_job_process_sync_health(job_id, admin=admin)
    return {
        "message": "Process sync safe repair complete",
        "job_id": job_id,
        "repairs_count": len(repairs),
        "repairs": repairs,
        "health": health,
    }

# Include the router in the main app after all routes have been registered.
# Important: FastAPI copies the APIRouter routes at include time, so this must stay at the end.
app.include_router(api_router)

# Run the application
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "server:app",
        host="0.0.0.0",
        port=PORT,
        reload=False,
        access_log=True
    )
